from __future__ import annotations

import re
import shutil
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "output" / "doc"
SOURCE = DOC_DIR / "校园物资智能管理系统设计与实现-正文压缩版.docx"
TARGET = DOC_DIR / "校园物资智能管理系统设计与实现-图表引用修复版.docx"
REPORT = DOC_DIR / "图表引用检查报告.md"

CAPTION_RE = re.compile(r"^(图|表)(\d+-\d+)\s+")
BODY_HEADING_RE = re.compile(r"^[1-7]\s")
STOP_TITLES = {"致  谢", "致谢", "参考文献", "附录 1 关键接口与测试补充说明", "附录"}


def body_paragraphs(doc: Document) -> list[tuple[int, Paragraph, str]]:
    items: list[tuple[int, Paragraph, str]] = []
    in_body = False
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name == "Heading 1":
            if text in STOP_TITLES:
                in_body = False
            elif BODY_HEADING_RE.match(text) or text == "结束语":
                in_body = True
        if in_body:
            items.append((index, paragraph, text))
    return items


def chapter_order_checks(captions: list[tuple[int, str, str, str]]) -> tuple[bool, list[str]]:
    groups: dict[tuple[str, int], list[int]] = defaultdict(list)
    issues: list[str] = []
    for _, kind, number, _ in captions:
        chapter_str, seq_str = number.split("-")
        groups[(kind, int(chapter_str))].append(int(seq_str))

    for (kind, chapter), seqs in sorted(groups.items()):
        seqs_sorted = sorted(seqs)
        expected = list(range(1, len(seqs_sorted) + 1))
        if seqs_sorted != expected or len(seqs_sorted) != len(set(seqs_sorted)):
            kind_name = "图" if kind == "图" else "表"
            issues.append(f"{kind_name}{chapter}章编号序列异常：实际为 {seqs_sorted}，期望为 {expected}")
    return (len(issues) == 0, issues)


def scan_document(doc: Document) -> dict[str, object]:
    items = body_paragraphs(doc)
    captions: list[tuple[int, str, str, str]] = []
    bad_phrases: list[str] = []
    unresolved: list[str] = []

    for index, _, text in items:
        if "如下图" in text or "如下表" in text:
            bad_phrases.append(f"[{index}] {text}")
        match = CAPTION_RE.match(text)
        if match:
            captions.append((index, match.group(1), match.group(2), text))

    original_refs: dict[str, str] = {}
    missing_refs: list[str] = []
    compliant_refs: list[str] = []
    body_texts = [(index, text) for index, _, text in items]

    for index, kind, number, _ in captions:
        if kind == "图":
            ref_pat = re.compile(rf"(如图{re.escape(number)}所示|见图{re.escape(number)}|图{re.escape(number)}所示)")
        else:
            ref_pat = re.compile(rf"(如表{re.escape(number)}所示|见表{re.escape(number)}|表{re.escape(number)}所示)")
        before = [(i, text) for i, text in body_texts if i < index]
        after = [(i, text) for i, text in body_texts if i > index]
        before_hit = next(((i, text) for i, text in reversed(before) if ref_pat.search(text)), None)
        after_hit = next(((i, text) for i, text in after if ref_pat.search(text)), None)
        label = f"{kind}{number}"
        if before_hit:
            original_refs[label] = "before"
            compliant_refs.append(label)
        elif after_hit:
            original_refs[label] = "after"
            missing_refs.append(label)
        else:
            original_refs[label] = "missing"
            missing_refs.append(label)

    numbering_ok, numbering_issues = chapter_order_checks(captions)

    return {
        "captions": captions,
        "figure_count": len([c for c in captions if c[1] == "图"]),
        "table_count": len([c for c in captions if c[1] == "表"]),
        "bad_phrases": bad_phrases,
        "missing_refs": missing_refs,
        "compliant_refs": compliant_refs,
        "numbering_ok": numbering_ok,
        "numbering_issues": numbering_issues,
        "unresolved": unresolved,
    }


def find_by_substring(doc: Document, needle: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text.strip():
            return paragraph
    raise RuntimeError(f"未找到锚点段落: {needle}")


def append_sentence(paragraph: Paragraph, sentence: str) -> None:
    current = paragraph.text.rstrip()
    if sentence in current:
        return
    paragraph.text = f"{current}{sentence}"


def apply_reference_fixes(doc: Document) -> list[str]:
    added: list[str] = []

    append_targets = [
        ("预警与统计需求包括低库存、库存积压、临期、过期和异常消耗预警", "具体功能需求见表3-1。", "表3-1"),
        ("可维护性体现在模块化组织和通用字段设计", "具体非功能需求见表3-2。", "表3-2"),
        ("申领审批流程从部门用户创建申领单开始", "调拨执行流程如图3-2所示。", "图3-2"),
        ("调拨执行流程强调跨仓协同", "预警处置流程如图3-3所示。", "图3-3"),
        ("系统采用 B/S 架构和前后端分离设计", "系统整体架构如图4-1所示。", "图4-1"),
        ("系统按实际代码目录和前端路由划分为认证授权", "系统功能结构如图4-2所示。", "图4-2"),
        ("预警、日志和通知支撑表包括 `warning_record`、`operation_log`、`login_log` 和 `notification`", "关键数据表及其作用见表5-1。", "表5-1"),
        ("本地演示由后端 8080 端口提供业务接口", "系统开发与运行环境如表6-1所示。", "表6-1"),
        ("库存模块由 `InventoryService` 实现", "库存查询页面展示如图6-3所示。", "图6-3"),
        ("预警扫描由 `WarningService.scan` 实现，包含低库存、积压、临期、过期和异常出库五类规则", "相关功能页面展示如图6-7所示。", "图6-7"),
        ("补货建议和移动平均预测由 SmartService 实现", "统计分析页面展示如图6-8所示。", "图6-8"),
        ("2026年4月26日重新执行 `mvn test`", "后端自动化测试执行情况如表7-1所示。", "表7-1"),
        ("前端执行 `npm run build` 后", "前端验证结果如表7-2所示。", "表7-2"),
        ("业务场景验证围绕申领审批、调拨执行、预警处理和认证续签展开", "典型业务场景验证说明如表7-3所示。", "表7-3"),
        ("tests/performance 下的 k6 脚本与基线记录采集于本地 screenshot profile 和隔离 H2 数据集", "本地 k6 基线记录见表7-4。", "表7-4"),
    ]

    for needle, sentence, label in append_targets:
        paragraph = find_by_substring(doc, needle)
        append_sentence(paragraph, sentence)
        added.append(label)

    insert_targets = [
        ("图3-1 申领审批闭环流程图", "申领审批闭环流程如图3-1所示。", "图3-1"),
        (
            "图5-1 用户角色与组织 E-R 图",
            "用户角色与组织关系如图5-1所示，库存与批次关系如图5-2所示，业务单据、预警与通知关系如图5-3所示。",
            ["图5-1", "图5-2", "图5-3"],
        ),
        (
            "图6-1 登录认证与令牌续签流程图",
            "登录认证与令牌续签流程如图6-1所示，系统登录界面如图6-2所示。",
            ["图6-1", "图6-2"],
        ),
        (
            "图6-4 部门用户申领界面",
            "部门用户申领界面如图6-4所示，调拨执行与候选仓排序流程如图6-5所示，调拨管理页面如图6-6所示。",
            ["图6-4", "图6-5", "图6-6"],
        ),
    ]

    for needle, text, labels in insert_targets:
        anchor = find_by_substring(doc, needle)
        anchor.insert_paragraph_before(text, style="Normal")
        if isinstance(labels, list):
            added.extend(labels)
        else:
            added.append(labels)

    # 去重并保留顺序
    seen: set[str] = set()
    ordered_added: list[str] = []
    for item in added:
        if item not in seen:
            seen.add(item)
            ordered_added.append(item)
    return ordered_added


def write_report(source_scan: dict[str, object], final_scan: dict[str, object], backup_path: Path, added_refs: list[str]) -> None:
    original_missing = source_scan["missing_refs"]
    original_compliant = source_scan["compliant_refs"]
    numbering_issues = source_scan["numbering_issues"]
    bad_phrases = source_scan["bad_phrases"]
    unresolved = final_scan["unresolved"]

    lines = [
        "# 图表引用检查报告",
        "",
        f"- 工作底稿：`{SOURCE.name}`",
        f"- 备份文件：`{backup_path.name}`",
        f"- 输出文件：`{TARGET.name}`",
        f"- 全文图数量：`{source_scan['figure_count']}`",
        f"- 全文表数量：`{source_scan['table_count']}`",
        "",
        "## 原本未被引用的图表",
        "",
    ]

    if original_missing:
        for label in original_missing:
            lines.append(f"- `{label}`")
    else:
        lines.append("- 无")

    lines.extend(
        [
            "",
            "## 已补充引用的图表",
            "",
        ]
    )

    if added_refs:
        for label in added_refs:
            lines.append(f"- `{label}`")
    else:
        lines.append("- 无新增")

    lines.extend(
        [
            "",
            "## 原本已合规引用的图表",
            "",
        ]
    )

    if original_compliant:
        for label in original_compliant:
            lines.append(f"- `{label}`")
    else:
        lines.append("- 无")

    lines.extend(
        [
            "",
            "## 编号检查",
            "",
        ]
    )

    if source_scan["numbering_ok"]:
        lines.append("- 未发现图号、表号重复、缺失或章节号不匹配问题。")
    else:
        for issue in numbering_issues:
            lines.append(f"- {issue}")

    lines.extend(
        [
            "",
            "## 不规范表达检查",
            "",
        ]
    )

    if bad_phrases:
        for item in bad_phrases:
            lines.append(f"- {item}")
    else:
        lines.append("- 未发现“如下图”“如下表”等不规范表达。")

    lines.extend(
        [
            "",
            "## 无法判断引用位置的图表",
            "",
        ]
    )

    if unresolved:
        for item in unresolved:
            lines.append(f"- {item}")
    else:
        lines.append("- 无。")

    REPORT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    backup_path = DOC_DIR / f"{SOURCE.stem}-图表引用修复前备份-{timestamp}{SOURCE.suffix}"
    shutil.copy2(SOURCE, backup_path)
    shutil.copy2(SOURCE, TARGET)

    source_scan = scan_document(Document(SOURCE))

    target_doc = Document(TARGET)
    added_refs = apply_reference_fixes(target_doc)
    target_doc.save(TARGET)

    final_scan = scan_document(Document(TARGET))
    if final_scan["missing_refs"]:
        raise RuntimeError(f"修复后仍有未前置引用的图表: {final_scan['missing_refs']}")
    write_report(source_scan, final_scan, backup_path, added_refs)

    print(f"[backup] {backup_path.name}")
    print(f"[output] {TARGET.name}")
    print(f"[figures] {source_scan['figure_count']}")
    print(f"[tables] {source_scan['table_count']}")
    print(f"[original-missing] {len(source_scan['missing_refs'])}")
    print(f"[added] {len(added_refs)}")
    print(f"[final-missing] {len(final_scan['missing_refs'])}")
    print(f"[numbering-ok] {final_scan['numbering_ok']}")
    print(f"[bad-phrases] {len(final_scan['bad_phrases'])}")


if __name__ == "__main__":
    main()
