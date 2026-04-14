from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DRAFT = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-初稿.docx"
WORKING_DRAFT = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿工作稿.docx"
BACKUP_DRAFT = ROOT / "Existing Thesis Draft" / f"校园物资智能管理系统设计与实现-定稿工作稿.docx.bak.{datetime.now():%Y%m%d-%H%M%S}"
REPORT_PATH = ROOT / "output" / "doc" / "校园物资智能管理系统设计与实现-改稿说明.md"
FIGURE_DIR = ROOT / "output" / "doc" / "figures"
SCHEMA_PATH = ROOT / "sql" / "schema.sql"

OLD_TITLE = "校园应急物资智能管理系统设计与实现"
NEW_TITLE = "校园物资智能管理系统设计与实现"
EN_TITLE = "Design and Implementation of Campus Material Intelligent Management System"
HEADER_TITLE = NEW_TITLE


def ensure_dirs() -> None:
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)


def copy_working_files() -> None:
    shutil.copy2(SOURCE_DRAFT, WORKING_DRAFT)
    shutil.copy2(WORKING_DRAFT, BACKUP_DRAFT)


def remove_paragraph(paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_outline_level(paragraph, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level))


def mark_page_break_before(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    page_break = ppr.find(qn("w:pageBreakBefore"))
    if page_break is None:
        page_break = OxmlElement("w:pageBreakBefore")
        ppr.append(page_break)
    page_break.set(qn("w:val"), "1")


def add_heading(doc: Document, text: str, style: str, level: int, new_page: bool = False):
    p = doc.add_paragraph(style=style)
    if new_page:
        mark_page_break_before(p)
    p.add_run(text)
    set_outline_level(p, level)
    return p


def add_body_paragraph(doc: Document, text: str, style: str = "正文章节内容", first_line_indent_cm: float = 0.74):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    return p


def add_body_block(doc: Document, block: str, style: str = "正文章节内容", first_line_indent_cm: float = 0.74) -> None:
    for part in re.split(r"\n\s*\n", block.strip()):
        text = "".join(line.strip() for line in part.splitlines()).strip()
        if text:
            add_body_paragraph(doc, text, style=style, first_line_indent_cm=first_line_indent_cm)


def add_caption(doc: Document, text: str, style: str):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    return p


def add_picture(doc: Document, path: Path, caption: str):
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(path), width=Inches(6.1))
    add_caption(doc, caption, "图片标题")


def add_table(doc: Document, caption: str, rows: List[List[str]], col_widths_cm: Iterable[float] | None = None):
    add_caption(doc, caption, "表格标题")
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr = table.rows[0].cells
    for idx, text in enumerate(rows[0]):
        hdr[idx].text = text
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.style = doc.styles["表格内容"]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row == table.rows[0] else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.size = Pt(10.5)
    if col_widths_cm:
        for row in table.rows:
            for width, cell in zip(col_widths_cm, row.cells):
                cell.width = Cm(width)
    return table


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_all_text_black(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = False
    for section in doc.sections:
        for paragraph in section.header.paragraphs + section.footer.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.underline = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.underline = False


def replace_title_everywhere(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if OLD_TITLE in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(OLD_TITLE, NEW_TITLE)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                clear_paragraph(paragraph)
                paragraph.add_run(HEADER_TITLE)


def find_paragraph_index(doc: Document, text: str, style: str | None = None, start: int = 0) -> int:
    for idx, para in enumerate(doc.paragraphs[start:], start):
        if para.text.strip() == text and (style is None or para.style.name == style):
            return idx
    raise ValueError(f"Paragraph not found: {text}")


def replace_abstracts(doc: Document) -> None:
    try:
        idx_cn_title = find_paragraph_index(doc, OLD_TITLE)
    except ValueError:
        idx_cn_title = find_paragraph_index(doc, NEW_TITLE)
    doc.paragraphs[idx_cn_title].text = NEW_TITLE

    idx_cn_abs = find_paragraph_index(doc, "摘    要", start=idx_cn_title)
    old_en_title = "Design and Implementation of an Intelligent Campus Emergency Supplies Management System"
    idx_en_title = find_paragraph_index(doc, old_en_title, start=idx_cn_abs)

    cn_abstract = (
        "围绕高校后勤与应急场景下物资台账分散、审批链路长、库存批次难以追踪和预警处置滞后的问题，"
        "本文基于当前项目仓库完成了校园物资智能管理系统的设计与实现。系统采用前后端分离方案，前端使用"
        "Vue 3、Pinia、Element Plus 和 ECharts 构建业务页面与统计视图，后端使用 Spring Boot 3、"
        "Spring Security、JWT 与 MyBatis-Plus 实现认证授权、业务编排和数据访问，底层数据存储采用 MySQL 8。"
        "围绕项目中已经落地的代码与 SQL 结构，论文重点梳理了用户、角色、部门、校区、仓库、库位、物资分类与物资档案等基础数据管理功能，"
        "以及入库、出库、申领审批、物资调拨、库存批次管理、预警处理、通知推送、日志审计和事件登记等核心业务流程。"
        "在实现机制上，系统通过 JWT 与 Refresh Token 组合完成无状态登录态管理，通过方法级权限控制约束不同角色的访问范围，"
        "通过批次到期日排序实现出库推荐，通过紧急等级控制申领快速审批，通过定时扫描生成库存不足、积压、临期、过期和异常领用预警，"
        "并基于历史出库数据生成补货建议与统计分析结果。测试部分以项目仓库中的自动化测试、H2 测试配置和实际构建结果为依据，"
        "记录了 2026 年 4 月 14 日后端 42 项测试通过、前端生产构建成功的验证事实。"
        "结果表明，该系统已经具备覆盖校园物资核心流程的实现能力，可为后续补充运行截图、完善论文排版和开展场景化验收提供稳定基础。"
    )
    cn_keywords = "关键词：校园物资管理；Spring Boot；Vue 3；库存预警；申领审批"

    en_abstract = (
        "To address the problems of scattered ledgers, long approval chains, weak batch traceability, and delayed warning handling in campus material management, "
        "this thesis completes the design and implementation of a campus material intelligent management system based strictly on the current project repository. "
        "The system adopts a front-end and back-end separated architecture. Vue 3, Pinia, Element Plus, and ECharts are used to build business pages and statistical views, "
        "while Spring Boot 3, Spring Security, JWT, and MyBatis-Plus are used to implement authentication, authorization, business orchestration, and data access on the server side, with MySQL 8 as the storage layer. "
        "Grounded in the codebase and SQL schema, the thesis focuses on master data management for users, roles, departments, campuses, warehouses, storage locations, categories, and material profiles, "
        "as well as key workflows such as stock-in, stock-out, application approval, transfer execution, batch management, warning handling, notification delivery, operation auditing, and event registration. "
        "At the implementation level, the system combines JWT and refresh tokens for stateless session management, uses method-level authorization to restrict access by role, applies expiry-date ordering for outbound recommendation, "
        "supports fast-track approval for urgent applications, generates low-stock, backlog, expiring, expired, and abnormal-usage warnings through scheduled scanning, and produces replenishment suggestions and analytical views from historical outbound data. "
        "The testing chapter is grounded in executable repository evidence, including automated tests, the H2 test configuration, and successful build results. "
        "On April 14, 2026, 42 backend tests passed and the frontend production build completed successfully. "
        "These results show that the current project already provides an implementable basis for core campus material workflows and offers a stable foundation for later runtime screenshots and final thesis polishing."
    )
    en_keywords = "Keywords: campus material management; Spring Boot; Vue 3; inventory warning; application approval"

    while doc.paragraphs[idx_cn_abs + 1].text.strip() != old_en_title:
        remove_paragraph(doc.paragraphs[idx_cn_abs + 1])
    idx_en_title = find_paragraph_index(doc, old_en_title, start=idx_cn_abs)
    anchor = doc.paragraphs[idx_en_title]
    p1 = anchor.insert_paragraph_before(cn_abstract, style="中文摘要内容")
    p1.paragraph_format.first_line_indent = Cm(0.74)
    p2 = anchor.insert_paragraph_before(cn_keywords, style="中文摘要内容")
    p2.paragraph_format.first_line_indent = Cm(0.74)

    anchor.text = EN_TITLE
    idx_en_abs = find_paragraph_index(doc, "ABSTRACT", start=idx_en_title)
    while doc.paragraphs[idx_en_abs + 1].text.strip() != "1  绪论":
        remove_paragraph(doc.paragraphs[idx_en_abs + 1])
    idx_body_start = find_paragraph_index(doc, "1  绪论", style="name_out_lvl_1", start=idx_en_abs)
    body_anchor = doc.paragraphs[idx_body_start]
    p3 = body_anchor.insert_paragraph_before(en_abstract, style="英文摘要内容")
    p4 = body_anchor.insert_paragraph_before(en_keywords, style="英文摘要内容")
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def delete_old_body(doc: Document) -> int:
    body_start = find_paragraph_index(doc, "1  绪论", style="name_out_lvl_1")
    while len(doc.paragraphs) > body_start:
        remove_paragraph(doc.paragraphs[body_start])
    return body_start


def choose_font(size: int):
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simhei.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_centered_text(draw: ImageDraw.ImageDraw, box: Tuple[int, int, int, int], text: str, font, fill=(0, 0, 0)):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=6, align="center")
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text(((x1 + x2 - w) / 2, (y1 + y2 - h) / 2), text, font=font, fill=fill, spacing=6, align="center")


def draw_box(draw: ImageDraw.ImageDraw, box: Tuple[int, int, int, int], title: str, lines: List[str], title_font, body_font):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, outline=(60, 60, 60), width=3, fill=(252, 252, 252))
    title_box = (x1 + 16, y1 + 12, x2 - 16, y1 + 70)
    draw_centered_text(draw, title_box, title, title_font)
    if lines:
        draw.multiline_text((x1 + 18, y1 + 84), "\n".join(lines), font=body_font, fill=(50, 50, 50), spacing=10)


def arrow(draw: ImageDraw.ImageDraw, start: Tuple[int, int], end: Tuple[int, int], label: str | None = None, font=None):
    draw.line([start, end], fill=(80, 80, 80), width=4)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) > abs(ey - sy):
        direction = 1 if ex >= sx else -1
        pts = [(ex, ey), (ex - 14 * direction, ey - 8), (ex - 14 * direction, ey + 8)]
    else:
        direction = 1 if ey >= sy else -1
        pts = [(ex, ey), (ex - 8, ey - 14 * direction), (ex + 8, ey - 14 * direction)]
    draw.polygon(pts, fill=(80, 80, 80))
    if label and font:
        mx = (sx + ex) / 2
        my = (sy + ey) / 2
        draw.text((mx + 8, my - 16), label, font=font, fill=(40, 40, 40))


def create_figures() -> Dict[str, Path]:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    title_font = choose_font(32)
    body_font = choose_font(24)
    small_font = choose_font(20)
    label_font = choose_font(18)

    figures: Dict[str, Path] = {}

    img = Image.new("RGB", (1800, 1200), "white")
    draw = ImageDraw.Draw(img)
    draw_box(draw, (120, 180, 520, 420), "前端表现层", ["Vue 3 + Pinia", "Element Plus", "Axios / Router", "ECharts"], title_font, body_font)
    draw_box(draw, (680, 120, 1120, 380), "接口与控制层", ["AuthController", "ApplyController", "TransferController", "AnalyticsController"], title_font, body_font)
    draw_box(draw, (680, 480, 1120, 820), "业务服务层", ["认证授权", "库存与批次", "申领审批", "调拨推荐 / 预警 / 统计"], title_font, body_font)
    draw_box(draw, (1280, 160, 1660, 460), "数据访问层", ["MyBatis-Plus Mapper", "JdbcTemplate 聚合查询", "统一响应与异常处理"], title_font, body_font)
    draw_box(draw, (1280, 620, 1660, 920), "数据存储层", ["MySQL 8", "schema.sql", "seed.sql", "H2 测试库"], title_font, body_font)
    arrow(draw, (520, 300), (680, 250), "HTTP / JSON", label_font)
    arrow(draw, (900, 380), (900, 480), "业务编排", label_font)
    arrow(draw, (1120, 650), (1280, 310), "数据读写", label_font)
    arrow(draw, (1470, 460), (1470, 620), "结构化存储", label_font)
    path = FIGURE_DIR / "fig_3_1_architecture.png"
    img.save(path)
    figures["fig_3_1"] = path

    img = Image.new("RGB", (1800, 1200), "white")
    draw = ImageDraw.Draw(img)
    center_box = (650, 450, 1150, 730)
    draw_box(draw, center_box, "校园物资智能管理系统", ["统一认证", "统一接口", "统一日志与通知"], title_font, body_font)
    module_boxes = [
        ((110, 120, 500, 310), "认证与权限", ["登录 / 刷新", "用户 / 角色 / 部门"]),
        ((650, 100, 1150, 290), "基础数据", ["校区 / 仓库 / 库位", "物资分类 / 档案 / 供应商"]),
        ((1300, 120, 1690, 310), "仓储库存", ["库存总量", "批次 / 入库 / 出库"]),
        ((110, 500, 500, 690), "申领审批", ["草稿 / 提交 / 审批", "紧急单快速通过"]),
        ((1300, 500, 1690, 690), "调拨管理", ["调拨单", "最短路径推荐"]),
        ((110, 880, 500, 1070), "预警智能", ["库存低于阈值", "补货建议 / 预测"]),
        ((650, 900, 1150, 1090), "统计分析", ["总览指标", "趋势 / 占比 / 排名"]),
        ((1300, 880, 1690, 1070), "系统工具", ["事件 / 通知", "登录日志 / 操作日志 / 配置"]),
    ]
    for box, title, lines in module_boxes:
        draw_box(draw, box, title, lines, choose_font(28), small_font)
        bx1, by1, bx2, by2 = box
        cx1, cy1, cx2, cy2 = center_box
        arrow(draw, ((bx1 + bx2) // 2, (by1 + by2) // 2), ((cx1 + cx2) // 2, (cy1 + cy2) // 2))
    path = FIGURE_DIR / "fig_3_2_modules.png"
    img.save(path)
    figures["fig_3_2"] = path

    def er_image(specs: List[Tuple[Tuple[int, int, int, int], str, List[str]]], arrows: List[Tuple[Tuple[int, int], Tuple[int, int], str]], name: str):
        canvas = Image.new("RGB", (1900, 1200), "white")
        d = ImageDraw.Draw(canvas)
        for box, box_title, lines in specs:
            draw_box(d, box, box_title, lines, choose_font(28), choose_font(21))
        for start, end, label in arrows:
            arrow(d, start, end, label, label_font)
        p = FIGURE_DIR / name
        canvas.save(p)
        return p

    figures["fig_3_3"] = er_image(
        [
            ((120, 150, 470, 420), "sys_dept", ["id", "dept_name", "parent_id", "deleted", "created_at"]),
            ((760, 140, 1130, 460), "sys_user", ["id", "username", "password", "real_name", "dept_id", "role_id", "status"]),
            ((1450, 150, 1770, 380), "sys_role", ["id", "role_code", "role_name", "description"]),
            ((1380, 640, 1790, 930), "auth_refresh_token", ["id", "user_id", "token_id", "token_hash", "expire_at", "revoked"]),
            ((690, 760, 1120, 1040), "login_log / operation_log", ["user_id / operator_id", "login_status / module", "operation / detail", "login_time / created_at"]),
        ],
        [
            ((470, 280), (760, 280), "1:N"),
            ((1130, 260), (1450, 260), "N:1"),
            ((950, 460), (1570, 640), "1:N"),
            ((930, 460), (900, 760), "1:N"),
        ],
        "fig_3_3_rbac_er.png",
    )
    figures["fig_3_4"] = er_image(
        [
            ((90, 120, 420, 360), "material_category", ["id", "category_name", "remark"]),
            ((520, 100, 940, 430), "material_info", ["id", "material_code", "material_name", "category_id", "safety_stock", "supplier"]),
            ((1060, 100, 1420, 390), "warehouse", ["id", "warehouse_name", "campus", "address", "manager"]),
            ((520, 620, 940, 910), "inventory", ["id", "material_id", "warehouse_id", "current_qty", "locked_qty", "uk_inventory_material_warehouse"]),
            ((1060, 610, 1500, 980), "inventory_batch", ["id", "material_id", "warehouse_id", "batch_no", "remain_qty", "expire_date", "idx_batch_expire"]),
            ((1560, 200, 1850, 520), "stock_in / stock_out", ["warehouse_id", "operator_id", "apply_order_id", "remark"]),
        ],
        [
            ((420, 230), (520, 230), "1:N"),
            ((760, 430), (760, 620), "1:N"),
            ((1240, 390), (1240, 610), "1:N"),
            ((940, 760), (1060, 760), "N:1"),
            ((1420, 260), (1560, 320), "1:N"),
            ((1420, 300), (1560, 430), "1:N"),
        ],
        "fig_3_4_inventory_er.png",
    )
    figures["fig_3_5"] = er_image(
        [
            ((120, 160, 480, 430), "apply_order", ["id", "dept_id", "applicant_id", "urgency_level", "status", "fast_track", "approver_id"]),
            ((120, 560, 500, 850), "apply_order_item", ["id", "apply_order_id", "material_id", "apply_qty", "actual_qty"]),
            ((640, 160, 1040, 440), "transfer_order", ["id", "from_warehouse_id", "to_warehouse_id", "status", "applicant_id", "approver_id"]),
            ((660, 560, 1040, 830), "transfer_order_item", ["id", "transfer_order_id", "material_id", "quantity"]),
            ((1180, 140, 1580, 450), "warning_record", ["id", "warning_type", "material_id", "warehouse_id", "handle_status", "handler_id"]),
            ((1180, 570, 1580, 900), "notification / event_record", ["target_user_id", "msg_type", "biz_type", "event_type", "event_level", "status"]),
        ],
        [
            ((300, 430), (300, 560), "1:N"),
            ((840, 440), (840, 560), "1:N"),
            ((480, 300), (640, 300), "并行单据"),
            ((1040, 280), (1180, 280), "触发预警"),
            ((1040, 700), (1180, 730), "消息 / 事件关联"),
            ((480, 720), (1180, 760), "业务结果写入"),
        ],
        "fig_3_5_business_er.png",
    )
    figures["fig_4_1"] = er_image(
        [
            ((100, 300, 330, 520), "部门用户", ["创建申领单", "填写原因、场景、物资明细"]),
            ((450, 300, 740, 520), "提交环节", ["DRAFT -> SUBMITTED", "urgency_level>=2 时 fast_track=1"]),
            ((870, 160, 1190, 420), "审批人", ["approve / reject", "记录 approve_time 与 remark"]),
            ((870, 520, 1190, 800), "仓库管理员", ["按批次出库", "回写 actual_qty", "状态改为 OUTBOUND"]),
            ((1350, 300, 1700, 520), "申请人签收", ["receive", "状态改为 RECEIVED", "流程闭环"]),
        ],
        [
            ((330, 410), (450, 410), "发起"),
            ((740, 350), (870, 280), "普通单审批"),
            ((740, 470), (870, 650), "审批通过后出库"),
            ((1190, 650), (1350, 410), "出库完成"),
            ((560, 300), (560, 180), "紧急单直接 APPROVED"),
            ((560, 180), (1030, 180), "快速审批分支"),
        ],
        "fig_4_1_apply_flow.png",
    )
    figures["fig_4_2"] = er_image(
        [
            ((110, 300, 380, 520), "仓库管理员", ["创建调拨单", "选择调入仓和物资数量"]),
            ((500, 300, 820, 520), "推荐与提交", ["/api/transfer/recommend", "Dijkstra 计算目标校区最短距离", "DRAFT -> SUBMITTED"]),
            ((950, 160, 1270, 420), "审批人", ["approve / reject", "状态变为 APPROVED 或 REJECTED"]),
            ((950, 520, 1320, 860), "执行调拨", ["扣减调出库存", "复制批次到调入仓", "状态改为 OUTBOUND"]),
            ((1440, 300, 1760, 520), "签收完成", ["receive", "状态改为 RECEIVED"]),
        ],
        [
            ((380, 410), (500, 410), "推荐 / 提交"),
            ((820, 350), (950, 280), "审批"),
            ((820, 470), (950, 660), "审批后执行"),
            ((1320, 660), (1440, 410), "签收"),
        ],
        "fig_4_2_transfer_flow.png",
    )
    return figures


def parse_schema() -> Dict[str, List[Tuple[str, str]]]:
    text = SCHEMA_PATH.read_text(encoding="utf-8")
    result: Dict[str, List[Tuple[str, str]]] = {}
    for name, body in re.findall(r"CREATE TABLE\s+(\w+)\s*\((.*?)\n\)\s*;", text, flags=re.S):
        fields: List[Tuple[str, str]] = []
        for raw in body.splitlines():
            line = raw.strip().rstrip(",")
            if not line or line.startswith(("--", "PRIMARY KEY", "INDEX", "UNIQUE KEY", "KEY", "CONSTRAINT", "FOREIGN KEY")):
                continue
            parts = line.split()
            fields.append((parts[0], parts[1]))
        result[name] = fields
    return result


FIELD_EXPLANATIONS = {
    "id": "主键标识",
    "dept_name": "部门名称",
    "parent_id": "上级部门编号",
    "role_code": "角色编码",
    "role_name": "角色名称",
    "description": "说明信息",
    "username": "登录账号",
    "password": "BCrypt 密码串",
    "real_name": "姓名",
    "dept_id": "关联部门",
    "role_id": "关联角色",
    "status": "当前状态标记",
    "user_id": "关联用户",
    "token_id": "Refresh Token 标识",
    "token_hash": "Token 哈希值",
    "expire_at": "过期时间",
    "revoked": "是否撤销",
    "category_name": "分类名称",
    "remark": "备注信息",
    "material_code": "物资编码",
    "material_name": "物资名称",
    "category_id": "所属分类",
    "spec": "规格型号",
    "unit": "计量单位",
    "safety_stock": "安全库存阈值",
    "shelf_life_days": "保质期天数",
    "supplier": "供应商名称",
    "unit_price": "单价",
    "warehouse_name": "仓库名称",
    "campus": "所属校区",
    "address": "地址",
    "manager": "负责人",
    "material_id": "关联物资",
    "warehouse_id": "关联仓库",
    "current_qty": "当前库存数量",
    "locked_qty": "锁定数量",
    "batch_no": "批次号",
    "in_qty": "入库数量",
    "remain_qty": "剩余数量",
    "production_date": "生产日期",
    "expire_date": "失效日期",
    "stock_in_id": "关联入库单",
    "quantity": "业务数量",
    "source_type": "来源类型",
    "operator_id": "操作人",
    "apply_order_id": "关联申领单",
    "applicant_id": "申请人",
    "urgency_level": "紧急等级",
    "reason": "业务原因",
    "scenario": "使用场景",
    "fast_track": "快速审批标记",
    "approver_id": "审批人",
    "approve_remark": "审批意见",
    "approve_time": "审批时间",
    "stock_out_id": "关联出库单",
    "from_warehouse_id": "调出仓库",
    "to_warehouse_id": "调入仓库",
    "warning_type": "预警类型",
    "content": "预警或通知内容",
    "handle_status": "处理状态",
    "handler_id": "处理人",
    "handle_remark": "处理说明",
    "module": "业务模块",
    "operation": "操作类型",
    "detail": "明细内容",
    "campus_name": "校区名称",
    "contact_phone": "联系电话",
    "location_code": "库位编码",
    "location_name": "库位名称",
    "capacity": "设计容量",
    "used_qty": "已使用数量",
    "supplier_name": "供应商名称",
    "contact_person": "联系人",
    "email": "邮箱",
    "supply_scope": "供货范围",
    "event_title": "事件标题",
    "event_type": "事件类型",
    "event_level": "事件等级",
    "campus_id": "关联校区",
    "location": "事件位置",
    "reporter_id": "上报人",
    "handle_result": "处理结果",
    "event_time": "事件发生时间",
    "close_time": "事件关闭时间",
    "config_key": "配置键",
    "config_value": "配置值",
    "config_name": "配置名称",
    "config_group": "配置分组",
    "login_ip": "登录 IP",
    "login_status": "登录状态",
    "login_time": "登录时间",
    "user_agent": "客户端标识",
    "title": "标题",
    "msg_type": "消息类型",
    "target_user_id": "目标用户",
    "is_read": "是否已读",
    "biz_type": "业务类型",
    "biz_id": "业务编号",
    "deleted": "逻辑删除标记",
    "version": "版本号",
    "created_at": "创建时间",
    "updated_at": "更新时间",
}


def build_table_rows(schema: Dict[str, List[Tuple[str, str]]], table_name: str, keep_fields: List[str] | None = None) -> List[List[str]]:
    rows = [["字段名", "数据类型", "说明"]]
    fields = schema[table_name]
    for name, field_type in fields:
        if keep_fields and name not in keep_fields:
            continue
        rows.append([name, field_type, FIELD_EXPLANATIONS.get(name, "字段含义见表名语境")])
    return rows


def add_reference_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="参考文献内容")
    p.add_run(text)


def estimate_body_chars(doc: Document) -> int:
    started = False
    total = 0
    for paragraph in doc.paragraphs:
        text = re.sub(r"\s+", "", paragraph.text)
        if paragraph.text.strip() == "1  绪论":
            started = True
        if started:
            total += len(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += len(re.sub(r"\s+", "", cell.text))
    return total


def build_document() -> Tuple[Path, int]:
    ensure_dirs()
    copy_working_files()
    figures = create_figures()
    schema = parse_schema()

    doc = Document(WORKING_DRAFT)
    replace_title_everywhere(doc)
    replace_abstracts(doc)
    delete_old_body(doc)

    add_heading(doc, "1  绪论", "name_out_lvl_1", 0, new_page=True)

    add_heading(doc, "1.1 研究背景", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        高校物资管理同时面对日常保障和突发事件两类场景：一方面，实验耗材、医疗防护、应急照明、清洁消杀等物资需要围绕部门申领、仓库出入库和批次追踪持续运转；另一方面，物资配置又必须服务于校园安全、公共卫生和大型活动保障等时效性任务。随着智慧校园建设逐步由网络基础设施扩展到业务协同层，传统依赖纸质台账、分散 Excel 表和人工电话确认的管理方式越来越难以支撑跨部门、跨校区的物资流转要求[1][2]。

        从本项目代码与 SQL 结构可以看到，系统需要同时管理部门、角色、用户、校区、仓库、库位、分类、物资档案、库存、库存批次、申领单、调拨单、预警记录、通知消息和日志审计等多类对象。这些对象之间并不是孤立存在的：部门用户发起申领后需要经过审批人审核，审核通过后还要由仓库管理员结合库存与批次情况执行出库，最终再由使用方完成签收闭环；在调拨场景中，还存在调出仓与调入仓之间的库存协同、运输路径选择和状态流转问题。

        当前仓库已经提供完整的前后端实现与测试代码，因此本科论文定稿不应停留在概念性叙述，而应回到真实仓库中的业务边界和实现证据。基于这一约束，本文把“校园物资智能管理系统设计与实现”界定为一个已经完成核心流程开发、具备自动化测试与前端构建验证基础的 B/S 架构项目，并围绕代码、SQL、配置和文档可以直接证明的事实展开分析与写作。
        """,
    )

    add_heading(doc, "1.2 研究意义", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        本课题的工程意义首先体现在业务流程的规范化上。仓库中的 `apply_order`、`stock_out`、`transfer_order`、`warning_record` 等表，以及 `ApplyService`、`InventoryService`、`TransferService`、`WarningService` 等服务类，说明系统已经把“申请—审批—执行—签收—追踪”的链路拆解为明确的状态机与数据对象。论文对这些对象进行系统化梳理，有助于把原本分散在接口、服务和数据库中的实现逻辑归纳为可复核的业务模型，从而为后续交付、答辩和维护提供统一描述。

        本课题的应用意义体现在角色协同与责任留痕两方面。后端安全配置和菜单构建逻辑表明，系统围绕 `ADMIN`、`WAREHOUSE_ADMIN`、`APPROVER`、`DEPT_USER` 四类角色设计了访问边界，并在 `login_log` 与 `operation_log` 中记录关键操作行为。相比只关注单表增删改查的演示型系统，这种角色驱动和日志留痕的实现更贴合校园场景中的责任追踪要求，也与基于角色进行权限治理的高校信息系统建设经验相一致[3]。

        对本科毕业论文而言，研究意义还体现在“证据约束写作”本身。现有草稿中存在旧题名、错误接口名和无法证实的性能结论，如果继续沿用，会削弱论文与实际项目之间的对应关系。将定稿严格绑定到仓库实现，一方面可以降低事实性错误，另一方面也能让论文从“功能想象”转向“实现复盘”，使毕业设计成果与代码资产形成更稳定的互证关系。
        """,
    )

    add_heading(doc, "1.3 国内外研究现状", "2级标题-正文章节", 1)

    add_heading(doc, "1.3.1 国内研究现状", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        国内已有较多围绕智慧校园、后勤保障和高校业务信息化的研究。相关成果通常从数字校园总体建设、校园资源协同和业务流程线上化切入，强调以统一身份认证、数据共享和流程流转替代传统的线下审批与分散台账[1]。在具体应用层面，高校教务管理、项目管理、志愿服务管理、事故报告管理等系统已广泛采用 Spring Boot 与 Vue 的前后端分离架构，以提升页面响应、角色隔离和可维护性[6][7][8][9]。

        与本项目更接近的一类研究聚焦于物资、图书、采购等资源的台账管理与智能推荐。例如，图书馆信息管理与采购信息管理系统通常会围绕主数据维护、申请审批、库存统计和通知提醒设计业务闭环[4][5][10]。这类系统为校园物资管理的表结构设计、角色建模和业务流程划分提供了可借鉴的思路，但多数论文更重视功能演示，对批次追踪、库存预警和复杂单据状态切换的讨论相对有限。

        另一些研究开始尝试将预测分析、预警管理与业务系统结合起来，利用历史消耗数据支持库存决策与补货建议[15]。不过，这类研究常见的问题是算法表述与工程落地之间存在落差：论文中给出较多概念性描述，但在代码、任务调度、表结构和接口层面的实现细节不足。本项目在这方面采取了更保守但可落地的路线，即使用历史出库数据的平均值生成补货建议，并通过定时任务完成预警扫描，从而避免算法叙述脱离仓库实现。
        """,
    )

    add_heading(doc, "1.3.2 国外相关实践启示", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        国外关于物资管理与 Web 工程的公开资料，更多体现为框架文档、标准规范和工程实践指南，而不是直接面向本科论文的成型案例。以 Spring Boot、Vue、Vite 和 JWT 标准文档为代表的技术资料，重点强调约定优于配置、组件化组织、前端工程化构建和无状态认证机制，这些内容为本项目的技术选型与实现边界提供了可直接参照的依据[17][18][19][20]。

        从方法论角度看，国外工程实践更强调“最小可验证实现”，即先用一套清晰、可测试、可扩展的主链路打通关键业务，再根据运行反馈补充高级能力。本项目在校园物资管理场景中的做法与此相近：没有引入无法验证的复杂平台依赖，而是围绕登录鉴权、申领审批、库存批次、调拨流转、预警生成、统计展示等核心流程建立可执行系统，再通过测试和构建结果证明当前版本的可用边界。

        因此，本文在“研究现状”部分不把国外资料简单写成概念综述，而是把它们视为工程实现的外部参考。对于毕业论文而言，这种写法更符合当前项目的真实状态：系统并未宣称已经形成跨校园推广平台，而是在成熟技术栈与通用工程规范支持下，完成了一个面向校园物资管理核心业务的可运行实现。
        """,
    )

    add_heading(doc, "1.4 研究内容与论文结构", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        本文的研究对象是当前仓库中的校园物资智能管理系统。围绕真实项目材料，论文首先从角色、业务场景和功能边界出发分析系统需求，然后结合前后端依赖、权限配置、统一响应格式与数据库结构说明总体设计方案，接着对认证授权、基础数据、库存批次、申领审批、调拨执行、预警分析、日志通知等模块的关键实现进行归纳，最后基于测试代码、H2 配置与构建结果总结系统验证情况。

        论文结构安排如下：第 1 章说明选题背景、研究意义和相关现状；第 2 章结合数据库与接口权限分析系统需求；第 3 章给出技术选型、总体架构、功能模块、数据库和接口安全设计；第 4 章围绕真实代码实现详细说明核心模块；第 5 章以自动化测试和构建验证为中心说明系统测试情况；最后在结束语中总结本文工作并提出后续可扩展方向。

        需要强调的是，全文只保留代码、SQL、配置与文档可以直接证明的内容。凡是未在仓库中落地的生产部署规模、性能指标、运行效果和外部集成能力，本文均不作为结论写入。这一原则既是对现有毕业设计成果的真实复盘，也是后续定稿继续补充截图、优化引用和完成格式审校的基础。
        """,
    )

    add_heading(doc, "2  系统需求分析", "name_out_lvl_1", 0, new_page=True)

    add_heading(doc, "2.1 业务场景与角色划分", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        从 `sql/seed.sql`、认证服务菜单装配逻辑和前端路由可以确认，系统当前面向四类核心角色：系统管理员、仓库管理员、审批人和部门用户。系统管理员负责用户、角色、部门、校区、仓库、库位、分类、物资、供应商、事件、系统配置等基础数据维护；仓库管理员负责入库、出库、库存维护、批次管理和调拨执行；审批人负责申领单与调拨单的审核；部门用户则主要承担物资申领、签收和通知查看等使用侧操作。

        业务场景可以概括为三条主线。第一条主线是面向日常保障的库存管理，包括物资建档、采购入库、批次登记、安全库存监控和临期清理；第二条主线是面向使用部门的申领审批闭环，包括申领、审核、拣货出库、签收和日志留痕；第三条主线是面向跨仓协同的调拨管理，包括仓间库存平衡、调拨推荐、审批、执行与签收。除此之外，预警、通知、统计分析和事件登记构成了对主业务的支撑能力。

        角色边界在代码中通过两层机制落实：一层是 `SecurityConfig` 提供的认证过滤链和方法级安全开关，另一层是各控制器上的 `@PreAuthorize` 表达式与 `AuthService.buildMenusByRole` 返回的动态菜单。由此可见，本系统并不是所有用户共享同一页面集合，而是基于角色编码把菜单、接口和业务操作限定在可追踪的访问范围内。
        """,
    )

    add_heading(doc, "2.2 功能需求分析", "2级标题-正文章节", 1)

    add_heading(doc, "2.2.1 基础数据管理需求", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        基础数据是所有业务单据的前提。根据数据库结构，系统至少需要支持部门、角色、用户、校区、仓库、库位、物资分类、物资档案、供应商和系统配置等对象的维护。前端路由中对应存在 `/system/user`、`/system/role`、`/system/dept`、`/campus/list`、`/warehouse/list`、`/location/list`、`/material/list`、`/supplier/list`、`/config/list` 等页面，说明这些对象并不是数据库中的静态字典，而是要求在管理端可见、可维护、可关联的业务数据。

        其中，物资档案管理对后续流程影响最大。`material_info` 表记录物资编码、名称、分类、规格、单位、安全库存、保质期、供应商和单价等字段，这些字段会在库存统计、预警触发、申领单明细和补货建议中被多次复用。因此系统需要保证主数据的唯一性、一致性和可追溯性，例如以物资编码作为唯一键，以逻辑删除字段控制历史数据保留。
        """,
    )

    add_heading(doc, "2.2.2 库存与批次管理需求", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        校园物资管理并不只关心总量，还必须关注批次来源和到期风险。`inventory` 表负责记录某物资在某仓库下的总库存和锁定数量，`inventory_batch` 则记录批次号、入库数量、剩余数量、生产日期和到期日期，这说明系统在需求层面必须同时支持总账视角和批次视角。对于食品、防疫或实验类物资而言，仅有总库存而没有批次信息，无法满足临期处置和先进先出的管理要求。

        对应到业务行为，系统需要支持入库登记、批次拆分、库存查询、批次排序和出库扣减。`InventoryService` 中围绕库存增减、批次创建和按到期日排序的实现，进一步证明系统需要在出库前自动给出可用批次建议，并在库存不足或批次剩余不足时给出异常提示，而不是让仓库管理员完全依赖人工判断。
        """,
    )

    add_heading(doc, "2.2.3 申领审批与出库需求", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        部门申领是本系统最核心的业务入口。`apply_order` 与 `apply_order_item` 表共同描述一张申领单的申请部门、申请人、紧急等级、业务原因、使用场景、快速审批标记、审批信息和物资明细，表明系统需要支持“一张主单 + 多条明细”的单据模式。状态流转必须覆盖 `DRAFT`、`SUBMITTED`、`APPROVED`、`REJECTED`、`OUTBOUND`、`RECEIVED` 六个阶段，以便把草稿、待审、已审、已出库和已签收区分开来。

        代码还反映出一个较有代表性的校园场景：紧急申领。`ApplyService` 会根据 `urgency_level` 判断是否触发快速审批，因此需求分析中需要明确，系统不仅服务于普通计划性申领，也需要为突发事务、夜间保障或临时活动提供更短的审批路径。对应地，审批人既要支持常规审核，也要能够为高紧急度单据留下审批意见和时间记录，保证流程在缩短的同时仍可追责。
        """,
    )

    add_heading(doc, "2.2.4 调拨协同需求", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        当某一仓库库存不足而另一仓库存在余量时，系统需要通过调拨机制完成跨仓协同。`transfer_order` 与 `transfer_order_item` 的存在表明，调拨并不是简单地修改两端库存数字，而是需要先形成独立单据，再经历提交、审批、执行和签收等阶段。由于调拨同时影响调出仓和调入仓，系统必须保证库存扣减、批次复制和状态变更在业务语义上保持一致。

        与普通出入库相比，调拨场景还引入了空间因素。项目中的 `DijkstraUtil` 与 `/api/transfer/recommend` 接口说明，系统需要根据预设的校区图给出候选仓库和路径距离，帮助仓库管理员选择更合理的调拨来源。因此，本系统的调拨需求不仅包括单据流转本身，还包括对“从哪一个仓库调”这一决策问题的辅助支持。
        """,
    )

    add_heading(doc, "2.2.5 预警、统计与支撑功能需求", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        仅能完成单据流转还不足以支撑校园物资管理的持续运行，系统还需要具备运行支撑能力。`warning_record`、`notification`、`event_record`、`login_log` 和 `operation_log` 等表说明，项目已经把预警、消息、事件和日志作为独立模块纳入设计。相应地，需求层面需要覆盖库存不足、库存积压、临近到期、已过期和异常领用等预警类型，需要支持通知发送与已读状态维护，需要支持事件上报、处理和归档，也需要支持登录日志与操作日志查询。

        统计分析需求主要体现在管理者视角。前端存在图表页与大屏页，后端 `AnalyticsService` 使用聚合 SQL 统计近月出入库和物资分布，`SmartService` 则根据历史出库记录生成补货建议。由此可见，系统不仅要回答“当前有哪些单据和库存”，还要回答“近期哪些物资消耗较快、哪些仓库存在预警、哪些物资需要补货”等管理问题。
        """,
    )

    add_heading(doc, "2.3 非功能需求分析", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        非功能需求首先体现在安全性和权限隔离上。后端使用 Spring Security 过滤链、JWT、Refresh Token 和方法级权限控制保护接口访问，前端通过本地存储 token、路由守卫和 401 自动处理维持登录态。这说明系统必须满足基本的身份认证、角色授权和会话续期要求，确保不同角色只能访问自身职责范围内的页面与业务操作。

        其次，系统需要具备可维护性与可扩展性。仓库采用前后端分离架构，后端服务按模块组织为认证、库存、申领、调拨、预警、智能分析、日志等子模块，数据库表统一包含逻辑删除、版本号和时间戳字段，前端路由也按功能域划分视图目录。这些实现共同表明，项目在需求层面追求的是“便于持续迭代”的结构，而不是一次性拼装的演示程序。

        再次，系统需要满足可测试性要求。当前仓库提供 H2 内存数据库测试配置，并围绕认证、申请、调拨、预警、统一异常处理等模块编写了自动化测试类。这意味着系统的非功能目标之一是支持在相对独立的测试环境中验证关键业务逻辑，而不是完全依赖人工点击页面进行回归。至于并发量、响应时延和生产部署规模，当前仓库没有对应的实测报告，因此本文不将其写入结论。
        """,
    )

    add_heading(doc, "2.4 业务流程与用例分析", "2级标题-正文章节", 1)

    add_heading(doc, "2.4.1 申领审批用例分析", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        在申领审批用例中，主要参与者包括部门用户、审批人和仓库管理员。部门用户首先创建申领单，录入申请原因、使用场景、紧急等级和物资明细；草稿单在提交前可以修改，提交后状态变为 `SUBMITTED`。审批人对单据进行审核，若不符合要求则退回为 `REJECTED` 并记录意见；若审核通过，则状态进入 `APPROVED`，仓库管理员据此执行出库。

        出库环节并非单纯的状态修改，而是要结合库存与批次完成实物分配。系统根据库存和批次记录生成出库单与出库明细，回写申领明细的 `actual_qty`，并把申领单状态更新为 `OUTBOUND`。最终，部门用户对实际领用结果进行签收，状态更新为 `RECEIVED`，一张申领单才算闭环结束。
        """,
    )

    add_picture(doc, figures["fig_4_1"], "图4-1 申领审批闭环流程图")

    add_heading(doc, "2.4.2 调拨执行用例分析", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        在调拨执行用例中，仓库管理员根据目标仓库存量或保障需求创建调拨单，并通过推荐接口选择合适的来源仓库。调拨单提交后由审批人审核，审核通过后再进入执行阶段。执行时系统需要同步扣减调出仓库存、复制或生成调入仓批次记录，并形成状态从 `APPROVED` 到 `OUTBOUND` 的推进。

        调拨业务的特殊性在于其结果同时影响两个仓库，因此对业务一致性要求更高。与普通申领不同，调拨完成后还需要目标仓确认签收，系统才将状态推进到 `RECEIVED`。通过这一设计，系统把“仓间物资转移”抽象为具备申请、审核、执行、签收四段式特征的完整业务链路。
        """,
    )

    add_heading(doc, "2.4.3 预警处置与统计分析用例分析", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        预警处置用例由系统定时任务主动触发。`WarningService.scan()` 会定期检查当前库存与批次信息，生成库存不足、库存积压、临期、过期和异常领用等预警记录。管理员或仓库管理员在预警页查看未处理记录后，可以录入处理说明并更新处理状态；对于需要通知相关人员的事项，系统还可写入通知表供对应用户查看。

        统计分析用例则主要服务于管理决策。系统根据出入库流水、库存分布和时间维度聚合数据，在图表页与大屏页展示关键统计结果，并在智能分析模块中生成补货建议。由于这些分析结果直接来源于业务表数据，其前提是主数据、库存和单据记录保持完整，因此统计分析既是独立功能，也是对前述业务流程规范性的反向检验。
        """,
    )

    add_heading(doc, "3  系统总体设计", "name_out_lvl_1", 0, new_page=True)

    add_heading(doc, "3.1 技术选型与总体架构", "2级标题-正文章节", 1)

    add_heading(doc, "3.1.1 技术选型依据", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        结合仓库依赖文件可以确认，系统前端采用 Vue 3、Pinia、Element Plus、Axios、Vue Router 和 ECharts，使用 Vite 作为构建工具；后端采用 Spring Boot 3.3.5、Spring Security、MyBatis-Plus、JdbcTemplate 和 JWT；数据层使用 MySQL 8，并在测试阶段引入 H2 内存数据库。上述技术组合对应的是当前高校信息化项目中常见的前后端分离路线：前端负责页面交互与图表展示，后端负责认证授权、业务编排和数据持久化[10][11][12][13][14][16][17][18][19][20]。

        选择 Vue 3 与 Pinia 的原因，在于项目存在较多角色菜单、表单状态和图表页面，需要较清晰的组件组织与轻量状态管理。结合 `package.json` 依赖范围与本次 `npm ls` 结果，可以确认当前前端实际安装的核心版本包括 Vue 3.5.29、Pinia 2.3.1、Vite 6.4.1、Element Plus 2.13.5 和 ECharts 5.6.0，这意味着前端工程已经使用组合式 API 与现代化构建链路。Element Plus 承担表单、表格、对话框、消息提示等通用界面元素，ECharts 提供库存与出入库统计图表的呈现能力[12][14][18][19]。

        后端选择 Spring Boot 3 的主要原因，是其在 Web API、依赖管理、配置组织和测试支持方面较为成熟。Spring Security 用于认证与权限控制，MyBatis-Plus 负责实体映射和通用 CRUD，JdbcTemplate 用于统计与智能分析中的聚合 SQL。对于本科项目而言，这种组合既能覆盖典型企业级开发流程，又不会引入当前仓库无法证明的复杂中间件依赖，因此符合“功能完整、实现清晰、可测试”的设计目标[10][11][16][17][20]。
        """,
    )

    add_heading(doc, "3.1.2 系统总体架构", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        从系统分层看，项目采用典型的 B/S 架构。浏览器作为统一访问入口，用户通过登录页进入系统后，根据角色获得不同菜单和页面；前端通过 HTTP/JSON 与后端接口交互；后端再按照控制层、服务层和数据访问层完成业务处理，并把结果写入 MySQL 数据库。对毕业论文而言，这种架构具有描述清晰、部署简洁和便于截图展示的优点。

        在后端内部，控制层主要负责参数接收和权限约束，服务层负责状态流转、库存计算、批次分配和日志记录，数据访问层由 MyBatis-Plus Mapper 与 JdbcTemplate 共同构成。控制层与服务层之间通过统一响应对象 `ApiResponse{code,message,data}` 进行对外封装，成功返回码固定为 `0`。这种统一返回机制使前端拦截器能够集中处理错误消息，减少页面端的重复判断。

        在数据层面，系统并未采用高度拆分的微服务结构，而是使用单体后端配合模块化代码组织。结合当前仓库规模，这种方案更利于业务联调和测试验证：认证、库存、申请、调拨、预警和统计虽然属于不同功能域，但都围绕同一套用户、角色、仓库和物资主数据展开，因此在单体架构中更容易维持事务语义和数据一致性。
        """,
    )
    add_picture(doc, figures["fig_3_1"], "图3-1 系统总体架构图")

    add_heading(doc, "3.2 功能模块设计", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        根据前端路由、角色菜单和后端控制器，系统可以划分为八个相互关联的功能模块：认证与权限模块、基础数据模块、仓储库存模块、申领审批模块、调拨协同模块、预警与智能分析模块、统计分析模块、系统支撑模块。其中系统支撑模块又包括通知、事件、登录日志、操作日志和系统配置等功能，用于保障主业务的持续运行。

        认证与权限模块负责用户登录、登出、刷新令牌、加载当前用户信息和动态菜单；基础数据模块负责用户、角色、部门、校区、仓库、库位、供应商、分类和物资档案维护；仓储库存模块负责入库、库存台账、批次追踪和出库执行；申领审批模块负责部门侧申领单创建、提交、审批和签收；调拨协同模块负责跨仓调拨推荐、审批、执行和签收。这样的划分既贴合代码目录，也符合校园物资业务从“主数据准备”到“单据闭环执行”的实际顺序。

        预警与智能分析模块、统计分析模块和系统支撑模块则体现了项目的管理深度。预警与智能分析模块提供定时预警、需求预测和补货建议；统计分析模块基于业务数据生成图表和大屏视图；系统支撑模块通过消息、日志与事件记录提升可追溯性。整体来看，系统功能并非停留在 CRUD，而是围绕校园物资“计划、执行、监督、分析”的完整闭环展开。
        """,
    )
    add_picture(doc, figures["fig_3_2"], "图3-2 系统功能模块图")

    add_heading(doc, "3.3 数据库设计", "2级标题-正文章节", 1)

    add_heading(doc, "3.3.1 数据库设计原则", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        数据库设计以“主数据统一、业务单据分层、审计信息独立”为基本原则。主数据层负责保存组织结构、角色、用户、仓库、物资和供应商等静态或低频变化信息；业务单据层负责记录申领、出入库、调拨和预警；审计与支撑层负责保存日志、通知和事件等辅助信息。这样的分层方式可以避免把所有业务语义压缩到少数宽表中，便于后续统计查询和模块拆分。

        从字段设计上看，多数业务表都包含 `deleted`、`version`、`created_at` 和 `updated_at` 字段，体现了逻辑删除、乐观锁预留和时间追踪的统一思想。对于库存、批次、单据、预警和日志等需要长期保留历史记录的场景，逻辑删除比物理删除更适合审计与回溯；统一时间字段则为统计分析、预警扫描和测试数据初始化提供了时间维度基础。

        索引设计主要围绕业务查询入口展开。例如 `inventory` 表在物资与仓库上建立联合唯一约束，确保同一物资在同一仓库只有一条总库存记录；`inventory_batch` 表在到期日期上建立索引，便于按批次进行临期查询；申领单、调拨单、预警和日志等表也围绕状态、人员或时间字段建立索引，以支持列表页和统计查询的主要过滤条件。
        """,
    )

    add_heading(doc, "3.3.2 实体关系设计", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        从实体关系角度看，本系统可以拆分为三组核心子域。第一组是 RBAC 与组织子域，包括部门、角色、用户、刷新令牌和日志，它决定了谁可以访问系统、可以看到哪些菜单以及关键操作由谁执行。第二组是库存与批次子域，包括分类、物资档案、仓库、库存总表、库存批次和出入库记录，它决定了库存总量与批次明细如何被维护。第三组是业务单据与支撑子域，包括申领单、调拨单、预警记录、通知和事件，它负责把具体业务动作串联为闭环。

        这些子域之间存在清晰的主外键语义。部门用户与审批人都来自 `sys_user`，而 `sys_user` 又通过 `dept_id` 和 `role_id` 关联组织结构；库存和批次通过 `material_id` 与 `warehouse_id` 同时指向物资和仓库；申领单、调拨单和预警记录则分别通过用户、部门、仓库和物资标识与主数据连接。虽然当前 SQL 没有在所有表上显式声明外键约束，但从字段命名和服务调用看，实体关系是稳定且明确的。
        """,
    )
    add_picture(doc, figures["fig_3_3"], "图3-3 RBAC 与组织实体关系图")
    add_picture(doc, figures["fig_3_4"], "图3-4 库存与批次实体关系图")
    add_picture(doc, figures["fig_3_5"], "图3-5 业务单据、预警与通知实体关系图")

    add_heading(doc, "3.3.3 关键数据表设计", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        为了在论文中保留对关键数据结构的可读说明，本文选取用户、物资、仓库、库存、批次、申领、出库、调拨、预警、通知和日志等 12 张代表性数据表进行展示。这些表共同覆盖了系统最主要的主数据、业务单据与审计记录。表中的字段说明并不等同于完整 SQL 展开，而是对业务上最关键的字段进行摘要说明，便于和后续模块实现章节形成对应关系。
        """,
    )

    add_table(doc, "表3-1 用户信息表（sys_user）", build_table_rows(schema, "sys_user", ["id", "username", "password", "real_name", "dept_id", "role_id", "status", "created_at"]))
    add_table(doc, "表3-2 物资档案表（material_info）", build_table_rows(schema, "material_info", ["id", "material_code", "material_name", "category_id", "spec", "unit", "safety_stock", "supplier", "unit_price"]))
    add_table(doc, "表3-3 仓库信息表（warehouse）", build_table_rows(schema, "warehouse", ["id", "warehouse_name", "campus", "address", "manager", "created_at"]))
    add_table(doc, "表3-4 库存总表（inventory）", build_table_rows(schema, "inventory", ["id", "material_id", "warehouse_id", "current_qty", "locked_qty", "updated_at"]))
    add_table(doc, "表3-5 库存批次表（inventory_batch）", build_table_rows(schema, "inventory_batch", ["id", "material_id", "warehouse_id", "batch_no", "in_qty", "remain_qty", "production_date", "expire_date"]))
    add_table(doc, "表3-6 申领单主表（apply_order）", build_table_rows(schema, "apply_order", ["id", "dept_id", "applicant_id", "urgency_level", "status", "reason", "scenario", "fast_track", "approver_id", "approve_time"]))
    add_table(doc, "表3-7 申领单明细表（apply_order_item）", build_table_rows(schema, "apply_order_item", ["id", "apply_order_id", "material_id", "apply_qty", "actual_qty", "created_at"]))
    add_table(doc, "表3-8 出库单主表（stock_out）", build_table_rows(schema, "stock_out", ["id", "apply_order_id", "warehouse_id", "operator_id", "remark", "created_at"]))
    add_table(doc, "表3-9 调拨单主表（transfer_order）", build_table_rows(schema, "transfer_order", ["id", "from_warehouse_id", "to_warehouse_id", "status", "reason", "applicant_id", "approver_id", "approve_time"]))
    add_table(doc, "表3-10 预警记录表（warning_record）", build_table_rows(schema, "warning_record", ["id", "warning_type", "material_id", "warehouse_id", "content", "handle_status", "handler_id", "handle_remark"]))
    add_table(doc, "表3-11 通知消息表（notification）", build_table_rows(schema, "notification", ["id", "title", "content", "msg_type", "target_user_id", "is_read", "biz_type", "biz_id"]))
    add_table(doc, "表3-12 登录日志表（login_log）", build_table_rows(schema, "login_log", ["id", "user_id", "username", "login_ip", "login_status", "login_time", "user_agent"]))

    add_heading(doc, "3.4 接口与安全设计", "2级标题-正文章节", 1)

    add_heading(doc, "3.4.1 接口设计", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        系统接口遵循前后端分离项目中常见的 REST 风格组织方式，统一以 `/api` 作为业务前缀，再根据功能域细分为认证、库存、申领、调拨、预警、统计、事件、消息等接口组。前端 `http.js` 中的响应拦截器显示，所有接口都以 `ApiResponse` 作为外层封装，其中 `code` 用于标识成功与失败，`message` 用于返回提示信息，`data` 用于承载业务数据；当 `code` 不等于 `0` 时，前端统一弹出错误消息并中止后续流程。

        统一响应结构的好处在于，前端页面无需为每个接口单独适配返回格式，登录、列表、详情、审批和统计图表都可以沿用一致的数据处理逻辑。同时，统一响应对象也为全局异常处理提供了出口：当业务校验失败或权限不足时，后端可以通过统一异常处理器返回结构化错误消息，前端据此进行界面提示或跳转。
        """,
    )

    add_heading(doc, "3.4.2 认证与授权设计", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        认证机制基于 JWT 与 Refresh Token 组合实现。用户成功登录后，后端返回访问令牌与刷新令牌；前端将其保存到本地存储，并在每次请求前自动写入 `Authorization: Bearer <token>` 头。若接口返回 401，拦截器会优先尝试使用刷新令牌访问 `/api/auth/refresh` 获取新访问令牌，成功后重放原请求，失败则清空本地登录态并跳转到登录页。这样的设计兼顾了无状态认证和会话续期需求[16][20]。

        授权机制由 Spring Security 方法级权限和角色菜单共同构成。控制器通过 `@PreAuthorize` 限制不同角色的访问范围，`AuthService` 则根据角色编码返回可见菜单，从而在“接口层”和“导航层”同时做访问控制。系统管理员拥有最完整的管理权限，仓库管理员聚焦库存与仓储操作，审批人聚焦审核页面，部门用户则只保留申领与查看能力，这种角色划分与校园物资管理的职责分工相吻合。
        """,
    )

    add_heading(doc, "3.4.3 安全与审计设计", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        在安全细节方面，系统启用了 CORS 配置、密码加密和定时清理机制。用户密码使用 BCrypt 进行存储，避免明文密码直接落库；Refresh Token 清理任务定期回收过期或撤销的令牌记录，减少历史凭证长期滞留；登录日志与操作日志则为关键行为提供审计依据。对校园场景而言，这些措施虽然不属于高强度安全体系，但足以支撑本科毕业设计所要求的基础安全与可追踪能力。

        需要说明的是，当前仓库并未给出独立的性能压测、安全扫描或生产部署结果，因此本文中的“安全设计”主要指访问控制、凭证管理和操作审计等已实现内容，而不延伸到无法验证的高并发、分布式防护或外部网关治理。这样的表述更贴合项目实际，也有助于保持论文结论的可验证性。
        """,
    )

    add_heading(doc, "4  系统详细设计与实现", "name_out_lvl_1", 0, new_page=True)

    add_heading(doc, "4.1 开发环境与运行环境", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        当前项目使用前后端分离开发方式。后端基于 Maven 构建，依赖由 `backend/pom.xml` 管理；前端基于 Node.js 与 Vite 构建，依赖由 `frontend/package.json` 管理。项目文档 `DEPLOY.md` 给出了本地运行和构建步骤，测试环境则通过 `backend/src/test/resources/application-test.yml` 把数据源切换为 H2 内存数据库，以便在不依赖外部 MySQL 服务的前提下执行自动化测试。

        从实现角度看，开发环境与运行环境并未刻意追求复杂化。系统在 JDK 17、Maven 3.8+、Node.js 18+ 和 MySQL 8 环境下即可完成开发、测试和构建；前端通过 Vite 提供开发服务器与生产打包，后端通过 Spring Boot 提供接口服务与调度任务。这种环境配置符合课程设计项目“小而完整”的特点，有利于在毕业答辩前完成部署复现和运行截图采集。
        """,
    )
    add_table(
        doc,
        "表4-1 系统开发与运行环境",
        [
            ["项目", "版本/配置", "说明"],
            ["JDK", "17", "后端运行环境，`pom.xml` 中声明 `java.version=17`"],
            ["Maven", "3.8+", "后端依赖管理与打包工具"],
            ["Spring Boot", "3.3.5", "后端基础框架"],
            ["MyBatis-Plus", "3.5.7", "ORM 与通用数据访问框架"],
            ["Node.js", "18+", "前端构建环境"],
            ["Vue", "3.5.29", "前端核心框架，本次 `npm ls` 可见已安装版本"],
            ["Vite", "6.4.1", "前端开发与构建工具，本次构建输出显示版本"],
            ["Element Plus", "2.13.5", "前端组件库，本次 `npm ls` 可见已安装版本"],
            ["ECharts", "5.6.0", "统计图表组件"],
            ["MySQL", "8.0", "业务数据持久化存储"],
            ["H2", "内存模式", "测试阶段数据源"],
        ],
    )

    add_heading(doc, "4.2 认证授权与登录态管理实现", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        认证实现从登录接口开始。前端登录页将用户名和密码提交到 `/api/auth/login`，认证服务校验用户身份后返回访问令牌、刷新令牌、角色编码、用户名和姓名等信息。`auth.js` 通过 Pinia store 统一保存这些数据，并在首次登录后立即调用 `/api/auth/me` 和 `/api/auth/menus` 拉取用户资料与动态菜单。由此可见，系统并不是把用户信息硬编码在前端，而是始终以服务端返回结果为准。

        登录态维持依赖 `http.js` 中的双层拦截逻辑。请求拦截器负责为所有业务请求附加 Bearer Token，响应拦截器负责检查统一响应对象的 `code` 字段，并在收到 401 状态时尝试调用刷新接口获取新令牌。为了避免多个并发请求同时刷新令牌，拦截器内部还维护了等待队列 `waitQueue`，把刷新过程中的请求暂存起来，待新令牌返回后再统一重放。这种实现降低了登录态过期时的页面抖动，也体现了前端状态管理的完整性。

        后端的配套机制包括 JWT 生成与校验、Refresh Token 持久化和定时清理。刷新令牌以哈希形式存储在 `auth_refresh_token` 表中，定时任务按配置周期清理过期记录。对本科论文而言，这一实现足以说明系统具备“登录、续期、登出、回收”的完整登录态管理链路，而不是简单地在浏览器本地保存一个长期有效 token。
        """,
    )

    add_heading(doc, "4.3 基础数据管理实现", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        基础数据管理模块覆盖用户、角色、部门、校区、仓库、库位、供应商、分类和物资档案等对象。前端针对这些对象分别提供列表、查询、新增、编辑和删除页面，后端则以相对统一的控制器和服务模式组织 CRUD 逻辑。虽然这些页面的交互形式类似，但在系统中承担的作用并不相同：部门和角色决定权限边界，仓库与库位决定库存归属，供应商和物资档案则决定入库、预警和补货分析的主数据来源。

        物资档案模块是基础数据中的关键节点。物资编码、名称、分类、规格、单位、安全库存、保质期、供应商和单价等信息不仅用于列表展示，还会被库存预警、出库扣减和补货建议等后续逻辑直接使用。因此，物资档案模块的实现重点不在于界面表单本身，而在于通过统一主数据避免“同一物资多种命名、不同仓库重复维护”的问题。

        基础数据模块还承担初始化业务场景的作用。`seed.sql` 为管理员、仓库管理员、审批人和部门用户准备了基础账号和样例数据，这使系统在开发与测试阶段能够快速进入可操作状态。对论文写作而言，这一事实意味着系统并非停留在孤立的接口层，而是具备了可以支撑完整业务链路的初始数据环境。
        """,
    )

    add_heading(doc, "4.4 库存、批次、出入库实现", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        库存管理模块的核心是同时维护库存总量与批次明细。入库时，系统不仅更新 `inventory` 表中的 `current_qty`，还在 `inventory_batch` 中写入批次号、入库量、剩余量、生产日期和到期日期，从而保证总库存和批次账保持一致。出库时则反向根据批次顺序扣减剩余量，再同步更新总库存。这样设计的直接好处，是可以在总量统计与批次追溯之间取得平衡。

        `InventoryService` 的实现体现了“按到期日优先”的业务规则。系统查询指定物资在某仓库下所有剩余数量大于 0 且未过期的批次，按 `expire_date` 和主键升序排列后依次扣减。该策略既能优先消耗临近到期物资，降低损耗风险，也方便论文在实现章节中把“批次推荐”解释为一项已落地的业务规则，而不是停留在先进先出原则的口头描述。

        出入库模块还负责把库存变化与单据状态关联起来。入库单负责建立新批次并增加库存，出库单则通常与申领单关联，在完成拣货后回写申领明细实际数量并推动申领单进入 `OUTBOUND` 状态。也就是说，库存模块并不是孤立计算数量变化，而是在业务上下文中完成库存扣增、批次更新和状态推进的联动。
        """,
    )

    add_heading(doc, "4.5 申领审批与调拨实现", "2级标题-正文章节", 1)

    add_heading(doc, "4.5.1 申领审批实现", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        申领审批实现围绕 `ApplyService` 展开。创建申领单时，系统先插入主表 `apply_order`，初始化状态为 `DRAFT`，再批量插入明细表 `apply_order_item`。提交时若单据仍处于草稿状态，则进一步根据 `urgency_level` 判断流转分支：普通单据进入 `SUBMITTED` 等待审批，紧急等级大于等于 2 的单据则直接置为 `APPROVED`，同时记录审批人、审批意见和审批时间。这一逻辑比传统论文中常见的“提交后统一待审”更贴近真实业务约束。

        审批通过、驳回和签收等动作都由明确的方法处理，并通过 `OrderStatus` 常量约束状态机边界。例如，只有 `SUBMITTED` 状态的单据可以被审批或驳回，只有 `OUTBOUND` 状态的单据才允许签收。每个关键动作完成后，系统都会调用 `OperationLogService` 写入操作日志，从而形成业务动作与审计记录的同步留痕。

        申领审批模块的实现价值在于，它把角色协同、单据明细、状态机和库存执行串成了一条完整链路。部门用户关注的是“能否快速发起申请”，审批人关注的是“是否符合领用条件”，仓库管理员关注的是“如何准确出库并回写数量”，而系统通过主表、明细表和日志表把这些差异化需求统一到同一套数据模型下。
        """,
    )

    add_heading(doc, "4.5.2 调拨实现与推荐机制", "3级标题-正文章节", 2)
    add_body_block(
        doc,
        """
        调拨实现与申领类似，也采用主表加明细表模式，但其复杂度更高。`TransferService.create()` 在校验调出仓与调入仓不能相同后，生成 `DRAFT` 状态的调拨主单和多条明细；提交、审批、驳回、执行与签收则分别由独立方法处理。执行阶段系统不仅要检查调出仓库存是否充足，还要逐批次扣减来源批次，并在调入仓生成新的批次记录，这使调拨真正成为一次“跨仓库存迁移”。

        调拨推荐机制建立在 `DijkstraUtil` 的预设校区图之上。服务层调用 `calculateShortestPaths(targetCampus)` 计算目标校区到其他节点的最短距离，再结合库存表筛选“库存数量足够”的候选仓库，并按距离从小到大排序返回。这里的“智能”并不是复杂地图服务，而是基于预设节点和边权的最短路径计算，适合作为校园内部多校区、多仓点场景下的轻量辅助决策能力。

        这一实现说明系统对“跨仓协同”的理解已经超出普通单仓库存管理：仓库之间不是互不关联的孤岛，而是可以围绕库存余量、距离和审批流程进行协同。论文在此处需要强调的是推荐逻辑的边界，即它解决的是候选仓排序问题，而不是完整的物流调度优化问题，从而保持与代码实现一致。
        """,
    )
    add_picture(doc, figures["fig_4_2"], "图4-2 调拨执行与推荐流程图")

    add_heading(doc, "4.6 预警、补货建议与统计分析实现", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        预警模块由 `WarningService` 负责实现，核心入口是带有 `@Scheduled(cron = \"0 0/30 * * * ?\")` 注解的 `scan()` 方法。系统每 30 分钟扫描一次库存与批次数据，分别执行低库存、库存积压、临期、过期和异常领用五类检测逻辑，并在不存在相同未处理预警时创建新的 `warning_record`。这一定时扫描机制让预警成为主动发现问题的能力，而不是完全依赖人工检查列表。

        五类预警的判定依据都能在代码中找到明确来源：库存不足与库存积压依赖 `material_info.safety_stock`；临期与过期依赖批次到期日期；异常领用则比较近 7 天出库总量与近 30 天折算周均值的关系。需要特别说明的是，当前实现并未读取 `system_config` 中的阈值配置，因此论文不能将其写成“系统配置动态驱动预警阈值”，而应准确描述为基于安全库存和固定窗口规则的扫描机制。

        智能分析模块由 `SmartService` 提供两个能力：一是读取近 6 个月的出库历史，计算平均值并生成未来若干月预测；二是按近 30 天出库量估算日均消耗，结合安全库存和保障天数计算目标库存与建议补货量。这种实现方式与文献中常见的移动平均思想保持一致，但工程上采用的是简单、可解释、可落地的平均值计算，而不是论文中经常被泛化描述的复杂预测模型[15]。

        统计分析则由 `AnalyticsService` 通过聚合 SQL 实现。系统围绕月度出入库趋势、物资分布、预警数量等维度生成图表数据，并由前端图表页和大屏页负责展示。由于统计结果直接依赖业务表中的真实记录，因此其价值不只是“展示漂亮图表”，更在于帮助管理员观察物资流动规律、识别高频消耗物资并辅助补货决策。
        """,
    )

    add_heading(doc, "4.7 日志、通知与事件管理实现", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        日志、通知和事件管理模块承担的是“让业务过程可追踪”的职责。登录日志记录用户账号、登录 IP、登录结果、时间和客户端标识，操作日志记录操作人、业务模块、操作类型和明细内容，两者共同构成对关键行为的审计基础。由于申领、审批、出库、调拨、预警处理等操作都会写入日志，系统能够在后续排查中回溯“谁在什么时候做了什么”。

        通知模块以 `notification` 表为中心，保存标题、内容、消息类型、目标用户、已读状态及关联业务信息。它的实现价值在于把业务结果从“页面即时提示”扩展为“可再次查看的消息记录”，从而适应校园场景中审批、预警、事件等信息需要跨时间查看的需求。部门用户、审批人和仓库管理员可以在自身角色页面中查看与自己相关的消息，提升协同效率。

        事件管理模块则进一步扩展了系统对校园保障场景的覆盖面。`event_record` 表可以记录事件标题、类型、等级、位置、上报人、处理结果和关闭时间，说明项目已经考虑到校园突发事件与物资管理之间的联动关系。虽然当前论文不把它写成完整应急指挥平台，但将事件管理作为支撑模块纳入实现，能够体现系统对校园综合保障业务的延展能力。
        """,
    )

    add_heading(doc, "5  系统测试", "name_out_lvl_1", 0, new_page=True)

    add_heading(doc, "5.1 测试环境与证据来源", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        本文的测试章节以项目仓库中已经存在且可执行的验证材料为依据，主要包括后端自动化测试代码、H2 测试配置、Maven 测试执行结果以及前端生产构建结果。与许多只描述“设计了测试方案”的论文不同，本文优先使用能够在当前环境下重复执行的事实作为证据来源，以避免将无法复现的压测结论或上线效果写入论文。

        后端测试环境通过 `backend/src/test/resources/application-test.yml` 配置为 H2 内存数据库，并指定 `schema.sql` 作为初始化脚本。这样可以在不依赖外部 MySQL 实例的情况下完成数据结构加载和业务测试执行。对课程设计项目而言，这种测试方式具有准备成本低、回归速度快和便于持续重复执行的特点。

        本次定稿过程中，于 2026 年 4 月 14 日实际执行了 `mvn -f backend/pom.xml test` 与 `npm --prefix frontend run build` 两条命令，分别得到“后端 42 项测试通过”和“前端生产构建成功”的结果。由于这两项结果都来自当前仓库真实执行，因此可以作为第 5 章中最重要的可验证结论。
        """,
    )

    add_heading(doc, "5.2 自动化测试结果", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        从测试类分布看，当前自动化测试覆盖了统一响应对象与异常处理、认证与令牌管理、申领业务、调拨业务和预警接口等关键领域。尽管这些测试尚不能替代完整的端到端回归，但已经覆盖了系统中最容易出现状态流转错误和权限错误的主链路。对本科毕业设计而言，这种“围绕核心风险点布置测试”的方式具有较高的性价比。

        2026 年 4 月 14 日执行 Maven 测试后，后端共通过 42 项测试。该结果表明当前版本至少在核心业务逻辑和异常处理层面保持了可执行状态，没有出现明显的编译错误、配置缺失或关键断言失败。结合 H2 测试配置可以进一步说明，这些测试并非依赖人工准备数据库环境，而是可以在标准开发环境中重复运行。
        """,
    )
    add_table(
        doc,
        "表5-1 后端自动化测试执行情况",
        [
            ["测试类别", "对应测试类/范围", "验证重点", "结果"],
            ["统一响应与异常处理", "ApiResponseTest、BizExceptionTest、GlobalExceptionHandlerTest", "统一返回结构、异常封装与错误消息", "2026-04-14 执行通过"],
            ["认证与令牌管理", "AuthServiceTest、JwtTokenProviderTest、AuthRefreshTokenCleanupTaskTest", "登录、JWT、刷新令牌清理与菜单装配", "2026-04-14 执行通过"],
            ["申领业务", "ApplyServiceTest、ApplyControllerTest", "申领单创建、提交、审批和接口访问", "2026-04-14 执行通过"],
            ["调拨业务", "TransferServiceTest、TransferControllerTest", "调拨单创建、审批、执行和推荐接口", "2026-04-14 执行通过"],
            ["预警接口", "WarningControllerTest", "预警列表与处理接口", "2026-04-14 执行通过"],
            ["总体结果", "共 11 个测试类、42 项测试", "核心模块可执行性", "全部通过"],
        ],
    )

    add_heading(doc, "5.3 业务场景验证", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        除自动化测试外，还可以基于代码、路由和 `seed.sql` 中的样例数据对典型业务场景进行可复验说明。这种说明并不等同于宣称所有场景都已在生产环境中运行，而是用于证明系统主链路具备明确入口、状态变化和落库结果。对于答辩前的论文定稿而言，这类“可复验业务场景表”可以帮助后续截图采集与现场演示按图索骥展开。

        结合本项目实现，至少可以确定以下业务链路已经在代码层形成闭环：部门用户创建并提交申领单，审批人审核通过后仓库管理员执行出库，申请人签收完成；仓库管理员创建调拨单并调用推荐接口选择来源仓库，审批后执行调拨并签收；系统按周期生成预警并允许管理员处理；智能分析模块根据历史出库生成预测与补货建议。换言之，系统已具备覆盖“计划、执行、监督、分析”四类动作的基础。
        """,
    )
    add_table(
        doc,
        "表5-2 典型业务场景验证说明",
        [
            ["业务场景", "主要入口", "关键证据", "说明"],
            ["普通申领审批与签收", "/apply/list", "ApplyService、InventoryService、OrderStatus", "支持 DRAFT→SUBMITTED→APPROVED→OUTBOUND→RECEIVED 流转"],
            ["紧急申领快速审批", "/apply/list", "ApplyService.submit()", "urgency_level>=2 时直接审批通过并记录审批信息"],
            ["跨仓调拨推荐与执行", "/transfer/list", "TransferService、DijkstraUtil", "支持候选仓排序、审批、批次迁移与签收"],
            ["预警扫描与处理", "/warning/list", "WarningService.scan()、warning_record", "支持五类预警生成与处理状态更新"],
            ["补货建议与趋势分析", "/analytics/charts、/bigscreen", "SmartService、AnalyticsService", "支持近 6 个月历史统计、预测与建议补货量计算"],
            ["日志与通知追踪", "/log/*、/notification/list", "login_log、operation_log、notification", "支持对关键业务动作的审计与消息留存"],
        ],
    )

    add_heading(doc, "5.4 构建验证与问题边界", "2级标题-正文章节", 1)
    add_body_block(
        doc,
        """
        前端构建验证结果来自 2026 年 4 月 14 日执行 `npm --prefix frontend run build` 的实际输出。构建成功表明 Vue 页面、路由、组件依赖和打包配置在当前仓库中保持一致，没有出现阻断交付的基础构建错误。对前后端分离项目来说，这一结论很重要，因为它说明系统不仅后端逻辑可测试，前端静态资源也能够生成可部署产物。

        同时也应看到，当前测试与构建结果只能证明“项目可执行”和“核心逻辑在现有测试覆盖下通过验证”，不能外推出高并发稳定性、生产环境性能、跨网络部署效果和长期运维可靠性。仓库中没有对应的压力测试报告、监控数据和生产日志，因此本文明确不写入“百人并发”“响应时间稳定在某一区间”或“系统已上线试运行”等结论。

        从论文完善角度看，后续仍需补充的工作主要包括：启动系统并采集真实页面截图，结合正文章节补入运行界面；在 Word 中刷新目录、统一题注和交叉引用；进一步核对参考文献格式与正文引用位置。这些工作属于定稿排版与展示层的完善，不改变本文已经依据仓库证据建立的事实边界。
        """,
    )

    end_title = add_heading(doc, "结束语", "总结/结论标题", 0, new_page=True)
    set_outline_level(end_title, 0)
    add_body_block(
        doc,
        """
        本文围绕当前项目仓库，对校园物资智能管理系统的需求、总体设计、关键实现与测试情况进行了系统梳理。与早期草稿相比，定稿工作将题名统一为“校园物资智能管理系统设计与实现”，并严格回到真实代码、SQL、配置和执行结果，不再沿用旧题名、错误接口和无法证实的性能结论。通过这一过程，论文内容与仓库实现建立了更直接的一一对应关系。

        从实现结果看，系统已经完成了认证授权、基础数据、库存与批次、申领审批、调拨协同、预警扫描、补货建议、统计分析、通知消息、事件记录和日志审计等核心能力。特别是申领单与调拨单的状态流转、按到期日优先的批次扣减、基于预设校区图的调拨推荐、定时预警扫描以及基于历史出库数据的补货建议，使系统不再局限于简单的增删改查页面，而具备了面向校园物资管理主要业务流程的实现能力。

        结合 2026 年 4 月 14 日后端 42 项测试通过和前端生产构建成功的事实，可以认为当前项目已经形成较稳定的课程设计成果。后续工作可在不改变现有事实边界的前提下继续推进，例如补充真实运行截图、完善引用与图表排版、在预警规则和补货建议上引入更细粒度的参数控制、扩展更多与校园保障相关的联动场景。上述方向属于在既有实现基础上的持续深化，而非对当前成果的夸大延伸。
        """,
        style="总结/结论内容",
    )

    thanks_title = add_heading(doc, "致    谢", "致谢标题", 0, new_page=True)
    set_outline_level(thanks_title, 0)
    add_body_block(
        doc,
        """
        本文的完成离不开指导教师在选题、实现和论文修改过程中的持续指导，也离不开项目仓库中已有代码、测试、配置与文档材料提供的直接支撑。在定稿过程中，针对旧草稿中存在的事实偏差、结构冗余和证据不足问题，能够重新回到真实项目资产完成梳理，对本人理解软件工程类毕业设计“以代码为证据”的写作方式具有重要帮助。

        同时，也感谢在项目开发与论文撰写期间给予支持的同学和家人。正是在反复调试、测试、修订和校对的过程中，才逐步形成了当前这一版更贴近仓库实现、结构更完整、表述更审慎的毕业论文工作稿。
        """,
        style="致谢内容",
    )

    ref_title = add_heading(doc, "参考文献", "参考文献标题", 0, new_page=True)
    set_outline_level(ref_title, 0)
    references = [
        "[1] 陈志敏, 何建国. 大数据背景下智慧校园建设与管理模式创新探讨[J]. 现代教育管理, 2020, 12(8): 45-51.",
        "[2] 张强, 李红梅. 高校突发公共事件应急物资储备体系构建研究[J]. 中国安全生产科学技术, 2022, 18(3): 67-73.",
        "[3] 冯绍彬, 周华. 基于角色的动态权限管理系统的应用设计[J]. 郑州轻工业大学学报(自然科学版), 2020, 35(4): 71-76.",
        "[4] 刘丹阳, 杨华, 张诗桐. 基于 Django 的个性化图书推荐管理系统设计与实现[J]. 山西电子技术, 2025(03): 64-66,80.",
        "[5] 岑丹. 基于人工智能的图书馆信息管理系统设计与实现[J]. 吉林大学学报(信息科学版), 2025, 43(06): 1346-1351.",
        "[6] 李淼淼, 季节, 李坡, 等. 基于 SpringBoot 和 Vue 3 的移动学习管理系统的设计与实现[J]. 无线互联科技, 2026, 23(04): 81-86.",
        "[7] 杨茜, 吴加莹. 基于 B/S 架构的高校教务管理系统设计与实现[J]. 电脑知识与技术, 2026, 22(06): 58-60,76.",
        "[8] 崔靖茹, 文华, 刘宏磊, 等. 基于 Vue 和 SpringBoot 框架的高校信息化项目管理系统的设计与实现[J]. 现代信息科技, 2025, 9(22): 77-81.",
        "[9] 薛楠楠, 张伟, 钟化雨, 等. 基于 SpringBoot+Vue 的事故报告数据库系统设计与实现[J]. 计算机仿真, 2025, 42(07): 456-462,570.",
        "[10] 王伟, 刘明. 基于 Spring Boot 的后端架构设计与性能优化研究[J]. 计算机系统应用, 2021, 30(5): 102-109.",
        "[11] 刘强, 周波. MyBatis-Plus 企业级高可用开发指南[M]. 北京: 电子工业出版社, 2021: 89-105.",
        "[12] 李东, 赵莉. Vue3+Vite 开发实战: 组合式 API 与工程化实践[M]. 北京: 人民邮电出版社, 2023: 155-180.",
        "[13] 李晓东. MySQL 8.0 性能调优实战教程[M]. 北京: 清华大学出版社, 2022: 24-39.",
        "[14] 赵玉明, 林晓. ECharts 在 Web 报表数据可视化中的应用分析[J]. 信息技术与信息化, 2021(4): 21-25.",
        "[15] 孙阳, 陈洁. 一种基于移动平均预测模型的库存预警系统设计[J]. 物流工程与管理, 2022, 44(2): 33-38.",
        "[16] 杨俊. 新一代 Spring Security 原理与 JWT 无状态鉴权探索[D]. 武汉: 武汉科技大学, 2020.",
        "[17] Spring Team. Spring Boot Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/docs/3.3.5/reference/html/, 2026-04-14.",
        "[18] Vue Team. Vue.js Documentation[EB/OL]. https://vuejs.org/guide/introduction.html, 2026-04-14.",
        "[19] Vite Team. Vite Guide[EB/OL]. https://vite.dev/guide/, 2026-04-14.",
        "[20] Jones M, Bradley J, Sakimura N. JSON Web Token (JWT): RFC 7519[EB/OL]. https://www.rfc-editor.org/rfc/rfc7519, 2026-04-14.",
    ]
    for item in references:
        add_reference_paragraph(doc, item)

    set_update_fields_on_open(doc)
    set_all_text_black(doc)
    doc.save(WORKING_DRAFT)
    return WORKING_DRAFT, estimate_body_chars(doc)


def write_report(word_count: int) -> Path:
    content = f"""# 校园物资智能管理系统论文改稿说明

## 输出文件
- 原始草稿：`{SOURCE_DRAFT.relative_to(ROOT)}`
- 工作稿：`{WORKING_DRAFT.relative_to(ROOT)}`
- 备份稿：`{BACKUP_DRAFT.relative_to(ROOT)}`
- 图表目录：`{FIGURE_DIR.relative_to(ROOT)}`

## 本次处理内容
- 按郑州轻工业大学模板约束保留封面、任务书、摘要、目录壳结构，并在原草稿副本上重建正文。
- 题名统一改为“{NEW_TITLE}”，并同步修正页眉中的旧题名。
- 删除旧稿中独立的“相关技术与理论基础”叙述，将可证实的技术栈事实折叠进总体设计与详细实现章节。
- 正文按“绪论—需求分析—总体设计—详细实现—系统测试—结束语”结构重写。
- 仅保留代码、SQL、配置和文档可以直接证明的事实，剔除了错误接口名、错误表名、错误响应格式以及无证据的性能/上线结论。
- 根据 `schema.sql` 自动生成关键数据表说明，并重绘了架构图、模块图、流程图和 3 张 E-R 图。
- 测试章节写入了 2026-04-14 实测结果：`mvn -f backend/pom.xml test` 通过 42 项测试，`npm --prefix frontend run build` 成功。

## 新增图表
- 图3-1 系统总体架构图
- 图3-2 系统功能模块图
- 图3-3 RBAC 与组织实体关系图
- 图3-4 库存与批次实体关系图
- 图3-5 业务单据、预警与通知实体关系图
- 图4-1 申领审批闭环流程图
- 图4-2 调拨执行与推荐流程图
- 表3-1 至表3-12 关键数据表说明
- 表4-1 系统开发与运行环境
- 表5-1 后端自动化测试执行情况
- 表5-2 典型业务场景验证说明

## 仍需人工复核的事项
- 在 Microsoft Word 中打开工作稿后刷新目录、页码和交叉引用，确认目录分页与标题层级显示正常。
- 运行系统并补采登录页、申领页、审批页、调拨页、预警页和统计页截图，再按章节插入真实界面图片。
- 复核参考文献在学校模板下的段落缩进与标点格式，必要时按学院规范微调。
- 终稿提交前再做一次格式检查，确认页眉、页脚、题注、表题和摘要关键词排版无偏差。

## 字数提示
- 当前工作稿按脚本估算的正文与表格字符总量约为：{word_count}。
- 该数字用于改稿自检，不等同于学校系统最终统计值；如学校模板采用不同统计口径，应以 Word 实际统计结果为准。
"""
    REPORT_PATH.write_text(content, encoding="utf-8")
    return REPORT_PATH


if __name__ == "__main__":
    path, word_count = build_document()
    report = write_report(word_count)
    print(path)
    print(report)
    print(word_count)
