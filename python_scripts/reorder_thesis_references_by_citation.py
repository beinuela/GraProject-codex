from pathlib import Path
import re
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿-无具体日期版.docx"
DST = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿-文献顺序修正版.docx"

CITATION_PATTERN = re.compile(r"\[(\d+(?:\s*[-,，]\s*\d+)*)\]")
REFERENCE_PATTERN = re.compile(r"^\[(\d+)\]\s*(.+)")


def replace_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def expand_citation_numbers(content):
    numbers = []
    for part in re.split(r"[,，]", content):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            start, end = [int(value.strip()) for value in part.split("-", 1)]
            numbers.extend(range(start, end + 1))
        else:
            numbers.append(int(part))
    return numbers


def collect_citation_order(document, body_start, reference_heading):
    order = []
    seen = set()
    for paragraph in document.paragraphs[body_start - 1 : reference_heading - 1]:
        for match in CITATION_PATTERN.finditer(paragraph.text):
            for number in expand_citation_numbers(match.group(1)):
                if number not in seen:
                    seen.add(number)
                    order.append(number)
    return order


def rewrite_citations_in_run_text(text, number_map):
    def replace_match(match):
        numbers = expand_citation_numbers(match.group(1))
        mapped = [str(number_map.get(number, number)) for number in numbers]
        if len(mapped) == 1:
            return f"[{mapped[0]}]"
        return "[" + ",".join(mapped) + "]"

    return CITATION_PATTERN.sub(replace_match, text)


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)

    shutil.copy2(SRC, DST)
    document = Document(DST)

    body_start = next(
        index for index, paragraph in enumerate(document.paragraphs, 1) if paragraph.text.strip() == "1 绪论"
    )
    reference_heading = next(
        index for index, paragraph in enumerate(document.paragraphs, 1) if paragraph.text.strip() == "参考文献"
    )

    citation_order = collect_citation_order(document, body_start, reference_heading)
    number_map = {old_number: new_number for new_number, old_number in enumerate(citation_order, 1)}

    reference_paragraphs = []
    references_by_old_number = {}
    for index, paragraph in enumerate(document.paragraphs[reference_heading:], reference_heading + 1):
        match = REFERENCE_PATTERN.match(paragraph.text.strip())
        if not match:
            continue
        old_number = int(match.group(1))
        reference_paragraphs.append(paragraph)
        references_by_old_number[old_number] = match.group(2).strip()

    missing = [number for number in citation_order if number not in references_by_old_number]
    if missing:
        raise ValueError(f"正文引用编号在参考文献中不存在: {missing}")

    # 正文引用只处理第 1 章至参考文献之前的正式正文，避免改动任务书中的“主要参考资料”。
    changed_runs = 0
    for paragraph in document.paragraphs[body_start - 1 : reference_heading - 1]:
        for run in paragraph.runs:
            updated = rewrite_citations_in_run_text(run.text, number_map)
            if updated != run.text:
                run.text = updated
                changed_runs += 1

    if len(reference_paragraphs) < len(citation_order):
        raise ValueError("文末参考文献段落数量少于正文引用数量")

    for new_number, old_number in enumerate(citation_order, 1):
        replace_paragraph_text(reference_paragraphs[new_number - 1], f"[{new_number}] {references_by_old_number[old_number]}")

    # 清空多余且未被正文引用的旧参考文献段落；当前文档通常不会进入该分支。
    for paragraph in reference_paragraphs[len(citation_order) :]:
        replace_paragraph_text(paragraph, "")

    document.save(DST)

    verified = Document(DST)
    new_order = collect_citation_order(verified, body_start, reference_heading)
    print(f"saved={DST}")
    print(f"old_to_new={number_map}")
    print(f"changed_citation_runs={changed_runs}")
    print(f"verified_first_citation_order={new_order}")


if __name__ == "__main__":
    main()
