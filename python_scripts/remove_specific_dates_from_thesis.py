from pathlib import Path
import re
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿-附录版式修正版.docx"
DST = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿-无具体日期版.docx"


TEXT_REPLACEMENTS = [
    (
        "在本轮定稿过程中，2026 年 4 月 21 日已经实际执行过",
        "在本轮定稿过程中，已实际执行过",
    ),
    (
        "2026 年 4 月 21 日执行 Maven 测试后，",
        "执行 Maven 测试后，",
    ),
    (
        "；同日执行前端构建与单元测试后，",
        "；执行前端构建与单元测试后，",
    ),
    (
        "前端构建验证来自 2026 年 4 月 14 日执行 npm --prefix frontend run build 的实际结果。",
        "前端构建验证来自 npm --prefix frontend run build 的实际执行结果。",
    ),
    (
        "结合 2026 年 4 月 21 日后端 49 项测试通过、",
        "结合后端 49 项测试通过、",
    ),
    (", 2026-04-14.", "."),
    ("，2026-04-14。", "。"),
    ("2026-04-21 执行通过", "执行通过"),
]


SPECIFIC_DATE_PATTERNS = [
    re.compile(r"\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日"),
    re.compile(r"\d{4}-\d{1,2}-\d{1,2}"),
]


def replace_text_preserve_first_run(paragraph, replacements):
    original = paragraph.text
    updated = original
    for old, new in replacements:
        updated = updated.replace(old, new)

    if updated == original:
        return False

    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)
    return True


def iter_table_paragraphs(document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def scan_specific_dates(document):
    matches = []
    for idx, paragraph in enumerate(document.paragraphs, 1):
        text = paragraph.text.strip()
        if text and any(pattern.search(text) for pattern in SPECIFIC_DATE_PATTERNS):
            matches.append(("paragraph", idx, text))

    for table_idx, table in enumerate(document.tables, 1):
        for row_idx, row in enumerate(table.rows, 1):
            for cell_idx, cell in enumerate(row.cells, 1):
                text = " ".join(cell.text.split())
                if text and any(pattern.search(text) for pattern in SPECIFIC_DATE_PATTERNS):
                    matches.append(("table", (table_idx, row_idx, cell_idx), text))
    return matches


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)

    shutil.copy2(SRC, DST)
    document = Document(DST)

    changed = 0
    for paragraph in document.paragraphs:
        if replace_text_preserve_first_run(paragraph, TEXT_REPLACEMENTS):
            changed += 1

    for paragraph in iter_table_paragraphs(document):
        if replace_text_preserve_first_run(paragraph, TEXT_REPLACEMENTS):
            changed += 1

    document.save(DST)

    verified = Document(DST)
    remaining = scan_specific_dates(verified)
    print(f"saved={DST}")
    print(f"changed_paragraphs_or_cells={changed}")
    print(f"remaining_specific_date_matches={len(remaining)}")
    for kind, location, text in remaining:
        print(f"{kind} {location}: {text}")


if __name__ == "__main__":
    main()
