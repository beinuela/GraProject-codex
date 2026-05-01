from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "output" / "doc"
SOURCE = DOC_DIR / "校园物资智能管理系统设计与实现-图表引用修复版.docx"
TARGET = DOC_DIR / "校园物资智能管理系统设计与实现-日期清理版.docx"
REPORT = DOC_DIR / "日期表述清理检查报告.md"

DATE_PATTERNS = [
    re.compile(r"\d{4}年\d{1,2}月\d{1,2}日"),
    re.compile(r"\d{1,2}月\d{1,2}日"),
    re.compile(r"\d{4}/\d{1,2}/\d{1,2}"),
    re.compile(r"\d{4}-\d{1,2}-\d{1,2}"),
]

BODY_HEADING_RE = re.compile(r"^[1-7]\s")
STOP_TITLES = {"致  谢", "致谢", "参考文献", "附录 1 关键接口与测试补充说明", "附录"}

BODY_REPLACEMENTS = {
    "2026年4月26日重新执行 `mvn test`，后端共 49 项测试通过，失败 0 项、错误 0 项、跳过 0 项。测试范围包括统一响应、业务异常、全局异常处理、登录刷新、限流、申领出库主链路、并发申领库存锁定、调拨执行、预警处理和分页接口。后端自动化测试执行情况如表7-1所示。":
        "经重新执行后端自动化测试，后端共 49 项测试通过，失败 0 项、错误 0 项、跳过 0 项。测试范围包括统一响应、业务异常、全局异常处理、登录刷新、限流、申领出库主链路、并发申领库存锁定、调拨执行、预警处理和分页接口。后端自动化测试执行情况如表7-1所示。"
}


def is_body_paragraph(index: int, paragraphs) -> bool:
    in_body = False
    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name == "Heading 1":
            if text in STOP_TITLES:
                in_body = False
            elif BODY_HEADING_RE.match(text) or text == "结束语":
                in_body = True
        if i == index:
            return in_body
    return False


def scan_dates(doc: Document) -> dict[str, list[str]]:
    body_hits: list[str] = []
    non_body_hits: list[str] = []
    table_hits: list[str] = []

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if any(p.search(text) for p in DATE_PATTERNS):
            item = f"P[{index}] {paragraph.style.name} :: {text}"
            if is_body_paragraph(index, doc.paragraphs):
                body_hits.append(item)
            else:
                non_body_hits.append(item)

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                text = cell.text.strip().replace("\n", " | ")
                if text and any(p.search(text) for p in DATE_PATTERNS):
                    table_hits.append(f"T[{table_index}][{row_index}][{cell_index}] :: {text}")

    return {
        "body": body_hits,
        "non_body": non_body_hits,
        "tables": table_hits,
    }


def apply_body_replacements(doc: Document) -> list[str]:
    changed: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in BODY_REPLACEMENTS:
            paragraph.text = BODY_REPLACEMENTS[text]
            changed.append(text)
    return changed


def write_report(
    backup_path: Path,
    before_scan: dict[str, list[str]],
    after_scan: dict[str, list[str]],
    changed: list[str],
) -> None:
    lines = [
        "# 日期表述清理检查报告",
        "",
        f"- 工作底稿：`{SOURCE.name}`",
        f"- 备份文件：`{backup_path.name}`",
        f"- 输出文件：`{TARGET.name}`",
        "",
        "## 扫描到的具体日期位置",
        "",
        "### 正文",
    ]

    if before_scan["body"]:
        for item in before_scan["body"]:
            lines.append(f"- {item}")
    else:
        lines.append("- 无")

    lines.extend(["", "### 非正文保留项"])
    if before_scan["non_body"]:
        for item in before_scan["non_body"]:
            lines.append(f"- {item}")
    else:
        lines.append("- 无")

    lines.extend(["", "### 表格"])
    if before_scan["tables"]:
        for item in before_scan["tables"]:
            lines.append(f"- {item}")
    else:
        lines.append("- 无")

    lines.extend(["", "## 已改写的正文段落", ""])
    if changed:
        for item in changed:
            lines.append(f"- {item}")
    else:
        lines.append("- 无")

    lines.extend(
        [
            "",
            "## 保留未改的日期清单及原因",
            "",
        ]
    )

    if before_scan["non_body"]:
        lines.append("- 封面与任务书中的日期属于学校模板信息，本轮按约定保留。")
    if any("[" in item for item in before_scan["non_body"] if "参考文献" not in item):
        pass
    if any("P[" in item for item in before_scan["non_body"]):
        lines.append("- 参考文献中的出版日期属于文献元数据，本轮按约定保留。")
    if before_scan["tables"]:
        lines.append("- 未发现需要清理的表格内过程性日期；若后续表格中出现正文性质时间痕迹，再单独处理。")

    lines.extend(["", "## 清理后正文残留检查", ""])
    if after_scan["body"]:
        for item in after_scan["body"]:
            lines.append(f"- 仍存在正文日期：{item}")
    else:
        lines.append("- 未发现正文中的具体日期残留。")

    REPORT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    backup_path = DOC_DIR / f"{SOURCE.stem}-日期清理前备份-{timestamp}{SOURCE.suffix}"
    shutil.copy2(SOURCE, backup_path)
    shutil.copy2(SOURCE, TARGET)

    before_scan = scan_dates(Document(SOURCE))

    target_doc = Document(TARGET)
    changed = apply_body_replacements(target_doc)
    target_doc.save(TARGET)

    after_scan = scan_dates(Document(TARGET))
    if after_scan["body"]:
        raise RuntimeError(f"清理后仍存在正文日期残留: {after_scan['body']}")

    write_report(backup_path, before_scan, after_scan, changed)

    print(f"[backup] {backup_path.name}")
    print(f"[output] {TARGET.name}")
    print(f"[before-body] {len(before_scan['body'])}")
    print(f"[before-non-body] {len(before_scan['non_body'])}")
    print(f"[changed] {len(changed)}")
    print(f"[after-body] {len(after_scan['body'])}")


if __name__ == "__main__":
    main()
