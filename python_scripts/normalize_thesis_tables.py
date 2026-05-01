from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿-图表重绘版.docx"
TARGET_DOCX = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿-图表重绘版-表格规范化.docx"

CAPTION_RE = re.compile(r"^表\d+-\d+\s+")

# A4 纵向正文区可用宽度约 16cm。
DB_TABLE_WIDTHS = [2.7, 2.2, 1.5, 2.6, 7.0]
SMART_TABLE_WIDTHS = [1.9, 2.8, 3.0, 2.2, 3.6, 2.5]
ENV_TABLE_WIDTHS = [2.4, 2.6, 11.0]
TEST_TABLE_WIDTHS = [2.4, 4.8, 4.4, 4.4]
CASE_TABLE_WIDTHS = [2.6, 2.8, 4.6, 6.0]


def twips_from_cm(value_cm: float) -> int:
    return int(round(value_cm * 567))


def get_or_add_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def remove_child(parent, tag: str) -> None:
    child = parent.find(qn(tag))
    if child is not None:
        parent.remove(child)


def set_table_width(table, width_cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = get_or_add_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:w"), str(twips_from_cm(width_cm)))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = get_or_add_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    remove_child(tbl_pr, "w:tblBorders")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = OxmlElement(f"w:{edge}")
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "8")
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), "000000")
        borders.append(item)
    tbl_pr.append(borders)


def set_table_cell_margins(table, top: int = 50, left: int = 70, bottom: int = 50, right: int = 70) -> None:
    tbl_pr = table._tbl.tblPr
    remove_child(tbl_pr, "w:tblCellMar")
    margins = OxmlElement("w:tblCellMar")
    for edge, size in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        item = OxmlElement(f"w:{edge}")
        item.set(qn("w:w"), str(size))
        item.set(qn("w:type"), "dxa")
        margins.append(item)
    tbl_pr.append(margins)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = get_or_add_child(tc_pr, "w:tcW")
    tc_w.set(qn("w:w"), str(twips_from_cm(width_cm)))
    tc_w.set(qn("w:type"), "dxa")


def clear_cell_flow_constraints(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for tag in ("w:textDirection", "w:noWrap", "w:tcFitText"):
        remove_child(tc_pr, tag)


def set_row_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_run_font(run, size_pt: float, bold: bool | None = None) -> None:
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")
    if bold is not None:
        run.bold = bold


def normalize_paragraph(paragraph, align, size_pt: float, bold: bool | None = None) -> None:
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    fmt.keep_together = True
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_font(run, size_pt, bold=bold)


def normalize_caption_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.0
    for run in paragraph.runs:
        set_run_font(run, 12, bold=True)


def smart_table_text(row_idx: int, col_idx: int, text: str) -> str:
    if row_idx == 0:
        return text
    if col_idx == 1 and len(text) > 24:
        return text.replace("、", "、\n").replace(", ", ",\n").replace(",", ",\n")
    if col_idx == 2 and len(text) > 20:
        return text.replace(" + ", " +\n").replace("，", "，\n").replace(", ", ",\n")
    if col_idx == 4 and "；" in text:
        return text.replace("；", "；\n")
    return text


def compact_long_list_text(col_idx: int, text: str) -> str:
    if col_idx in (1, 2) and len(text) > 28 and "、" in text:
        return text.replace("、", "、\n")
    return text


def style_table(
    table,
    widths_cm: list[float],
    body_font_pt: float,
    header_font_pt: float,
    column_alignments,
    transform_text=None,
    min_row_height_cm: float = 0.72,
) -> None:
    total_width = sum(widths_cm)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_layout_fixed(table)
    set_table_width(table, total_width)
    set_table_borders(table)
    set_table_cell_margins(table)

    for row_idx, row in enumerate(table.rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(min_row_height_cm)
        set_row_cant_split(row)
        if row_idx == 0:
            set_row_header(row)

        for col_idx, cell in enumerate(row.cells):
            clear_cell_flow_constraints(cell)
            set_cell_width(cell, widths_cm[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if transform_text is not None:
                new_text = transform_text(row_idx, col_idx, cell.text.strip())
                if new_text != cell.text:
                    cell.text = new_text

            target_align = column_alignments[col_idx]
            for paragraph in cell.paragraphs:
                normalize_paragraph(
                    paragraph,
                    target_align if row_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER,
                    header_font_pt if row_idx == 0 else body_font_pt,
                    bold=True if row_idx == 0 else None,
                )


def normalize_database_tables(doc: Document) -> None:
    for idx in range(12):
        style_table(
            doc.tables[idx],
            widths_cm=DB_TABLE_WIDTHS,
            body_font_pt=9.5,
            header_font_pt=10,
            column_alignments=[
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.LEFT,
            ],
        )


def normalize_smart_table(doc: Document) -> None:
    style_table(
        doc.tables[12],
        widths_cm=SMART_TABLE_WIDTHS,
        body_font_pt=8.5,
        header_font_pt=9.5,
        column_alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
        ],
        transform_text=smart_table_text,
        min_row_height_cm=0.82,
    )


def normalize_environment_table(doc: Document) -> None:
    style_table(
        doc.tables[13],
        widths_cm=ENV_TABLE_WIDTHS,
        body_font_pt=9.5,
        header_font_pt=10,
        column_alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
        ],
        min_row_height_cm=0.78,
    )


def normalize_test_tables(doc: Document) -> None:
    style_table(
        doc.tables[14],
        widths_cm=TEST_TABLE_WIDTHS,
        body_font_pt=9.2,
        header_font_pt=10,
        column_alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
        transform_text=lambda r, c, t: compact_long_list_text(c, t) if r > 0 else t,
        min_row_height_cm=0.8,
    )
    style_table(
        doc.tables[15],
        widths_cm=CASE_TABLE_WIDTHS,
        body_font_pt=9.2,
        header_font_pt=10,
        column_alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
        ],
        transform_text=lambda r, c, t: compact_long_list_text(c, t) if r > 0 else t,
        min_row_height_cm=0.8,
    )


def normalize_captions(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if CAPTION_RE.match(paragraph.text.strip()):
            normalize_caption_paragraph(paragraph)


def build_doc() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"未找到源文档: {SOURCE_DOCX}")
    shutil.copy2(SOURCE_DOCX, TARGET_DOCX)
    doc = Document(TARGET_DOCX)
    normalize_captions(doc)
    normalize_database_tables(doc)
    normalize_smart_table(doc)
    normalize_environment_table(doc)
    normalize_test_tables(doc)
    doc.save(TARGET_DOCX)
    print(f"生成完成: {TARGET_DOCX}")


if __name__ == "__main__":
    build_doc()
