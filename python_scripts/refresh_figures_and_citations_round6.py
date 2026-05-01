from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import audit_template_and_redraw_with_drawio_round5 as r5


ROOT = Path(__file__).resolve().parents[1]
WORKING_DRAFT = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿工作稿.docx"
BACKUP_PATH = ROOT / "Existing Thesis Draft" / f"校园物资智能管理系统设计与实现-定稿工作稿.docx.bak.{datetime.now():%Y%m%d-%H%M%S}-figure-citation-unify"
REPORT_PATH = ROOT / "output" / "doc" / "校园物资智能管理系统设计与实现-改稿说明-第六轮图风格统一与文献落点补引.md"


def rect(node_id: str, label: str, x: int, y: int, size: str = "xl", font_size: int = 16):
    return r5.bw_node(node_id, label, "service", x, y, "shape=rectangle;rounded=0", size=size, font_size=font_size)


def proc(node_id: str, label: str, x: int, y: int, size: str = "xl", font_size: int = 16):
    return r5.bw_node(node_id, label, "process", x, y, size=size, font_size=font_size)


def term(node_id: str, label: str, x: int, y: int, size: str = "xl", font_size: int = 16):
    return r5.bw_node(node_id, label, "terminal", x, y, size=size, font_size=font_size)


def decision(node_id: str, label: str, x: int, y: int, size: str = "large", font_size: int = 16):
    return r5.bw_node(node_id, label, "decision", x, y, size=size, font_size=font_size)


def edge(src: str, dst: str, label: str | None = None, end_arrow: str = "block"):
    item = {
        "from": src,
        "to": dst,
        "type": "primary",
        "style": {
            "strokeColor": "#000000",
            "endArrow": end_arrow,
            "strokeWidth": 1.4,
            "fontColor": "#000000",
            "fontSize": 12,
        },
    }
    if label:
        item["label"] = label
    return item


def make_architecture_spec():
    nodes = [
        rect("browser", "浏览器访问", 470, 40),
        rect("frontend", "前端表示层", 470, 180),
        rect("router", "路由与状态", 470, 330),
        rect("gateway", "接口访问层", 470, 490),
        rect("controller", "控制层", 120, 690),
        rect("service", "服务层", 470, 690),
        rect("persistence", "数据访问层", 820, 690),
        rect("db", "MySQL 数据库", 470, 890),
    ]
    edges = [
        edge("browser", "frontend"),
        edge("frontend", "router"),
        edge("router", "gateway"),
        edge("gateway", "controller"),
        edge("gateway", "service"),
        edge("gateway", "persistence"),
        edge("controller", "service"),
        edge("service", "persistence"),
        edge("persistence", "db"),
    ]
    return {
        "meta": {
            "profile": "academic-paper",
            "theme": "academic",
            "layout": "vertical",
            "routing": "orthogonal",
            "canvas": "980x1080",
            "title": "系统总体架构图",
            "description": "校园物资智能管理系统的 B/S 分层架构",
        },
        "nodes": nodes,
        "edges": edges,
    }


def make_modules_spec():
    nodes = [
        rect("core", "校园物资智能管理系统", 470, 390, font_size=17),
        rect("auth", "认证与权限", 110, 110),
        rect("base", "基础数据", 470, 90),
        rect("inventory", "仓储库存", 830, 110),
        rect("apply", "申领审批", 110, 390),
        rect("transfer", "调拨协同", 830, 390),
        rect("warning", "预警与补货", 110, 690),
        rect("analytics", "统计分析", 470, 730),
        rect("support", "系统支撑", 830, 690),
    ]
    edges = [
        edge("core", "auth", end_arrow="none"),
        edge("core", "base", end_arrow="none"),
        edge("core", "inventory", end_arrow="none"),
        edge("core", "apply", end_arrow="none"),
        edge("core", "transfer", end_arrow="none"),
        edge("core", "warning", end_arrow="none"),
        edge("core", "analytics", end_arrow="none"),
        edge("core", "support", end_arrow="none"),
    ]
    return {
        "meta": {
            "profile": "academic-paper",
            "theme": "academic",
            "layout": "star",
            "routing": "straight",
            "canvas": "980x920",
            "title": "系统功能模块图",
            "description": "校园物资智能管理系统的功能模块划分",
        },
        "nodes": nodes,
        "edges": edges,
    }


def make_apply_flow_spec():
    nodes = [
        term("start", "创建申领单", 330, 40),
        proc("draft", "保存草稿", 330, 145),
        proc("submit", "提交申领单", 330, 250),
        decision("urgent", "紧急等级≥2？", 330, 360),
        proc("approve", "审批人审核", 110, 500),
        proc("fast", "快速审批", 550, 500),
        decision("pass", "审批通过？", 330, 635),
        proc("reject", "驳回并记录意见", 580, 635),
        proc("stockout", "库存匹配\n批次出库", 330, 770),
        proc("receive", "部门用户签收", 330, 885),
        term("end", "状态为 RECEIVED", 330, 995),
    ]
    edges = [
        edge("start", "draft"),
        edge("draft", "submit"),
        edge("submit", "urgent"),
        edge("urgent", "approve", "否"),
        edge("urgent", "fast", "是"),
        edge("approve", "pass"),
        edge("fast", "pass"),
        edge("pass", "reject", "否"),
        edge("pass", "stockout", "是"),
        edge("stockout", "receive"),
        edge("receive", "end"),
    ]
    return {
        "meta": {
            "profile": "academic-paper",
            "theme": "academic",
            "layout": "vertical",
            "routing": "orthogonal",
            "canvas": "760x1060",
            "title": "申领审批闭环流程图",
            "description": "部门用户、审批人和仓库管理员围绕申领单状态推进的闭环流程",
        },
        "nodes": nodes,
        "edges": edges,
    }


def make_transfer_flow_spec():
    nodes = [
        term("start", "创建调拨单", 330, 40),
        proc("recommend", "选择目标仓\n参考候选来源仓", 330, 150),
        proc("submit", "提交调拨单", 330, 270),
        proc("approve", "审批人审核", 330, 390),
        decision("pass", "审批通过？", 330, 515),
        proc("reject", "驳回并结束", 580, 515),
        proc("execute", "执行调拨\n同步两端库存", 330, 660),
        proc("receive", "调入仓签收", 330, 790),
        term("end", "状态为 RECEIVED", 330, 910),
    ]
    edges = [
        edge("start", "recommend"),
        edge("recommend", "submit"),
        edge("submit", "approve"),
        edge("approve", "pass"),
        edge("pass", "reject", "否"),
        edge("pass", "execute", "是"),
        edge("execute", "receive"),
        edge("receive", "end"),
    ]
    return {
        "meta": {
            "profile": "academic-paper",
            "theme": "academic",
            "layout": "vertical",
            "routing": "orthogonal",
            "canvas": "760x980",
            "title": "调拨执行流程图",
            "description": "调拨单从创建、提交、审批到执行和签收的状态推进流程",
        },
        "nodes": nodes,
        "edges": edges,
    }


def make_warning_flow_spec():
    nodes = [
        term("start", "触发定时扫描", 330, 40),
        proc("scan", "检查库存与批次", 330, 150),
        decision("risk", "发现异常？", 330, 280),
        proc("record", "生成预警记录", 330, 410),
        proc("view", "查看待处理预警", 330, 535),
        proc("handle", "填写说明\n更新状态", 330, 660),
        proc("notify", "按需发送通知", 330, 785),
        term("end", "进入统计分析", 330, 910),
        proc("noop", "无预警\n结束扫描", 580, 280),
    ]
    edges = [
        edge("start", "scan"),
        edge("scan", "risk"),
        edge("risk", "record", "是"),
        edge("risk", "noop", "否"),
        edge("record", "view"),
        edge("view", "handle"),
        edge("handle", "notify"),
        edge("notify", "end"),
    ]
    return {
        "meta": {
            "profile": "academic-paper",
            "theme": "academic",
            "layout": "vertical",
            "routing": "orthogonal",
            "canvas": "760x980",
            "title": "预警处置流程图",
            "description": "定时扫描、预警生成、人工处置与通知联动流程图",
        },
        "nodes": nodes,
        "edges": edges,
    }


def make_auth_flow_spec():
    nodes = [
        term("login", "提交账号密码", 330, 40),
        proc("verify", "校验用户信息", 330, 145),
        proc("issue", "签发令牌", 330, 255),
        proc("init", "保存令牌并加载菜单", 330, 380),
        proc("api", "携带 Token 访问接口", 330, 500),
        decision("unauth", "收到 401？", 330, 625),
        proc("refresh", "调用刷新接口", 330, 755),
        decision("rotated", "刷新成功？", 330, 885),
        proc("retry", "更新令牌并重放请求", 120, 1010),
        proc("logout", "清理令牌并返回登录页", 540, 1010),
        term("done", "继续访问页面", 120, 1125),
    ]
    edges = [
        edge("login", "verify"),
        edge("verify", "issue"),
        edge("issue", "init"),
        edge("init", "api"),
        edge("api", "unauth"),
        edge("unauth", "refresh", "是"),
        edge("unauth", "done", "否"),
        edge("refresh", "rotated"),
        edge("rotated", "retry", "是"),
        edge("rotated", "logout", "否"),
        edge("retry", "done"),
    ]
    return {
        "meta": {
            "profile": "academic-paper",
            "theme": "academic",
            "layout": "vertical",
            "routing": "orthogonal",
            "canvas": "760x1180",
            "title": "登录认证与令牌续签流程图",
            "description": "登录、令牌刷新和失败回退链路",
        },
        "nodes": nodes,
        "edges": edges,
    }


def make_transfer_recommend_flow_spec():
    nodes = [
        term("start", "发起调拨需求", 330, 40),
        proc("input", "输入目标校区\n物资与数量", 330, 145),
        proc("calc", "计算最短路径", 330, 255),
        proc("candidate", "筛选候选来源仓", 330, 365),
        proc("rank", "按距离排序\n返回推荐结果", 330, 475),
        proc("submit", "选择来源仓\n提交调拨单", 330, 600),
        decision("approve", "审批通过？", 330, 730),
        proc("reject", "驳回并结束", 580, 730),
        proc("execute", "执行调拨\n迁移批次", 330, 870),
        proc("receive", "调入仓签收", 330, 985),
        term("end", "状态为 RECEIVED", 330, 1095),
    ]
    edges = [
        edge("start", "input"),
        edge("input", "calc"),
        edge("calc", "candidate"),
        edge("candidate", "rank"),
        edge("rank", "submit"),
        edge("submit", "approve"),
        edge("approve", "reject", "否"),
        edge("approve", "execute", "是"),
        edge("execute", "receive"),
        edge("receive", "end"),
    ]
    return {
        "meta": {
            "profile": "academic-paper",
            "theme": "academic",
            "layout": "vertical",
            "routing": "orthogonal",
            "canvas": "760x1160",
            "title": "调拨执行与推荐流程图",
            "description": "调拨推荐、审批与执行流程图",
        },
        "nodes": nodes,
        "edges": edges,
    }


def update_body_paragraph(paragraph, text: str):
    r5.clear_paragraph_content(paragraph)
    paragraph.style = "正文章节内容"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    r5.set_run_fonts(run, "宋体", "Times New Roman", 12, bold=False)


def render_figures():
    figure_specs = {
        "fig_3_1_architecture": make_architecture_spec(),
        "fig_3_2_modules": make_modules_spec(),
        "fig_2_1_apply_flow": make_apply_flow_spec(),
        "fig_2_2_transfer_flow": make_transfer_flow_spec(),
        "fig_2_3_warning_flow": make_warning_flow_spec(),
        "fig_4_1_auth_flow": make_auth_flow_spec(),
        "fig_4_2_transfer_recommend_flow": make_transfer_recommend_flow_spec(),
    }
    rendered = {}
    for stem, spec in figure_specs.items():
        rendered[stem] = r5.render_bundle(spec, stem)
    return rendered


def replace_caption_figure(doc: Document, caption_text: str, image_path: Path):
    caption = r5.find_paragraph(doc, caption_text)
    picture = doc.paragraphs[r5.paragraph_index(doc, caption) - 1]
    r5.replace_picture_paragraph(picture, image_path)
    r5.set_caption(caption, caption_text)


def collect_used_refs(doc: Document) -> list[int]:
    started = False
    used: set[int] = set()
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt == "1 绪论":
            started = True
        if txt == "参考文献":
            break
        if not started:
            continue
        for m in re.finditer(r"\[(\d+)\]", txt):
            used.add(int(m.group(1)))
    return sorted(used)


def main():
    shutil.copy2(WORKING_DRAFT, BACKUP_PATH)

    figures = render_figures()

    doc = Document(WORKING_DRAFT)
    replace_caption_figure(doc, "图3-1 系统总体架构图", figures["fig_3_1_architecture"])
    replace_caption_figure(doc, "图3-2 系统功能模块图", figures["fig_3_2_modules"])
    replace_caption_figure(doc, "图2-1 申领审批闭环流程图", figures["fig_2_1_apply_flow"])
    replace_caption_figure(doc, "图2-2 调拨执行流程图", figures["fig_2_2_transfer_flow"])
    replace_caption_figure(doc, "图2-3 预警处置流程图", figures["fig_2_3_warning_flow"])
    replace_caption_figure(doc, "图4-1 登录认证与令牌续签流程图", figures["fig_4_1_auth_flow"])
    replace_caption_figure(doc, "图4-5 调拨执行与推荐流程图", figures["fig_4_2_transfer_recommend_flow"])

    auth_para = None
    for para in doc.paragraphs:
        if para.text.strip().startswith("授权控制分布在两个层面。"):
            auth_para = para
            break
    if auth_para is None:
        raise ValueError("Failed to locate authorization paragraph for citation update")
    update_body_paragraph(
        auth_para,
        "授权控制分布在两个层面。其一是 SecurityConfig 打开的方法级安全能力和控制器上的 PreAuthorize 表达式，其二是 AuthService.buildMenusByRole 返回的菜单集合。接口层和导航层一起工作后，系统管理员、仓库管理员、审批人和部门用户才真正表现出不同的页面与操作边界[3]。",
    )

    r5.set_all_text_black(doc)
    doc.save(WORKING_DRAFT)

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    used_refs = collect_used_refs(doc)
    REPORT_PATH.write_text(
        "\n".join(
            [
                "# 校园物资智能管理系统论文改稿说明（第六轮图风格统一与文献落点补引）",
                "",
                "## 本轮处理范围",
                f"- working draft：`{WORKING_DRAFT.relative_to(ROOT)}`",
                f"- fresh backup：`{BACKUP_PATH.relative_to(ROOT)}`",
                "- 不改正文结构，不改参考文献表顺序，只统一工程图视觉风格并补正文文献落点。",
                "",
                "## 图件处理",
                "- 使用 `$drawio` 重画 `图3-1 系统总体架构图`、`图3-2 系统功能模块图`，统一为黑白学术风格。",
                "- 同步重画 `图2-1`、`图2-2`、`图2-3`、`图4-1`、`图4-5` 五张流程图，收紧画布留白并与 E-R 图保持同一套黑白线条风格。",
                "- drawio sidecar 已写入 `output/doc/figures/drawio/`，便于后续继续微调。",
                "",
                "## 文献引用处理",
                "- 审核正文后确认，原来只有 `[3] 基于角色的动态权限管理系统的应用设计` 未在正文章节中被引用。",
                "- 本轮将 `[3]` 补入 `3.4.2 认证与授权设计` 的授权边界论述段，与角色菜单、接口权限和方法级安全能力的内容直接对应。",
                f"- 处理后，正文章节中已覆盖引用的参考文献编号为：`{', '.join(str(i) for i in used_refs)}`。",
                "",
                "## 交付说明",
                "- 如需终版打印，仍建议在 Word 中执行一次 `Ctrl+A`、`F9` 做目录和题注刷新，再做分页终检。",
            ]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
