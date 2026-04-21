from __future__ import annotations

import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
import xml.etree.ElementTree as ET

import yaml
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from playwright.sync_api import sync_playwright


ROOT = Path(__file__).resolve().parents[1]
WORKING_DRAFT = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿工作稿.docx"
ORIGINAL_DRAFT = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-初稿.docx"
FIGURE_DIR = ROOT / "output" / "doc" / "figures"
DRAWIO_BUNDLE_DIR = FIGURE_DIR / "drawio"
REPORT_PATH = ROOT / "output" / "doc" / "校园物资智能管理系统设计与实现-改稿说明-第五轮模板审查与drawio重绘.md"
BACKUP_PATH = ROOT / "Existing Thesis Draft" / f"校园物资智能管理系统设计与实现-定稿工作稿.docx.bak.{datetime.now():%Y%m%d-%H%M%S}-template-drawio-review"
DRAWIO_CLI = ROOT / ".codex" / "skills" / "drawio" / "scripts" / "cli.js"
TEMPLATE_DOC = ROOT / "muban" / "7 软件学院本科毕业设计（论文）模板 2026.docx"
GUIDE_DOC = ROOT / "muban" / "4 软件学院本科毕业设计（论文）撰写规范.docx"

TITLE = "校园物资智能管理系统设计与实现"


def find_paragraph(doc: Document, *candidates: str):
    wanted = {c.strip() for c in candidates}
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in wanted:
            return paragraph
    raise ValueError(f"Paragraph not found: {candidates}")


def find_paragraph_or_none(doc: Document, *candidates: str):
    wanted = {c.strip() for c in candidates}
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in wanted:
            return paragraph
    return None


def paragraph_index(doc: Document, paragraph) -> int:
    for idx, item in enumerate(doc.paragraphs):
        if item._element is paragraph._element:
            return idx
    raise ValueError("Paragraph is not part of the current document paragraph list")


def clear_paragraph_content(paragraph) -> None:
    element = paragraph._element
    for child in list(element):
        if child.tag == qn("w:pPr"):
            continue
        element.remove(child)


def set_run_fonts(run, east_asia: str, ascii_font: str, size_pt: float, bold: bool | None = None):
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def configure_style(style, east_asia: str, ascii_font: str, size_pt: float, bold: bool = False,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent_cm: float | None = None,
                    line_spacing: float = 1.5, page_break_before: bool = False):
    style.font.name = ascii_font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia)

    pfmt = style.paragraph_format
    pfmt.alignment = alignment
    pfmt.space_before = Pt(0)
    pfmt.space_after = Pt(0)
    pfmt.line_spacing = line_spacing
    pfmt.first_line_indent = Cm(first_line_indent_cm) if first_line_indent_cm is not None else None
    pfmt.page_break_before = page_break_before


def mark_page_break_before(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    page_break = p_pr.find(qn("w:pageBreakBefore"))
    if page_break is None:
        page_break = OxmlElement("w:pageBreakBefore")
        p_pr.append(page_break)
    page_break.set(qn("w:val"), "1")


def set_outline_level(paragraph, level: int):
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def set_heading_text(paragraph, text: str, style_name: str, level: int):
    clear_paragraph_content(paragraph)
    paragraph.style = style_name
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    mark_page_break_before(paragraph)
    set_outline_level(paragraph, level)
    run = paragraph.add_run(text)
    set_run_fonts(run, "黑体", "Times New Roman", 16, bold=False)


def set_caption(paragraph, text: str):
    clear_paragraph_content(paragraph)
    paragraph.style = "图片标题"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_fonts(run, "宋体", "Times New Roman", 10.5, bold=True)


def replace_picture_paragraph(paragraph, image_path: Path):
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(6.1))


def add_body_before(anchor, text: str):
    paragraph = anchor.insert_paragraph_before(text, style="正文章节内容")
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_picture_before(anchor, image_path: Path, caption: str):
    picture_paragraph = anchor.insert_paragraph_before()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(image_path), width=Inches(6.1))
    caption_paragraph = anchor.insert_paragraph_before(caption, style="图片标题")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return picture_paragraph, caption_paragraph


def ensure_lead_before_picture(doc: Document, caption_text: str, lead_text: str):
    caption = find_paragraph(doc, caption_text)
    picture = doc.paragraphs[paragraph_index(doc, caption) - 1]
    prev_idx = paragraph_index(doc, picture) - 1
    if prev_idx >= 0 and doc.paragraphs[prev_idx].text.strip() == lead_text.strip():
        return
    add_body_before(picture, lead_text)


def set_all_text_black(doc: Document):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = False
    for section in doc.sections:
        for paragraph in section.header.paragraphs + section.footer.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.underline = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.underline = False


def configure_document_styles(doc: Document):
    configure_style(doc.styles["name_out_lvl_1"], "黑体", "Times New Roman", 16, False, WD_ALIGN_PARAGRAPH.CENTER)
    configure_style(doc.styles["1级标题-正文章节"], "黑体", "Times New Roman", 16, False, WD_ALIGN_PARAGRAPH.LEFT, None, 1.5, True)
    configure_style(doc.styles["2级标题-正文章节"], "黑体", "Times New Roman", 15, False, WD_ALIGN_PARAGRAPH.LEFT)
    configure_style(doc.styles["3级标题-正文章节"], "黑体", "Times New Roman", 14, False, WD_ALIGN_PARAGRAPH.LEFT)
    configure_style(doc.styles["正文章节内容"], "宋体", "Times New Roman", 12, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 0.74)
    configure_style(doc.styles["中文摘要内容"], "宋体", "Times New Roman", 12, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 0.74)
    configure_style(doc.styles["英文摘要内容"], "宋体", "Times New Roman", 12, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 0.74)
    configure_style(doc.styles["图片标题"], "宋体", "Times New Roman", 10.5, True, WD_ALIGN_PARAGRAPH.CENTER)
    configure_style(doc.styles["表格标题"], "宋体", "Times New Roman", 10.5, True, WD_ALIGN_PARAGRAPH.CENTER)

    toc1 = doc.styles["toc 1"].paragraph_format
    toc1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc1.space_before = Pt(0)
    toc1.space_after = Pt(0)
    toc1.line_spacing = 1.5
    toc2 = doc.styles["toc 2"].paragraph_format
    toc2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc2.space_before = Pt(0)
    toc2.space_after = Pt(0)
    toc2.line_spacing = 1.5


def configure_sections(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)
        section.start_type = WD_SECTION_START.NEW_PAGE

    if len(doc.sections) > 1:
        body_section = doc.sections[1]
        header = body_section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clear_paragraph_content(header)
        header_run = header.add_run(TITLE)
        set_run_fonts(header_run, "宋体", "Times New Roman", 9, bold=False)

        footer = body_section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            set_run_fonts(run, "宋体", "Times New Roman", 9, bold=False)


def svg_dimensions(svg_path: Path) -> tuple[int, int]:
    root = ET.fromstring(svg_path.read_text(encoding="utf-8"))
    width = root.attrib.get("width", "1600")
    height = root.attrib.get("height", "1000")

    def parse_dim(value: str) -> int:
        digits = "".join(ch for ch in value if ch.isdigit())
        return int(digits) if digits else 1000

    return parse_dim(width), parse_dim(height)


def svg_to_png(svg_path: Path, png_path: Path) -> None:
    width, height = svg_dimensions(svg_path)
    html = (
        "<!doctype html><html><head><meta charset='utf-8'>"
        "<style>html,body{margin:0;background:#fff;}img{display:block;margin:0 auto;}</style>"
        f"</head><body><img src='{svg_path.as_uri()}' width='{width}' height='{height}'></body></html>"
    )
    with tempfile.TemporaryDirectory() as tmp:
        html_path = Path(tmp) / "preview.html"
        html_path.write_text(html, encoding="utf-8")
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page(viewport={"width": width + 80, "height": min(height + 80, 1400)}, device_scale_factor=2)
            page.goto(html_path.as_uri(), wait_until="load")
            page.screenshot(path=str(png_path), full_page=True)
            browser.close()


def render_bundle(spec: dict, stem: str) -> Path:
    DRAWIO_BUNDLE_DIR.mkdir(parents=True, exist_ok=True)
    spec_path = DRAWIO_BUNDLE_DIR / f"{stem}.yaml"
    svg_path = DRAWIO_BUNDLE_DIR / f"{stem}.svg"
    png_path = FIGURE_DIR / f"{stem}.png"
    spec_path.write_text(yaml.safe_dump(spec, allow_unicode=True, sort_keys=False), encoding="utf-8")

    subprocess.run(
        ["node", str(DRAWIO_CLI), str(spec_path), str(svg_path), "--validate", "--write-sidecars"],
        cwd=str(ROOT),
        text=True,
        encoding="utf-8",
        errors="ignore",
        check=True,
    )
    svg_to_png(svg_path, png_path)
    return png_path


def bw_node(node_id: str, label: str, node_type: str, x: int, y: int, shape: str | None = None, size: str = "medium", font_size: int = 13):
    style = {
        "fillColor": "#FFFFFF",
        "strokeColor": "#000000",
        "fontColor": "#000000",
        "strokeWidth": 1.6,
        "fontSize": font_size,
    }
    if shape:
        style["shape"] = shape
    return {
        "id": node_id,
        "label": label,
        "type": node_type,
        "size": size,
        "position": {"x": x, "y": y},
        "style": style,
    }


def bw_edge(src: str, dst: str, label: str | None = None, label_pos: str = "center"):
    edge = {
        "from": src,
        "to": dst,
        "type": "primary",
        "style": {
            "strokeColor": "#000000",
            "endArrow": "none",
            "strokeWidth": 1.4,
            "fontColor": "#000000",
            "fontSize": 12,
        },
    }
    if label:
        edge["label"] = label
        edge["labelPosition"] = label_pos
    return edge


def make_rbac_er_spec():
    nodes = [
        bw_node("dept", "部门\\nsys_dept", "service", 120, 260, "shape=rectangle;rounded=0"),
        bw_node("user", "用户\\nsys_user", "service", 700, 250, "shape=rectangle;rounded=0"),
        bw_node("role", "角色\\nsys_role", "service", 1280, 260, "shape=rectangle;rounded=0"),
        bw_node("token", "刷新令牌\\nauth_refresh_token", "service", 360, 680, "shape=rectangle;rounded=0"),
        bw_node("log", "登录日志\\nlogin_log", "service", 1040, 680, "shape=rectangle;rounded=0"),
        bw_node("rel_dept_user", "隶属", "decision", 410, 265, font_size=12),
        bw_node("rel_user_role", "拥有", "decision", 985, 265, font_size=12),
        bw_node("rel_user_token", "签发", "decision", 530, 520, font_size=12),
        bw_node("rel_user_log", "产生", "decision", 860, 520, font_size=12),
    ]
    attrs = [
        ("dept_id", "部门编号", 40, 90), ("dept_name", "部门名称", 150, 40), ("dept_parent", "上级部门", 280, 70), ("dept_ver", "版本号", 290, 150),
        ("user_id", "用户编号", 620, 60), ("user_name", "用户账号", 760, 20), ("user_real", "姓名", 910, 60), ("user_status", "状态", 980, 150), ("user_dept", "dept_id", 570, 150), ("user_role", "role_id", 880, 150),
        ("role_id", "角色编号", 1190, 80), ("role_code", "角色编码", 1310, 40), ("role_name", "角色名称", 1445, 80), ("role_desc", "角色说明", 1510, 165),
        ("token_id", "令牌编号", 160, 540), ("token_user", "user_id", 290, 500), ("token_hash", "token_hash", 510, 500), ("token_expire", "expire_at", 600, 570), ("token_revoked", "revoked", 600, 655),
        ("log_user", "user_id", 860, 510), ("log_name", "username", 1035, 475), ("log_ip", "login_ip", 1190, 505), ("log_status", "登录状态", 1245, 595), ("log_time", "login_time", 1150, 675),
    ]
    for node_id, label, x, y in attrs:
        nodes.append(bw_node(node_id, label, "user", x, y, size="small", font_size=11))

    edges = [
        bw_edge("dept", "rel_dept_user", "1", "end"),
        bw_edge("rel_dept_user", "user", "N", "end"),
        bw_edge("role", "rel_user_role", "1", "end"),
        bw_edge("rel_user_role", "user", "N", "end"),
        bw_edge("user", "rel_user_token", "1", "end"),
        bw_edge("rel_user_token", "token", "N", "end"),
        bw_edge("user", "rel_user_log", "1", "end"),
        bw_edge("rel_user_log", "log", "N", "end"),
    ]
    for entity, attr_prefix in [("dept", "dept_"), ("user", "user_"), ("role", "role_"), ("token", "token_"), ("log", "log_")]:
        for attr_id, *_ in [a for a in attrs if a[0].startswith(attr_prefix)]:
            edges.append(bw_edge(entity, attr_id))

    return {
        "meta": {
            "profile": "academic-paper",
            "theme": "academic",
            "layout": "horizontal",
            "routing": "straight",
            "canvas": "1600x980",
            "title": "RBAC 与组织 E-R图",
            "description": "用户、部门、角色、刷新令牌与登录日志的 Chen 风格 E-R 图",
        },
        "nodes": nodes,
        "edges": edges,
    }


def make_inventory_er_spec():
    nodes = [
        bw_node("category", "分类\\nmaterial_category", "service", 90, 250, "shape=rectangle;rounded=0"),
        bw_node("material", "物资档案\\nmaterial_info", "service", 550, 250, "shape=rectangle;rounded=0"),
        bw_node("warehouse", "仓库\\nwarehouse", "service", 1060, 250, "shape=rectangle;rounded=0"),
        bw_node("inventory", "库存\\ninventory", "service", 450, 650, "shape=rectangle;rounded=0"),
        bw_node("batch", "库存批次\\ninventory_batch", "service", 990, 650, "shape=rectangle;rounded=0"),
        bw_node("rel_cat_material", "归类", "decision", 305, 255, font_size=12),
        bw_node("rel_material_inventory", "形成", "decision", 500, 485, font_size=12),
        bw_node("rel_wh_inventory", "存放", "decision", 815, 485, font_size=12),
        bw_node("rel_inventory_batch", "细分", "decision", 720, 650, font_size=12),
    ]
    attrs = [
        ("cat_id", "id", 10, 90), ("cat_name", "category_name", 110, 40), ("cat_remark", "remark", 220, 90),
        ("mat_id", "id", 405, 70), ("mat_code", "material_code", 550, 20), ("mat_name", "material_name", 700, 70), ("mat_safety", "safety_stock", 735, 165), ("mat_spec", "spec/unit", 360, 160),
        ("wh_id", "id", 940, 70), ("wh_name", "warehouse_name", 1075, 20), ("wh_campus", "campus", 1220, 70), ("wh_manager", "manager", 1245, 160),
        ("inv_id", "id", 320, 500), ("inv_material", "material_id", 355, 785), ("inv_wh", "warehouse_id", 520, 785), ("inv_qty", "current_qty", 610, 535), ("inv_locked", "locked_qty", 635, 710),
        ("batch_id", "id", 905, 500), ("batch_no", "batch_no", 1035, 475), ("batch_in", "in_qty", 1175, 520), ("batch_remain", "remain_qty", 1200, 635), ("batch_expire", "expire_date", 1125, 760),
    ]
    for node_id, label, x, y in attrs:
        nodes.append(bw_node(node_id, label, "user", x, y, size="small", font_size=11))

    edges = [
        bw_edge("category", "rel_cat_material", "1", "end"),
        bw_edge("rel_cat_material", "material", "N", "end"),
        bw_edge("material", "rel_material_inventory", "1", "end"),
        bw_edge("rel_material_inventory", "inventory", "N", "end"),
        bw_edge("warehouse", "rel_wh_inventory", "1", "end"),
        bw_edge("rel_wh_inventory", "inventory", "N", "end"),
        bw_edge("inventory", "rel_inventory_batch", "1", "end"),
        bw_edge("rel_inventory_batch", "batch", "N", "end"),
    ]
    for entity, attr_prefix in [("category", "cat_"), ("material", "mat_"), ("warehouse", "wh_"), ("inventory", "inv_"), ("batch", "batch_")]:
        for attr_id, *_ in [a for a in attrs if a[0].startswith(attr_prefix)]:
            edges.append(bw_edge(entity, attr_id))

    return {
        "meta": {
            "profile": "academic-paper",
            "theme": "academic",
            "layout": "horizontal",
            "routing": "straight",
            "canvas": "1500x980",
            "title": "库存与批次 E-R图",
            "description": "分类、物资档案、仓库、库存总表与库存批次之间的 E-R 图",
        },
        "nodes": nodes,
        "edges": edges,
    }


def make_business_er_spec():
    nodes = [
        bw_node("apply", "申领单\\napply_order", "service", 110, 250, "shape=rectangle;rounded=0"),
        bw_node("item", "申领明细\\napply_order_item", "service", 590, 250, "shape=rectangle;rounded=0"),
        bw_node("stockout", "出库单\\nstock_out", "service", 1110, 250, "shape=rectangle;rounded=0"),
        bw_node("transfer", "调拨单\\ntransfer_order", "service", 360, 660, "shape=rectangle;rounded=0"),
        bw_node("warning", "预警记录\\nwarning_record", "service", 840, 660, "shape=rectangle;rounded=0"),
        bw_node("notify", "通知消息\\nnotification", "service", 1290, 660, "shape=rectangle;rounded=0"),
        bw_node("rel_apply_item", "包含", "decision", 350, 255, font_size=12),
        bw_node("rel_apply_stock", "生成", "decision", 850, 255, font_size=12),
        bw_node("rel_transfer_warning", "触发", "decision", 610, 540, font_size=12),
        bw_node("rel_warning_notify", "发送", "decision", 1065, 660, font_size=12),
    ]
    attrs = [
        ("apply_id", "id", 25, 85), ("apply_dept", "dept_id", 110, 35), ("apply_applicant", "applicant_id", 235, 70), ("apply_status", "status", 270, 155),
        ("item_id", "id", 505, 95), ("item_order", "apply_order_id", 590, 40), ("item_material", "material_id", 725, 95), ("item_qty", "apply_qty", 760, 175),
        ("out_id", "id", 1025, 95), ("out_apply", "apply_order_id", 1115, 40), ("out_wh", "warehouse_id", 1255, 95), ("out_operator", "operator_id", 1295, 175),
        ("transfer_id", "id", 265, 500), ("transfer_from", "from_warehouse_id", 320, 810), ("transfer_to", "to_warehouse_id", 475, 810), ("transfer_status", "status", 505, 585),
        ("warn_id", "id", 750, 500), ("warn_type", "warning_type", 845, 470), ("warn_status", "handle_status", 980, 515), ("warn_handler", "handler_id", 1025, 610),
        ("notify_id", "id", 1205, 515), ("notify_title", "title", 1290, 470), ("notify_target", "target_user_id", 1435, 525), ("notify_read", "is_read", 1450, 625),
    ]
    for node_id, label, x, y in attrs:
        nodes.append(bw_node(node_id, label, "user", x, y, size="small", font_size=11))

    edges = [
        bw_edge("apply", "rel_apply_item", "1", "end"),
        bw_edge("rel_apply_item", "item", "N", "end"),
        bw_edge("apply", "rel_apply_stock", "1", "end"),
        bw_edge("rel_apply_stock", "stockout", "1", "end"),
        bw_edge("transfer", "rel_transfer_warning", "1", "end"),
        bw_edge("rel_transfer_warning", "warning", "N", "end"),
        bw_edge("warning", "rel_warning_notify", "1", "end"),
        bw_edge("rel_warning_notify", "notify", "N", "end"),
    ]
    for entity, attr_prefix in [("apply", "apply_"), ("item", "item_"), ("stockout", "out_"), ("transfer", "transfer_"), ("warning", "warn_"), ("notify", "notify_")]:
        for attr_id, *_ in [a for a in attrs if a[0].startswith(attr_prefix)]:
            edges.append(bw_edge(entity, attr_id))

    return {
        "meta": {
            "profile": "academic-paper",
            "theme": "academic",
            "layout": "horizontal",
            "routing": "straight",
            "canvas": "1600x980",
            "title": "业务单据、预警与通知 E-R图",
            "description": "申领单、出库单、调拨单、预警记录与通知消息之间的 E-R 图",
        },
        "nodes": nodes,
        "edges": edges,
    }


def make_apply_flow_spec():
    def node(node_id, label, node_type, x, y):
        return bw_node(node_id, label, node_type, x, y, font_size=12)

    return {
        "meta": {"profile": "academic-paper", "theme": "academic", "layout": "vertical", "routing": "orthogonal", "canvas": "1100x980", "title": "申领审批闭环流程图", "description": "部门用户、审批人和仓库管理员围绕申领单状态推进的闭环流程"},
        "nodes": [
            node("start", "部门用户创建申领单", "terminal", 390, 40),
            node("draft", "保存草稿", "process", 390, 145),
            node("submit", "提交申领单", "process", 390, 255),
            node("urgent", "紧急等级≥2？", "decision", 390, 370),
            node("approve", "审批人审核", "process", 180, 500),
            node("fast", "快速审批", "process", 600, 500),
            node("passed", "审批通过？", "decision", 390, 620),
            node("reject", "驳回并记录意见", "process", 650, 620),
            node("stockout", "仓库管理员结合库存与批次出库", "process", 390, 750),
            node("receive", "部门用户签收", "process", 390, 860),
            node("end", "状态更新为 RECEIVED", "terminal", 390, 965),
        ],
        "edges": [
            {"from": "start", "to": "draft", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "draft", "to": "submit", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "submit", "to": "urgent", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "urgent", "to": "approve", "type": "primary", "label": "否", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "urgent", "to": "fast", "type": "primary", "label": "是", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "approve", "to": "passed", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "fast", "to": "passed", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "passed", "to": "reject", "type": "primary", "label": "否", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "passed", "to": "stockout", "type": "primary", "label": "是", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "stockout", "to": "receive", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "receive", "to": "end", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
        ],
    }


def make_transfer_flow_spec():
    def node(node_id, label, node_type, x, y):
        return bw_node(node_id, label, node_type, x, y, font_size=12)

    return {
        "meta": {"profile": "academic-paper", "theme": "academic", "layout": "vertical", "routing": "orthogonal", "canvas": "1100x980", "title": "调拨执行流程图", "description": "调拨单从创建、提交、审批到执行和签收的状态推进流程"},
        "nodes": [
            node("start", "仓库管理员创建调拨单", "terminal", 390, 40),
            node("recommend", "选择目标仓并参考候选来源仓", "process", 390, 150),
            node("submit", "提交调拨单", "process", 390, 265),
            node("approve", "审批人审核", "process", 390, 380),
            node("pass", "审批通过？", "decision", 390, 495),
            node("reject", "驳回并结束", "process", 660, 495),
            node("execute", "执行调拨并同步两端库存", "process", 390, 635),
            node("receive", "调入仓确认签收", "process", 390, 760),
            node("end", "状态更新为 RECEIVED", "terminal", 390, 880),
        ],
        "edges": [
            {"from": "start", "to": "recommend", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "recommend", "to": "submit", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "submit", "to": "approve", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "approve", "to": "pass", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "pass", "to": "reject", "type": "primary", "label": "否", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "pass", "to": "execute", "type": "primary", "label": "是", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "execute", "to": "receive", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "receive", "to": "end", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
        ],
    }


def make_warning_flow_spec():
    def node(node_id, label, node_type, x, y):
        return bw_node(node_id, label, node_type, x, y, font_size=12)

    return {
        "meta": {"profile": "academic-paper", "theme": "academic", "layout": "vertical", "routing": "orthogonal", "canvas": "1200x980", "title": "预警处置流程图", "description": "定时扫描、预警生成、人工处置与通知联动的流程图"},
        "nodes": [
            node("start", "定时任务触发扫描", "terminal", 420, 40),
            node("scan", "检查库存与批次数据", "process", 420, 155),
            node("risk", "发现低库存 / 积压 / 临期 / 过期 / 异常领用？", "decision", 420, 290),
            node("record", "生成 warning_record", "process", 420, 430),
            node("view", "管理员查看待处理预警", "process", 420, 560),
            node("handle", "填写处理说明并更新 HANDLE_STATUS", "process", 420, 690),
            node("notify", "按需要写入 notification", "process", 420, 820),
            node("end", "结束并进入统计分析", "terminal", 420, 930),
            node("noop", "无预警，结束本轮扫描", "process", 780, 290),
        ],
        "edges": [
            {"from": "start", "to": "scan", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "scan", "to": "risk", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "risk", "to": "noop", "type": "primary", "label": "否", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "risk", "to": "record", "type": "primary", "label": "是", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "record", "to": "view", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "view", "to": "handle", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "handle", "to": "notify", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "notify", "to": "end", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
        ],
    }


def make_auth_flow_spec():
    def node(node_id, label, node_type, x, y):
        return bw_node(node_id, label, node_type, x, y, font_size=12)

    return {
        "meta": {"profile": "academic-paper", "theme": "academic", "layout": "vertical", "routing": "orthogonal", "canvas": "1280x1180", "title": "登录认证与令牌续签流程图", "description": "登录、刷新令牌轮换与失败回退链路的流程图"},
        "nodes": [
            node("login", "用户提交账号密码", "terminal", 420, 40),
            node("verify", "AuthService 校验用户信息", "process", 420, 150),
            node("issue", "签发 access token + refresh token", "process", 420, 260),
            node("init", "前端保存令牌并加载用户 / 菜单", "process", 420, 375),
            node("api", "携带 Bearer Token 访问接口", "process", 420, 490),
            node("unauth", "返回 code=401？", "decision", 420, 610),
            node("refresh", "调用 /api/auth/refresh", "process", 420, 745),
            node("rotated", "刷新成功并轮换 refresh token？", "decision", 420, 870),
            node("retry", "更新本地令牌并重放原请求", "process", 190, 995),
            node("logout", "清理本地令牌并回到登录页", "process", 660, 995),
            node("done", "继续访问业务页面", "terminal", 190, 1110),
        ],
        "edges": [
            {"from": "login", "to": "verify", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "verify", "to": "issue", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "issue", "to": "init", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "init", "to": "api", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "api", "to": "unauth", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "unauth", "to": "refresh", "type": "primary", "label": "是", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "unauth", "to": "done", "type": "primary", "label": "否", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "refresh", "to": "rotated", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "rotated", "to": "retry", "type": "primary", "label": "是", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "rotated", "to": "logout", "type": "primary", "label": "否", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "retry", "to": "done", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
        ],
    }


def make_transfer_recommend_flow_spec():
    def node(node_id, label, node_type, x, y):
        return bw_node(node_id, label, node_type, x, y, font_size=12)

    return {
        "meta": {"profile": "academic-paper", "theme": "academic", "layout": "vertical", "routing": "orthogonal", "canvas": "1280x1180", "title": "调拨执行与推荐流程图", "description": "目标校区推荐候选仓、审批与执行联动流程图"},
        "nodes": [
            node("start", "仓库管理员发起调拨需求", "terminal", 420, 40),
            node("input", "输入目标校区、物资与数量", "process", 420, 150),
            node("calc", "DijkstraUtil 计算最短路径", "process", 420, 265),
            node("candidate", "筛选库存充足的候选来源仓", "process", 420, 380),
            node("rank", "按距离排序返回推荐结果", "process", 420, 495),
            node("submit", "选择来源仓并提交调拨单", "process", 420, 610),
            node("approve", "审批通过？", "decision", 420, 735),
            node("reject", "驳回并结束", "process", 700, 735),
            node("execute", "执行调出、调入与批次迁移", "process", 420, 875),
            node("receive", "调入仓签收", "process", 420, 995),
            node("end", "状态更新为 RECEIVED", "terminal", 420, 1105),
        ],
        "edges": [
            {"from": "start", "to": "input", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "input", "to": "calc", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "calc", "to": "candidate", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "candidate", "to": "rank", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "rank", "to": "submit", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "submit", "to": "approve", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "approve", "to": "reject", "type": "primary", "label": "否", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "approve", "to": "execute", "type": "primary", "label": "是", "style": {"strokeColor": "#000000", "endArrow": "block", "fontColor": "#000000"}},
            {"from": "execute", "to": "receive", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
            {"from": "receive", "to": "end", "type": "primary", "style": {"strokeColor": "#000000", "endArrow": "block"}},
        ],
    }


def generate_all_figures():
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    figure_map = {
        "fig_3_3_rbac_er": make_rbac_er_spec(),
        "fig_3_4_inventory_er": make_inventory_er_spec(),
        "fig_3_5_business_er": make_business_er_spec(),
        "fig_2_1_apply_flow": make_apply_flow_spec(),
        "fig_2_2_transfer_flow": make_transfer_flow_spec(),
        "fig_2_3_warning_flow": make_warning_flow_spec(),
        "fig_4_1_auth_flow": make_auth_flow_spec(),
        "fig_4_2_transfer_recommend_flow": make_transfer_recommend_flow_spec(),
    }
    rendered = {}
    for stem, spec in figure_map.items():
        rendered[stem] = render_bundle(spec, stem)
    return rendered


def replace_existing_caption_figure(doc: Document, old_caption: str, new_caption: str, image_path: Path):
    caption_paragraph = find_paragraph(doc, old_caption, new_caption)
    replace_picture_paragraph(doc.paragraphs[paragraph_index(doc, caption_paragraph) - 1], image_path)
    set_caption(caption_paragraph, new_caption)


def rename_caption_only(doc: Document, old_caption: str, new_caption: str):
    caption_paragraph = find_paragraph(doc, old_caption, new_caption)
    set_caption(caption_paragraph, new_caption)


def insert_new_figure_before(doc: Document, anchor_text: str, lead_text: str, caption: str, image_path: Path):
    existing_caption = find_paragraph_or_none(doc, caption)
    if existing_caption is not None:
        replace_picture_paragraph(doc.paragraphs[paragraph_index(doc, existing_caption) - 1], image_path)
        set_caption(existing_caption, caption)
        return

    anchor = find_paragraph(doc, anchor_text)
    add_body_before(anchor, lead_text)
    add_picture_before(anchor, image_path, caption)


def update_figures(doc: Document, figures: dict[str, Path]):
    replace_existing_caption_figure(doc, "图4-1 申领审批闭环流程图", "图2-1 申领审批闭环流程图", figures["fig_2_1_apply_flow"])
    ensure_lead_before_picture(doc, "图2-1 申领审批闭环流程图", "申领审批的主要参与者、状态推进与出库签收关系可概括为图2-1所示。")

    insert_new_figure_before(doc, "2.4.3 预警处置与统计分析用例分析", "调拨执行用例中从推荐候选仓到签收完成的状态推进关系可概括为图2-2所示。", "图2-2 调拨执行流程图", figures["fig_2_2_transfer_flow"])
    insert_new_figure_before(doc, "3 系统总体设计", "预警生成、人工处置与通知联动之间的关系可概括为图2-3所示。", "图2-3 预警处置流程图", figures["fig_2_3_warning_flow"])
    insert_new_figure_before(doc, "为了直观展示系统登录入口与统一认证界面，本节补充运行截图如图4-3所示。", "认证授权与登录态维持的请求处理链路如图4-1所示。", "图4-1 登录认证与令牌续签流程图", figures["fig_4_1_auth_flow"])

    rename_caption_only(doc, "图4-3 系统登录界面", "图4-2 系统登录界面")
    rename_caption_only(doc, "图4-4 部门用户申领界面", "图4-3 部门用户申领界面")
    rename_caption_only(doc, "图4-5 审批与出库执行界面", "图4-4 审批与出库执行界面")
    replace_existing_caption_figure(doc, "图4-2 调拨执行与推荐流程图", "图4-5 调拨执行与推荐流程图", figures["fig_4_2_transfer_recommend_flow"])
    replace_existing_caption_figure(doc, "图3-3 RBAC 与组织实体图", "图3-3 RBAC 与组织 E-R图", figures["fig_3_3_rbac_er"])
    replace_existing_caption_figure(doc, "图3-4 库存与批次实体图", "图3-4 库存与批次 E-R图", figures["fig_3_4_inventory_er"])
    replace_existing_caption_figure(doc, "图3-5 业务单据、预警与通知实体图", "图3-5 业务单据、预警与通知 E-R图", figures["fig_3_5_business_er"])

    lead_updates = {
        "为了直观展示系统登录入口与统一认证界面，本节补充运行截图如图4-3所示。": "为了直观展示系统登录入口与统一认证界面，本节补充运行截图如图4-2所示。",
        "部门用户发起物资申请时，可以在申领页面直接录入原因、场景和物资明细，运行界面如图4-4所示。": "部门用户发起物资申请时，可以在申领页面直接录入原因、场景和物资明细，运行界面如图4-3所示。",
        "审批人与仓库管理员分别在审批列表和出库页面完成业务闭环，组合运行界面如图4-5所示。": "审批人与仓库管理员分别在审批列表和出库页面完成业务闭环，组合运行界面如图4-4所示。",
    }
    for old_text, new_text in lead_updates.items():
        paragraph = find_paragraph_or_none(doc, old_text, new_text)
        if paragraph is not None:
            clear_paragraph_content(paragraph)
            paragraph.style = "正文章节内容"
            run = paragraph.add_run(new_text)
            set_run_fonts(run, "宋体", "Times New Roman", 12, bold=False)


def normalize_headings(doc: Document):
    heading_map = {
        "1  绪论": "1 绪论",
        "2  系统需求分析": "2 系统需求分析",
        "3  系统总体设计": "3 系统总体设计",
        "4  系统详细设计与实现": "4 系统详细设计与实现",
        "5  系统测试": "5 系统测试",
    }
    for old_text, new_text in heading_map.items():
        paragraph = find_paragraph(doc, old_text, new_text)
        set_heading_text(paragraph, new_text, "1级标题-正文章节", 0)

    for text in ["摘    要", "ABSTRACT"]:
        paragraph = find_paragraph(doc, text)
        paragraph.style = "name_out_lvl_1"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_fonts(run, "黑体", "Times New Roman", 16, bold=False)


def write_report(image_count: int, table_count: int):
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(
        "\n".join(
            [
                "# 校园物资智能管理系统论文改稿说明（第五轮模板审查与 drawio 重绘）",
                "",
                "## 本轮依据",
                f"- 学校模板：`{TEMPLATE_DOC.relative_to(ROOT)}`",
                f"- 撰写规范：`{GUIDE_DOC.relative_to(ROOT)}`",
                f"- working draft：`{WORKING_DRAFT.relative_to(ROOT)}`",
                f"- protected original：`{ORIGINAL_DRAFT.relative_to(ROOT)}`",
                f"- fresh backup：`{BACKUP_PATH.relative_to(ROOT)}`",
                "",
                "## 模板审查后修正的热点",
                "- 将正文第 1 至第 5 章标题统一改为“三号黑体、左顶格、1.5 倍行距、章节另起页”，并把旧的双空格章标题改为单空格。",
                "- 将二级、三级标题统一改为左顶格黑体，去掉原来“两端对齐”带来的标题观感偏差。",
                "- 将正文、中文摘要、英文摘要统一为“小四号宋体 + Times New Roman 数字/字母、1.5 倍行距、首行缩进 2 字符”。",
                "- 按模板统一表题和图题为图下/表上居中、5 号加粗，保持全文黑色文本。",
                "- 校正正文部分页边距、页眉页脚距离，并保留正文页眉为论文题目、页脚为居中页码。",
                "- 修正第二章中错误的图号：原“图4-1 申领审批闭环流程图”已收口为“图2-1”。",
                "",
                "## drawio 图件处理",
                "- 使用 `$drawio` 重画了 `图3-3` 至 `图3-5` 三张黑白 E-R 图，统一为 Chen 风格：实体矩形、关系菱形、属性椭圆、黑白线条。",
                "- 在论文中新增 `图2-2 调拨执行流程图`、`图2-3 预警处置流程图`、`图4-1 登录认证与令牌续签流程图`。",
                "- 重画了 `图4-2 调拨执行与推荐流程图`，使其与新增流程图保持一致的黑白 drawio 风格。",
                "",
                "## 交付结果",
                f"- 当前图片总数：`{image_count}`",
                f"- 当前表格总数：`{table_count}`",
                f"- 图件目录：`{FIGURE_DIR.relative_to(ROOT)}`",
                f"- drawio 侧产物目录：`{DRAWIO_BUNDLE_DIR.relative_to(ROOT)}`",
                "",
                "## 仍需本机 Word 复核",
                "- 打开 working draft 后刷新目录、题注和页码域。",
                "- 重点检查第 2 章新增流程图和第 3 章 3 张 E-R 图的分页位置。",
                "- 如导师对黑白图宽度有偏好，可在 Word 中把单张图宽微调到 5.8 至 6.1 英寸范围内。",
            ]
        ),
        encoding="utf-8",
    )


def main():
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(WORKING_DRAFT, BACKUP_PATH)

    figures = generate_all_figures()

    doc = Document(WORKING_DRAFT)
    configure_document_styles(doc)
    configure_sections(doc)
    normalize_headings(doc)
    update_figures(doc, figures)
    set_all_text_black(doc)
    doc.save(WORKING_DRAFT)

    saved = Document(WORKING_DRAFT)
    write_report(len(saved.inline_shapes), len(saved.tables))
    print(WORKING_DRAFT)
    print(BACKUP_PATH)
    print(REPORT_PATH)
    print("images", len(saved.inline_shapes))
    print("tables", len(saved.tables))


if __name__ == "__main__":
    main()
