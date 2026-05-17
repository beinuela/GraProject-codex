from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]


FIGURE_MAP = {
    "4.1 系统架构设计": [
        ("output/doc/figures/fig_3_1_architecture.png", "图4-1 系统总体架构图", 5.2),
    ],
    "4.2 系统关键类设计": [
        ("output/doc/figures/fig_4_2_key_classes.png", "图4-2 系统关键类设计图", 5.7),
    ],
    "4.3 系统功能设计": [
        ("output/doc/figures/fig_3_2_modules.png", "图4-3 系统功能模块图", 5.4),
        ("output/doc/figures/fig_4_3_front_backend_flow.png", "图4-4 前后端交互流程图", 5.2),
    ],
    "4.3.3 申领审批设计": [
        ("output/doc/figures/fig_2_1_apply_flow.png", "图4-5 申领审批闭环流程图", 5.1),
    ],
    "4.3.4 调拨管理设计": [
        ("output/doc/figures/fig_2_2_transfer_flow.png", "图4-6 调拨执行流程图", 5.1),
    ],
    "4.3.5 预警与 AI 分析设计": [
        ("output/doc/figures/fig_2_3_warning_flow.png", "图4-7 预警处置流程图", 5.1),
    ],
    "4.4.2 实体关系设计": [
        ("output/doc/figures/fig_3_3_rbac_er.png", "图4-8 权限组织 E-R 图", 5.25),
        ("output/doc/figures/fig_3_4_inventory_er.png", "图4-9 库存与批次 E-R 图", 5.25),
        ("output/doc/figures/fig_3_5_business_er.png", "图4-10 业务单据、预警与通知 E-R 图", 5.25),
    ],
    "5.3 用户基本功能实现": [
        ("output/doc/figures/fig_4_1_auth_flow.png", "图5-1 登录认证与令牌续签流程图", 5.2),
        ("output/doc/runtime-screenshots/fig_6_2_login.png", "图5-2 系统登录界面", 5.2),
    ],
    "5.4 库存、申领与调拨功能实现": [
        ("output/doc/runtime-screenshots/fig_6_3_apply.png", "图5-3 部门用户申领界面", 5.2),
        ("output/doc/runtime-screenshots/fig_6_4_inventory.png", "图5-4 库存查询界面", 5.2),
        ("output/doc/runtime-screenshots/fig_6_6_transfer.png", "图5-5 调拨管理界面", 5.2),
        ("output/doc/figures/fig_6_10_stock_out_flow.png", "图5-6 出库业务流程图", 5.0),
    ],
    "5.5 预警、AI 分析与统计功能实现": [
        ("output/doc/runtime-screenshots/fig_6_7_warning.png", "图5-7 预警管理界面", 5.2),
        ("output/doc/runtime-screenshots/fig_6_8_dashboard.png", "图5-8 运营总览界面", 5.2),
        ("output/doc/runtime-screenshots/fig_6_9_analytics.png", "图5-9 统计分析界面", 5.2),
        ("output/doc/figures/fig_6_11_analytics_flow.png", "图5-10 统计数据流转流程图", 5.0),
    ],
}


def is_table_separator(line: str) -> bool:
    stripped = line.strip().replace(" ", "")
    if not stripped.startswith("|"):
        return False
    core = stripped.strip("|")
    return all(part and set(part) <= {"-", ":"} for part in core.split("|"))


def parse_blocks(text: str):
    blocks = []
    lines = text.splitlines()
    current_para: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("#"):
            if current_para:
                blocks.append(("p", "\n".join(current_para).strip()))
                current_para = []
            level = len(line) - len(line.lstrip("#"))
            blocks.append((f"h{level}", line[level:].strip()))
            i += 1
            continue
        if line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            if current_para:
                blocks.append(("p", "\n".join(current_para).strip()))
                current_para = []
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            blocks.append(("table", table_lines))
            continue
        if not line.strip():
            if current_para:
                blocks.append(("p", "\n".join(current_para).strip()))
                current_para = []
            i += 1
            continue
        current_para.append(line)
        i += 1
    if current_para:
        blocks.append(("p", "\n".join(current_para).strip()))
    return blocks


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def set_run_font(run, east_asia: str, size: int, bold: bool = False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def make_doc(md_path: Path, out_path: Path, include_figures: bool):
    blocks = parse_blocks(md_path.read_text(encoding="utf-8"))
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    current_heading = None
    pending_figures = []
    inserted_for_heading: set[str] = set()

    def add_captioned_figure(img_rel: str, caption: str, width_inches: float):
        img = ROOT / img_rel
        if not img.exists():
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(img), width=Inches(width_inches))
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = c.add_run(caption)
        set_run_font(caption_run, "宋体", 12, bold=True)

    def flush_pending_if_needed():
        nonlocal current_heading, pending_figures
        if include_figures and current_heading and pending_figures and current_heading not in inserted_for_heading:
            for img_rel, caption, width_inches in pending_figures:
                add_captioned_figure(img_rel, caption, width_inches)
            inserted_for_heading.add(current_heading)

    for kind, content in blocks:
        if kind.startswith("h"):
            flush_pending_if_needed()

        if kind == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            set_run_font(run, "黑体", 18, bold=True)
            current_heading = content
            pending_figures = FIGURE_MAP.get(content, [])
        elif kind == "h2":
            if doc.paragraphs:
                doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            set_run_font(run, "黑体", 16, bold=True)
            current_heading = content
            pending_figures = FIGURE_MAP.get(content, [])
        elif kind == "h3":
            p = doc.add_paragraph()
            run = p.add_run(content)
            set_run_font(run, "黑体", 14, bold=True)
            current_heading = content
            pending_figures = FIGURE_MAP.get(content, [])
        elif kind == "h4":
            p = doc.add_paragraph()
            run = p.add_run(content)
            set_run_font(run, "黑体", 12, bold=True)
            current_heading = content
            pending_figures = FIGURE_MAP.get(content, [])
        elif kind == "p":
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(content.replace("**", ""))
            set_run_font(run, "宋体", 12)
            flush_pending_if_needed()
        elif kind == "table":
            rows = [split_table_row(line) for line in content]
            header = rows[0]
            body = rows[2:]
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            for idx, value in enumerate(header):
                hdr_cells[idx].text = value
                hdr_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for para in hdr_cells[idx].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        set_run_font(run, "宋体", 12, bold=True)
            for row in body:
                cells = table.add_row().cells
                for idx, value in enumerate(row):
                    cells[idx].text = value
                    cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for para in cells[idx].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in para.runs:
                            set_run_font(run, "宋体", 12)
            doc.add_paragraph()

    flush_pending_if_needed()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input_md", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--with-figures", action="store_true")
    args = parser.parse_args()
    make_doc(args.input_md, args.output_docx, include_figures=args.with_figures)


if __name__ == "__main__":
    main()
