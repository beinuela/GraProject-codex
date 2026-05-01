from __future__ import annotations

import shutil
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿工作稿.docx"
TARGET_DOCX = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿-图表重绘版.docx"

FIGURE_DIR = ROOT / "output" / "doc" / "figures"

IMAGE_REPLACEMENTS = {
    "word/media/image40.png": FIGURE_DIR / "fig_2_1_apply_flow.png",
    "word/media/image41.png": FIGURE_DIR / "fig_2_2_transfer_flow.png",
    "word/media/image42.png": FIGURE_DIR / "fig_2_3_warning_flow.png",
    "word/media/image45.png": FIGURE_DIR / "fig_3_1_architecture.png",
    "word/media/image46.png": FIGURE_DIR / "fig_3_2_modules.png",
    "word/media/image6.png": FIGURE_DIR / "fig_3_3_rbac_er.png",
    "word/media/image7.png": FIGURE_DIR / "fig_3_4_inventory_er.png",
    "word/media/image8.png": FIGURE_DIR / "fig_3_5_business_er.png",
    "word/media/image47.png": FIGURE_DIR / "fig_4_1_auth_flow.png",
    "word/media/image44.png": FIGURE_DIR / "fig_4_2_transfer_recommend_flow.png",
}

SOURCE_INDEX = [
    (
        "图2-1 申领审批闭环流程图",
        "output/doc/figures/drawio/fig_2_1_apply_flow.drawio",
        "output/doc/figures/fig_2_1_apply_flow.png",
    ),
    (
        "图2-2 调拨执行流程图",
        "output/doc/figures/drawio/fig_2_2_transfer_flow.drawio",
        "output/doc/figures/fig_2_2_transfer_flow.png",
    ),
    (
        "图2-3 预警处置流程图",
        "output/doc/figures/drawio/fig_2_3_warning_flow.drawio",
        "output/doc/figures/fig_2_3_warning_flow.png",
    ),
    (
        "图3-1 系统总体架构图",
        "output/doc/figures/drawio/fig_3_1_architecture.drawio",
        "output/doc/figures/fig_3_1_architecture.png",
    ),
    (
        "图3-2 系统功能模块图",
        "output/doc/figures/drawio/fig_3_2_modules.drawio",
        "output/doc/figures/fig_3_2_modules.png",
    ),
    (
        "图3-3 RBAC 与组织 E-R图",
        "output/doc/figures/drawio/fig_3_3_rbac_er.drawio",
        "output/doc/figures/fig_3_3_rbac_er.png",
    ),
    (
        "图3-4 库存与批次 E-R图",
        "output/doc/figures/drawio/fig_3_4_inventory_er.drawio",
        "output/doc/figures/fig_3_4_inventory_er.png",
    ),
    (
        "图3-5 业务单据、预警与通知 E-R图",
        "output/doc/figures/drawio/fig_3_5_business_er.drawio",
        "output/doc/figures/fig_3_5_business_er.png",
    ),
    (
        "图4-1 登录认证与令牌续签流程图",
        "output/doc/figures/drawio/fig_4_1_auth_flow.drawio",
        "output/doc/figures/fig_4_1_auth_flow.png",
    ),
    (
        "图4-5 调拨执行与推荐流程图",
        "output/doc/figures/drawio/fig_4_2_transfer_recommend_flow.drawio",
        "output/doc/figures/fig_4_2_transfer_recommend_flow.png",
    ),
]


def choose_style_name(doc: Document, *candidates: str) -> str | None:
    for name in candidates:
        try:
            doc.styles[name]
            return name
        except KeyError:
            continue
    return None


def add_paragraph(doc: Document, text: str, style_name: str | None) -> None:
    paragraph = doc.add_paragraph()
    if style_name:
        paragraph.style = style_name
    paragraph.add_run(text)


def appendix_exists(doc: Document) -> bool:
    return any(paragraph.text.strip().startswith("附录A 图表绘制规范与源文件索引") for paragraph in doc.paragraphs)


def append_appendix(doc_path: Path) -> None:
    doc = Document(doc_path)
    if appendix_exists(doc):
        doc.save(doc_path)
        return

    heading_style = choose_style_name(doc, "参考文献标题", "总结/结论标题", "Normal")
    subheading_style = choose_style_name(doc, "2级标题-正文章节", "参考文献标题", "Normal")
    body_style = choose_style_name(doc, "正文章节内容", "总结/结论内容", "Normal")

    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    add_paragraph(doc, "附录A 图表绘制规范与源文件索引", heading_style)
    add_paragraph(doc, "A.1 图表重绘说明", subheading_style)
    add_paragraph(
        doc,
        "本次重绘针对论文中的关键结构图、流程图与 E-R 图进行表达层规范化处理，统一版式、节点形态、连线方式和关系展示规则，不改变原有业务流程逻辑、数据库实体语义以及正文中的章节编号、图号和图题含义。",
        body_style,
    )
    add_paragraph(
        doc,
        "结构图与流程图统一采用自上而下或分层分组的论文版式，主路径居中，分支清晰，连线统一采用垂直或水平的正交折线，避免斜线交叉、文字压线和线条穿过图形。流程图中的开始和结束节点使用椭圆，处理步骤使用圆角矩形，判断条件使用菱形，并明确标注“是/否”分支。",
        body_style,
    )
    add_paragraph(
        doc,
        "E-R 图统一采用论文风格的实体关系表达方式。实体框、属性字段、主键和外键标识、关系菱形以及基数标注均按统一视觉规范重排，重点优化实体对齐、字段密度、阅读顺序和连线清晰度，确保在论文打印和答辩投屏场景下都具备较好的可读性。",
        body_style,
    )
    add_paragraph(doc, "A.2 图形建模约定", subheading_style)
    add_paragraph(
        doc,
        "流程图中的主路径优先沿页面纵向向下延伸，判断节点的“是/否”分支采用左右分流设计；E-R 图中的实体名称与表名保持对应，字段顺序以当前系统模型为准，关系和基数表达遵循“一对多、多对一和多角色关联可读优先”的原则。",
        body_style,
    )
    add_paragraph(doc, "A.3 图号与源文件索引", subheading_style)
    for figure_name, drawio_path, image_path in SOURCE_INDEX:
        add_paragraph(
            doc,
            f"{figure_name}：源文件 {drawio_path}；论文插图 {image_path}。",
            body_style,
        )

    doc.save(doc_path)


def replace_docx_media(doc_path: Path) -> None:
    missing = [str(path) for path in IMAGE_REPLACEMENTS.values() if not path.exists()]
    if missing:
        raise FileNotFoundError(f"缺少待替换图片: {missing}")

    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        temp_path = Path(tmp_file.name)

    try:
        with ZipFile(doc_path, "r") as source_zip, ZipFile(temp_path, "w", compression=ZIP_DEFLATED) as target_zip:
            for item in source_zip.infolist():
                if item.filename in IMAGE_REPLACEMENTS:
                    target_zip.writestr(item, IMAGE_REPLACEMENTS[item.filename].read_bytes())
                else:
                    target_zip.writestr(item, source_zip.read(item.filename))
        shutil.move(str(temp_path), str(doc_path))
    finally:
        if temp_path.exists():
            temp_path.unlink()


def build_delivery_doc() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"未找到论文工作稿: {SOURCE_DOCX}")
    shutil.copy2(SOURCE_DOCX, TARGET_DOCX)
    append_appendix(TARGET_DOCX)
    replace_docx_media(TARGET_DOCX)
    print(f"生成完成: {TARGET_DOCX}")


if __name__ == "__main__":
    build_delivery_doc()
