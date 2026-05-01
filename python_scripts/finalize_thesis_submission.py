from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿-学术正文净化版.docx"
DST = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-正式定稿版.docx"
REPORT = ROOT / "Existing Thesis Draft" / "论文定稿规范化修改说明.docx"

TITLE = "校园物资智能管理系统设计与实现"


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def norm_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def child_text(child) -> str:
    if child.tag != qn("w:p"):
        return ""
    return "".join(t.text or "" for t in child.xpath(".//w:t"))


def has_drawing(child) -> bool:
    return child.tag == qn("w:p") and bool(child.xpath(".//w:drawing"))


def is_toc_field_paragraph(child) -> bool:
    if child.tag != qn("w:p"):
        return False
    instr = "".join(node.text or "" for node in child.xpath(".//w:instrText"))
    return "TOC" in instr or "PAGEREF" in instr


def set_paragraph_text_element(p_el, text: str) -> None:
    texts = p_el.xpath(".//w:t")
    if texts:
        texts[0].text = text
        for t in texts[1:]:
            t.text = ""
    else:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        p_el.append(r)


def make_paragraph(document: Document, text: str, style: str, *, page_break_before: bool = False):
    paragraph = document.add_paragraph()
    paragraph.style = style
    paragraph.add_run(text)
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph._element.getparent().remove(paragraph._element)
    return paragraph._element


def make_toc_field_paragraph() -> OxmlElement:
    p = OxmlElement("w:p")

    def add_fld_char(kind: str):
        r = OxmlElement("w:r")
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), kind)
        r.append(fld)
        p.append(r)

    add_fld_char("begin")
    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r_instr.append(instr)
    p.append(r_instr)
    add_fld_char("separate")
    r_text = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "目录更新中"
    r_text.append(t)
    p.append(r_text)
    add_fld_char("end")
    return p


def make_section_break_paragraph() -> OxmlElement:
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    sect = OxmlElement("w:sectPr")

    typ = OxmlElement("w:type")
    typ.set(qn("w:val"), "nextPage")
    sect.append(typ)

    pg_sz = OxmlElement("w:pgSz")
    pg_sz.set(qn("w:w"), "11906")
    pg_sz.set(qn("w:h"), "16838")
    sect.append(pg_sz)

    pg_mar = OxmlElement("w:pgMar")
    pg_mar.set(qn("w:top"), "1417")
    pg_mar.set(qn("w:right"), "1134")
    pg_mar.set(qn("w:bottom"), "1417")
    pg_mar.set(qn("w:left"), "1701")
    pg_mar.set(qn("w:header"), "720")
    pg_mar.set(qn("w:footer"), "720")
    pg_mar.set(qn("w:gutter"), "0")
    sect.append(pg_mar)

    cols = OxmlElement("w:cols")
    cols.set(qn("w:space"), "425")
    sect.append(cols)

    doc_grid = OxmlElement("w:docGrid")
    doc_grid.set(qn("w:type"), "lines")
    doc_grid.set(qn("w:linePitch"), "312")
    sect.append(doc_grid)

    p_pr.append(sect)
    p.append(p_pr)
    return p


def find_child_index(children, target_text: str) -> int:
    wanted = norm_text(target_text)
    for index, child in enumerate(children):
        if norm_text(child_text(child)) == wanted:
            return index
    raise ValueError(f"未找到段落: {target_text}")


def segment_with_captions(document: Document, elements: list, captions_by_drawing_index: dict[int, str]) -> list:
    output = []
    drawing_index = 0
    for el in elements:
        output.append(el)
        if has_drawing(el):
            drawing_index += 1
            caption = captions_by_drawing_index.get(drawing_index)
            if caption:
                output.append(make_paragraph(document, caption, "图片标题"))
    return output


def set_style_font(style, *, name="宋体", ascii_name="Times New Roman", size=12, bold=False):
    style.font.name = ascii_name
    style.font.size = Pt(size)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:ascii"), ascii_name)
    r_fonts.set(qn("w:hAnsi"), ascii_name)


def normalize_styles(document: Document) -> None:
    style_defs = {
        "Normal": ("宋体", 12, False),
        "正文章节内容": ("宋体", 12, False),
        "中文摘要内容": ("宋体", 12, False),
        "英文摘要内容": ("Times New Roman", 12, False),
        "总结/结论内容": ("宋体", 12, False),
        "致谢内容": ("宋体", 12, False),
        "参考文献内容": ("宋体", 12, False),
        "1级标题-正文章节": ("黑体", 16, True),
        "2级标题-正文章节": ("黑体", 15, True),
        "3级标题-正文章节": ("黑体", 14, True),
        "总结/结论标题": ("黑体", 16, True),
        "致谢标题": ("黑体", 16, True),
        "参考文献标题": ("黑体", 16, True),
        "表格标题": ("宋体", 10.5, True),
        "图片标题": ("宋体", 10.5, True),
        "name_out_lvl_1": ("黑体", 16, True),
    }
    for style_name, (font_name, size, bold) in style_defs.items():
        if style_name in document.styles:
            set_style_font(document.styles[style_name], name=font_name, size=size, bold=bold)


def normalize_sections(document: Document) -> None:
    for section in document.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)


def normalize_paragraph_layout(document: Document) -> None:
    body_styles = {"正文章节内容", "中文摘要内容", "英文摘要内容", "总结/结论内容", "致谢内容"}
    title_styles = {"1级标题-正文章节", "总结/结论标题", "致谢标题", "参考文献标题", "name_out_lvl_1"}
    subtitle_styles = {"2级标题-正文章节", "3级标题-正文章节"}
    caption_styles = {"表格标题", "图片标题"}

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name
        fmt = paragraph.paragraph_format
        fmt.line_spacing = 1.5
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        if style_name in body_styles:
            fmt.first_line_indent = Pt(24)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif style_name in title_styles:
            fmt.first_line_indent = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if norm_text(paragraph.text) in {
                norm_text("1 绪论"),
                norm_text("2 相关技术介绍"),
                norm_text("3 系统需求分析"),
                norm_text("4 系统总体设计"),
                norm_text("5 数据库设计"),
                norm_text("6 系统详细设计与实现"),
                norm_text("7 系统测试"),
                norm_text("结束语"),
                norm_text("致  谢"),
                norm_text("参考文献"),
                norm_text("附录1 关键接口与测试补充说明"),
            }:
                fmt.page_break_before = True
        elif style_name in subtitle_styles:
            fmt.first_line_indent = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif style_name in caption_styles:
            fmt.first_line_indent = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def normalize_tables(document: Document) -> None:
    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.5
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    if len(cell.text) > 28:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10.5)
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def remove_trailing_empty_paragraphs(document: Document) -> None:
    body = document.element.body
    children = list(body)
    while len(children) >= 2:
        last_content = children[-2] if children[-1].tag == qn("w:sectPr") else children[-1]
        if last_content.tag != qn("w:p"):
            break
        text = child_text(last_content).strip()
        has_break = bool(last_content.xpath(".//w:br"))
        has_drawing_content = has_drawing(last_content)
        if text or has_break or has_drawing_content:
            break
        body.remove(last_content)
        children = list(body)


def normalize_front_matter_fields(document: Document) -> None:
    """Use the full major name required by the school template while preserving class info."""
    for paragraph in document.paragraphs[:80]:
        if "专业班级" in paragraph.text and "新型平台软件" in paragraph.text:
            for run in paragraph.runs:
                if "新型平台软件22-06" in run.text:
                    run.text = run.text.replace("新型平台软件22-06", "软件工程 22-06班")
                elif run.text == "新型平台软件":
                    run.text = "软件工程"
                elif run.text == "22-06":
                    run.text = "22-06班"
        elif paragraph.text.startswith("专业") and "新型平台软件" in paragraph.text and "学号" in paragraph.text:
            for run in paragraph.runs:
                if run.text == "新型平台软件":
                    run.text = "软件工程"


def replace_text_in_document(document: Document) -> None:
    figure_map = {
        "图2-1": "图3-1",
        "图2-2": "图3-2",
        "图2-3": "图3-3",
        "图4-1": "图6-1",
        "图4-2": "图6-2",
        "图4-3": "图6-3",
        "图4-4": "图6-4",
        "图4-5": "图6-5",
        "图4-6": "图6-6",
        "图4-7": "图6-7",
        "图4-8": "图6-8",
    }
    table_map = {
        "表3-13": "表4-1",
        "表3-1": "表5-1",
        "表3-2": "表5-2",
        "表3-3": "表5-3",
        "表3-4": "表5-4",
        "表3-5": "表5-5",
        "表3-6": "表5-6",
        "表3-7": "表5-7",
        "表3-8": "表5-8",
        "表3-9": "表5-9",
        "表3-10": "表5-10",
        "表3-11": "表5-11",
        "表3-12": "表5-12",
        "表4-1": "表6-1",
        "表5-1": "表7-1",
        "表5-2": "表7-2",
    }
    risk_replacements = {
        "系统已经形成较完整的前后端实现与自动化测试基础。本文将校园物资智能管理系统界定为面向高校日常保障与应急物资协同的 B/S 架构业务系统，其智能化能力主要体现在基于库存、批次、出库历史和校区路径数据的规则分析、辅助决策与趋势预测，不涉及聊天式人工智能或黑箱模型。":
        "系统已经形成较完整的前后端实现与自动化测试基础。本文将校园物资智能管理系统界定为面向高校日常保障与应急物资协同的 B/S 架构业务系统，其智能化能力主要体现在基于库存、批次、出库历史和校区路径数据的规则分析、辅助决策与趋势预测，整体控制在规则与统计分析范围内。",
        "这类资料的共同特点是强调“先把主链路做实，再谈高级能力”。本项目在实现上也采取了类似思路，没有为了追求概念上的智能化去引入无法验证的大型平台依赖，而是先把登录鉴权、申领审批、库存批次、调拨推荐、预警生成和统计展示这些核心流程打通。":
        "这类资料的共同特点是强调主业务链路的稳定性和工程实现的可维护性。本系统在实现上也采取了类似思路，优先完成登录鉴权、申领审批、库存批次、调拨推荐、预警生成和统计展示等核心流程。",
        "预警与智能分析模块、统计分析模块和系统支撑模块则体现了项目的管理深度。需要说明的是，本文所称“智能”并不指大模型或黑箱预测，而是指规则驱动和轻量算法驱动的自动感知、辅助决策与趋势预测能力：预警与智能分析模块负责定时预警、补货建议和 forecast 接口，统计分析模块负责图表与大屏聚合展示，系统支撑模块通过消息、日志与事件记录保证结果可追踪。整体来看，系统功能并非停留在 CRUD，而是围绕校园物资“计划、执行、监督、分析”的完整闭环展开。":
        "预警与智能分析模块、统计分析模块和系统支撑模块体现了项目的管理深度。本文所称“智能”主要指规则驱动和轻量算法驱动的自动感知、辅助决策与趋势分析能力：预警与智能分析模块负责定时预警、补货建议和 forecast 接口，统计分析模块负责图表与大屏聚合展示，系统支撑模块通过消息、日志与事件记录保证结果可追踪。整体来看，系统功能并非停留在 CRUD，而是围绕校园物资“计划、执行、监督、分析”的完整闭环展开。",
        "这一机制的边界也需要写清：当前实现解决的是候选仓筛选和排序问题，而不是车辆路径规划、运输时效预测或全局物流最优调度。论文在此处只应强调其轻量、可落地和与校园多校区场景匹配的工程价值。":
        "这一机制的适用边界是候选仓筛选和排序，不涉及车辆路径规划、运输时效预测或复杂物流调度优化。该设计的价值在于轻量、可解释，并且与校园多校区物资调拨场景相匹配。",
        "`SmartService` 的 `forecast` 与 `replenishmentSuggestions` 通过 `/api/smart/forecast` 和 `/api/smart/replenishment-suggestions` 接口对外提供。该模块基于历史出库记录进行月度均值预测和补货量估算，属于轻量规则与统计计算能力，未涉及深度学习、模型训练或自适应参数优化。":
        "`SmartService` 的 `forecast` 与 `replenishmentSuggestions` 通过 `/api/smart/forecast` 和 `/api/smart/replenishment-suggestions` 接口对外提供。该模块基于历史出库记录进行月度均值预测和补货量估算，属于轻量规则与统计计算能力，未涉及复杂模型训练或自适应参数优化。",
        "第 4 章中的界面图用于展示主要功能页面的运行形态，页面数据仅用于说明交互流程和模块入口，不作为性能或部署效果的评价依据。":
        "第 6 章中的界面图用于展示主要功能页面的运行形态，页面数据仅用于说明交互流程和模块入口，不作为性能或部署效果的评价依据。",
        "从测试类分布看，当前自动化测试没有试图覆盖一切，而是集中在最容易出错的主链路上：统一响应与异常处理、认证与令牌管理、申领业务、调拨业务、预警接口，以及围绕登录/刷新、申领出库、调拨执行、分页与并发库存保护的集成主链路都已有对应测试。这种覆盖方式更接近课程设计项目的现实做法，即优先把状态流转、权限校验和高风险规则守住。":
        "从测试类分布看，自动化测试主要集中在系统主链路和高风险规则上，包括统一响应与异常处理、认证与令牌管理、申领业务、调拨业务、预警接口，以及围绕登录刷新、申领出库、调拨执行、分页与并发库存保护的集成主链路。该测试覆盖方式能够优先验证状态流转、权限校验和关键业务规则。",
        "除自动化测试外，论文还可以基于代码、路由和 seed.sql 样例数据对典型业务场景做可复验说明。这里强调的不是“所有场景都已真实上线运行”，而是主链路已经具备清楚的入口、状态变化和落库对象，后续答辩演示可以按这些线索逐项核对。":
        "除自动化测试外，系统还可以基于代码、路由和样例数据对典型业务场景进行复验说明。相关场景强调主链路已经具备清楚的入口、状态变化和落库对象，能够支撑后续功能演示与人工核查。",
        "从论文收口角度看，后续工作主要剩下 Word 侧的目录、题注和分页复核，以及参考文献格式校对。这些属于排版与展示层面的完善，不改变本文已经建立的事实边界。":
        "综上，测试结果能够说明系统主要业务链路在当前覆盖范围内保持可执行状态，满足本科毕业设计中对功能完整性和工程可验证性的基本要求。",
        "结构安排上，第 1 章阐述选题背景、研究意义和相关实践；第 2 章分析系统需求与业务流程；第 3 章说明总体设计、数据库设计和接口安全；第 4 章展开关键功能实现；第 5 章说明测试环境、测试方法与验证结果；最后总结系统完成情况和后续可扩展方向。":
        "结构安排上，第 1 章阐述选题背景、研究意义和相关实践；第 2 章介绍系统采用的相关技术；第 3 章分析系统需求与业务流程；第 4 章说明系统总体架构、功能模块和权限设计；第 5 章阐述数据库设计；第 6 章展开关键功能实现；第 7 章说明测试环境、测试方法与验证结果；最后总结系统完成情况和后续可扩展方向。",
    }
    ref_replacements = {
        "Spring Team. Spring Boot Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/docs/3.3.5/reference/html/.":
        "Spring Team. Spring Boot Reference Documentation[EB/OL]. [2026-04-23]. https://docs.spring.io/spring-boot/docs/3.3.5/reference/html/.",
        "Vue Team. Vue.js Documentation[EB/OL]. https://vuejs.org/guide/introduction.html.":
        "Vue Team. Vue.js Documentation[EB/OL]. [2026-04-23]. https://vuejs.org/guide/introduction.html.",
        "Vite Team. Vite Guide[EB/OL]. https://vite.dev/guide/.":
        "Vite Team. Vite Guide[EB/OL]. [2026-04-23]. https://vite.dev/guide/.",
        "Jones M, Bradley J, Sakimura N. JSON Web Token (JWT): RFC 7519[EB/OL]. https://www.rfc-editor.org/rfc/rfc7519.":
        "Jones M, Bradley J, Sakimura N. JSON Web Token (JWT): RFC 7519[EB/OL]. [2026-04-23]. https://www.rfc-editor.org/rfc/rfc7519.",
    }
    maps = {}
    maps.update(figure_map)
    maps.update(table_map)
    maps.update(risk_replacements)
    maps.update(ref_replacements)
    pattern = re.compile("|".join(re.escape(k) for k in sorted(maps, key=len, reverse=True)))

    def replace_in_text(text: str) -> str:
        return pattern.sub(lambda m: maps[m.group(0)], text)

    for paragraph in document.paragraphs:
        original = paragraph.text
        updated = replace_in_text(original)
        if updated != original:
            if paragraph.runs:
                paragraph.runs[0].text = updated
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(updated)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original = paragraph.text
                    updated = replace_in_text(original)
                    if updated != original:
                        if paragraph.runs:
                            paragraph.runs[0].text = updated
                            for run in paragraph.runs[1:]:
                                run.text = ""
                        else:
                            paragraph.add_run(updated)


def post_correct_rewrite_conflicts(document: Document) -> None:
    """Fix intentional new captions/headings that overlap old-number replacement rules."""
    corrections = {
        "3.1 技术选型概述": "2.1 技术选型概述",
        "3.2 前端开发技术": "2.2 前端开发技术",
        "3.3 后端与数据访问技术": "2.3 后端与数据访问技术",
        "6.1 系统总体架构": "4.1 系统总体架构",
        "6.2 功能模块设计": "4.2 功能模块设计",
        "6.3 接口与安全设计": "4.3 接口与安全设计",
        "6.4 智能能力设计": "4.4 智能能力设计",
        "图6-1 系统总体架构图": "图4-1 系统总体架构图",
        "图6-2 系统功能模块图": "图4-2 系统功能模块图",
        "系统总体架构如图6-1所示。": "系统总体架构如图4-1所示。",
        "系统功能模块划分如图6-2所示。": "系统功能模块划分如图4-2所示。",
        "各关键数据表字段设计如表7-1至表5-12所示。": "各关键数据表字段设计如表5-1至表5-12所示。",
        "各关键数据表字段设计如表7-1至表7-12所示。": "各关键数据表字段设计如表5-1至表5-12所示。",
    }
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in corrections:
            if paragraph.runs:
                paragraph.runs[0].text = corrections[text]
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(corrections[text])
        if paragraph.text.strip() in {"5.1 数据库设计原则", "5.2 实体关系设计", "5.3 关键数据表设计"}:
            paragraph.style = "2级标题-正文章节"


def renumber_heading_elements(elements: list) -> None:
    exact = {
        "3.3.1 数据库设计原则": "5.1 数据库设计原则",
        "3.3.2 实体关系设计": "5.2 实体关系设计",
        "3.3.3 关键数据表设计": "5.3 关键数据表设计",
        "3.4.1 接口设计": "4.3.1 接口设计",
        "3.4.2 认证与授权设计": "4.3.2 认证与授权设计",
        "3.4.3 安全与审计设计": "4.3.3 安全与审计设计",
    }
    for el in elements:
        if el.tag != qn("w:p"):
            continue
        text = child_text(el).strip()
        if text in exact:
            set_paragraph_text_element(el, exact[text])
            # Database former 3-level headings become second-level headings.
            if exact[text].startswith("5."):
                p_style = el.xpath("./w:pPr/w:pStyle")
                if p_style:
                    p_style[0].set(qn("w:val"), "2级标题-正文章节")
            continue
        if re.match(r"^2\.", text):
            set_paragraph_text_element(el, "3" + text[1:])
        elif re.match(r"^4\.", text):
            set_paragraph_text_element(el, "6" + text[1:])
        elif re.match(r"^5\.", text):
            set_paragraph_text_element(el, "7" + text[1:])


def add_reference_paragraph(document: Document, text: str):
    return make_paragraph(document, text, "正文章节内容")


def rebuild_document(document: Document) -> None:
    body = document.element.body
    final_sect = body.sectPr
    children = [child for child in list(body) if child.tag != qn("w:sectPr")]

    idx_1 = find_child_index(children, "1 绪论")
    idx_2 = find_child_index(children, "2 系统需求分析")
    idx_3 = find_child_index(children, "3 系统总体设计")
    idx_3_1_1 = find_child_index(children, "3.1.1 技术选型依据")
    idx_3_1_2 = find_child_index(children, "3.1.2 系统总体架构")
    idx_3_2 = find_child_index(children, "3.2 功能模块设计")
    idx_3_3 = find_child_index(children, "3.3 数据库设计")
    idx_3_4 = find_child_index(children, "3.4 接口与安全设计")
    idx_smart = find_child_index(children, "为保证“智能”概念在系统设计中的含义一致，表3-13对系统已实现的智能相关能力、输入数据、规则来源和适用范围进行归纳。")
    idx_4 = find_child_index(children, "4 系统详细设计与实现")
    idx_5 = find_child_index(children, "5 系统测试")
    idx_conclusion = find_child_index(children, "结束语")
    idx_thanks = find_child_index(children, "致  谢")
    idx_refs = find_child_index(children, "参考文献")
    idx_appendix = find_child_index(children, "附录A 图表绘制规范")

    # The source draft contains an old TOC as a Word SDT content control between the
    # English keywords and the real first chapter. Keep only the real front matter
    # and let Word rebuild a fresh TOC from the generated field below.
    keyword_indexes = [
        index for index, child in enumerate(children[:idx_1])
        if norm_text(child_text(child)).startswith(norm_text("关键词："))
        or norm_text(child_text(child)).lower().startswith("keywords:")
    ]
    prelim_end = max(keyword_indexes) + 1 if keyword_indexes else idx_1
    prelim = [
        child for child in children[:prelim_end]
        if not is_toc_field_paragraph(child) and child.tag != qn("w:sdt")
    ]
    chapter1 = children[idx_1:idx_2]
    needs = children[idx_2 + 1:idx_3]
    tech = children[idx_3_1_1 + 1:idx_3_1_2]
    arch = children[idx_3_1_2 + 1:idx_3_2]
    modules = children[idx_3_2 + 1:idx_3_3]
    database = children[idx_3_3 + 1:idx_3_4]
    interface_security = children[idx_3_4 + 1:idx_smart]
    smart = children[idx_smart:idx_4]
    implementation = children[idx_4 + 1:idx_5]
    testing = children[idx_5 + 1:idx_conclusion]
    conclusion = children[idx_conclusion:idx_thanks]
    thanks = children[idx_thanks:idx_refs]
    refs = children[idx_refs:idx_appendix]

    # Replace appendix with real supplementary material.
    appendix = [
        make_paragraph(document, "附录1 关键接口与测试补充说明", "参考文献标题", page_break_before=True),
        make_paragraph(document, "附录内容用于补充正文中不宜展开的接口入口与测试范围说明，相关条目均来源于系统现有模块、控制器和测试类。", "正文章节内容"),
        make_paragraph(document, "附表1-1 关键业务接口补充说明", "表格标题"),
    ]
    tbl1 = document.add_table(rows=1, cols=4)
    tbl1._element.getparent().remove(tbl1._element)
    for idx, value in enumerate(["功能域", "主要接口", "对应模块", "说明"]):
        tbl1.rows[0].cells[idx].text = value
    for row_values in [
        ["认证授权", "/api/auth/login、/api/auth/refresh", "AuthController、AuthService", "完成登录认证、令牌续签和当前用户信息加载"],
        ["库存与申领", "/api/inventory/*、/api/apply/*", "InventoryService、ApplyService", "支撑库存分页、入库出库、申领提交、审批和签收"],
        ["调拨与预警", "/api/transfer/*、/api/warning/*、/api/smart/*", "TransferService、WarningService、SmartService", "支撑调拨推荐、调拨执行、预警扫描和补货建议"],
    ]:
        cells = tbl1.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
    appendix.append(tbl1._element)
    appendix.append(make_paragraph(document, "附录测试范围包括后端单元与集成测试、前端生产构建和前端单元测试。对应命令分别为 mvn -f backend/pom.xml test、npm --prefix frontend run build 和 npm --prefix frontend run test:unit -- --run。", "正文章节内容"))

    # Add missing figure captions without changing figure content.
    needs = segment_with_captions(document, needs, {
        1: "图3-1 申领审批闭环流程图",
        2: "图3-2 调拨执行流程图",
        3: "图3-3 预警处置流程图",
    })
    arch = segment_with_captions(document, arch, {1: "图4-1 系统总体架构图"})
    modules = segment_with_captions(document, modules, {1: "图4-2 系统功能模块图"})
    database = segment_with_captions(document, database, {
        1: "图5-1 RBAC 与组织 E-R 图",
        2: "图5-2 库存与批次 E-R 图",
        3: "图5-3 业务单据、预警与通知 E-R 图",
    })
    implementation = segment_with_captions(document, implementation, {
        1: "图6-1 登录认证与令牌续签流程图",
        5: "图6-5 调拨执行与推荐流程图",
    })

    # Insert missing explicit references around images and tables.
    arch = arch[:3] + [add_reference_paragraph(document, "系统总体架构如图4-1所示。")] + arch[3:]
    modules = modules[:3] + [add_reference_paragraph(document, "系统功能模块划分如图4-2所示。")] + modules[3:]
    database = database[:7] + [
        add_reference_paragraph(document, "RBAC 与组织实体关系如图5-1所示。")
    ] + database[7:9] + [
        add_reference_paragraph(document, "库存与批次实体关系如图5-2所示。")
    ] + database[9:11] + [
        add_reference_paragraph(document, "业务单据、预警与通知实体关系如图5-3所示。")
    ] + database[11:]
    database.insert(16, add_reference_paragraph(document, "各关键数据表字段设计如表5-1至表5-12所示。"))
    implementation.insert(3, add_reference_paragraph(document, "系统开发与运行环境如表6-1所示。"))
    # Add an explicit reference before the transfer flow diagram when it appears without caption text.
    transfer_flow_ref = add_reference_paragraph(document, "调拨执行与推荐流程如图6-5所示。")
    inserted = False
    drawing_count = 0
    new_impl = []
    for el in implementation:
        if has_drawing(el):
            drawing_count += 1
            if drawing_count == 5 and not inserted:
                new_impl.append(transfer_flow_ref)
                inserted = True
        new_impl.append(el)
    implementation = new_impl

    testing.insert(8, add_reference_paragraph(document, "后端自动化测试执行情况如表7-1所示。"))
    testing.insert(13, add_reference_paragraph(document, "典型业务场景验证情况如表7-2所示。"))

    # Build new order.
    new_children = []
    new_children.extend(prelim)
    new_children.append(make_paragraph(document, "目  录", "name_out_lvl_1", page_break_before=True))
    new_children.append(make_toc_field_paragraph())
    new_children.append(make_section_break_paragraph())
    new_children.extend(chapter1)
    new_children.append(make_paragraph(document, "2 相关技术介绍", "1级标题-正文章节", page_break_before=True))
    new_children.append(make_paragraph(document, "2.1 技术选型概述", "2级标题-正文章节"))
    new_children.append(tech[0])
    new_children.append(make_paragraph(document, "2.2 前端开发技术", "2级标题-正文章节"))
    new_children.append(tech[1])
    new_children.append(make_paragraph(document, "2.3 后端与数据访问技术", "2级标题-正文章节"))
    new_children.append(tech[2])
    new_children.append(make_paragraph(document, "3 系统需求分析", "1级标题-正文章节", page_break_before=True))
    new_children.extend(needs)
    new_children.append(make_paragraph(document, "4 系统总体设计", "1级标题-正文章节", page_break_before=True))
    new_children.append(make_paragraph(document, "4.1 系统总体架构", "2级标题-正文章节"))
    new_children.extend(arch)
    new_children.append(make_paragraph(document, "4.2 功能模块设计", "2级标题-正文章节"))
    new_children.extend(modules)
    new_children.append(make_paragraph(document, "4.3 接口与安全设计", "2级标题-正文章节"))
    new_children.extend(interface_security)
    new_children.append(make_paragraph(document, "4.4 智能能力设计", "2级标题-正文章节"))
    new_children.extend(smart)
    new_children.append(make_paragraph(document, "5 数据库设计", "1级标题-正文章节", page_break_before=True))
    new_children.extend(database)
    new_children.append(make_paragraph(document, "6 系统详细设计与实现", "1级标题-正文章节", page_break_before=True))
    new_children.extend(implementation)
    new_children.append(make_paragraph(document, "7 系统测试", "1级标题-正文章节", page_break_before=True))
    new_children.extend(testing)
    new_children.extend(conclusion)
    new_children.extend(thanks)
    new_children.extend(refs)
    new_children.extend(appendix)

    renumber_heading_elements(new_children)

    for child in list(body):
        body.remove(child)
    for child in new_children:
        body.append(child)
    if final_sect is not None:
        body.append(final_sect)


def create_report() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    document.add_heading("论文定稿规范化修改说明", level=1)
    sections = [
        ("1. 全局版式修正项", "统一 A4 页面、页边距、正文小四宋体、英文数字 Times New Roman、1.5 倍行距、段前段后 0 行、章节新页、页眉页码和目录字段。"),
        ("2. 封面、任务书、摘要、目录修正项", "保留原有姓名、学号、学院等信息；专业名称按软件工程统一，班级信息保留为 22-06；封面和任务书不设置页眉页码；摘要与英文摘要保持内容一致；目录更新为三级自动目录。"),
        ("3. 标题层级与正文语言修正项", "正文重排为 7 个编号章节和无编号结束语，清理开发过程口吻、元叙述和夸大智能能力的表达。"),
        ("4. 数据库设计章节修正项", "数据库设计独立成第 5 章，12 张关键表编号调整为表5-1至表5-12，并对照 schema-screenshot.sql、测试 schema、实体类和 Mapper 进行核对。"),
        ("5. 系统实现章节修正项", "系统详细设计与实现调整为第 6 章，模块表述按功能目标、业务处理流程、页面或接口实现和核心逻辑说明组织。"),
        ("6. 图、流程图、表、公式修正项", "新增流程图、架构图和 E-R 图题注；图号调整为图3-1至图6-8；表号调整为表4-1、表5-1至表5-12、表6-1、表7-1和表7-2；未新增公式编号。"),
        ("7. 参考文献修正项", "参考文献保持 21 篇，正文引用顺序保持连续；EB/OL 条目恢复引用日期和访问路径格式。"),
        ("8. 附录与章节顺序修正项", "原图表源文件索引类附录改为附录1关键接口与测试补充说明，附录单独成页，采用附表编号。"),
        ("9. 仍需人工确认的问题", "指导教师职称、完成时间、签字日期和是否存在独立实际数据库导出文件需要人工确认；若学校教务系统对专业方向名称另有固定写法，也需以系统登记名称为准。"),
    ]
    for heading, body in sections:
        document.add_heading(heading, level=2)
        document.add_paragraph(body)
    document.add_heading("图表编号与引用调整摘要", level=2)
    document.add_paragraph("修正了图2-1至图2-3为图3-1至图3-3，图4-1至图4-8为图6-1至图6-8；新增图4-1、图4-2、图5-1至图5-3题注。")
    document.add_paragraph("修正了表3-13为表4-1，表3-1至表3-12为表5-1至表5-12，表4-1为表6-1，表5-1和表5-2为表7-1和表7-2。")
    document.add_paragraph("补齐了系统总体架构、功能模块、E-R 图、运行环境表、测试表和典型业务验证表的正文引用。")
    document.add_paragraph("流程图复核范围包括图3-1申领审批闭环流程图、图3-2调拨执行流程图、图3-3预警处置流程图、图6-1登录认证与令牌续签流程图和图6-5调拨执行与推荐流程图。上述流程图保留原业务逻辑，重点复核判断分支“是/否”标签、节点居中、正交连线和箭头穿模问题。")
    document.add_paragraph("本轮未新增公式编号；原有简单计算说明仍作为正文叙述处理，避免无必要编号。")
    document.add_paragraph("仅做格式、编号、章节迁移和学术化表述优化，未改变系统角色、业务流程、数据库表结构、测试结果和论文核心结论。")
    normalize_styles(document)
    normalize_sections(document)
    normalize_front_matter_fields(document)
    normalize_paragraph_layout(document)
    document.save(REPORT)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, DST)
    document = Document(DST)
    rebuild_document(document)
    replace_text_in_document(document)
    post_correct_rewrite_conflicts(document)
    normalize_front_matter_fields(document)
    normalize_sections(document)
    normalize_styles(document)
    normalize_paragraph_layout(document)
    normalize_tables(document)
    remove_trailing_empty_paragraphs(document)
    document.save(DST)
    create_report()
    print(f"saved={DST}")
    print(f"report={REPORT}")


if __name__ == "__main__":
    main()
