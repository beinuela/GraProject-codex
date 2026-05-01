from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿-表格摘要附录修订版.docx"
TARGET_DOCX = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿-附录版式修正版.docx"

APPENDIX_TITLE = "附录A 图表绘制规范与源文件索引"
APPENDIX_SUBHEADINGS = ["A.1 图表重绘说明", "A.2 图形建模约定", "A.3 图号与源文件索引"]
APPENDIX_BODY_TEXTS = [
    "本次重绘针对论文中的关键结构图、流程图与 E-R 图进行表达层规范化处理，统一版式、节点形态、连线方式和关系展示规则，不改变原有业务流程逻辑、数据库实体语义以及正文中的章节编号、图号和图题含义。",
    "结构图与流程图统一采用自上而下或分层分组的论文版式，主路径居中，分支清晰，连线统一采用垂直或水平的正交折线，避免斜线交叉、文字压线和线条穿过图形。流程图中的开始和结束节点使用椭圆，处理步骤使用圆角矩形，判断条件使用菱形，并明确标注“是/否”分支。",
    "E-R 图统一采用论文风格的实体关系表达方式。实体框、属性字段、主键和外键标识、关系菱形以及基数标注均按统一视觉规范重排，重点优化实体对齐、字段密度、阅读顺序和连线清晰度，确保在论文打印和答辩投屏场景下都具备较好的可读性。",
    "流程图中的主路径优先沿页面纵向向下延伸，判断节点的“是/否”分支采用左右分流设计；E-R 图中的实体名称与表名保持对应，字段顺序以当前系统模型为准，关系和基数表达遵循“一对多、多对一和多角色关联可读优先”的原则。",
]


def remove_child(parent, tag: str) -> None:
    child = parent.find(qn(tag))
    if child is not None:
        parent.remove(child)


def get_or_add_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_run_font(run, size_pt: float, east_asia_font: str, ascii_font: str = "Times New Roman", bold: bool | None = None) -> None:
    run.font.size = Pt(size_pt)
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    if bold is not None:
        run.bold = bold


def set_paragraph_alignment_xml(paragraph, value: str) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    remove_child(p_pr, "w:jc")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), value)
    p_pr.append(jc)


def set_first_line_chars(paragraph, chars: int = 2) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    remove_child(p_pr, "w:ind")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLineChars"), str(chars * 100))
    ind.set(qn("w:firstLine"), "420")
    p_pr.append(ind)


def clear_first_line_indent(paragraph) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    remove_child(p_pr, "w:ind")


def find_paragraph_index(doc: Document, exact_text: str) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == exact_text:
            return idx
    raise ValueError(f"未找到段落: {exact_text}")


def insert_blank_paragraph_after(paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def format_appendix_title(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_alignment_xml(paragraph, "center")
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    clear_first_line_indent(paragraph)
    if not paragraph.runs:
        paragraph.add_run(paragraph.text)
    for run in paragraph.runs:
        set_run_font(run, 15, east_asia_font="黑体", bold=True)


def format_appendix_subheading(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_alignment_xml(paragraph, "left")
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    clear_first_line_indent(paragraph)
    for run in paragraph.runs:
        set_run_font(run, 12, east_asia_font="黑体", bold=True)


def format_appendix_body(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_alignment_xml(paragraph, "both")
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    set_first_line_chars(paragraph, 2)
    for run in paragraph.runs:
        set_run_font(run, 12, east_asia_font="宋体")


def format_blank_line(paragraph) -> None:
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_alignment_xml(paragraph, "left")
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    clear_first_line_indent(paragraph)


def format_appendix_table(table) -> None:
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [1.7, 3.2, 5.3, 5.8]
    tbl_pr = table._tbl.tblPr
    layout = get_or_add_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_w = get_or_add_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:w"), str(int(round(sum(widths) * 567))))
    tbl_w.set(qn("w:type"), "dxa")

    for row_idx, row in enumerate(table.rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.78)
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if row_idx == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))

        for col_idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            for tag in ("w:textDirection", "w:noWrap", "w:tcFitText"):
                remove_child(tc_pr, tag)
            tc_w = get_or_add_child(tc_pr, "w:tcW")
            tc_w.set(qn("w:w"), str(int(round(widths[col_idx] * 567))))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Cm(widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 or col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_alignment_xml(paragraph, "center" if row_idx == 0 or col_idx == 0 else "left")
                fmt = paragraph.paragraph_format
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)
                fmt.line_spacing = 1.2
                clear_first_line_indent(paragraph)
                for run in paragraph.runs:
                    set_run_font(run, 10.5 if row_idx == 0 else 9.5, east_asia_font="宋体", bold=True if row_idx == 0 else None)


def build_doc() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"未找到源文档: {SOURCE_DOCX}")
    shutil.copy2(SOURCE_DOCX, TARGET_DOCX)
    doc = Document(TARGET_DOCX)

    title_idx = find_paragraph_index(doc, APPENDIX_TITLE)
    title_para = doc.paragraphs[title_idx]
    format_appendix_title(title_para)

    if title_idx + 1 >= len(doc.paragraphs) or doc.paragraphs[title_idx + 1].text.strip():
        blank_para = insert_blank_paragraph_after(title_para)
    else:
        blank_para = doc.paragraphs[title_idx + 1]
    format_blank_line(blank_para)

    for heading in APPENDIX_SUBHEADINGS:
        format_appendix_subheading(doc.paragraphs[find_paragraph_index(doc, heading)])

    for body_text in APPENDIX_BODY_TEXTS:
        format_appendix_body(doc.paragraphs[find_paragraph_index(doc, body_text)])

    format_appendix_table(doc.tables[-1])

    doc.save(TARGET_DOCX)
    print(f"生成完成: {TARGET_DOCX}")


if __name__ == "__main__":
    build_doc()
