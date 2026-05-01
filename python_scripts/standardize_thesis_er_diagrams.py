from __future__ import annotations

import hashlib
import json
import re
import shutil
import subprocess
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from tempfile import NamedTemporaryFile
from xml.etree import ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile

import yaml
from docx import Document
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "output" / "doc"
FIGURE_DIR = DOC_DIR / "figures"
DRAWIO_DIR = FIGURE_DIR / "drawio"
DRAWIO_CLI = ROOT / ".codex" / "skills" / "drawio" / "scripts" / "cli.js"
PLAYWRIGHT_MODULE = ROOT / "frontend" / "node_modules" / "playwright"

SOURCE_DOCX = DOC_DIR / "校园物资智能管理系统设计与实现-E-R图中文规范优化版.docx"
TARGET_DOCX = DOC_DIR / "校园物资智能管理系统设计与实现-E-R图人工精修版.docx"
REPORT_PATH = DOC_DIR / "E-R图严格布局优化与中文化修改说明.md"

STOP_TITLES = {"致  谢", "致谢", "参考文献", "附录 1 关键接口与测试补充说明", "附录"}
BODY_HEADING_RE = re.compile(r"^[1-7]\s")
CAPTION_RE = re.compile(r"^图5-[1-3]\s+.+E-R 图$")


@dataclass(frozen=True)
class EntityData:
    name: str
    table: str
    attributes: tuple[str, ...]


@dataclass(frozen=True)
class Blueprint:
    canvas: str
    title: str
    description: str
    entity_positions: dict[str, tuple[int, int]]
    relation_positions: dict[str, tuple[int, int]]
    attribute_positions: dict[str, list[tuple[int, int]]]
    relation_edges: list[tuple[str, str, str]]


LINE = (0, 0, 0)
BG = (255, 255, 255)
RENDER_SCALE = 2


def choose_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


BLUEPRINTS: dict[str, Blueprint] = {
    "fig_3_3_rbac_er": Blueprint(
        canvas="1760x760",
        title="用户角色与组织 E-R 图",
        description="用户、部门、角色、刷新令牌与登录日志的实体关系图",
        entity_positions={
            "dept": (180, 210),
            "user": (880, 200),
            "role": (1580, 210),
            "token": (560, 500),
            "log": (1200, 500),
        },
        relation_positions={
            "rel_dept_user": (500, 240),
            "rel_user_role": (1245, 240),
            "rel_user_token": (720, 390),
            "rel_user_log": (1040, 390),
        },
        attribute_positions={
            "dept": [(70, 95), (200, 30), (360, 90), (380, 180)],
            "user": [(760, 70), (910, 20), (1070, 70), (1180, 160), (660, 160)],
            "role": [(1470, 95), (1590, 30), (1740, 90), (1735, 180)],
            "token": [(390, 420), (540, 360), (700, 420), (790, 560)],
            "log": [(1040, 420), (1200, 360), (1370, 420), (1460, 560)],
        },
        relation_edges=[
            ("dept", "rel_dept_user", "1"),
            ("rel_dept_user", "user", "n"),
            ("user", "rel_user_role", "n"),
            ("rel_user_role", "role", "1"),
            ("user", "rel_user_token", "1"),
            ("rel_user_token", "token", "n"),
            ("user", "rel_user_log", "1"),
            ("rel_user_log", "log", "n"),
        ],
    ),
    "fig_3_4_inventory_er": Blueprint(
        canvas="1900x1280",
        title="库存与批次 E-R 图",
        description="分类、物资档案、仓库、库存和库存批次的实体关系图",
        entity_positions={
            "category": (280, 210),
            "material": (930, 230),
            "warehouse": (1600, 230),
            "inventory": (930, 610),
            "batch": (930, 1000),
        },
        relation_positions={
            "rel_cat_material": (610, 250),
            "rel_material_inventory": (930, 420),
            "rel_wh_inventory": (1320, 430),
            "rel_inventory_batch": (930, 820),
        },
        attribute_positions={
            "category": [(150, 100), (300, 35), (470, 105)],
            "material": [(760, 115), (930, 55), (1110, 120), (690, 350), (1170, 350)],
            "warehouse": [(1460, 110), (1610, 40), (1790, 120), (1810, 265)],
            "inventory": [(760, 520), (620, 780), (770, 840), (1240, 790), (1260, 650)],
            "batch": [(760, 930), (1140, 930), (720, 1120), (1150, 1120), (940, 1230)],
        },
        relation_edges=[
            ("category", "rel_cat_material", "1"),
            ("rel_cat_material", "material", "n"),
            ("material", "rel_material_inventory", "1"),
            ("rel_material_inventory", "inventory", "n"),
            ("warehouse", "rel_wh_inventory", "1"),
            ("rel_wh_inventory", "inventory", "n"),
            ("inventory", "rel_inventory_batch", "1"),
            ("rel_inventory_batch", "batch", "n"),
        ],
    ),
    "fig_3_5_business_er": Blueprint(
        canvas="1960x1200",
        title="业务单据、预警与通知 E-R 图",
        description="申领、出库、调拨、预警和通知的实体关系图",
        entity_positions={
            "apply": (260, 250),
            "item": (950, 500),
            "stockout": (1640, 250),
            "transfer": (260, 900),
            "warning": (980, 900),
            "notify": (1660, 900),
        },
        relation_positions={
            "rel_apply_item": (590, 420),
            "rel_apply_stock": (950, 110),
            "rel_transfer_warning": (620, 900),
            "rel_warning_notify": (1320, 900),
        },
        attribute_positions={
            "apply": [(150, 130), (290, 45), (470, 110), (90, 255)],
            "item": [(780, 390), (950, 300), (1130, 390), (1160, 540)],
            "stockout": [(1510, 130), (1670, 40), (1835, 110), (1905, 220)],
            "transfer": [(150, 780), (260, 1070), (500, 1070), (90, 900)],
            "warning": [(850, 780), (980, 700), (1170, 780), (1060, 1070)],
            "notify": [(1530, 780), (1670, 700), (1860, 800), (1895, 930)],
        },
        relation_edges=[
            ("apply", "rel_apply_item", "1"),
            ("rel_apply_item", "item", "n"),
            ("apply", "rel_apply_stock", "1"),
            ("rel_apply_stock", "stockout", "1"),
            ("transfer", "rel_transfer_warning", "1"),
            ("rel_transfer_warning", "warning", "n"),
            ("warning", "rel_warning_notify", "1"),
            ("rel_warning_notify", "notify", "n"),
        ],
    ),
}

FALLBACK_SOURCE_MODELS: dict[str, dict[str, dict[str, object]]] = {
    "fig_3_3_rbac_er": {
        "entities": {
            "dept": EntityData("部门", "", ("部门编号", "部门名称", "上级部门", "版本号")),
            "user": EntityData("用户", "", ("用户编号", "用户账号", "姓名", "状态", "部门编号 / 角色编号")),
            "role": EntityData("角色", "", ("角色编号", "角色编码", "角色名称", "角色说明")),
            "token": EntityData("刷新令牌", "", ("令牌编号", "用户编号", "令牌摘要", "到期时间 / 吊销状态")),
            "log": EntityData("登录日志", "", ("用户编号", "用户账号", "登录地址", "登录状态 / 登录时间")),
        },
        "relations": {
            "rel_dept_user": "隶属",
            "rel_user_role": "拥有",
            "rel_user_token": "签发",
            "rel_user_log": "产生",
        },
    },
    "fig_3_4_inventory_er": {
        "entities": {
            "category": EntityData("物资分类", "", ("编号", "分类名称", "备注")),
            "material": EntityData("物资档案", "", ("编号", "物资编码", "物资名称", "安全库存", "规格单位")),
            "warehouse": EntityData("仓库", "", ("编号", "仓库名称", "所属校区", "负责人")),
            "inventory": EntityData("库存", "", ("编号", "物资编号", "仓库编号", "当前库存", "锁定库存")),
            "batch": EntityData("库存批次", "", ("编号", "批次号", "入库数量", "剩余数量", "过期日期")),
        },
        "relations": {
            "rel_cat_material": "归类",
            "rel_material_inventory": "形成",
            "rel_wh_inventory": "存放",
            "rel_inventory_batch": "细分",
        },
    },
    "fig_3_5_business_er": {
        "entities": {
            "apply": EntityData("申领单", "", ("编号", "部门编号", "申请人编号", "状态")),
            "item": EntityData("申领明细", "", ("编号", "申领单编号", "物资编号", "申请数量")),
            "stockout": EntityData("出库单", "", ("编号", "申领单编号", "仓库编号", "操作人编号")),
            "transfer": EntityData("调拨单", "", ("编号", "调出仓库编号", "调入仓库编号", "状态")),
            "warning": EntityData("预警记录", "", ("编号", "预警类型", "处理状态", "处理人编号")),
            "notify": EntityData("通知消息", "", ("编号", "消息标题", "目标用户编号", "已读状态")),
        },
        "relations": {
            "rel_apply_item": "包含",
            "rel_apply_stock": "生成",
            "rel_transfer_warning": "触发",
            "rel_warning_notify": "发送",
        },
    },
}

MEDIA_REPLACEMENTS = {
    "word/media/image7.png": FIGURE_DIR / "fig_3_3_rbac_er.png",
    "word/media/image8.png": FIGURE_DIR / "fig_3_4_inventory_er.png",
    "word/media/image9.png": FIGURE_DIR / "fig_3_5_business_er.png",
}

EDGE_OVERRIDES: dict[str, dict[tuple[str, str], dict[str, object]]] = {
    "fig_3_3_rbac_er": {
        ("user", "rel_user_token"): {
            "waypoints": [{"x": 760, "y": 330}],
            "style": {"exitX": 0.38, "exitY": 1, "entryX": 0.5, "entryY": 0},
        },
        ("user", "rel_user_log"): {
            "waypoints": [{"x": 1040, "y": 330}],
            "style": {"exitX": 0.62, "exitY": 1, "entryX": 0.5, "entryY": 0},
        },
        ("rel_user_token", "token"): {
            "style": {"exitX": 0.5, "exitY": 1, "entryX": 0.82, "entryY": 0},
        },
        ("rel_user_log", "log"): {
            "style": {"exitX": 0.5, "exitY": 1, "entryX": 0.18, "entryY": 0},
        },
    },
    "fig_3_4_inventory_er": {
        ("category", "rel_cat_material"): {
            "style": {"exitX": 1, "exitY": 0.56, "entryX": 0.0, "entryY": 0.5},
        },
        ("rel_cat_material", "material"): {
            "style": {"exitX": 1, "exitY": 0.5, "entryX": 0, "entryY": 0.5},
        },
        ("material", "rel_material_inventory"): {
            "style": {"exitX": 0.5, "exitY": 1, "entryX": 0.5, "entryY": 0},
        },
        ("rel_material_inventory", "inventory"): {
            "style": {"exitX": 0.5, "exitY": 1, "entryX": 0.5, "entryY": 0},
        },
        ("warehouse", "rel_wh_inventory"): {
            "waypoints": [{"x": 1450, "y": 430}],
            "style": {"exitX": 0.0, "exitY": 0.62, "entryX": 1, "entryY": 0.5},
        },
        ("rel_wh_inventory", "inventory"): {
            "waypoints": [{"x": 1120, "y": 430}, {"x": 1120, "y": 620}],
            "style": {"exitX": 0.0, "exitY": 0.6, "entryX": 1, "entryY": 0.36},
        },
        ("inventory", "rel_inventory_batch"): {
            "style": {"exitX": 0.5, "exitY": 1, "entryX": 0.5, "entryY": 0},
        },
        ("rel_inventory_batch", "batch"): {
            "style": {"exitX": 0.5, "exitY": 1, "entryX": 0.5, "entryY": 0},
        },
    },
    "fig_3_5_business_er": {
        ("apply", "rel_apply_item"): {
            "style": {"exitX": 1, "exitY": 0.72, "entryX": 0.0, "entryY": 0.5},
        },
        ("rel_apply_item", "item"): {
            "style": {"exitX": 1, "exitY": 0.5, "entryX": 0.0, "entryY": 0.58},
        },
        ("apply", "rel_apply_stock"): {
            "style": {"exitX": 1, "exitY": 0.38, "entryX": 0.0, "entryY": 0.5},
        },
        ("rel_apply_stock", "stockout"): {
            "style": {"exitX": 1, "exitY": 0.5, "entryX": 0.0, "entryY": 0.38},
        },
        ("transfer", "rel_transfer_warning"): {
            "style": {"exitX": 1, "exitY": 0.52, "entryX": 0.0, "entryY": 0.5},
        },
        ("rel_transfer_warning", "warning"): {
            "style": {"exitX": 1, "exitY": 0.5, "entryX": 0.0, "entryY": 0.52},
        },
        ("warning", "rel_warning_notify"): {
            "style": {"exitX": 1, "exitY": 0.5, "entryX": 0.0, "entryY": 0.5},
        },
        ("rel_warning_notify", "notify"): {
            "style": {"exitX": 1, "exitY": 0.5, "entryX": 0.0, "entryY": 0.5},
        },
    },
}

EXPECTED_CAPTIONS = (
    "图5-1 用户角色与组织 E-R 图",
    "图5-2 库存与批次 E-R 图",
    "图5-3 业务单据、预警与通知 E-R 图",
)
EXPECTED_REFERENCE_SENTENCE = "用户角色与组织关系如图5-1所示，库存与批次关系如图5-2所示，业务单据、预警与通知关系如图5-3所示。"
ASCII_LABEL_RE = re.compile(r"[A-Za-z_]")


def body_paragraphs(doc: Document) -> list[tuple[int, str]]:
    items: list[tuple[int, str]] = []
    in_body = False
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name == "Heading 1":
            if text in STOP_TITLES:
                in_body = False
            elif BODY_HEADING_RE.match(text) or text == "结束语":
                in_body = True
        if in_body:
            items.append((index, text))
    return items


def parse_entity_model(stem: str, blueprint: Blueprint) -> tuple[dict[str, EntityData], dict[str, str]]:
    fallback = FALLBACK_SOURCE_MODELS[stem]
    entities = dict(fallback["entities"])
    relations = dict(fallback["relations"])

    for entity_id, entity in entities.items():
        if len(entity.attributes) > len(blueprint.attribute_positions[entity_id]):
            raise ValueError(f"{stem} 的实体 {entity_id} 属性数量超过预设布局容量。")

    return entities, relations


def bw_node(
    node_id: str,
    label: str,
    node_type: str,
    x: int,
    y: int,
    *,
    shape: str | None = None,
    size: str = "medium",
    font_size: int = 13,
) -> dict[str, object]:
    style: dict[str, object] = {
        "fillColor": "#FFFFFF",
        "strokeColor": "#000000",
        "fontColor": "#000000",
        "strokeWidth": 1.6,
        "fontSize": font_size,
    }
    if shape:
        style["shape"] = shape
    return {
        "id": node_id,
        "label": label,
        "type": node_type,
        "size": size,
        "position": {"x": x, "y": y},
        "style": style,
    }


def bw_edge(
    src: str,
    dst: str,
    label: str | None = None,
    *,
    waypoints: list[dict[str, int]] | None = None,
    style_updates: dict[str, object] | None = None,
) -> dict[str, object]:
    edge: dict[str, object] = {
        "from": src,
        "to": dst,
        "type": "primary",
        "style": {
            "strokeColor": "#000000",
            "endArrow": "none",
            "strokeWidth": 1.4,
            "fontColor": "#000000",
            "fontSize": 12,
        },
    }
    if style_updates:
        edge["style"].update(style_updates)
    if label is not None:
        edge["label"] = label
        edge["labelPosition"] = "end"
    if waypoints:
        edge["waypoints"] = waypoints
    return edge


def attribute_size(label: str) -> str:
    if len(label) >= 14:
        return "large"
    if len(label) >= 9:
        return "medium"
    return "small"


def build_er_spec(stem: str, blueprint: Blueprint) -> dict[str, object]:
    entities, relations = parse_entity_model(stem, blueprint)
    nodes: list[dict[str, object]] = []
    edges: list[dict[str, object]] = []

    for entity_id, pos in blueprint.entity_positions.items():
        entity = entities[entity_id]
        entity_label = entity.name if not entity.table else f"{entity.name}\n{entity.table}"
        nodes.append(
            bw_node(
                entity_id,
                entity_label,
                "service",
                pos[0],
                pos[1],
                shape="shape=rectangle;rounded=0;whiteSpace=wrap;html=1;spacing=8",
                size="xl",
                font_size=13,
            )
        )

        for index, attribute_label in enumerate(entity.attributes):
            attr_x, attr_y = blueprint.attribute_positions[entity_id][index]
            attr_id = f"{entity_id}_attr_{index + 1}"
            nodes.append(
                bw_node(
                    attr_id,
                    attribute_label,
                    "user",
                    attr_x,
                    attr_y,
                    size=attribute_size(attribute_label),
                    font_size=11,
                )
            )
            edges.append(bw_edge(entity_id, attr_id))

    for relation_id, pos in blueprint.relation_positions.items():
        nodes.append(
            bw_node(
                relation_id,
                relations[relation_id],
                "decision",
                pos[0],
                pos[1],
                size="medium",
                font_size=12,
            )
        )

    edge_overrides = EDGE_OVERRIDES.get(stem, {})
    for src, dst, cardinality in blueprint.relation_edges:
        override = edge_overrides.get((src, dst), {})
        edges.append(
            bw_edge(
                src,
                dst,
                cardinality,
                waypoints=override.get("waypoints"),
                style_updates=override.get("style"),
            )
        )

    return {
        "meta": {
            "profile": "academic-paper",
            "theme": "academic",
            "layout": "horizontal",
            "routing": "straight",
            "canvas": blueprint.canvas,
            "title": blueprint.title,
            "description": blueprint.description,
        },
        "nodes": nodes,
        "edges": edges,
        "modules": [],
    }


def parse_canvas_size(canvas: str) -> tuple[int, int]:
    width_text, height_text = canvas.lower().split("x", 1)
    return int(width_text), int(height_text)


def text_size(draw: ImageDraw.ImageDraw, text: str, font) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def center_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font) -> None:
    width, height = text_size(draw, text, font)
    x1, y1, x2, y2 = box
    draw.text(((x1 + x2 - width) / 2, (y1 + y2 - height) / 2), text, font=font, fill=LINE)


def node_kind(node_id: str, spec: dict[str, object]) -> str:
    for node in spec["nodes"]:
        if node["id"] == node_id:
            if node["id"].startswith("rel_"):
                return "relation"
            if "_attr_" in node["id"]:
                return "attribute"
            return "entity"
    raise KeyError(node_id)


def build_manual_node_layout(spec: dict[str, object]) -> dict[str, dict[str, float]]:
    canvas_w, canvas_h = parse_canvas_size(spec["meta"]["canvas"])
    dummy = Image.new("RGB", (10, 10), BG)
    draw = ImageDraw.Draw(dummy)
    font_entity = choose_font(28, bold=True)
    font_attr = choose_font(21)
    font_rel = choose_font(24, bold=True)
    layout: dict[str, dict[str, float]] = {}
    for node in spec["nodes"]:
        nid = node["id"]
        kind = "relation" if nid.startswith("rel_") else "attribute" if "_attr_" in nid else "entity"
        label = str(node.get("label", ""))
        cx = node["position"]["x"] * RENDER_SCALE
        cy = node["position"]["y"] * RENDER_SCALE
        if kind == "entity":
            width = 210 * RENDER_SCALE
            height = 96 * RENDER_SCALE
        elif kind == "relation":
            width = 110 * RENDER_SCALE
            height = 74 * RENDER_SCALE
        else:
            text_w, text_h = text_size(draw, label, font_attr)
            width = max(text_w + 36, 92 * RENDER_SCALE)
            height = max(text_h + 22, 46 * RENDER_SCALE)
        layout[nid] = {
            "kind": kind,
            "cx": cx,
            "cy": cy,
            "w": width,
            "h": height,
            "x1": cx - width / 2,
            "y1": cy - height / 2,
            "x2": cx + width / 2,
            "y2": cy + height / 2,
        }
    layout["_canvas"] = {"w": canvas_w * RENDER_SCALE, "h": canvas_h * RENDER_SCALE}
    return layout


def point_on_node_border(node: dict[str, float], toward: tuple[float, float], fraction: tuple[float, float] | None = None) -> tuple[float, float]:
    if fraction is not None:
        fx, fy = fraction
        return (node["x1"] + node["w"] * fx, node["y1"] + node["h"] * fy)

    cx, cy = node["cx"], node["cy"]
    dx = toward[0] - cx
    dy = toward[1] - cy
    if dx == 0 and dy == 0:
        return (cx, cy)

    if node["kind"] == "entity":
        scale = min((node["w"] / 2) / abs(dx) if dx else float("inf"), (node["h"] / 2) / abs(dy) if dy else float("inf"))
        return (cx + dx * scale, cy + dy * scale)
    if node["kind"] == "relation":
        scale = 1 / ((abs(dx) / (node["w"] / 2)) + (abs(dy) / (node["h"] / 2)))
        return (cx + dx * scale, cy + dy * scale)

    denom = ((dx * dx) / ((node["w"] / 2) ** 2)) + ((dy * dy) / ((node["h"] / 2) ** 2))
    scale = 1 / (denom ** 0.5)
    return (cx + dx * scale, cy + dy * scale)


def polyline_point_and_normal(points: list[tuple[float, float]], ratio: float) -> tuple[tuple[float, float], tuple[float, float]]:
    if len(points) < 2:
        return points[0], (0.0, -1.0)
    segments: list[tuple[tuple[float, float], tuple[float, float], float]] = []
    total = 0.0
    for start, end in zip(points, points[1:]):
        length = ((end[0] - start[0]) ** 2 + (end[1] - start[1]) ** 2) ** 0.5
        segments.append((start, end, length))
        total += length
    if total == 0:
        return points[0], (0.0, -1.0)
    target = total * ratio
    traversed = 0.0
    for start, end, length in segments:
        if traversed + length >= target:
            inner = (target - traversed) / length if length else 0.0
            px = start[0] + (end[0] - start[0]) * inner
            py = start[1] + (end[1] - start[1]) * inner
            dx = end[0] - start[0]
            dy = end[1] - start[1]
            seg_len = (dx * dx + dy * dy) ** 0.5 or 1.0
            return (px, py), (-dy / seg_len, dx / seg_len)
        traversed += length
    last_start, last_end, _ = segments[-1]
    dx = last_end[0] - last_start[0]
    dy = last_end[1] - last_start[1]
    seg_len = (dx * dx + dy * dy) ** 0.5 or 1.0
    return last_end, (-dy / seg_len, dx / seg_len)


def draw_manual_er_png(spec: dict[str, object], png_path: Path) -> None:
    layout = build_manual_node_layout(spec)
    canvas_w = int(layout["_canvas"]["w"])
    canvas_h = int(layout["_canvas"]["h"])
    image = Image.new("RGB", (canvas_w, canvas_h), BG)
    draw = ImageDraw.Draw(image)
    font_entity = choose_font(28, bold=True)
    font_attr = choose_font(21)
    font_rel = choose_font(24, bold=True)
    font_card = choose_font(20, bold=True)

    for node in spec["nodes"]:
        node_box = layout[node["id"]]
        box = (int(node_box["x1"]), int(node_box["y1"]), int(node_box["x2"]), int(node_box["y2"]))
        label = str(node.get("label", ""))
        if node_box["kind"] == "entity":
            draw.rectangle(box, outline=LINE, width=4, fill=BG)
            center_text(draw, box, label, font_entity)
        elif node_box["kind"] == "relation":
            x1, y1, x2, y2 = box
            diamond = [((x1 + x2) // 2, y1), (x2, (y1 + y2) // 2), ((x1 + x2) // 2, y2), (x1, (y1 + y2) // 2)]
            draw.polygon(diamond, outline=LINE, width=4, fill=BG)
            center_text(draw, box, label, font_rel)
        else:
            draw.ellipse(box, outline=LINE, width=4, fill=BG)
            center_text(draw, box, label, font_attr)

    for edge in spec["edges"]:
        src = layout[edge["from"]]
        dst = layout[edge["to"]]
        waypoints = [(point["x"] * RENDER_SCALE, point["y"] * RENDER_SCALE) for point in edge.get("waypoints", [])]
        style = edge.get("style", {})
        src_fraction = None
        dst_fraction = None
        if "exitX" in style and "exitY" in style:
            src_fraction = (float(style["exitX"]), float(style["exitY"]))
        if "entryX" in style and "entryY" in style:
            dst_fraction = (float(style["entryX"]), float(style["entryY"]))
        start_target = waypoints[0] if waypoints else (dst["cx"], dst["cy"])
        end_source = waypoints[-1] if waypoints else (src["cx"], src["cy"])
        points = [point_on_node_border(src, start_target, src_fraction)]
        points.extend(waypoints)
        points.append(point_on_node_border(dst, end_source, dst_fraction))
        draw.line(points, fill=LINE, width=4)

        label = str(edge.get("label", "")).strip()
        if not label:
            continue
        src_kind = src["kind"]
        dst_kind = dst["kind"]
        ratio = 0.18 if src_kind == "entity" and dst_kind == "relation" else 0.82 if src_kind == "relation" and dst_kind == "entity" else 0.5
        (px, py), normal = polyline_point_and_normal(points, ratio)
        label_w, label_h = text_size(draw, label, font_card)
        lx = px + normal[0] * 22 - label_w / 2
        ly = py + normal[1] * 22 - label_h / 2
        draw.rectangle((lx - 8, ly - 6, lx + label_w + 8, ly + label_h + 6), fill=BG)
        draw.text((lx, ly), label, font=font_card, fill=LINE)

    image.save(png_path)


def svg_dimensions(svg_path: Path) -> tuple[int, int]:
    root = ET.fromstring(svg_path.read_text(encoding="utf-8"))

    def parse_dimension(value: str | None, fallback: int) -> int:
        if not value:
            return fallback
        match = re.search(r"(\d+(?:\.\d+)?)", value)
        return int(float(match.group(1))) if match else fallback

    width = parse_dimension(root.attrib.get("width"), 1600)
    height = parse_dimension(root.attrib.get("height"), 1000)
    if "viewBox" in root.attrib:
        _, _, view_w, view_h = [float(part) for part in root.attrib["viewBox"].split()]
        width = max(width, int(view_w))
        height = max(height, int(view_h))
    return width, height


def svg_to_png(svg_path: Path, png_path: Path) -> None:
    width, height = svg_dimensions(svg_path)
    html = (
        "<!doctype html><html><head><meta charset='utf-8'>"
        "<style>html,body{margin:0;background:#fff;}img{display:block;margin:0 auto;}</style>"
        f"</head><body><img id='diagram' src='{svg_path.as_uri()}' width='{width}' height='{height}'></body></html>"
    )
    html_path = svg_path.with_suffix(".preview.html")
    html_path.write_text(html, encoding="utf-8")

    node_script = f"""
const {{ chromium }} = require({json.dumps(str(PLAYWRIGHT_MODULE))});
(async () => {{
  const browser = await chromium.launch();
  const page = await browser.newPage({{
    viewport: {{ width: {width + 80}, height: {min(height + 80, 1400)} }},
    deviceScaleFactor: 2
  }});
  await page.goto({json.dumps(html_path.as_uri())}, {{ waitUntil: 'load' }});
  await page.waitForFunction(() => {{
    const img = document.getElementById('diagram');
    return img && img.complete && img.naturalWidth > 0;
  }});
  await page.screenshot({{ path: {json.dumps(str(png_path))}, fullPage: true }});
  await browser.close();
}})().catch(err => {{
  console.error(err);
  process.exit(1);
}});
"""
    try:
        subprocess.run(
            ["node", "-"],
            input=node_script,
            text=True,
            cwd=str(ROOT),
            check=True,
        )
    finally:
        if html_path.exists():
            html_path.unlink()


def render_bundle(spec: dict[str, object], stem: str) -> Path:
    DRAWIO_DIR.mkdir(parents=True, exist_ok=True)
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)

    yaml_path = DRAWIO_DIR / f"{stem}.yaml"
    drawio_path = DRAWIO_DIR / f"{stem}.drawio"
    svg_path = DRAWIO_DIR / f"{stem}.svg"
    png_path = FIGURE_DIR / f"{stem}.png"

    yaml_path.write_text(yaml.safe_dump(spec, allow_unicode=True, sort_keys=False), encoding="utf-8")

    subprocess.run(
        ["node", str(DRAWIO_CLI), str(yaml_path), str(drawio_path), "--validate", "--write-sidecars"],
        cwd=str(ROOT),
        text=True,
        encoding="utf-8",
        errors="ignore",
        check=True,
    )
    subprocess.run(
        ["node", str(DRAWIO_CLI), str(yaml_path), str(svg_path), "--validate", "--write-sidecars"],
        cwd=str(ROOT),
        text=True,
        encoding="utf-8",
        errors="ignore",
        check=True,
    )
    draw_manual_er_png(spec, png_path)
    return png_path


def assert_node_labels_are_chinese(spec: dict[str, object], stem: str) -> None:
    offenders: list[str] = []
    for node in spec.get("nodes", []):
        label = str(node.get("label", ""))
        if label and ASCII_LABEL_RE.search(label):
            offenders.append(f"{node['id']}={label!r}")
    if offenders:
        raise RuntimeError(f"{stem} 仍存在英文或中英混排标签: {offenders}")


def md5_bytes(data: bytes) -> str:
    return hashlib.md5(data).hexdigest()


def md5_file(path: Path) -> str:
    return md5_bytes(path.read_bytes())


def replace_docx_media(doc_path: Path) -> None:
    missing = [str(path) for path in MEDIA_REPLACEMENTS.values() if not path.exists()]
    if missing:
        raise FileNotFoundError(f"缺少待替换图片: {missing}")

    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        temp_path = Path(tmp_file.name)

    try:
        with ZipFile(doc_path, "r") as source_zip, ZipFile(temp_path, "w", compression=ZIP_DEFLATED) as target_zip:
            for item in source_zip.infolist():
                replacement = MEDIA_REPLACEMENTS.get(item.filename)
                if replacement is None:
                    target_zip.writestr(item, source_zip.read(item.filename))
                else:
                    target_zip.writestr(item, replacement.read_bytes())
        shutil.move(str(temp_path), str(doc_path))
    finally:
        if temp_path.exists():
            temp_path.unlink()


def verify_docx(doc_path: Path) -> dict[str, object]:
    doc = Document(doc_path)
    body = body_paragraphs(doc)
    captions = [text for _, text in body if CAPTION_RE.match(text)]
    caption_alignments: dict[str, int | None] = {}
    image_alignments: dict[str, int | None] = {}

    def paragraph_has_drawing(paragraph) -> bool:
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        return bool(paragraph._element.findall(".//w:drawing", ns))

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text in EXPECTED_CAPTIONS:
            caption_alignments[text] = paragraph.alignment
            image_alignment = None
            for prev_index in range(index - 1, -1, -1):
                candidate = doc.paragraphs[prev_index]
                if paragraph_has_drawing(candidate):
                    image_alignment = candidate.alignment
                    break
                if candidate.text.strip():
                    continue
            image_alignments[text] = image_alignment

    reference_sentence_present = any(text == EXPECTED_REFERENCE_SENTENCE for _, text in body)

    media_hashes: dict[str, str] = {}
    with ZipFile(doc_path) as archive:
        for media_name in MEDIA_REPLACEMENTS:
            media_hashes[media_name] = md5_bytes(archive.read(media_name))

    return {
        "captions": captions,
        "caption_alignments": caption_alignments,
        "image_alignments": image_alignments,
        "reference_sentence_present": reference_sentence_present,
        "media_hashes": media_hashes,
    }


def write_report(backup_path: Path, verification: dict[str, object]) -> None:
    optimized = [
        ("图5-1 用户角色与组织 E-R 图", "完成中文化，并按人工精修方式重排用户、角色、部门与日志/令牌支路"),
        ("图5-2 库存与批次 E-R 图", "完成中文化，并重点清理库存附近的主干线与属性线冲突"),
        ("图5-3 业务单据、预警与通知 E-R 图", "完成中文化，并重绘申领/出库与调拨/预警/通知两个模块"),
    ]
    lines = [
        "# E-R 图严格布局优化与中文化修改说明",
        "",
        f"- 工作底稿：`{SOURCE_DOCX.name}`",
        f"- 备份文件：`{backup_path.name}`",
        f"- 输出文件：`{TARGET_DOCX.name}`",
        "",
        "## 修改结果",
        "",
        "- 共检查 `3` 张 E-R 图。",
        "- 共优化 `3` 张 E-R 图。",
        "",
        "### 重点布局调整",
        "",
    ]

    for caption, layout_change in optimized:
        lines.append(f"- `{caption}`：已完成 {layout_change}。")

    lines.extend(
        [
            "",
            "### 中文化处理",
            "",
            "- `图5-1`、`图5-2`、`图5-3` 的实体名、属性名和联系名已全部改为中文。",
            "- 已移除实体框中的英文表名，不再保留中文与英文表名混排形式。",
            "- 已将字段类英文标签统一转换为规范中文，例如“编号、仓库名称、当前库存、锁定库存、操作人编号、目标用户编号”等。",
            "",
            "### 图题、图号与正文引用",
            "",
            "- 图题与图号未新增也未改号，仍保持 `图5-1` 至 `图5-3` 连续编号。",
            f"- 正文引用校验结果：`{'已保留原有正确引用' if verification['reference_sentence_present'] else '需要人工复核'}`。",
            "- 本轮未补充新的图题、图号或正文引用文本。",
            "",
            "### 排版优化",
            "",
            "- 保持 Chen 风格不变，但最终论文内嵌图片已改为人工级几何重绘，不再依赖原先的 SVG 截图连线效果。",
            "- 已重点处理关系主干交叉、长斜线、仓库和库存附近的穿框观感，以及属性椭圆过远和局部拥挤问题。",
            "- 图片段落和图题段落结构保持原位，仅替换内嵌图片内容。",
            "",
            "### 图源更新路径",
            "",
            "- `output/doc/figures/fig_3_3_rbac_er.png`",
            "- `output/doc/figures/fig_3_4_inventory_er.png`",
            "- `output/doc/figures/fig_3_5_business_er.png`",
            "- `output/doc/figures/drawio/fig_3_3_rbac_er.{yaml,spec.yaml,drawio,svg,arch.json}`",
            "- `output/doc/figures/drawio/fig_3_4_inventory_er.{yaml,spec.yaml,drawio,svg,arch.json}`",
            "- `output/doc/figures/drawio/fig_3_5_business_er.{yaml,spec.yaml,drawio,svg,arch.json}`",
            "",
            "### 媒体替换校验",
            "",
        ]
    )

    for media_name, media_hash in sorted(verification["media_hashes"].items()):
        figure_name = MEDIA_REPLACEMENTS[media_name].name
        lines.append(f"- `{media_name}` -> `{figure_name}`，MD5=`{media_hash}`。")

    lines.extend(
        [
            "",
            "### 视觉复核提示",
            "",
            "- 当前环境未提供 Word 页面级渲染工具，建议在本地 Word 中打开输出文件，对第 5 章分页以及 `图5-2`、`图5-3` 的页内缩放观感做一次最终视觉复核。",
        ]
    )

    REPORT_PATH.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    if not DRAWIO_CLI.exists():
        raise FileNotFoundError(DRAWIO_CLI)
    if not PLAYWRIGHT_MODULE.exists():
        raise FileNotFoundError(PLAYWRIGHT_MODULE)

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    backup_path = DOC_DIR / f"{SOURCE_DOCX.stem}-人工精修前备份-{timestamp}{SOURCE_DOCX.suffix}"
    shutil.copy2(SOURCE_DOCX, backup_path)
    shutil.copy2(SOURCE_DOCX, TARGET_DOCX)

    rendered: dict[str, Path] = {}
    for stem, blueprint in BLUEPRINTS.items():
        spec = build_er_spec(stem, blueprint)
        assert_node_labels_are_chinese(spec, stem)
        rendered[stem] = render_bundle(spec, stem)

    replace_docx_media(TARGET_DOCX)
    verification = verify_docx(TARGET_DOCX)

    if tuple(verification["captions"]) != EXPECTED_CAPTIONS:
        raise RuntimeError(f"E-R 图题检查失败: {verification['captions']}")
    if not verification["reference_sentence_present"]:
        raise RuntimeError("正文中未找到预期的 E-R 图引用语句。")
    for caption in EXPECTED_CAPTIONS:
        if verification["caption_alignments"].get(caption) != 1:
            raise RuntimeError(f"图题未保持居中: {caption}")
        if verification["image_alignments"].get(caption) != 1:
            raise RuntimeError(f"图片段落未保持居中: {caption}")
    for media_name, figure_path in MEDIA_REPLACEMENTS.items():
        if verification["media_hashes"][media_name] != md5_file(figure_path):
            raise RuntimeError(f"内嵌媒体哈希不匹配: {media_name} -> {figure_path.name}")

    write_report(backup_path, verification)

    print(f"[backup] {backup_path.name}")
    print(f"[output] {TARGET_DOCX.name}")
    for stem, path in rendered.items():
        print(f"[figure] {stem} -> {path.name}")
    for media_name, media_hash in sorted(verification['media_hashes'].items()):
        print(f"[media] {media_name} {media_hash}")


if __name__ == "__main__":
    main()
