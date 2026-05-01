from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿工作稿.docx"
BACKUP_PATH = (
    ROOT
    / "Existing Thesis Draft"
    / f"校园物资智能管理系统设计与实现-定稿工作稿.docx.bak.{datetime.now():%Y%m%d-%H%M%S}-ch5-table-layout"
)

TABLE_51_HEADER = ["测试类别", "对应测试类/范围", "验证重点", "结果"]
TABLE_52_HEADER = ["业务场景", "主要入口", "关键证据", "说明"]


def set_cell_width(cell, width) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width.twips)))


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    table.autofit = False


def set_table_grid(table, widths) -> None:
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width.twips)))
        tbl_grid.append(grid_col)


def set_table_widths(table, widths) -> None:
    set_table_layout_fixed(table)
    set_table_grid(table, widths)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_cell_margins(table, top=55, start=70, bottom=55, end=70) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tbl_cell_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_first_row_as_header(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def format_table_text(table, body_size=9.0, header_size=9.5) -> None:
    for row_idx, row in enumerate(table.rows):
        set_row_cant_split(row)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(header_size if row_idx == 0 else body_size)
                    if row_idx == 0:
                        run.bold = True


def find_paragraph(document: Document, text: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def table_header_text(table) -> list[str]:
    return [cell.text.strip() for cell in table.rows[0].cells]


def find_table(document: Document, header: list[str]):
    for table in document.tables:
        if table_header_text(table) == header:
            return table
    raise ValueError(f"Table not found: {header}")


def format_caption(paragraph, page_break_before: bool) -> None:
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)


def main() -> None:
    BACKUP_PATH.write_bytes(DOCX_PATH.read_bytes())
    document = Document(str(DOCX_PATH))

    table_51 = find_table(document, TABLE_51_HEADER)
    table_52 = find_table(document, TABLE_52_HEADER)

    set_table_widths(table_51, [Cm(3.15), Cm(5.55), Cm(3.75), Cm(2.55)])
    set_table_widths(table_52, [Cm(3.05), Cm(2.55), Cm(4.0), Cm(5.5)])
    set_table_cell_margins(table_51)
    set_table_cell_margins(table_52)
    mark_first_row_as_header(table_51)
    mark_first_row_as_header(table_52)
    format_table_text(table_51, body_size=8.5, header_size=9.0)
    format_table_text(table_52, body_size=8.5, header_size=9.0)

    format_caption(find_paragraph(document, "表5-1 后端自动化测试执行情况"), page_break_before=True)
    format_caption(find_paragraph(document, "表5-2 典型业务场景验证说明"), page_break_before=True)

    document.save(str(DOCX_PATH))
    print(f"Updated: {DOCX_PATH}")
    print(f"Backup: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
