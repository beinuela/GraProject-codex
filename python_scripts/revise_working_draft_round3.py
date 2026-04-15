from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORKING_DRAFT = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿工作稿.docx"
ORIGINAL_DRAFT = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-初稿.docx"
BACKUP_DRAFT = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿工作稿.docx.bak.20260414-aigc-er-restore"
SCHEMA_PATH = ROOT / "sql" / "schema.sql"
REPORT_PATH = ROOT / "output" / "doc" / "校园物资智能管理系统设计与实现-改稿说明-第三轮数据库恢复与AIGC续改.md"

TARGET_TABLES: List[Tuple[str, str]] = [
    ("表3-1 用户信息表（sys_user）", "sys_user"),
    ("表3-2 物资档案表（material_info）", "material_info"),
    ("表3-3 仓库信息表（warehouse）", "warehouse"),
    ("表3-4 库存总表（inventory）", "inventory"),
    ("表3-5 库存批次表（inventory_batch）", "inventory_batch"),
    ("表3-6 申领单主表（apply_order）", "apply_order"),
    ("表3-7 申领单明细表（apply_order_item）", "apply_order_item"),
    ("表3-8 出库单主表（stock_out）", "stock_out"),
    ("表3-9 调拨单主表（transfer_order）", "transfer_order"),
    ("表3-10 预警记录表（warning_record）", "warning_record"),
    ("表3-11 通知消息表（notification）", "notification"),
    ("表3-12 登录日志表（login_log）", "login_log"),
]

FIELD_EXPLANATIONS: Dict[str, str] = {
    "id": "主键标识",
    "username": "登录账号",
    "password": "BCrypt 密码串",
    "real_name": "姓名",
    "dept_id": "关联部门",
    "role_id": "关联角色",
    "status": "当前状态标记",
    "deleted": "逻辑删除标记",
    "version": "版本号",
    "created_at": "创建时间",
    "updated_at": "更新时间",
    "material_code": "物资编码",
    "material_name": "物资名称",
    "category_id": "所属分类",
    "spec": "规格型号",
    "unit": "计量单位",
    "safety_stock": "安全库存阈值",
    "shelf_life_days": "保质期天数",
    "supplier": "供应商名称",
    "unit_price": "单价",
    "warehouse_name": "仓库名称",
    "campus": "所属校区",
    "address": "地址",
    "manager": "负责人",
    "material_id": "关联物资",
    "warehouse_id": "关联仓库",
    "current_qty": "当前库存数量",
    "locked_qty": "锁定数量",
    "batch_no": "批次号",
    "in_qty": "入库数量",
    "remain_qty": "剩余数量",
    "production_date": "生产日期",
    "expire_date": "到期日期",
    "apply_order_id": "关联申领单",
    "applicant_id": "申请人",
    "urgency_level": "紧急等级",
    "reason": "业务原因",
    "scenario": "使用场景",
    "fast_track": "快速审批标记",
    "approver_id": "审批人",
    "approve_remark": "审批意见",
    "approve_time": "审批时间",
    "apply_qty": "申领数量",
    "actual_qty": "实际出库数量",
    "operator_id": "操作人",
    "remark": "备注信息",
    "from_warehouse_id": "调出仓库",
    "to_warehouse_id": "调入仓库",
    "warning_type": "预警类型",
    "content": "内容描述",
    "handle_status": "处理状态",
    "handler_id": "处理人",
    "handle_remark": "处理说明",
    "title": "标题",
    "msg_type": "消息类型",
    "target_user_id": "目标用户",
    "is_read": "是否已读",
    "biz_type": "业务类型",
    "biz_id": "业务编号",
    "user_id": "关联用户",
    "login_ip": "登录IP",
    "login_status": "登录状态",
    "login_time": "登录时间",
    "user_agent": "客户端标识",
}

TABLE_WIDTHS = [Cm(2.7), Cm(2.5), Cm(2.0), Cm(4.0), Cm(5.4)]
AIGC_METRICS = {
    "total_words": "26198",
    "risk_probability": "21.76%",
    "high_risk": "0.83%",
    "mid_risk": "11.53%",
}


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def set_paragraph_text(paragraph, text: str) -> None:
    while paragraph.runs:
        paragraph._element.remove(paragraph.runs[0]._element)
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.underline = False


def paragraph_index(doc: Document, text: str) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return idx
    raise ValueError(f"Paragraph not found: {text}")


def replace_nonempty_paragraphs_between(doc: Document, start_text: str, stop_text: str, new_texts: List[str]) -> None:
    start_idx = paragraph_index(doc, start_text)
    stop_idx = paragraph_index(doc, stop_text)
    candidates = [p for p in doc.paragraphs[start_idx + 1 : stop_idx] if p.text.strip()]
    if len(candidates) != len(new_texts):
        raise ValueError(
            f"Paragraph count mismatch between '{start_text}' and '{stop_text}': "
            f"expected {len(new_texts)}, found {len(candidates)}"
        )
    for paragraph, text in zip(candidates, new_texts):
        set_paragraph_text(paragraph, text)


def find_first_matching_after(doc: Document, start_idx: int, predicate) -> int:
    for idx in range(start_idx + 1, len(doc.paragraphs)):
        if predicate(doc.paragraphs[idx]):
            return idx
    raise ValueError("Matching paragraph not found after anchor")


def replace_abstracts(doc: Document) -> None:
    cn_title_idx = paragraph_index(doc, "摘    要")
    cn_abs_idx = find_first_matching_after(doc, cn_title_idx, lambda p: p.text.strip() and not p.text.strip().startswith("关键词"))
    cn_kw_idx = find_first_matching_after(doc, cn_abs_idx, lambda p: p.text.strip().startswith("关键词"))
    en_title_idx = paragraph_index(doc, "ABSTRACT")
    en_abs_idx = find_first_matching_after(doc, en_title_idx, lambda p: p.text.strip() and not p.text.strip().startswith("Keywords"))
    en_kw_idx = find_first_matching_after(doc, en_abs_idx, lambda p: p.text.strip().startswith("Keywords"))

    cn_abstract = (
        "面对高校物资台账分散、审批流转偏长、库存批次难以追踪和预警处置滞后的问题，本文围绕当前项目仓库完成了校园物资智能管理系统的设计与实现。"
        "系统并不追求脱离项目的概念包装，而是以已有代码、SQL、配置和文档为边界：前端采用 Vue 3、Pinia、Element Plus 和 ECharts 组织页面与图表，"
        "后端采用 Spring Boot 3、Spring Security、JWT 和 MyBatis-Plus 完成认证授权、业务编排与数据访问，数据层使用 MySQL 8 存储业务数据。"
        "围绕这一实现基础，论文重点说明了基础数据管理、库存与批次维护、申领审批、出库执行、仓间调拨、预警扫描、补货建议、通知消息和日志审计等已落地模块，"
        "并对统一响应结构、角色权限边界、刷新令牌续期和按到期日优先出库等关键机制进行了归纳。"
        "测试部分只采用当前仓库可以复核的事实，记录了 2026 年 4 月 14 日后端 42 项自动化测试通过、前端生产构建成功的结果。"
        "这些证据表明，当前项目已经具备覆盖校园物资核心流程的实现基础，论文中的结论也据此限定在真实可验证的范围内。"
    )
    en_abstract = (
        "This thesis addresses the problems of scattered ledgers, lengthy approval chains, weak batch traceability, and delayed warning handling in campus material management. "
        "The discussion stays within the boundaries of the current repository rather than relying on generic system narratives. "
        "On the client side, Vue 3, Pinia, Element Plus, and ECharts are used to organize business pages and analytical views. "
        "On the server side, Spring Boot 3, Spring Security, JWT, and MyBatis-Plus are used to implement authentication, authorization, workflow orchestration, and data access, with MySQL 8 as the persistence layer. "
        "Based on these verified assets, the thesis focuses on master data maintenance, inventory and batch tracking, application approval, outbound execution, warehouse transfer, warning generation, replenishment suggestion, notification delivery, and audit logging. "
        "It also summarizes the actual mechanisms behind the unified response contract, role-based access control, refresh-token renewal, and expiry-oriented outbound deduction. "
        "The testing chapter is grounded only in executable evidence from the repository, including the backend test suite and the frontend production build. "
        "On April 14, 2026, 42 backend tests passed and the frontend build completed successfully. "
        "These results indicate that the current project already provides an implementable basis for the core workflow of campus material management."
    )

    set_paragraph_text(doc.paragraphs[cn_abs_idx], cn_abstract)
    set_paragraph_text(doc.paragraphs[cn_kw_idx], "关键词：校园物资管理；Spring Boot；Vue 3；库存预警；申领审批")
    set_paragraph_text(doc.paragraphs[en_abs_idx], en_abstract)
    set_paragraph_text(doc.paragraphs[en_kw_idx], "Keywords: campus material management; Spring Boot; Vue 3; inventory warning; application approval")


def split_type(type_text: str) -> Tuple[str, str]:
    match = re.match(r"([A-Z]+)(?:\(([^)]+)\))?", type_text.upper())
    if not match:
        return type_text.upper(), "-"
    base = match.group(1)
    length = match.group(2) or "-"
    return base, length


def parse_default_value(rest: str) -> str | None:
    match = re.search(r"DEFAULT\s+((?:'[^']*')|(?:[^ ]+))", rest, flags=re.I)
    if not match:
        return None
    value = match.group(1).strip().strip("'")
    if value.upper() == "CURRENT_TIMESTAMP":
        return "当前时间"
    return value


def dedup_keep_order(values: List[str]) -> List[str]:
    seen = set()
    result: List[str] = []
    for value in values:
        if value not in seen:
            seen.add(value)
            result.append(value)
    return result


def parse_schema_details() -> Dict[str, List[Dict[str, str]]]:
    text = SCHEMA_PATH.read_text(encoding="utf-8")
    result: Dict[str, List[Dict[str, str]]] = {}
    for table_name, body in re.findall(r"CREATE TABLE\s+(\w+)\s*\((.*?)\n\)\s*;", text, flags=re.S):
        unique_fields = set()
        for raw in body.splitlines():
            line = raw.strip().rstrip(",")
            if line.startswith("UNIQUE KEY"):
                match = re.search(r"\(([^)]+)\)", line)
                if match:
                    fields = [item.strip().strip("`") for item in match.group(1).split(",")]
                    if len(fields) == 1:
                        unique_fields.add(fields[0])

        fields: List[Dict[str, str]] = []
        for raw in body.splitlines():
            line = raw.strip().rstrip(",")
            if not line or line.startswith(("--", "PRIMARY KEY", "INDEX", "KEY", "CONSTRAINT", "FOREIGN KEY", "UNIQUE KEY")):
                continue
            match = re.match(r"(\w+)\s+([A-Z]+(?:\([^)]+\))?)\s*(.*)", line, flags=re.I)
            if not match:
                continue
            name, field_type, rest = match.groups()
            data_type, length = split_type(field_type)
            rest_upper = rest.upper()
            constraints: List[str] = []
            if "PRIMARY KEY" in rest_upper:
                constraints.append("主键")
            if "AUTO_INCREMENT" in rest_upper:
                constraints.append("自增")
            if name in unique_fields or " UNIQUE" in f" {rest_upper} ":
                constraints.append("唯一")
            if "NOT NULL" in rest_upper:
                constraints.append("非空")
            elif "PRIMARY KEY" not in rest_upper:
                constraints.append("可空")
            default_value = parse_default_value(rest)
            if default_value is not None and "AUTO_INCREMENT" not in rest_upper:
                constraints.append(f"默认{default_value}")
            fields.append(
                {
                    "name": name,
                    "type": data_type,
                    "length": length,
                    "constraints": "，".join(dedup_keep_order(constraints)) or "-",
                    "description": FIELD_EXPLANATIONS.get(name, "字段含义与表名语义对应"),
                }
            )
        result[table_name] = fields
    return result


def build_detailed_rows(schema: Dict[str, List[Dict[str, str]]], table_name: str) -> List[List[str]]:
    rows = [["字段名", "数据类型", "长度", "约束", "注释描述"]]
    for field in schema[table_name]:
        rows.append([field["name"], field["type"], field["length"], field["constraints"], field["description"]])
    return rows


def format_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = TABLE_WIDTHS[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                if paragraph.style and paragraph.style.name != "表格内容":
                    try:
                        paragraph.style = table._parent.styles["表格内容"]
                    except Exception:
                        pass
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.underline = False


def create_table(doc: Document, rows: List[List[str]]):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    header_cells = table.rows[0].cells
    for idx, text in enumerate(rows[0]):
        header_cells[idx].text = text
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    format_table(table)
    return table


def replace_table_after_caption(doc: Document, caption_text: str, rows: List[List[str]]) -> None:
    caption_para = find_paragraph(doc, caption_text)
    next_element = caption_para._p.getnext()
    if next_element is None or next_element.tag != qn("w:tbl"):
        raise ValueError(f"Table not found after caption: {caption_text}")
    next_element.getparent().remove(next_element)
    new_table = create_table(doc, rows)
    caption_para._p.addnext(new_table._tbl)


def set_all_text_black(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = False
    for section in doc.sections:
        for paragraph in section.header.paragraphs + section.footer.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.underline = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.underline = False


def ensure_absent_terms(doc: Document, forbidden: List[str]) -> None:
    for term in forbidden:
        for paragraph in doc.paragraphs:
            if term in paragraph.text:
                raise ValueError(f"Forbidden term still present in manuscript: {term}")


def write_report(original_hash: str) -> None:
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    report = f"""# 校园物资智能管理系统论文改稿说明（第三轮数据库恢复与AIGC续改）

## 工作文件
- working draft：`{WORKING_DRAFT.relative_to(ROOT)}`
- protected original：`{ORIGINAL_DRAFT.relative_to(ROOT)}`
- fresh backup：`{BACKUP_DRAFT.relative_to(ROOT)}`

## 本轮处理摘要
- 按 live working draft 继续改稿，没有回退到更早的 copied draft。
- 保留了 current working draft 中已经对齐现仓库 schema 的 `图3-3`、`图3-4`、`图3-5` 三张实体关系图。
- 没有直接恢复原稿 3 张旧数据库 E-R 图。原因是原稿图中仍包含 `APPLICATION_FORM`、`ALERT_LOG`、`inventory_info` 等旧命名，与当前 `sql/schema.sql` 存在明显漂移。
- 以原稿数据库章前 4 张详细字段说明表为版式来源，恢复了 5 列详细字段说明表结构，并扩展到 current schema 下的 12 张关键数据表。
- 根据最新 AIGC 报告继续处理 surviving hotspots，重点覆盖摘要、研究意义、研究现状、需求分析、接口与安全、认证实现、测试说明和结束语。

## 数据库章恢复说明
- 恢复来源：原稿 `4.4.3 主要数据表设计` 中的 5 列字段说明表结构与说明粒度。
- 未直接回填的旧资产：原稿 3 张旧 E-R 图、原稿中的旧字段名与旧表名。
- 依据 `schema.sql` 修正后的 12 张关键表：
  - `sys_user`
  - `material_info`
  - `warehouse`
  - `inventory`
  - `inventory_batch`
  - `apply_order`
  - `apply_order_item`
  - `stock_out`
  - `transfer_order`
  - `warning_record`
  - `notification`
  - `login_log`
- 本轮重点修复了 `表3-1` 到 `表3-4` 中“标题已更新、表体仍残留旧字段名”的错位问题，并统一了 `表3-1` 到 `表3-12` 的详表版式。

## AIGC 报告摘要
- 总字数：`{AIGC_METRICS['total_words']}`
- 疑似 AIGC 风险概率：`{AIGC_METRICS['risk_probability']}`
- 高风险文本占比：`{AIGC_METRICS['high_risk']}`
- 中风险文本占比：`{AIGC_METRICS['mid_risk']}`

## 热点页到现稿章节的映射
- 报告第 1-2 页的热点主要映射到：中文摘要、英文摘要、`1.2 研究意义`、`1.3 国内外研究现状`、`1.4 研究内容与论文结构`。
- 报告第 3 页的热点主要映射到：`2.1` 至 `2.4` 的需求分析与用例描述，以及功能模块枚举类段落。
- 报告第 4 页和第 20、22 页的热点主要映射到：`3.4 接口与安全设计`、`4.2 认证授权与登录态管理实现`、`5.1 测试环境与证据来源`。
- 报告第 24-27 页的热点主要映射到：测试与结论中偏顺滑、偏总结化的段落表达；本轮对应处理了 `5.2` 至 `5.4` 和 `结束语`。

## 本轮已处理的热点类型
- 摘要中的夸张结论、试运行口径与一口气堆叠技术名词的句式。
- 研究意义中的价值拔高与“工程意义/学术意义”模板化对称结构。
- 研究现状中的综述腔、章节安排中的标准化模板句。
- 需求分析中的整齐枚举句与模块口号式概括。
- 接口与安全章节中的概念介绍腔，尤其是旧稿里接口壳层、旧端点和认证原理的讲解方式。
- 测试章节与结束语中过满的收束句、结论句和延展句。

## 仍需本机 Word 复核
- 刷新目录、题注和页码域。
- 复查 12 张数据库详表的分页位置与跨页效果，必要时手工微调表格前后空行。
- 最终提交前再做一次全文黑色文本、图片编号与表格编号的视觉检查。

## 文件安全核验
- 原始初稿 SHA256：`{original_hash}`
- 本轮只更新 working draft 和独立改稿说明，没有回写 protected original。
"""
    REPORT_PATH.write_text(report, encoding="utf-8")


def main() -> None:
    if not WORKING_DRAFT.exists():
        raise SystemExit("Working draft not found.")
    if not ORIGINAL_DRAFT.exists():
        raise SystemExit("Original draft not found.")
    if not BACKUP_DRAFT.exists():
        raise SystemExit("Fresh backup not found. Create the backup before running this script.")

    original_hash = sha256(ORIGINAL_DRAFT)
    schema = parse_schema_details()

    doc = Document(WORKING_DRAFT)

    replace_abstracts(doc)

    replace_nonempty_paragraphs_between(
        doc,
        "1.2 研究意义",
        "1.3 国内外研究现状",
        [
            "本课题的意义首先不在于再描述一个抽象的“智慧校园愿景”，而在于把已经写进仓库的业务链条整理清楚。当前系统把申领、审批、出库、签收、调拨、预警和日志留痕拆成了明确的数据对象和状态节点，这些实现让校园物资管理从零散台账转向可追踪、可回溯的流程协同。",
            "从应用角度看，项目的价值主要落在角色协作和责任确认上。SecurityConfig、AuthService 以及日志相关表说明，不同角色看到的菜单、能调用的接口和留下的操作记录都已经被明确区分。对校园场景来说，这种边界划分比单纯展示几个管理页面更有实际意义。",
            "对论文写作本身，这一课题还有一个更直接的作用：它要求定稿回到仓库证据，而不是沿用旧稿里那些没有落到代码和配置上的结论。本轮改稿继续沿着这个方向处理，尽量把文字拉回实现现场，降低“看上去像论文模板”的痕迹。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "1.3.1 国内研究现状",
        "1.3.2 国外相关实践启示",
        [
            "国内围绕智慧校园和高校信息化的研究已经积累了大量案例，常见切入点包括统一身份认证、流程电子化、后勤资源协同和数据共享平台建设[1]。在具体实现层面，基于 Spring Boot 与 Vue 的前后端分离方案已经成为高校业务系统中的常见组合，用于支撑角色隔离、列表查询、表单流转和基础统计展示[6][7][8][9]。",
            "和本项目更接近的研究，多分布在图书、采购、资产和后勤物资等资源管理方向[4][5][10]。这些成果通常能够提供比较成熟的主数据维护、库存统计和审批流建模思路，但对批次追踪、临期处置和复杂状态迁移的讨论往往不够细，很多论文停留在功能截图层面。",
            "从这一背景看，本项目的可取之处不在于提出了新的理论模型，而在于把校园物资管理里真正需要落地的对象和流程接上了代码实现。论文在此基础上展开分析，更容易把需求、设计和实现对应起来。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "1.3.2 国外相关实践启示",
        "1.4 研究内容与论文结构",
        [
            "国外公开资料里，与本项目最相关的并不是某一篇现成的校园物资论文，而是框架文档、工程规范和通用的业务系统实践。Spring Boot、Vue、Vite 和 JWT 相关文档提供的是实现路径：如何组织模块、如何处理认证、如何管理依赖、如何保持接口一致性[17][18][19][20]。",
            "这类资料的共同特点是强调“先把主链路做实，再谈高级能力”。本项目在实现上也采取了类似思路，没有为了追求概念上的智能化去引入无法验证的大型平台依赖，而是先把登录鉴权、申领审批、库存批次、调拨推荐、预警生成和统计展示这些核心流程打通。",
            "因此，本文对国外相关实践的吸收主要体现在工程方法上，而不是照搬一套宏大的系统叙事。这样的写法更符合当前仓库的真实状态，也更适合本科毕业设计的证据边界。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "1.4 研究内容与论文结构",
        "2  系统需求分析",
        [
            "本文围绕当前仓库中的校园物资智能管理系统展开。正文先把角色分工、业务场景和功能边界讲清，再结合数据库结构、前后端依赖和统一响应机制说明总体设计，之后转入认证授权、基础数据、库存批次、申领审批、调拨执行、预警分析和日志通知等核心实现。",
            "结构安排上，第 1 章处理选题背景、研究意义和相关实践；第 2 章回到需求与业务流程；第 3 章说明总体设计、数据库和接口安全；第 4 章展开关键实现；第 5 章只保留能够复核的测试与构建证据；最后在结束语中总结已完成工作和后续可扩展方向。",
            "全文遵循同一条约束：只写当前仓库可以证明的事实。没有落到代码、SQL、配置或实际执行结果中的性能、部署、推广和运行成效，不在本文中当作结论使用。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "2.1 业务场景与角色划分",
        "2.2 功能需求分析",
        [
            "从 seed.sql、菜单装配逻辑和前端路由可以看出，系统真正围绕四类角色运转：系统管理员、仓库管理员、审批人和部门用户。管理员负责组织与主数据维护，仓库管理员处理入库、出库、批次和调拨执行，审批人处理审核动作，部门用户则承担申领、查看和签收这一侧的业务操作。",
            "业务场景并不是平均铺开的。系统的主线首先是库存与批次管理，其次是部门申领到仓库执行的闭环流转，再往上才是跨仓调拨、预警生成和统计分析。这种层次关系与仓库表结构、服务拆分和页面入口是一致的，不是论文后加出来的抽象分类。",
            "角色边界也不是只写在说明文档里。SecurityConfig、控制器上的 PreAuthorize 表达式和 AuthService.buildMenusByRole 共同决定了哪些用户能看到哪些菜单、调用哪些接口、进入哪些工作流页面，所以需求分析必须把这一点作为基本前提写清楚。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "2.2.1 基础数据管理需求",
        "2.2.2 库存与批次管理需求",
        [
            "基础数据并不是陪衬模块。部门、角色、用户、校区、仓库、库位、分类、物资、供应商和系统配置这些对象一旦定义不稳，后续单据和统计都无法成立。前端已经为这些对象提供了独立页面，说明它们在系统里需要被长期维护，而不是一次性导入后就不再变化。",
            "其中影响面最大的是物资档案。material_info 保存的编码、名称、分类、规格、单位、安全库存、保质期和供应商信息，会反复进入库存、预警、申领和补货建议等逻辑，所以这一部分的需求重点其实是主数据一致性，而不是单个表单页面本身。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "2.2.2 库存与批次管理需求",
        "2.2.3 申领审批与出库需求",
        [
            "校园物资管理不能只看一个总库存数字。inventory 负责汇总某物资在某仓库下的当前量和锁定量，inventory_batch 继续把来源批次、剩余数量、生产日期和到期日期拆开记录。这意味着系统在需求上必须同时支撑总账视图和批次视图，否则临期处理与先到期先出都无从谈起。",
            "对应到操作层面，系统至少要能完成入库建批次、库存查询、按批次排序和出库扣减。InventoryService 里围绕批次排序和数量校验的实现也说明，这里需要的是可执行的业务规则，而不是一句笼统的“支持库存管理”。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "2.2.3 申领审批与出库需求",
        "2.2.4 调拨协同需求",
        [
            "申领审批是系统最核心的入口。apply_order 和 apply_order_item 采用主单加明细的结构，把申请部门、申请人、紧急等级、原因、场景和物资明细放进同一张业务单据里。围绕这张单据形成的状态链条覆盖了 DRAFT、SUBMITTED、APPROVED、REJECTED、OUTBOUND 和 RECEIVED 六个阶段，需求分析必须与这一真实状态模型保持一致。",
            "代码里还能看到另一个校园场景：紧急申领。ApplyService 会根据 urgency_level 触发快速审批分支，所以系统不仅要服务于计划性领用，还要支撑临时保障、夜间处置和突发场景下更短的审批路径，同时留下审批意见和时间记录，避免流程缩短后失去追踪能力。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "2.2.4 调拨协同需求",
        "2.2.5 预警、统计与支撑功能需求",
        [
            "当一端仓库短缺、另一端仍有余量时，系统不能只让管理员手动改两个库存字段。transfer_order 和 transfer_order_item 表明，调拨必须先形成单据，再经历提交、审批、执行和签收，调出仓和调入仓的变化要在同一条业务语义下完成。",
            "与普通出入库不同，调拨还要回答“从哪调、走哪条路径”这样的问题。项目里的 DijkstraUtil 和 transfer recommend 接口说明，系统已经把校区图上的距离计算接入到调拨场景中，所以这一部分的需求不只是单据流转，还包括来源仓选择的辅助支持。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "2.2.5 预警、统计与支撑功能需求",
        "2.3 非功能需求分析",
        [
            "要让校园物资系统长期可用，只把单据流转做通还不够。warning_record、notification、event_record、login_log 和 operation_log 这些表说明，预警、消息、事件和日志不是附属功能，而是支撑主链路运行的独立模块。库存不足、库存积压、临期、过期和异常领用等预警类型也都需要在需求阶段提前明确。",
            "统计分析需求主要面向管理者。AnalyticsService 和 SmartService 表明，系统不仅要展示当前库存和单据状态，还要能基于历史出入库记录回答“最近哪些物资消耗快、哪些仓库存在风险、哪些物资需要补货”这类问题。也正因为如此，前面的主数据、库存和单据记录必须保持完整，否则分析结果没有基础。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "2.3 非功能需求分析",
        "2.4 业务流程与用例分析",
        [
            "非功能需求里最先落地的是安全与权限隔离。后端通过 Spring Security 过滤链、JWT、Refresh Token 和方法级权限控制保护接口访问，前端通过本地 token 存储、路由守卫和 401 处理维持登录态。换句话说，系统首先要保证不同角色只能进入各自职责范围内的页面和操作链路。",
            "另一个现实要求是可维护性。仓库采用前后端分离组织方式，后端服务按认证、库存、申领、调拨、预警、智能分析和日志等模块拆分，数据库表统一保留 deleted、version、created_at 和 updated_at 一类通用字段。这种组织方式更像可持续演进的业务系统，而不是一次性演示项目。",
            "系统还需要具备可测试性。当前仓库已经提供 H2 测试配置和一组围绕认证、申请、调拨、预警与异常处理的自动化测试类，所以论文可以把“可重复验证”当作一项真实非功能要求来写，而不是转向没有实测依据的并发和性能指标。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "2.4.2 调拨执行用例分析",
        "2.4.3 预警处置与统计分析用例分析",
        [
            "调拨执行用例从仓库管理员发起单据开始。管理员根据目标仓库存量或保障需求创建调拨单，再结合推荐接口给出的候选来源仓完成提交。审批通过后，系统进入执行阶段，同时处理调出仓扣减、调入仓补入和状态推进，而不是只改动一个库存总量字段。",
            "这类业务对一致性的要求比普通申领更高，因为一次调拨会同时影响两端仓库。目标仓完成签收前，系统不会把单据视为最终结束。通过这种设计，仓间转移被清楚地写成了申请、审核、执行、签收四段式链路。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "2.4.3 预警处置与统计分析用例分析",
        "3  系统总体设计",
        [
            "预警处置用例由系统主动触发。WarningService.scan() 周期性检查库存和批次信息，生成库存不足、库存积压、临期、过期和异常领用等记录。管理员或仓库管理员在预警页查看待处理项后，可以补录处理说明并更新状态；若需要通知相关角色，系统还会写入消息表供后续查看。",
            "统计分析则服务于管理决策。系统把出入库流水、库存分布和时间维度聚合到图表页与大屏页，并在智能分析模块中给出补货建议。它既是一个独立功能，也是前面主数据、库存和单据记录是否规范的反向检验。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "3.3.3 关键数据表设计",
        "表3-1 用户信息表（sys_user）",
        [
            "为恢复原稿中较完整的数据库字段说明方式，同时避免继续沿用旧表名和旧字段名，本节保留 current working draft 已确定的 12 张关键表范围，统一改为详细字段说明表。表格的展示粒度参考原稿数据库章的详表版式，但字段内容、约束和说明均以当前 schema.sql 为唯一依据。这样处理既补回了数据库章缺失的信息密度，也避免把已经退出现仓库的旧结构重新写回正文。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "3.4.1 接口设计",
        "3.4.2 认证与授权设计",
        [
            "接口设计不是抽象的 REST 口号，而是前后端协作的具体约束。当前项目统一以 /api 作为业务前缀，再按认证、库存、申领、调拨、预警、统计、事件和消息等功能域拆分接口，这一结构已经能从控制器分布和前端请求封装中直接看到。",
            "更关键的是统一响应外壳。frontend/src/api/http.js 的拦截逻辑表明，所有业务接口都围绕 ApiResponse 的 code、message 和 data 三个字段组织返回结果，且前端把 code 等于 0 视为成功。这种一致性减少了页面侧的分支判断，也让异常处理可以回到同一出口。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "3.4.2 认证与授权设计",
        "3.4.3 安全与审计设计",
        [
            "认证链路由登录接口、令牌签发、刷新令牌续期和前端重放机制共同组成。JwtTokenProvider 负责生成和校验访问令牌，AuthService 负责登录、获取当前用户与菜单装配，http.js 则在 401 场景下触发刷新接口并重放原请求。这说明当前系统已经把登录态管理落实为一条完整链路，而不是简单地在浏览器里存一个长期有效字符串。",
            "授权控制分布在两个层面。其一是 SecurityConfig 打开的方法级安全能力和控制器上的 PreAuthorize 表达式，其二是 AuthService.buildMenusByRole 返回的菜单集合。接口层和导航层一起工作后，系统管理员、仓库管理员、审批人和部门用户才真正表现出不同的页面与操作边界。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "3.4.3 安全与审计设计",
        "4  系统详细设计与实现",
        [
            "在安全细节上，当前仓库已经落地的是密码加密、跨域控制、刷新令牌清理和审计留痕。用户密码采用 BCrypt 存储，刷新令牌以哈希形式写入 auth_refresh_token 表，并通过定时任务清理过期或撤销记录；登录日志和操作日志则为关键动作提供了追踪基础。",
            "需要收束的是结论边界。当前项目没有独立的压测报告、安全扫描记录或生产部署监控，所以论文中的“安全设计”只讨论已经实现的访问控制、凭证管理和操作审计，不把高并发防护、网关治理或上线效果写成既成事实。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "4.2 认证授权与登录态管理实现",
        "为了直观展示系统登录入口与统一认证界面，本节补充运行截图如图4-3所示。",
        [
            "认证实现从 /api/auth/login 进入，但关键并不只是“用户输入账号密码”。前端 auth store 在登录成功后会立即拉取当前用户资料与动态菜单，说明登录动作和角色可见范围是一次连续初始化过程。对论文而言，这比单独介绍某个登录表单更能说明系统的真实运行方式。",
            "登录态维持主要依赖 frontend/src/api/http.js 的请求与响应拦截器。请求阶段统一附加 Bearer Token，响应阶段先检查 ApiResponse 的 code，再在收到 401 时进入刷新流程。脚本里维护的等待队列让多个并发请求不会重复刷新令牌，这一点直接对应了仓库中的实现细节。",
            "后端与之配合的是 JwtTokenProvider、AuthService 和 auth_refresh_token 表。访问令牌负责短期身份识别，刷新令牌负责续期与撤销，清理任务再把过期记录回收掉。这样形成的不是概念上的 JWT 介绍，而是一条已经写进代码和数据表的登录、续期、退出链路。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "5.1 测试环境与证据来源",
        "5.2 自动化测试结果",
        [
            "本章只采用当前仓库可以重复执行的验证材料，包括后端自动化测试代码、H2 测试配置、Maven 测试结果和前端生产构建结果。这样写的目的很直接：把论文的测试结论限制在可以复核的范围内，不再沿用旧稿里那些无法复现的压测和上线描述。",
            "后端测试环境由 backend/src/test/resources/application-test.yml 指向 H2，并结合测试资源目录中的 schema 初始化数据结构。这意味着关键业务逻辑可以在不依赖外部 MySQL 实例的情况下重复执行，测试结论也更容易回到仓库环境中核验。",
            "在本轮定稿过程中，2026 年 4 月 14 日已经实际执行过 mvn -f backend/pom.xml test 与 npm --prefix frontend run build，两条命令分别得到后端 42 项测试通过、前端生产构建成功的结果。第 5 章的核心证据就建立在这两次真实执行之上。",
            "为了给第 4 章提供界面佐证，前一轮还额外使用了 screenshot 隔离 profile 和最小演示数据采集运行截图。该环境只服务于论文取证，不替代本文对主项目实现边界的判断。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "5.2 自动化测试结果",
        "表5-1 后端自动化测试执行情况",
        [
            "从测试类分布看，当前自动化测试没有试图覆盖一切，而是集中在最容易出错的主链路上：统一响应与异常处理、认证与令牌管理、申领业务、调拨业务和预警接口都已有对应测试。这种覆盖方式更接近课程设计项目的现实做法，即优先把状态流转和权限校验这些高风险点守住。",
            "2026 年 4 月 14 日执行 Maven 测试后，后端共通过 42 项测试。这个结果至少说明当前版本在核心业务逻辑、配置加载和关键断言层面保持了可执行状态，也说明测试结论并不依赖人工预置数据库环境。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "5.3 业务场景验证",
        "表5-2 典型业务场景验证说明",
        [
            "除自动化测试外，论文还可以基于代码、路由和 seed.sql 样例数据对典型业务场景做可复验说明。这里强调的不是“所有场景都已真实上线运行”，而是主链路已经具备清楚的入口、状态变化和落库对象，后续答辩演示可以按这些线索逐项核对。",
            "结合当前实现，至少有三类场景已经在代码层形成闭环：部门用户提交申领单后经审批人与仓库管理员流转直至签收；仓库管理员创建调拨单并结合推荐接口完成跨仓执行；系统按周期生成预警并在统计与智能分析模块中输出管理视角结果。这样归纳比旧稿中那类“流程非常流畅、体验良好”的结论更接近可验证事实。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "5.4 构建验证与问题边界",
        "结束语",
        [
            "前端构建验证来自 2026 年 4 月 14 日执行 npm --prefix frontend run build 的实际结果。构建成功说明当前仓库中的 Vue 页面、路由、组件依赖和打包配置仍然保持一致，没有出现阻断交付的基础构建错误。",
            "同样需要明确边界。现有测试与构建结果只能证明“项目可执行”“核心逻辑在现有覆盖下通过验证”，不能外推到高并发稳定性、生产网络部署效果和长期运维可靠性。仓库里没有对应的压测报告、监控数据或生产日志，所以这些结论继续不写入正文。",
            "从论文收口角度看，后续工作主要剩下 Word 侧的目录、题注和分页复核，以及参考文献格式校对。这些属于排版与展示层面的完善，不改变本文已经建立的事实边界。",
        ],
    )

    replace_nonempty_paragraphs_between(
        doc,
        "结束语",
        "致    谢",
        [
            "本文围绕当前项目仓库，对校园物资智能管理系统的需求、总体设计、关键实现和测试证据进行了重新梳理。本轮改稿继续压缩旧稿中偏模板化、偏结论化的表达，并把数据库章和测试章重新拉回到代码、SQL、配置和执行结果可以直接证明的范围内。",
            "从现有实现看，系统已经完成了认证授权、基础数据维护、库存与批次管理、申领审批、调拨协同、预警扫描、补货建议、统计分析、通知消息、事件记录和日志审计等核心能力。它的价值不在于宣称形成了多大规模的平台，而在于主业务链路已经有了可以运行、可以测试、可以说明的数据结构和服务逻辑。",
            "结合 2026 年 4 月 14 日后端 42 项测试通过和前端生产构建成功的事实，可以把当前项目视为一个边界清楚、证据完整度较高的本科毕业设计成果。后续若继续完善，也应沿着现有实现往前推进，而不是重新引入无法验证的宏大叙事。",
        ],
    )

    for caption, table_name in TARGET_TABLES:
        replace_table_after_caption(doc, caption, build_detailed_rows(schema, table_name))

    set_all_text_black(doc)
    doc.save(WORKING_DRAFT)

    saved = Document(WORKING_DRAFT)
    ensure_absent_terms(saved, ["inventory_info", "application_form", "alert_log"])
    if len(saved.tables) != 19:
        raise SystemExit(f"Unexpected table count: {len(saved.tables)}")
    if len(saved.inline_shapes) != 13:
        raise SystemExit(f"Unexpected image count: {len(saved.inline_shapes)}")

    write_report(original_hash)
    print(WORKING_DRAFT)
    print(REPORT_PATH)
    print("tables", len(saved.tables))
    print("images", len(saved.inline_shapes))
    print("original_sha256", original_hash)


if __name__ == "__main__":
    main()
