from __future__ import annotations

import re
import shutil
from collections import OrderedDict
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "output" / "doc"
SOURCE = DOC_DIR / "校园物资智能管理系统设计与实现-正文扩写定稿版.docx"
TARGET = DOC_DIR / "校园物资智能管理系统设计与实现-正文压缩版.docx"
REPORT = DOC_DIR / "正文压缩说明.md"


REPLACEMENTS = [
    (96, "校园物资管理同时具有行政管理、仓储管理和服务保障属性", "校园物资管理兼具行政管理、仓储管理和服务保障属性。高校物资种类分散、使用部门多、消耗节奏受教学安排和突发事件影响明显。若仍依赖纸质台账或分散 Excel，管理人员难以及时掌握全校库存余量、批次效期和申领进度，因此需要以信息化手段实现物资流转闭环[1]。"),
    (97, "系统并非单一库存查询工具，而是围绕物资档案、仓库、库存、批次、申领、审批、出库、调拨、预警和统计分析构建的校园物资管理业务系统", "本系统围绕物资档案、仓库、库存、批次、申领、审批、出库、调拨、预警和统计分析构建。sql/schema.sql 中的 material_info、warehouse、inventory、inventory_batch、apply_order、transfer_order 和 warning_record 等表说明，系统不仅记录数量，还保留来源、去向、状态、处理人与时间信息，为流程追踪、批次核对和库存预警提供数据基础。"),
    (99, "高校日常运行离不开实验耗材、办公用品、医疗防护、清洁消杀、应急照明和设备备件等物资", "高校运行涉及实验耗材、办公用品、医疗防护、清洁消杀和设备备件等多类物资。传统管理方式容易出现数据更新滞后、审批记录不完整和库存批次不透明等问题，难以支撑跨部门协同和风险提醒[1][2]。"),
    (103, "本系统的实践意义首先体现在库存数据的统一维护", "系统的应用价值体现在库存统一维护和业务状态可追踪。管理员通过入库、出库和批次记录同步维护库存数据，部门用户能够查看申领从草稿、提交、审批到签收的状态变化，审批人员也可以依据单据明细和审批意见进行追溯。"),
    (104, "本系统的工程意义在于把前后端分离、权限控制、事务处理、批次扣减和预警扫描等技术组织成一个完整的管理系统", "系统的工程价值在于将前后端分离、权限控制、事务处理、批次扣减和预警扫描组织为完整的管理闭环。相较一般信息化系统，本课题更强调库存锁定、FEFO 扣减、候选仓排序和规则化预警等与校园物资业务直接相关的实现细节。"),
    (107, "本课题的工程重点在于将需求分析、数据库设计、前后端分离、权限控制、事务处理和测试验证结合起来", "本课题的工程重点在于将需求分析、数据库设计、前后端分离、权限控制、事务处理和测试验证结合起来，形成面向校园物资管理的完整业务系统。系统围绕 sys_user、material_info、inventory_batch、apply_order、transfer_order 和 warning_record 等真实数据对象组织功能。"),
    (110, "国内研究中，与本系统最接近的是高校、医疗和后勤物资管理系统的设计与实现类成果", "国内相关研究多以高校、医疗和后勤物资管理为对象，通常将入库、领用、审核和库存监控作为核心模块[1][2][4]。这些成果说明校园物资系统应覆盖申请、审批、仓库、批次、预警和统计，而不是停留在库存台账层面。"),
    (112, "国外或外文资料中，基于 Spring Boot、Vue 和 MySQL 的 Web 系统开发、安全设计和管理信息系统实现也有相关研究", "国外和外文资料更多从 Web 系统工程实践与安全实现角度提供参考[10][11][12][13]。相关研究支持本文对分层架构、数据库持久化和认证授权的讨论，但本系统的实现范围仅限于 JWT、refresh token、BCrypt、CORS、请求限流和安全响应头等已落地内容。"),
    (115, "与本课题更接近的研究集中在物资流通、后勤物资、实验室危化品和库存监控等方向", "与本课题更接近的研究集中在物资流通、后勤物资和实验室物资管理等方向。相关成果为入库、领用、审批和库存追踪提供了流程建模参考，但本文仍以当前项目的代码、数据库和页面实现为依据。"),
    (116, "国内近年的智能管理研究也开始关注预警、流程自动化和数据分析", "近年来也有研究将预警、自动化流程和数据分析引入物资管理[21][22]。相比之下，本系统采用固定阈值、时间窗口、移动平均和候选排序等轻量方法，更符合本科毕业设计的实现规模。"),
    (121, "本文主要研究内容包括：分析校园物资管理的角色和业务流程", "本文围绕校园物资管理系统的需求分析、总体设计、数据库设计、详细实现和系统测试展开。全文共 7 章：第 1 章介绍研究背景、意义与现状；第 2 章说明本系统采用的关键技术；第 3 至第 6 章分别阐述需求分析、总体设计、数据库设计和详细实现；第 7 章给出测试结果，最后给出结束语、致谢、参考文献和附录。"),
    (124, "技术选型的依据来自 backend/pom.xml、frontend/package.json 和后端配置文件", "技术选型以 backend/pom.xml、frontend/package.json 和配置文件为依据。后端采用 Spring Boot 3.3.5、Spring Security、MyBatis-Plus、JdbcTemplate、JJWT、Bucket4j、Caffeine 等组件，支撑认证授权、事务处理、数据访问、限流和缓存；前端采用 Vue 3.5、Vite 6、Pinia、Vue Router、Element Plus、Axios 和 ECharts，负责页面组织、状态管理、接口调用与图表展示；Vitest 与 Playwright 用于验证前端功能。"),
    (127, "根据 frontend/package.json 和 backend/pom.xml，系统实际使用的核心技术包括", "本系统的技术组合围绕“业务闭环”和“数据可追踪”两项目标展开。前端负责表单输入、列表筛选、状态展示和图表渲染，后端负责认证、状态流转、事务一致性和聚合查询，MySQL 负责持久化物资、库存、单据、日志和通知等数据。"),
    (129, "前端工程采用 Vite 作为开发与构建工具", "前端以 Vue 3 组件化页面为核心，Vue Router 管理登录、基础数据、库存、申领、调拨、预警、统计、日志和通知等路由，Pinia 保存登录状态和菜单信息，Axios 统一封装请求。该组合满足管理系统多页面、状态联动和接口复用的需求。"),
    (130, "Element Plus 在本系统中主要承担表单、表格、对话框、分页、标签和消息提示等管理系统常见交互", "Element Plus 为本系统提供表单、表格、对话框、分页和消息提示等通用交互组件，ECharts 负责库存占比、出入库趋势、部门排行、效期统计和仓库分布等图表展示，相关页面数据均来自后端统计接口。"),
    (134, "后端采用 Controller、Service、Mapper 和数据库表的分层结构", "后端按 Controller、Service、Mapper 和 Entity 分层组织。Controller 负责接收 REST 请求与响应封装，Service 负责业务规则、状态控制和事务处理，Mapper 通过 MyBatis-Plus 完成实体映射，JdbcTemplate 负责统计分析和补货建议中的聚合查询。"),
    (135, "MyBatis-Plus 在项目中承担基础 CRUD、分页查询、逻辑删除和乐观锁更新等常见数据访问工作", "MyBatis-Plus 主要承担基础 CRUD、分页查询、逻辑删除和乐观锁更新。BaseEntity 统一提供 deleted、version、created_at 和 updated_at 字段，使多数业务表具备软删除、版本控制和审计时间能力。"),
    (137, "数据层使用 MySQL 8 作为业务数据库，测试和截图场景使用 H2 模拟 MySQL 模式", "MySQL 8 用作正式业务数据库，H2 用于测试和截图环境。系统索引主要覆盖用户角色、批次出库、预警查询、日志时间和通知时间等高频访问场景。"),
    (139, "认证授权由 Spring Security、JWT 和自定义过滤器共同完成", "认证授权由 Spring Security、JWT 和自定义过滤器实现。SecurityConfig 放行登录、刷新令牌和接口文档等少数地址，JwtAuthenticationFilter 解析 Authorization 头并构建登录上下文，AuthService 负责用户名、密码和账号状态校验。"),
    (140, "安全设计还包括 refresh token 持久化与轮换", "refresh token 采用持久化与轮换策略。auth_refresh_token 记录 token_id、token_hash、expire_at 和 revoked 等字段，刷新时校验令牌类型、哈希、有效期和撤销状态，成功后撤销旧令牌并签发新令牌。"),
    (145, "系统角色以 AuthService.buildMenusByRole 和初始化数据为准", "系统角色以 AuthService.buildMenusByRole 和初始化数据为准，分为管理员 ADMIN、仓储管理员 WAREHOUSE_ADMIN、审批人员 APPROVER 和部门用户 DEPT_USER。管理员负责全局配置与基础数据维护，仓储管理员负责仓库、库存、出入库、调拨和预警，审批人员负责申领、调拨和预警审核，部门用户负责申领与签收。"),
    (146, "从业务边界看，系统支持校园物资从建档、入库、库存汇总、申领审批、出库、签收、调拨、预警到统计分析的完整闭环", "从业务边界看，系统覆盖物资建档、入库、库存汇总、申领审批、出库、签收、调拨、预警和统计分析等环节。当前实现没有独立配送任务表和配送角色，相关动作主要体现在申领出库签收和调拨执行接收流程中。"),
    (149, "基础数据需求包括部门、校区、仓库、库位、供应商、物资分类和物资档案维护", "基础数据需求包括部门、校区、仓库、库位、供应商、物资分类和物资档案维护。物资档案中的安全库存、保质期、仓库和校区等字段会直接影响预警、补货建议、库存分布和调拨推荐。"),
    (150, "库存业务需求包括库存查询、批次查询、入库登记、出库登记和盘点调整", "库存业务需求包括库存查询、批次查询、入库、出库和盘点调整。入库需要记录仓库、来源类型、物资明细、批次、数量和效期并同步更新库存汇总；出库需要校验可用库存、按 FEFO 扣减批次并回写出库单与申领状态；盘点调整当前主要用于库存汇总数量修正。"),
    (151, "申领审批需求体现部门用户和审批人员之间的业务协作", "申领审批需求体现为部门用户创建申领、系统锁定可满足需求的仓库库存、审批人员执行通过或驳回、仓储人员完成出库、部门用户完成签收。ApplyService 中的 DRAFT、SUBMITTED、APPROVED、OUTBOUND、RECEIVED 和 REJECTED 状态构成申领流程的核心依据。"),
    (152, "调拨需求解决不同仓库之间库存分布不均的问题", "调拨需求用于解决跨仓库存分布不均。调拨单需经过创建、提交、审批、执行和接收等状态转换，执行时同时更新来源仓和目标仓的库存与批次；候选仓推荐仅基于固定校区图和库存数量排序。"),
    (153, "预警与统计需求包括低库存、库存积压、临期、过期和异常消耗预警", "预警与统计需求包括低库存、库存积压、临期、过期和异常消耗预警，以及库存占比、出入库趋势、部门排行、效期统计、仓库分布和应急消耗等分析主题。系统还通过通知、登录日志和操作日志提供提醒与追溯能力。"),
    (156, "基础数据是系统运行的前提", "基础数据是库存、申领、调拨和预警的前提。物资编码、分类、安全库存、保质期、仓库和库位等信息会被多个业务模块重复引用，因此需要保证主数据完整和一致。"),
    (158, "校园物资管理不能只依赖一个总库存数字", "库存批次管理要求同时保存汇总库存和批次余量、生产日期、过期日期等信息，以支持 FEFO 出库和临期、过期预警。"),
    (160, "申领流程要求部门用户先创建单据，再提交进入审批", "申领流程要求部门用户先创建并提交单据，系统按可用库存选择仓库并锁定库存，审批通过后进入出库环节，最终由部门用户签收形成闭环。"),
    (162, "调拨流程用于处理跨仓库存不平衡", "调拨流程需要记录调出仓、调入仓、物资和数量，并通过提交、审批、执行和接收等状态节点反映跨仓协同过程。"),
    (164, "预警需求包括低库存、库存积压、临期、过期和异常出库", "预警与统计功能用于发现库存风险和展示物资流转情况，为管理人员提供规则化提醒和数据支持。"),
    (166, "安全性需求要求系统能够区分未登录、登录失败、账号禁用、令牌失效和权限不足等情况", "安全性要求系统能够区分未登录、登录失败、账号禁用、令牌失效和权限不足等情况。后端通过 Spring Security、JWT、BCrypt 和 refresh token 轮换控制访问边界，前端在无 token 时跳转登录页并限制管理员页面访问。"),
    (167, "一致性需求主要出现在库存相关流程", "一致性要求主要体现在库存相关流程。入库、出库、申领锁定和调拨执行都会同时影响多张表，因此关键 Service 需要在事务中完成库存校验、批次扣减、锁定释放和单据回写。"),
    (168, "可维护性需求体现在模块划分和通用字段设计", "可维护性体现在模块化组织和通用字段设计。后端按 auth、rbac、material、inventory、apply、transfer、warning、analytics 等模块划分，前端按 views、router、api 和 store 组织，多数业务表统一保留 deleted、version、created_at 和 updated_at 字段。"),
    (182, "系统采用典型 B/S 架构和前后端分离模式", "系统采用 B/S 架构和前后端分离模式。浏览器访问 Vue 前端，前端通过 Axios 调用 Spring Boot REST 接口，后端在 Spring Security 过滤链完成认证后进入 Controller、Service、Mapper 和 MySQL，实现页面、业务和数据持久化的分层处理。"),
    (183, "后端分层并不意味着系统采用微服务", "当前项目为单体 Spring Boot 应用，各模块通过包结构和 Service 分层划分职责。Actuator、Prometheus 指标和 Sentry 依赖主要用于工程观测与异常采集，不涉及微服务拆分和分布式部署。"),
    (188, "系统功能模块按照实际代码目录和前端路由划分", "系统按实际代码目录和前端路由划分为认证授权、RBAC、基础数据、库存出入库、申领审批、调拨管理、预警中心、统计分析、通知消息、事件管理和日志审计等模块，模块之间通过业务单据与数据库主键关联。"),
    (192, "接口设计遵循资源路径和业务动作相结合的方式", "接口设计遵循统一前缀与业务动作结合的方式。认证接口集中在 /api/auth，库存、申领、调拨、预警和统计分别使用独立资源路径；列表查询采用分页参数，提交、审批、驳回、执行、接收和处理等状态动作通过 submit、approve、reject、execute、receive 和 handle 等路径表达。"),
    (193, "如图4-3所示，前端页面到后端数据库的交互流程包括页面输入、请求封装、令牌携带、接口分发、业务处理、数据访问和页面刷新", "如图4-3所示，前后端交互链路包括页面输入、请求封装、JWT 鉴权、Controller 分发、Service 处理、Mapper 访问数据库和结果回传。前端根据返回数据更新 Pinia 状态、表格和 ECharts 图表，这一链路是各业务模块共享的基本执行模式。"),
    (197, "系统接口统一以 `/api` 为前缀", "系统接口统一以 /api 为前缀，登录、刷新、当前用户和菜单由 AuthController 提供，申领、库存、调拨、预警、统计、通知和日志由对应 Controller 提供。统一响应使用 ApiResponse，分页结果使用 PageQuery 和 PageResult。"),
    (198, "安全设计包括三层边界", "安全设计包含登录认证、JWT 鉴权和角色边界三层控制。用户通过用户名和 BCrypt 密码获取令牌，后端从 Authorization 头解析身份，方法级权限与菜单控制共同限制不同角色的可用功能，refresh token 以 hash 形式持久化并在刷新后轮换。"),
    (200, "本系统的智能能力设计重点是“可解释”和“可复核”", "系统的智能能力以规则可解释和结果可复核为前提。低库存、积压、临期、过期和异常消耗预警分别基于安全库存、库存倍数、批次效期和近 7 天与近 30 天出库量对比实现；补货建议结合安全库存、当前库存和近 30 天出库量计算建议数量；移动平均预测基于近 6 个月出库历史生成趋势参考。"),
    (201, "调拨推荐的实现重点在于为候选仓选择提供辅助依据", "调拨推荐用于辅助候选仓选择。TransferService.recommendTransfer 仅在满足数量需求的仓库中结合 DijkstraUtil 计算的校区距离排序，不涉及路径优化、车辆调度或实时交通决策。"),
    (206, "数据库设计以业务闭环和数据可追溯为基本原则", "数据库设计以业务闭环和数据可追溯为原则。用户、角色和部门承载组织与权限；物资、仓库、校区和库位构成基础数据；库存、批次、入出库、申领和调拨承载核心业务；预警、通知和日志提供支撑能力。各表之间主要通过业务字段形成逻辑关联。"),
    (207, "多数业务表使用 `id` 作为自增主键", "多数业务表使用 id 主键，deleted 表示逻辑删除，version 用于版本控制，created_at 和 updated_at 用于审计追踪。索引围绕用户角色、库存唯一约束、批次效期、申领与调拨状态、预警筛选以及日志通知分页等高频查询场景设置。"),
    (210, "系统实体关系可分为三组", "系统实体关系可分为权限组织、物资库存和业务单据三组，分别支撑身份认证、库存批次管理以及申领调拨与审计闭环。"),
    (218, "用户权限相关表包括 `sys_dept`、`sys_role`、`sys_user` 和 `auth_refresh_token`", "用户权限相关表包括 sys_dept、sys_role、sys_user 和 auth_refresh_token。它们分别保存部门层级、角色编码、用户账号状态以及 refresh token 的 token_id、token_hash、过期时间和撤销状态，用于支撑登录认证、角色菜单和令牌轮换。"),
    (219, "物资基础数据表包括 material_category、material_info、supplier、warehouse、campus 和 storage_location", "物资主数据表包括 material_category、material_info、supplier、warehouse、campus 和 storage_location。material_info 通过 material_code 唯一约束保存物资分类、安全库存和保质期参数；warehouse、campus 和 storage_location 共同描述仓库空间及其负责人信息。"),
    (220, "库存与批次表包括 `inventory` 和 `inventory_batch`", "inventory 与 inventory_batch 分别保存物资在仓库维度的汇总库存、锁定数量以及批次号、剩余数量、生产日期和过期日期。出库与调拨均按未过期且有余量的批次排序扣减，因此 inventory_batch 是 FEFO 规则的关键载体。"),
    (221, "入库与出库单据表包括 `stock_in`、`stock_in_item`、`stock_out` 和 `stock_out_item`", "stock_in、stock_in_item、stock_out 和 stock_out_item 构成入出库主从单据结构。主表记录仓库、来源类型、操作人和关联单据，明细表记录物资、批次和数量，用于支撑单据查询和数量统计。"),
    (222, "申领相关表包括 apply_order 和 apply_order_item", "apply_order 与 apply_order_item 用于保存部门、申请人、紧急等级、状态、锁定仓库、审批信息以及物资申请数量、实发数量等字段，是申领状态控制的核心数据来源。"),
    (223, "调拨相关表包括 `transfer_order` 和 `transfer_order_item`", "transfer_order 与 transfer_order_item 保存调出仓、调入仓、审批信息和调拨数量。调拨执行不仅推进单据状态，还会同步影响来源仓和目标仓的库存汇总与批次记录。"),
    (224, "预警、日志和通知支撑表包括 `warning_record`、`operation_log`、`login_log` 和 `notification`", "warning_record、operation_log、login_log 和 notification 分别用于保存预警内容、操作轨迹、登录行为和站内消息，使系统具备问题提醒、业务追溯和运行支撑能力。"),
    (227, "系统索引围绕实际查询和状态过滤设计", "索引与约束围绕实际查询场景设计。sys_user.username、sys_role.role_code、material_info.material_code 以及 inventory(material_id, warehouse_id) 使用唯一约束保证关键对象不重复；idx_batch_outbound_pick 支撑按物资、仓库和效期查找批次；预警、日志和通知则按类型、状态、时间和用户建立索引以支持分页筛选。"),
    (230, "开发环境与运行环境以项目文件为准", "开发与运行环境以项目文件为准。后端使用 Java 17、Spring Boot 3.3.5 和 Maven，默认由 application.yml 激活 dev profile，测试环境使用 H2 与 application-test.yml；前端使用 Node.js、Vue 3、Vite、Element Plus 和 ECharts，Vitest 与 Playwright 用于自动化验证。"),
    (231, "系统在本地演示时通常由后端 8080 端口提供业务接口", "本地演示由后端 8080 端口提供业务接口，管理端点通过 18080 暴露观测信息。当前验证范围限于本地开发、自动化测试和性能基线场景，不对应线上运行规模。"),
    (234, "登录流程从前端登录页开始", "登录时，前端调用 /api/auth/login，后端 AuthService.login 读取用户名、IP 和 user_agent，从 sys_user 查询用户并使用 BCrypt 校验密码与账号状态，认证通过后生成 access token 和 refresh token，将 refresh token 的 token_id、token_hash、过期时间和撤销状态写入 auth_refresh_token，并同步记录登录日志。"),
    (235, "请求鉴权由 Spring Security 过滤链完成", "请求鉴权由 Spring Security 过滤链完成。SecurityConfig 放行登录、刷新令牌和接口文档等少数地址，其余业务接口需要登录；前端通过 Axios 在 Authorization 头中携带 token，JwtAuthenticationFilter 解析令牌后将用户身份放入安全上下文。"),
    (236, "refresh token 采用持久化和轮换策略", "refresh token 采用单次使用和轮换策略。刷新时需要同时校验令牌类型、token_id、哈希、撤销状态和有效期，校验通过后撤销旧令牌并签发新的 token 对，以降低令牌重复使用风险。"),
    (237, "角色菜单由后端 `AuthService.buildMenusByRole` 统一生成", "角色菜单由 AuthService.buildMenusByRole 统一生成，前端根据返回的用户信息和菜单渲染导航，并对安全策略等页面追加管理员角色限制，从而保持页面可见范围与后端权限边界一致。"),
    (245, "用户与角色管理由 `RbacController` 和 `RbacService` 支撑", "用户与角色管理由 RbacController 和 RbacService 支撑。新增用户必须填写密码并通过 BCrypt 加密保存，修改用户时未提交新密码则保留原密码；角色表保存角色编码和名称，部门表保存部门层级，相关维护操作均写入操作日志。"),
    (246, "基础数据模块还包括校区、仓库、库位、供应商、物资分类和物资档案", "基础数据模块覆盖校区、仓库、库位、供应商、物资分类和物资档案，对应 MaterialController、WarehouseController、StorageLocationController 等接口。其主要职责是维护支撑库存、申领、调拨和预警的主数据。"),
    (247, "物资档案维护的重点是物资编码唯一、分类归属、安全库存和效期参数", "物资档案维护重点是物资编码唯一、分类归属、安全库存和效期参数。安全库存驱动低库存与积压规则，保质期和批次过期日期支撑临期提醒，仓库与校区信息服务于库存分布和调拨推荐。"),
    (250, "入库管理由前端 `StockInView.vue`、后端 `/api/inventory/stock-in` 接口和 `InventoryService.stockIn` 方法支撑", "入库管理由 StockInView.vue、/api/inventory/stock-in 接口和 InventoryService.stockIn 方法实现。用户填写仓库、来源类型和物资明细后，系统创建 stock_in 主记录与 stock_in_item 明细，未提供批次号时由后端按时间生成批次号。"),
    (251, "入库过程的关键不是保存单据本身，而是同步库存汇总和批次余量", "入库的关键在于同步库存汇总和批次余量。系统先获取或创建 inventory 汇总记录，再增加 current_qty 并写入 inventory_batch；整个过程在事务中完成，避免单据保存成功但库存未同步更新。"),
    (253, "出库管理由 `/api/inventory/stock-out` 接口和 `InventoryService.stockOut` 方法实现", "出库由 /api/inventory/stock-out 接口和 InventoryService.stockOut 方法实现。普通出库可以独立创建，申领出库则需要校验申领单状态、锁定仓库和明细实发数量，避免超出申请范围或从错误仓库出库。"),
    (254, "出库扣减采用 FEFO 思路，即优先扣减过期日期更早且未过期的批次", "出库扣减采用 FEFO 规则。系统按 expire_date 和 id 升序查询未过期且有余量的批次，按出库数量逐个扣减批次剩余量；若批次总量不足则抛出业务异常。扣减成功后同步减少库存汇总和申领锁定数量，并回写实际出库数量。"),
    (255, "出库业务流程如图6-10所示", "出库完成后系统写入 stock_out_item 明细，关联申领时把申领单状态更新为 OUTBOUND，并在库存低于安全库存时生成低库存预警。"),
    (261, "申领单创建由部门用户发起", "部门用户创建申领单时，前端提交物资、数量、紧急等级、申请原因和场景说明，ApplyService.create 生成 apply_order 主表和 apply_order_item 明细，初始状态为 DRAFT。草稿机制便于用户先保存明细再提交审批。"),
    (262, "提交申领时，系统会汇总所有明细物资的申请数量", "提交申领时，系统汇总各物资需求，在库存表中选择能够一次性满足需求的仓库，并将对应数量写入 locked_qty、仓库 id 写入 reserved_warehouse_id。若库存不足或没有仓库可满足需求，提交会失败。"),
    (263, "审批流程严格依赖状态", "申领审批严格依赖状态流转。普通申领由 SUBMITTED 经审批进入 APPROVED，驳回时释放锁定库存；紧急等级达到阈值时可走快速通道并自动记录审批信息。仓库出库后单据进入 OUTBOUND，部门用户签收后进入 RECEIVED 并释放剩余锁定量。"),
    (265, "调拨模块用于处理不同仓库之间的物资调剂", "调拨单需要填写调出仓、调入仓和明细，状态依次经历 DRAFT、SUBMITTED、APPROVED、OUTBOUND、RECEIVED，驳回状态为 REJECTED。该状态集合对应当前代码中的完整调拨流程。"),
    (266, "调拨执行会对库存和批次产生实质影响", "调拨执行时，系统先校验来源仓库存，再按 FEFO 顺序扣减来源仓批次，同时在目标仓生成带有调拨标记的镜像批次，并同步更新两侧库存汇总。整个过程位于事务中，确保批次搬移和库存变更一致。"),
    (267, "候选仓排序接口为调拨申请提供参考", "候选仓排序接口先筛选库存满足需求的仓库，再根据固定校区图计算距离并排序，为调拨申请提供参考，不涉及物流优化算法或实时地图调度。"),
    (277, "预警扫描由 `WarningService.scan` 统一触发", "预警扫描由 WarningService.scan 统一触发，既支持定时任务每 30 分钟执行，也支持页面手动扫描。扫描结果会记录生成数量和耗时，预警列表支持按类型和处理状态筛选，处理动作会保存处理人和备注。"),
    (278, "低库存和积压预警依赖 `inventory` 与 `material_info`", "系统实现低库存、积压、临期、过期和异常消耗五类预警。低库存与积压分别基于安全库存阈值和倍数判断，临期与过期根据批次效期判断，异常消耗比较近 7 天出库量与近 30 天折算周均值。"),
    (279, "预警去重通过 `createWarningIfAbsent` 完成", "预警去重通过 createWarningIfAbsent 完成。系统在生成新预警前会按类型、物资、仓库和未处理状态检查是否已有记录，避免同一问题重复堆积。"),
    (280, "统计分析模块由 `AnalyticsController` 提供统一入口", "统计分析由 AnalyticsController 提供总览、库存占比、出入库趋势、部门排行、效期统计、仓库分布和应急消耗等接口，前端 AnalyticsView.vue 使用 ECharts 渲染结果，运营总览和大屏页面复用部分统计数据。"),
    (281, "补货建议由 `SmartService.replenishmentSuggestions` 实现", "补货建议由 SmartService.replenishmentSuggestions 实现。系统结合当前库存、安全库存、近 30 天出库量和保障天数计算目标库存与建议补货量，结果主要作为仓储管理员的补货参考。"),
    (282, "移动平均预测由 SmartService.forecast 实现", "移动平均预测由 SmartService.forecast 实现。系统按月汇总指定物资近 6 个月的出库量，并以历史月均值作为未来若干月的趋势参考。"),
    (299, "日志模块分为登录日志和操作日志", "日志模块分为登录日志和操作日志。认证服务负责记录登录用户、IP、状态、时间和 user_agent，业务 Service 在创建、提交、审批、出入库、调拨和预警处理后写入操作日志，便于管理员按模块或时间追踪关键操作。"),
    (300, "通知模块保存系统消息和业务关联信息", "通知模块通过 notification 表保存标题、内容、消息类型、目标用户、已读状态和业务关联信息，前端提供列表、未读数量、标记已读、全部已读和删除能力，当前仅实现站内消息。"),
    (301, "系统配置和事件记录模块属于支撑能力", "系统配置与事件记录模块用于保存运行参数以及校园突发事件的标题、类型、等级、地点、描述、处理状态、上报人和处理结果，可为系统维护和事件处置提供数据支撑。"),
    (305, "测试环境由后端单元/集成测试、前端构建、前端单测和 E2E 流程测试组成", "测试环境由后端单元/集成测试、前端构建、前端单测、E2E 流程测试和本地性能基线组成。后端使用 Maven、JUnit、Spring Boot Test 与 H2，前端使用 Vite、Vitest 和 Playwright，各项测试共同覆盖业务校验、构建完整性和核心流程联调。"),
    (306, "系统测试阶段重新执行了 mvn test、npm run build、npm run test:unit 和 npm run test:e2e", "系统测试阶段重新执行了 mvn test、npm run build、npm run test:unit 和 npm run test:e2e。结果显示后端 49 项测试全部通过，前端构建通过，前端单测 3 个文件、8 项通过，E2E 4 个场景通过。"),
    (309, "后端自动化测试覆盖公共响应、业务异常、全局异常处理、JWT 令牌、认证服务、refresh token 清理、申领服务、申领控制器、调拨服务、调拨控制器、预警控制器和核心流程集成测试", "后端自动化测试覆盖统一响应、业务异常、JWT 令牌、认证服务、refresh token 清理、申领服务、调拨服务、预警控制器和核心流程集成测试。CoreFlowIntegrationTest 在真实 Spring Boot 上下文中验证了登录、库存查询、申领审批、库存扣减和调拨执行等组合流程。"),
    (310, "申领相关测试重点验证状态流转和库存锁定", "申领相关测试重点验证状态流转和库存锁定，调拨测试验证同仓调拨被拒绝、审批状态合法和执行后库存变化，认证测试验证密码校验、账号状态、token 生成、刷新失败和清理逻辑。"),
    (315, "前端构建验证了 Vue 单文件组件、路由懒加载、Element Plus、ECharts 和样式资源能够被 Vite 正常打包", "前端构建验证了 Vue 单文件组件、路由懒加载、Element Plus、ECharts 和样式资源能够被 Vite 正常打包。构建输出覆盖登录、运营总览、库存、入出库、申领、调拨、预警、统计、日志和通知等页面资源。"),
    (316, "前端单测主要覆盖 HTTP 请求封装和登录页面交互", "前端单测主要覆盖 HTTP 请求封装和登录页面交互，E2E 则覆盖管理员查看库存、部门用户创建申领单、仓储管理员创建调拨单和处理预警四个典型流程，能够验证页面工程和核心业务闭环。"),
    (317, "权限测试在 E2E 和后端日志中也有所体现", "权限测试也体现在 E2E 和后端日志中。部门用户访问无权接口时会记录权限拒绝异常，但页面流程仍可继续完成预期动作，说明权限边界控制已生效。"),
    (321, "业务场景验证主要围绕系统核心流程", "业务场景验证围绕申领审批、调拨执行、预警处理和认证续签展开，重点检查单据状态变化、来源仓扣减与目标仓增加、批次镜像生成、手动扫描处理以及旧 refresh token 失效等行为。"),
    (324, "项目提供登录、库存列表、预警列表和操作日志列表等接口的本地性能基线脚本", "项目提供登录、库存列表、预警列表和操作日志列表等接口的本地性能基线脚本，并在 docs/performance-baseline.md 中保留了对应测试材料。性能章节用于说明系统具备接口级基线验证能力。"),
    (325, "综合测试结果看，系统核心功能能够在本地环境下完成自动化验证", "综合测试结果表明，系统核心功能能够在本地环境下完成自动化验证，满足毕业设计演示和核心流程验证要求，但不足以推导大规模上线运行结论。"),
    (326, "仓库中提供了 tests/performance 下的 k6 脚本和 docs/performance-baseline.md 基线记录", "tests/performance 下的 k6 脚本与基线记录采集于本地 screenshot profile 和隔离 H2 数据集，场景包括登录、库存分页、预警分页和操作日志分页，仅适用于相同环境下的基线对比。"),
    (329, "本文围绕校园物资管理场景，完成了基于 Spring Boot 与 Vue 的前后端分离系统设计与实现", "本文围绕校园物资管理场景，完成了基于 Spring Boot 与 Vue 的前后端分离系统设计与实现。系统建立了组织、角色、物资、仓库、库存、批次、单据、预警、日志和通知等核心数据结构，能够支撑物资建档、入库、库存查询、申领审批、出库签收、仓间调拨、预警处理和统计分析等业务流程。"),
    (330, "本系统的主要工程特点包括", "系统的主要工程特点在于：采用 JWT 与角色菜单实现认证授权，采用库存汇总与批次表共同表达库存状态，出库和调拨按 FEFO 规则扣减批次，申领阶段锁定库存并在驳回或签收后释放剩余锁定量，预警与统计模块为管理人员提供规则化辅助能力。"),
    (331, "系统仍存在可完善空间", "系统仍有进一步完善空间，包括扩展移动端或扫码能力、细化采购与报废等协同流程、在积累更多真实数据后研究更细粒度的需求预测，以及在真实运行环境中补充压测、备份恢复和运维监控方案。"),
]

SKIP_REPLACEMENTS = {127, 139, 193, 207, 218, 224, 246, 251, 278, 331}

CHAPTER_ORDER = [
    "1 绪论",
    "2 相关技术介绍",
    "3 系统需求分析",
    "4 系统总体设计",
    "5 数据库设计",
    "6 系统详细设计与实现",
    "7 系统测试",
    "结束语",
]


def count_text(text: str) -> int:
    return len(re.sub(r"\s+", "", text))


def chapter_stats(doc: Document) -> OrderedDict[str, int]:
    stats: OrderedDict[str, int] = OrderedDict()
    current = None
    active = False
    stop_titles = {"致  谢", "致谢", "参考文献", "附录 1 关键接口与测试补充说明", "附录"}

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name == "Heading 1":
            if text in stop_titles:
                current = None
                active = False
                continue
            if re.match(r"^[1-7]\s", text) or text == "结束语":
                current = text
                stats[current] = 0
                active = True
        if active and current:
            stats[current] += count_text(text)
    return stats


def ensure_anchor(paragraphs, index: int, anchor: str, filename: str) -> None:
    text = paragraphs[index].text.strip()
    if anchor not in text:
        raise RuntimeError(f"{filename} 段落 {index} 未命中锚点: {anchor}\n实际内容: {text}")


def apply_changes(doc_path: Path) -> None:
    doc = Document(doc_path)
    paragraphs = doc.paragraphs

    for index, anchor, new_text in REPLACEMENTS:
        if index in SKIP_REPLACEMENTS:
            continue
        ensure_anchor(paragraphs, index, anchor, doc_path.name)
        paragraphs[index].text = new_text

    doc.save(doc_path)


def write_report(before_stats: OrderedDict[str, int], after_stats: OrderedDict[str, int]) -> None:
    before_total = sum(before_stats.values())
    after_total = sum(after_stats.values())
    chapter_labels = {
        "1 绪论": "绪论",
        "2 相关技术介绍": "相关技术介绍",
        "3 系统需求分析": "系统需求分析",
        "4 系统总体设计": "系统总体设计",
        "5 数据库设计": "数据库设计",
        "6 系统详细设计与实现": "系统详细设计与实现",
        "7 系统测试": "系统测试",
        "结束语": "结束语",
    }

    lines = [
        "# 正文压缩说明",
        "",
        f"- 源文件：`{SOURCE.name}`",
        f"- 输出文件：`{TARGET.name}`",
        f"- 统计口径：按正文段落统计，范围为“1 绪论”至“结束语”，不含封面、任务书、摘要、目录、参考文献、致谢、附录；统计字符时去除空白符，未计入表格单元格文字。",
        f"- 压缩前正文约：`{before_total}` 字",
        f"- 压缩后正文约：`{after_total}` 字",
        "",
        "## 各章节压缩情况",
        "",
    ]

    for chapter in CHAPTER_ORDER:
        before = before_stats.get(chapter, 0)
        after = after_stats.get(chapter, 0)
        ratio = 0 if before == 0 else round((before - after) / before * 100, 1)
        lines.append(f"- {chapter_labels[chapter]}：由约 `{before}` 字压缩至约 `{after}` 字，压缩约 `{ratio}%`。")

    lines.extend(
        [
            "",
            "## 主要删减类型",
            "",
            "- 删除绪论、技术介绍和测试章节中的重复背景、通用技术优势和重复结论。",
            "- 合并需求分析、总体设计与详细实现中对同一业务流程的重复解释。",
            "- 压缩页面级操作描述和接口罗列，保留关键业务逻辑、状态流转和核心数据结构。",
            "- 将技术章节从框架介绍式写法压缩为“技术在本系统中的作用”导向表达。",
            "- 压缩测试章节中与表格重复的文字说明，仅保留环境、方法、关键结果和边界结论。",
            "",
            "## 保留原因",
            "",
            "- 申领审批、库存批次、调拨执行、预警扫描、补货建议和统计分析等段落保留了核心实现逻辑，因为它们直接对应代码、数据库表和测试结果。",
            "- 图表引用、章节层级、表名、字段名、接口路径和关键技术名称均保持不变，以避免破坏论文结构和证据链。",
            "- 测试结果中的通过数量、k6 基线描述和安全边界说明保留了事实性内容，避免压缩后削弱可验证性。",
        ]
    )

    REPORT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    shutil.copy2(SOURCE, TARGET)

    before_doc = Document(SOURCE)
    before_stats = chapter_stats(before_doc)

    apply_changes(TARGET)

    after_doc = Document(TARGET)
    after_stats = chapter_stats(after_doc)

    write_report(before_stats, after_stats)

    print(f"[source] {SOURCE.name}")
    print(f"[target] {TARGET.name}")
    print(f"[before] {sum(before_stats.values())}")
    print(f"[after] {sum(after_stats.values())}")
    for chapter in CHAPTER_ORDER:
        print(f"{chapter}\t{before_stats.get(chapter, 0)}\t{after_stats.get(chapter, 0)}")


if __name__ == "__main__":
    main()
