from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

from build_thesis_working_draft import disable_update_fields_on_open, normalize_front_matter_order


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DOCX = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿工作稿.docx"


def main() -> int:
    parser = argparse.ArgumentParser(description="Fix the front-matter block order of the thesis DOCX.")
    parser.add_argument("docx_path", nargs="?", type=Path, default=DEFAULT_DOCX)
    args = parser.parse_args()

    docx_path = args.docx_path.resolve()
    doc = Document(docx_path)
    before_images = len(doc.inline_shapes)
    before_tables = len(doc.tables)

    normalize_front_matter_order(doc)
    disable_update_fields_on_open(doc)

    tmp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
    doc.save(tmp_path)
    tmp_path.replace(docx_path)

    saved = Document(docx_path)
    print(docx_path)
    print("images", len(saved.inline_shapes), "tables", len(saved.tables))
    print("same_images", len(saved.inline_shapes) == before_images)
    print("same_tables", len(saved.tables) == before_tables)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
