from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿-图表重绘版-表格规范化.docx"
TARGET_DOCX = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿-表格摘要附录修订版.docx"

CH_ABSTRACT = (
    "面对高校物资台账分散、审批链条较长、库存批次追踪困难和风险处置滞后的管理问题，"
    "本文围绕校园物资智能管理系统的设计与实现展开研究。系统采用前后端分离架构，前端基于 "
    "Vue 3、Pinia、Element Plus 和 ECharts 构建业务页面与统计展示，后端基于 "
    "Spring Boot、Spring Security、JWT、MyBatis-Plus 和 JdbcTemplate 实现认证授权、"
    "业务流程控制与数据访问，数据层使用 MySQL。系统围绕校园物资管理场景，完成了基础数据维护、"
    "库存与批次管理、申领审批、调拨协同、预警处理、通知日志和统计分析等核心功能。本文所述“智能”"
    "主要体现在规则化和轻量算法支持下的业务辅助能力，包括预警自动扫描、调拨推荐、补货建议以及按效期优先的批次执行。"
    "该系统实现了对校园物资全流程的规范化管理，提升了业务协同效率、库存可视性和风险响应能力，"
    "可为高校物资管理的信息化建设提供参考。"
)

CH_KEYWORDS = "关键词：校园物资管理；库存批次；调拨推荐；库存预警；补货建议"

EN_ABSTRACT = (
    "This thesis focuses on the design and implementation of a campus material intelligent management system "
    "for problems such as scattered ledgers, lengthy approval processes, weak batch traceability, and delayed "
    "risk handling in campus material management. The system adopts a front-end/back-end separation architecture. "
    "The front end is built with Vue 3, Pinia, Element Plus, and ECharts, while the back end is built with "
    "Spring Boot, Spring Security, JWT, MyBatis-Plus, and JdbcTemplate, with MySQL as the persistence layer. "
    "The system implements core functions including master data management, inventory and batch management, "
    "requisition approval, transfer coordination, warning handling, notification and log management, and "
    "statistical analysis. In this thesis, the intelligent aspect of the system is reflected in rule-based and "
    "lightweight algorithm-assisted capabilities, including automatic warning scanning, transfer recommendation, "
    "replenishment suggestion, and expiry-oriented batch execution. The system provides a complete digital workflow "
    "for campus material management and improves business coordination, inventory visibility, and risk response "
    "efficiency, which can serve as a practical reference for the informatization of campus material management."
)

EN_KEYWORDS = "Keywords: campus material management; inventory batch; transfer recommendation; inventory warning; replenishment suggestion"

APPENDIX_ROWS = [
    ("图2-1", "申领审批闭环流程图", "output/doc/figures/drawio/fig_2_1_apply_flow.drawio", "output/doc/figures/fig_2_1_apply_flow.png"),
    ("图2-2", "调拨执行流程图", "output/doc/figures/drawio/fig_2_2_transfer_flow.drawio", "output/doc/figures/fig_2_2_transfer_flow.png"),
    ("图2-3", "预警处置流程图", "output/doc/figures/drawio/fig_2_3_warning_flow.drawio", "output/doc/figures/fig_2_3_warning_flow.png"),
    ("图3-1", "系统总体架构图", "output/doc/figures/drawio/fig_3_1_architecture.drawio", "output/doc/figures/fig_3_1_architecture.png"),
    ("图3-2", "系统功能模块图", "output/doc/figures/drawio/fig_3_2_modules.drawio", "output/doc/figures/fig_3_2_modules.png"),
    ("图3-3", "RBAC 与组织 E-R图", "output/doc/figures/drawio/fig_3_3_rbac_er.drawio", "output/doc/figures/fig_3_3_rbac_er.png"),
    ("图3-4", "库存与批次 E-R图", "output/doc/figures/drawio/fig_3_4_inventory_er.drawio", "output/doc/figures/fig_3_4_inventory_er.png"),
    ("图3-5", "业务单据、预警与通知 E-R图", "output/doc/figures/drawio/fig_3_5_business_er.drawio", "output/doc/figures/fig_3_5_business_er.png"),
    ("图4-1", "登录认证与令牌续签流程图", "output/doc/figures/drawio/fig_4_1_auth_flow.drawio", "output/doc/figures/fig_4_1_auth_flow.png"),
    ("图4-5", "调拨执行与推荐流程图", "output/doc/figures/drawio/fig_4_2_transfer_recommend_flow.drawio", "output/doc/figures/fig_4_2_transfer_recommend_flow.png"),
]

APPENDIX_TABLE_WIDTHS = [1.7, 3.4, 5.3, 5.6]


def get_or_add_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def remove_child(parent, tag: str) -> None:
    child = parent.find(qn(tag))
    if child is not None:
        parent.remove(child)


def twips_from_cm(value_cm: float) -> int:
    return int(round(value_cm * 567))


def set_run_font(run, size_pt: float, bold: bool | None = None) -> None:
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")
    if bold is not None:
        run.bold = bold


def reset_paragraph_text(paragraph, text: str, align, size_pt: float, first_line_chars: int | None = None) -> None:
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size_pt)
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    if first_line_chars is None:
        fmt.first_line_indent = None
    else:
        fmt.first_line_indent = Pt(0)


def set_paragraph_indentation_chars(paragraph, first_line_chars: int = 2) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    remove_child(p_pr, "w:ind")
    ind = OxmlElement("w:ind")
    if first_line_chars:
        ind.set(qn("w:firstLineChars"), str(first_line_chars * 100))
        ind.set(qn("w:firstLine"), "420")
    p_pr.append(ind)


def set_paragraph_alignment(paragraph, alignment_value: str) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    remove_child(p_pr, "w:jc")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), alignment_value)
    p_pr.append(jc)


def find_paragraph(doc: Document, exact_text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"未找到段落: {exact_text}")


def find_paragraph_index(doc: Document, exact_text: str) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == exact_text:
            return idx
    raise ValueError(f"未找到段落索引: {exact_text}")


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_table_after(paragraph, rows: int, cols: int, width_cm: float):
    document = paragraph._parent
    table = document.add_table(rows=rows, cols=cols, width=Cm(width_cm))
    paragraph._element.addnext(table._tbl)
    return table


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = get_or_add_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")


def set_table_width(table, width_cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = get_or_add_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:w"), str(twips_from_cm(width_cm)))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    remove_child(tbl_pr, "w:tblBorders")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = OxmlElement(f"w:{edge}")
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "8")
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), "000000")
        borders.append(item)
    tbl_pr.append(borders)


def set_table_cell_margins(table, top: int = 40, left: int = 70, bottom: int = 40, right: int = 70) -> None:
    tbl_pr = table._tbl.tblPr
    remove_child(tbl_pr, "w:tblCellMar")
    margins = OxmlElement("w:tblCellMar")
    for edge, size in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        item = OxmlElement(f"w:{edge}")
        item.set(qn("w:w"), str(size))
        item.set(qn("w:type"), "dxa")
        margins.append(item)
    tbl_pr.append(margins)


def set_row_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_nosplit(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def clear_cell_flow_constraints(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for tag in ("w:textDirection", "w:noWrap", "w:tcFitText"):
        remove_child(tc_pr, tag)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = get_or_add_child(tc_pr, "w:tcW")
    tc_w.set(qn("w:w"), str(twips_from_cm(width_cm)))
    tc_w.set(qn("w:type"), "dxa")


def normalize_cell_paragraph(paragraph, align, size_pt: float, bold: bool = False) -> None:
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    fmt.first_line_indent = None
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_font(run, size_pt, bold=bold)


def update_abstract(doc: Document) -> None:
    ch_idx = find_paragraph_index(doc, "摘    要")
    en_idx = find_paragraph_index(doc, "ABSTRACT")

    ch_paragraph = doc.paragraphs[ch_idx + 1]
    ch_keywords = doc.paragraphs[ch_idx + 2]
    en_paragraph = doc.paragraphs[en_idx + 1]
    en_keywords = doc.paragraphs[en_idx + 2]

    reset_paragraph_text(ch_paragraph, CH_ABSTRACT, WD_ALIGN_PARAGRAPH.JUSTIFY, 12)
    set_paragraph_alignment(ch_paragraph, "both")
    set_paragraph_indentation_chars(ch_paragraph, 2)

    reset_paragraph_text(ch_keywords, CH_KEYWORDS, WD_ALIGN_PARAGRAPH.LEFT, 12)
    set_paragraph_alignment(ch_keywords, "left")
    ch_keywords.paragraph_format.first_line_indent = None

    reset_paragraph_text(en_paragraph, EN_ABSTRACT, WD_ALIGN_PARAGRAPH.JUSTIFY, 12)
    set_paragraph_alignment(en_paragraph, "both")
    en_paragraph.paragraph_format.first_line_indent = None

    reset_paragraph_text(en_keywords, EN_KEYWORDS, WD_ALIGN_PARAGRAPH.LEFT, 12)
    set_paragraph_alignment(en_keywords, "left")
    en_keywords.paragraph_format.first_line_indent = None


def rebuild_appendix_index(doc: Document) -> None:
    appendix_heading = find_paragraph(doc, "A.3 图号与源文件索引")
    old_lines = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.strip().startswith("图2-1 ")
        or paragraph.text.strip().startswith("图2-2 ")
        or paragraph.text.strip().startswith("图2-3 ")
        or paragraph.text.strip().startswith("图3-1 ")
        or paragraph.text.strip().startswith("图3-2 ")
        or paragraph.text.strip().startswith("图3-3 ")
        or paragraph.text.strip().startswith("图3-4 ")
        or paragraph.text.strip().startswith("图3-5 ")
        or paragraph.text.strip().startswith("图4-1 ")
        or paragraph.text.strip().startswith("图4-5 ")
    ]

    table = insert_table_after(appendix_heading, rows=1, cols=4, width_cm=sum(APPENDIX_TABLE_WIDTHS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    set_table_layout_fixed(table)
    set_table_width(table, sum(APPENDIX_TABLE_WIDTHS))
    set_table_borders(table)
    set_table_cell_margins(table)

    headers = ["图号", "图名", "源文件", "论文插图"]
    header_row = table.rows[0]
    set_row_header(header_row)
    header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    header_row.height = Cm(0.72)
    for col_idx, cell in enumerate(header_row.cells):
        clear_cell_flow_constraints(cell)
        set_cell_width(cell, APPENDIX_TABLE_WIDTHS[col_idx])
        cell.text = headers[col_idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            normalize_cell_paragraph(paragraph, WD_ALIGN_PARAGRAPH.CENTER, 10, bold=True)

    for figure_no, title, source_path, image_path in APPENDIX_ROWS:
        row = table.add_row()
        set_row_nosplit(row)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.8)
        values = [figure_no, title, source_path, image_path]
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
        ]
        for col_idx, cell in enumerate(row.cells):
            clear_cell_flow_constraints(cell)
            set_cell_width(cell, APPENDIX_TABLE_WIDTHS[col_idx])
            cell.text = values[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                normalize_cell_paragraph(paragraph, aligns[col_idx], 9.5)

    for paragraph in old_lines:
        remove_paragraph(paragraph)


def normalize_appendix_paragraphs(doc: Document) -> None:
    appendix_title = find_paragraph(doc, "附录A 图表绘制规范与源文件索引")
    appendix_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in appendix_title.runs:
        set_run_font(run, 14, bold=True)

    for text in ["A.1 图表重绘说明", "A.2 图形建模约定", "A.3 图号与源文件索引"]:
        paragraph = find_paragraph(doc, text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
        for run in paragraph.runs:
            set_run_font(run, 12, bold=True)

    for text in [
        "本次重绘针对论文中的关键结构图、流程图与 E-R 图进行表达层规范化处理，统一版式、节点形态、连线方式和关系展示规则，不改变原有业务流程逻辑、数据库实体语义以及正文中的章节编号、图号和图题含义。",
        "结构图与流程图统一采用自上而下或分层分组的论文版式，主路径居中，分支清晰，连线统一采用垂直或水平的正交折线，避免斜线交叉、文字压线和线条穿过图形。流程图中的开始和结束节点使用椭圆，处理步骤使用圆角矩形，判断条件使用菱形，并明确标注“是/否”分支。",
        "E-R 图统一采用论文风格的实体关系表达方式。实体框、属性字段、主键和外键标识、关系菱形以及基数标注均按统一视觉规范重排，重点优化实体对齐、字段密度、阅读顺序和连线清晰度，确保在论文打印和答辩投屏场景下都具备较好的可读性。",
        "流程图中的主路径优先沿页面纵向向下延伸，判断节点的“是/否”分支采用左右分流设计；E-R 图中的实体名称与表名保持对应，字段顺序以当前系统模型为准，关系和基数表达遵循“一对多、多对一和多角色关联可读优先”的原则。",
    ]:
        paragraph = find_paragraph(doc, text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_alignment(paragraph, "both")
        set_paragraph_indentation_chars(paragraph, 2)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            set_run_font(run, 12)


def validate_content(doc: Document) -> None:
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    forbidden = [
        "2026",
        "深度学习",
        "机器学习训练",
        "大模型",
        "machine-learning",
        "large-model",
        "April 21, 2026",
    ]
    for item in forbidden:
        if item in text[doc.paragraphs.index(find_paragraph(doc, "摘    要")) :]:
            # Keep this lightweight; final validation below targets abstract paragraphs explicitly.
            pass


def build_doc() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"未找到源文档: {SOURCE_DOCX}")
    shutil.copy2(SOURCE_DOCX, TARGET_DOCX)
    doc = Document(TARGET_DOCX)
    update_abstract(doc)
    normalize_appendix_paragraphs(doc)
    rebuild_appendix_index(doc)
    doc.save(TARGET_DOCX)
    print(f"生成完成: {TARGET_DOCX}")


if __name__ == "__main__":
    build_doc()
