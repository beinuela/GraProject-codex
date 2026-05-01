from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "output" / "doc"
SOURCE = DOC_DIR / "校园物资智能管理系统设计与实现-日期清理版.docx"
TARGET = DOC_DIR / "校园物资智能管理系统设计与实现-参考文献引用修复版.docx"
REPORT = DOC_DIR / "参考文献引用检查与修复报告.md"

BODY_HEADING_RE = re.compile(r"^[1-7]\s")
REFERENCE_ENTRY_RE = re.compile(r"^\[(\d+)\]\s+")
CITATION_RE = re.compile(r"\[(\d+(?:\s*[-,，]\s*\d+)*)\]")
FULL_WIDTH_CITATION_RE = re.compile(r"［\d+(?:\s*[-,，]\s*\d+)*］")
STOP_TITLES = {"致  谢", "致谢", "参考文献", "附录 1 关键接口与测试补充说明", "附录"}


@dataclass(frozen=True)
class CitationAddition:
    needle: str
    sentence: str
    references: tuple[int, ...]
    location: str


ADDITIONS = (
    CitationAddition(
        needle="MySQL 8 用作正式业务数据库，H2 用于测试和截图环境。",
        sentence="基于 Java 与 MySQL 的管理信息系统实践表明，此类技术路线能够较好支撑业务数据持久化与查询组织[14]。",
        references=(14,),
        location="第2章 2.3 后端与数据访问技术中 MySQL 数据库角色说明段",
    ),
    CitationAddition(
        needle="索引与约束围绕实际查询场景设计。",
        sentence="相关索引设计与查询优化思路可参照 MySQL 应用实践与性能优化研究[17][18]。",
        references=(17, 18),
        location="第5章 5.4 索引与约束设计段",
    ),
)


def body_paragraphs(doc: Document) -> list[tuple[int, Paragraph, str]]:
    items: list[tuple[int, Paragraph, str]] = []
    in_body = False
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name == "Heading 1":
            if text in STOP_TITLES:
                in_body = False
            elif BODY_HEADING_RE.match(text) or text == "结束语":
                in_body = True
        if in_body:
            items.append((index, paragraph, text))
    return items


def expand_citation_numbers(content: str) -> list[int]:
    numbers: list[int] = []
    for part in re.split(r"[,，]", content):
        token = part.strip()
        if not token:
            continue
        if "-" in token:
            start_str, end_str = [value.strip() for value in token.split("-", 1)]
            start = int(start_str)
            end = int(end_str)
            numbers.extend(range(start, end + 1))
        else:
            numbers.append(int(token))
    return numbers


def collect_reference_entries(doc: Document) -> dict[int, str]:
    refs: dict[int, str] = {}
    in_references = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text == "参考文献":
            in_references = True
            continue
        if not in_references:
            continue
        match = REFERENCE_ENTRY_RE.match(text)
        if not match:
            if paragraph.style.name == "Heading 1":
                break
            continue
        refs[int(match.group(1))] = text
    return refs


def scan_body_citations(doc: Document) -> dict[str, object]:
    items = body_paragraphs(doc)
    cited_numbers: set[int] = set()
    first_positions: dict[int, int] = {}
    citations_by_paragraph: list[tuple[int, str]] = []
    format_issues: list[str] = []

    for index, _, text in items:
        paragraph_numbers: list[int] = []
        for match in CITATION_RE.finditer(text):
            numbers = expand_citation_numbers(match.group(1))
            paragraph_numbers.extend(numbers)
            for number in numbers:
                cited_numbers.add(number)
                first_positions.setdefault(number, index)
        if paragraph_numbers:
            citations_by_paragraph.append((index, text))

        if FULL_WIDTH_CITATION_RE.search(text):
            format_issues.append(f"P[{index}] 存在全角引用括号：{text}")
        if "[ " in text or " ]" in text:
            format_issues.append(f"P[{index}] 存在方括号空格：{text}")

    return {
        "cited_numbers": cited_numbers,
        "first_positions": first_positions,
        "citations_by_paragraph": citations_by_paragraph,
        "format_issues": format_issues,
    }


def find_by_substring(doc: Document, needle: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise RuntimeError(f"未找到锚点段落: {needle}")


def append_sentence(paragraph: Paragraph, sentence: str) -> None:
    current = paragraph.text.rstrip()
    if sentence in current:
        return
    joiner = "" if current.endswith(("。", "！", "？")) else " "
    paragraph.add_run(f"{joiner}{sentence}")


def apply_additions(doc: Document) -> list[CitationAddition]:
    applied: list[CitationAddition] = []
    for item in ADDITIONS:
        paragraph = find_by_substring(doc, item.needle)
        if all(f"[{number}]" in paragraph.text for number in item.references):
            continue
        append_sentence(paragraph, item.sentence)
        applied.append(item)
    return applied


def build_report(
    backup_path: Path,
    before_refs: dict[int, str],
    before_scan: dict[str, object],
    after_scan: dict[str, object],
    applied: list[CitationAddition],
    dangling_before: list[int],
    dangling_after: list[int],
) -> None:
    ref_numbers = sorted(before_refs)
    cited_before = sorted(before_scan["cited_numbers"])
    cited_after = sorted(after_scan["cited_numbers"])
    missing_before = [number for number in ref_numbers if number not in before_scan["cited_numbers"]]
    missing_after = [number for number in ref_numbers if number not in after_scan["cited_numbers"]]
    format_issues = before_scan["format_issues"]

    lines = [
        "# 参考文献引用检查与修复报告",
        "",
        f"- 工作底稿：`{SOURCE.name}`",
        f"- 备份文件：`{backup_path.name}`",
        f"- 输出文件：`{TARGET.name}`",
        f"- 参考文献总数：`{len(before_refs)}`",
        "",
        "## 正文引用情况",
        "",
        f"- 修复前已引用文献编号：`{cited_before}`",
        f"- 修复后已引用文献编号：`{cited_after}`",
        f"- 修复前未在正文中引用的文献编号：`{missing_before}`",
        f"- 修复后未在正文中引用的文献编号：`{missing_after}`",
        "",
        "## 已补充引用的文献",
        "",
    ]

    if applied:
        for item in applied:
            refs = "、".join(f"[{number}]" for number in item.references)
            lines.append(f"- {refs}：{item.location}")
    else:
        lines.append("- 无新增补引。")

    lines.extend(
        [
            "",
            "## 仍建议删除或人工确认的文献",
            "",
        ]
    )

    if missing_after:
        lines.append(f"- 仍未引用：`{missing_after}`，需要人工判断是否保留。")
    else:
        lines.append("- 无。22 条文献均已在正文中出现至少一次。")

    lines.extend(
        [
            "",
            "## 正文引用格式检查",
            "",
        ]
    )

    if format_issues:
        lines.append("- 存在以下格式问题：")
        for issue in format_issues:
            lines.append(f"  - {issue}")
    else:
        lines.append("- 正文引用统一采用方括号数字标注。连续单引形式如 `[6][7][8]` 予以保留，未做风格性改写。")

    lines.extend(
        [
            "",
            "## 悬空引用检查",
            "",
            f"- 修复前正文引用了但文末未列出的编号：`{dangling_before}`",
            f"- 修复后正文引用了但文末未列出的编号：`{dangling_after}`",
            "",
            "## 编号策略说明",
            "",
            "- 本轮采用最小修复策略，保留文末参考文献现有编号以及正文大部分引用编号，不按首次出现顺序重排全文。当前编号与文末列表一一对应，但首次出现顺序并非严格升序，此为有意保留结果。",
        ]
    )

    REPORT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    source_doc = Document(SOURCE)
    before_refs = collect_reference_entries(source_doc)
    if not before_refs:
        raise RuntimeError("未能从工作底稿中提取参考文献条目。")

    before_scan = scan_body_citations(source_doc)
    dangling_before = sorted(number for number in before_scan["cited_numbers"] if number not in before_refs)

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    backup_path = DOC_DIR / f"{SOURCE.stem}-参考文献修复前备份-{timestamp}{SOURCE.suffix}"
    shutil.copy2(SOURCE, backup_path)
    shutil.copy2(SOURCE, TARGET)

    target_doc = Document(TARGET)
    applied = apply_additions(target_doc)
    target_doc.save(TARGET)

    verified_doc = Document(TARGET)
    after_refs = collect_reference_entries(verified_doc)
    after_scan = scan_body_citations(verified_doc)
    dangling_after = sorted(number for number in after_scan["cited_numbers"] if number not in after_refs)
    missing_after = sorted(number for number in after_refs if number not in after_scan["cited_numbers"])

    if dangling_after:
        raise RuntimeError(f"修复后仍存在正文悬空引用: {dangling_after}")
    if missing_after:
        raise RuntimeError(f"修复后仍存在未在正文中引用的文献: {missing_after}")

    build_report(
        backup_path=backup_path,
        before_refs=before_refs,
        before_scan=before_scan,
        after_scan=after_scan,
        applied=applied,
        dangling_before=dangling_before,
        dangling_after=dangling_after,
    )

    print(f"[backup] {backup_path.name}")
    print(f"[output] {TARGET.name}")
    print(f"[references] {len(before_refs)}")
    print(f"[before-cited] {sorted(before_scan['cited_numbers'])}")
    print(f"[after-cited] {sorted(after_scan['cited_numbers'])}")
    print(f"[applied] {[item.references for item in applied]}")
    print(f"[dangling-before] {dangling_before}")
    print(f"[dangling-after] {dangling_after}")


if __name__ == "__main__":
    main()
