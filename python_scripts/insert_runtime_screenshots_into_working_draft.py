from __future__ import annotations

import hashlib
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ORIGINAL_DRAFT = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-初稿.docx"
WORKING_DRAFT = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿工作稿.docx"
SCREENSHOT_DIR = ROOT / "output" / "playwright" / "thesis-runtime"
ROUND2_REPORT = ROOT / "output" / "doc" / "校园物资智能管理系统设计与实现-改稿说明-第二轮截图补充.md"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def add_body_before(anchor, text: str, style: str = "正文章节内容"):
    paragraph = anchor.insert_paragraph_before(text, style=style)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    return paragraph


def add_picture_before(anchor, image_path: Path, caption: str):
    picture_paragraph = anchor.insert_paragraph_before()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(image_path), width=Inches(6.1))
    caption_paragraph = anchor.insert_paragraph_before(caption, style="图片标题")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return picture_paragraph, caption_paragraph


def set_all_text_black(doc: Document):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = False
    for section in doc.sections:
        for paragraph in section.header.paragraphs + section.footer.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.underline = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.underline = False


def ensure_no_old_title(doc: Document):
    old_terms = ["校园应急物资智能管理系统", "Campus Emergency Material System"]
    for paragraph in doc.paragraphs:
        for term in old_terms:
            if term in paragraph.text:
                raise ValueError(f"Legacy title text still present in working draft: {term}")


def insert_section_screens(doc: Document):
    add_screenshot_block(
        doc,
        stop_heading="4.3 基础数据管理实现",
        lead="为了直观展示系统登录入口与统一认证界面，本节补充运行截图如图4-3所示。",
        image_path=SCREENSHOT_DIR / "fig_4_3_login.png",
        caption="图4-3 系统登录界面",
    )

    add_screenshot_block(
        doc,
        stop_heading="4.5.2 调拨实现与推荐机制",
        lead="部门用户发起物资申请时，可以在申领页面直接录入原因、场景和物资明细，运行界面如图4-4所示。",
        image_path=SCREENSHOT_DIR / "fig_4_4_apply_dept.png",
        caption="图4-4 部门用户申领界面",
    )
    add_screenshot_block(
        doc,
        stop_heading="4.5.2 调拨实现与推荐机制",
        lead="审批人与仓库管理员分别在审批列表和出库页面完成业务闭环，组合运行界面如图4-5所示。",
        image_path=SCREENSHOT_DIR / "fig_4_5_approval_stockout_combo.png",
        caption="图4-5 审批与出库执行界面",
    )

    add_screenshot_block(
        doc,
        stop_heading="4.6 预警、补货建议与统计分析实现",
        lead="仓库管理员在调拨页面可直接发起跨仓调拨，并调用推荐结果辅助选择来源仓库，运行界面如图4-6所示。",
        image_path=SCREENSHOT_DIR / "fig_4_6_transfer_recommend.png",
        caption="图4-6 调拨推荐界面",
    )

    add_screenshot_block(
        doc,
        stop_heading="4.7 日志、通知与事件管理实现",
        lead="预警模块在列表中同时展示预警类型、处理状态和处置入口，真实运行界面如图4-7所示。",
        image_path=SCREENSHOT_DIR / "fig_4_7_warning.png",
        caption="图4-7 预警管理界面",
    )
    add_screenshot_block(
        doc,
        stop_heading="4.7 日志、通知与事件管理实现",
        lead="统计分析模块通过图表卡片集中展示库存、趋势和排行结果，真实运行界面如图4-8所示。",
        image_path=SCREENSHOT_DIR / "fig_4_8_analytics.png",
        caption="图4-8 统计分析界面",
    )

    note_anchor = find_paragraph(doc, "5.2 自动化测试结果")
    add_body_before(
        note_anchor,
        "为补充第4章的运行界面证据，本轮在不接入本地 MySQL 主库的前提下，新增了基于 H2 文件库的 screenshot 演示 profile，并使用最小演示数据启动系统后采集截图。该环境仅用于论文取证和界面展示，不替代正文中基于主项目代码与测试结果形成的事实结论。",
    )


def add_screenshot_block(doc: Document, stop_heading: str, lead: str, image_path: Path, caption: str):
    if not image_path.exists():
        raise FileNotFoundError(f"Screenshot not found: {image_path}")
    anchor = find_paragraph(doc, stop_heading)
    add_body_before(anchor, lead)
    add_picture_before(anchor, image_path, caption)


def count_inline_shapes(doc: Document) -> int:
    return len(doc.inline_shapes)


def latest_continue_backup() -> Path | None:
    backups = sorted(
        (ROOT / "Existing Thesis Draft").glob("*定稿工作稿.docx.bak.*-continue-screens"),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    return backups[0] if backups else None


def write_report(original_hash: str, backup_path: Path, image_count: int):
    screenshots = [
        "图4-3 系统登录界面",
        "图4-4 部门用户申领界面",
        "图4-5 审批与出库执行界面",
        "图4-6 调拨推荐界面",
        "图4-7 预警管理界面",
        "图4-8 统计分析界面",
    ]
    text = f"""# 校园物资智能管理系统论文改稿说明（第二轮截图补充）

## 工作文件
- working draft：`{WORKING_DRAFT.relative_to(ROOT)}`
- protected original：`{ORIGINAL_DRAFT.relative_to(ROOT)}`
- fresh backup：`{backup_path.relative_to(ROOT)}`
- screenshot output：`{SCREENSHOT_DIR.relative_to(ROOT)}`

## 本轮完成内容
- 保留当前 working draft 既有正文、7 张工程图和 19 张表，只做增量插图与承接文字修改。
- 修正登录页旧英文副标题，保证截图与当前论文题名一致。
- 基于 `screenshot` 隔离演示 profile 启动真实系统，补采 6 张运行截图，并生成 1 张审批/出库组合图。
- 将图4-3至图4-8 按正文锚点插入 `4.2`、`4.5.1`、`4.5.2`、`4.6` 对应位置。
- 在 `5.1 测试环境与证据来源` 后补入隔离演示环境说明，明确截图环境不接入主库。

## 截图清单
{chr(10).join(f"- {item}" for item in screenshots)}

## 验证摘要
- 原始初稿 SHA256：`{original_hash}`
- 当前 working draft 图片总数：`{image_count}`
- 预期图片总数：`13`

## 仍需本地 Word 复核
- 打开 working draft 后刷新目录、页码与所有题注域。
- 检查截图分页是否与学校模板页边距一致，必要时手动微调图片宽度或前后空行。
- 最终提交前，再做一次全文黑色文本与图题连续编号检查。
"""
    ROUND2_REPORT.write_text(text, encoding="utf-8")


def main():
    original_hash = sha256(ORIGINAL_DRAFT)
    backup_path = latest_continue_backup()
    if backup_path is None:
        raise SystemExit("Missing continue-screens backup for the working draft.")
    doc = Document(WORKING_DRAFT)
    ensure_no_old_title(doc)
    insert_section_screens(doc)
    set_all_text_black(doc)
    doc.save(WORKING_DRAFT)
    image_count = count_inline_shapes(Document(WORKING_DRAFT))
    write_report(original_hash, backup_path, image_count)
    print(WORKING_DRAFT)
    print(ROUND2_REPORT)
    print(image_count)


if __name__ == "__main__":
    main()
