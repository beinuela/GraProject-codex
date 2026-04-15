from __future__ import annotations

import argparse
import math
import shutil
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
WORKING_DRAFT = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿工作稿.docx"
DEFAULT_OUTPUT_DRAFT = ROOT / "Existing Thesis Draft" / "校园物资智能管理系统设计与实现-定稿工作稿-图表重绘版.docx"
FIGURE_DIR = ROOT / "output" / "doc" / "figures"
REPORT_PATH = ROOT / "output" / "doc" / "校园物资智能管理系统设计与实现-改稿说明-第四轮图表重绘.md"
BACKUP_PATH = ROOT / "Existing Thesis Draft" / f"校园物资智能管理系统设计与实现-定稿工作稿.docx.bak.{datetime.now():%Y%m%d-%H%M%S}-redraw-figures"


def choose_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


FONT_TITLE = choose_font(34, bold=True)
FONT_BOX_TITLE = choose_font(28, bold=True)
FONT_BOX_BODY = choose_font(21)
FONT_ENTITY = choose_font(24, bold=True)
FONT_ATTR = choose_font(18)


def draw_centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill=(0, 0, 0), spacing: int = 6):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text(((x1 + x2 - w) / 2, (y1 + y2 - h) / 2), text, font=font, fill=fill, spacing=spacing, align="center")


def draw_module_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, body_lines: list[str]) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill="white", outline="black", width=3)
    draw_centered(draw, (x1 + 16, y1 + 12, x2 - 16, y1 + 62), title, FONT_BOX_TITLE)
    body_text = "\n".join(body_lines)
    draw.multiline_text((x1 + 20, y1 + 76), body_text, font=FONT_BOX_BODY, fill="black", spacing=8)


def line_between(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="black", width=3)


def render_module_figure(path: Path) -> None:
    img = Image.new("RGB", (1800, 1180), "white")
    draw = ImageDraw.Draw(img)

    center = (710, 425, 1090, 685)
    boxes = [
        ((110, 90, 460, 250), "认证与权限", ["登录 / 刷新", "用户 / 角色 / 部门"]),
        ((725, 70, 1075, 230), "基础数据", ["校区 / 仓库 / 库位", "分类 / 物资 / 供应商"]),
        ((1340, 90, 1690, 250), "仓储库存", ["库存总量", "批次 / 入库 / 出库"]),
        ((110, 415, 460, 575), "申领审批", ["草稿 / 提交 / 审批", "紧急单快速通过"]),
        ((1340, 415, 1690, 575), "调拨管理", ["调拨单", "候选仓推荐"]),
        ((110, 760, 460, 920), "预警智能", ["库存不足 / 积压", "临期 / 异常领用"]),
        ((725, 780, 1075, 940), "统计分析", ["趋势 / 占比 / 排名", "补货建议"]),
        ((1340, 760, 1690, 920), "系统工具", ["事件 / 通知", "登录 / 操作日志"]),
    ]

    draw_module_box(draw, center, "校园物资智能管理系统", ["统一认证", "统一接口", "统一日志与通知"])
    cx1, cy1, cx2, cy2 = center
    c_top = ((cx1 + cx2) // 2, cy1)
    c_bottom = ((cx1 + cx2) // 2, cy2)
    c_left = (cx1, (cy1 + cy2) // 2)
    c_right = (cx2, (cy1 + cy2) // 2)
    c_lt = (cx1 + 40, cy1 + 25)
    c_rt = (cx2 - 40, cy1 + 25)
    c_lb = (cx1 + 50, cy2 - 25)
    c_rb = (cx2 - 50, cy2 - 25)

    for box, title, body in boxes:
        draw_module_box(draw, box, title, body)

    line_between(draw, ((110 + 460) // 2, 250), c_lt)
    line_between(draw, ((725 + 1075) // 2, 230), c_top)
    line_between(draw, ((1340 + 1690) // 2, 250), c_rt)
    line_between(draw, (460, (415 + 575) // 2), c_left)
    line_between(draw, (1340, (415 + 575) // 2), c_right)
    line_between(draw, ((110 + 460) // 2, 760), c_lb)
    line_between(draw, ((725 + 1075) // 2, 780), c_bottom)
    line_between(draw, ((1340 + 1690) // 2, 760), c_rb)

    img.save(path)


def draw_entity_panel(draw: ImageDraw.ImageDraw, area: tuple[int, int, int, int], entity_label: str, attrs: list[str]) -> None:
    x1, y1, x2, y2 = area
    cx = (x1 + x2) / 2
    panel_height = y2 - y1
    rect_w = min(230, int((x2 - x1) * 0.38))
    rect_h = 54
    rect_y = int(y1 + panel_height * 0.62)
    rect = (int(cx - rect_w / 2), rect_y, int(cx + rect_w / 2), rect_y + rect_h)
    draw.rectangle(rect, fill="white", outline="black", width=3)
    draw_centered(draw, rect, entity_label, FONT_ENTITY)

    n = len(attrs)
    rx = min(220, int((x2 - x1) * 0.34))
    ry = min(170, int((y2 - y1) * 0.34))
    angles = [math.radians(170 - i * 140 / (max(n - 1, 1))) for i in range(n)]
    node_w = 154
    node_h = 52
    origin = (cx, rect_y + 2)
    for angle, attr in zip(angles, attrs):
        ex = int(origin[0] + rx * math.cos(angle))
        ey = int(origin[1] - ry * math.sin(angle))
        oval = (ex - node_w // 2, ey - node_h // 2, ex + node_w // 2, ey + node_h // 2)
        draw.ellipse(oval, fill="white", outline="black", width=3)
        draw.line([(cx, rect_y), (ex, ey + node_h // 2 - 6)], fill="black", width=2)
        draw_centered(draw, oval, attr, FONT_ATTR)


def render_entity_figure(path: Path, panels: Iterable[tuple[tuple[int, int, int, int], str, list[str]]]) -> None:
    img = Image.new("RGB", (1900, 1240), "white")
    draw = ImageDraw.Draw(img)
    for area, label, attrs in panels:
        draw_entity_panel(draw, area, label, attrs)
    img.save(path)


def clear_paragraph_content(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag.endswith("}pPr"):
            continue
        p.remove(child)


def replace_picture_paragraph(paragraph, image_path: Path) -> None:
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.1))


def set_caption(paragraph, text: str) -> None:
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)


def paragraph_index(doc: Document, text: str) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return idx
    raise ValueError(f"Paragraph not found: {text}")


def replace_figures_in_doc(source_path: Path, target_path: Path) -> None:
    doc = Document(source_path)

    caption_map = {
        "图3-2 系统功能模块图": ("图3-2 系统功能模块图", FIGURE_DIR / "fig_3_2_modules.png"),
        "图3-3 RBAC 与组织实体关系图": ("图3-3 RBAC 与组织实体图", FIGURE_DIR / "fig_3_3_rbac_entity.png"),
        "图3-4 库存与批次实体关系图": ("图3-4 库存与批次实体图", FIGURE_DIR / "fig_3_4_inventory_entity.png"),
        "图3-5 业务单据、预警与通知实体关系图": ("图3-5 业务单据、预警与通知实体图", FIGURE_DIR / "fig_3_5_business_entity.png"),
    }

    for old_caption, (new_caption, image_path) in caption_map.items():
        idx = paragraph_index(doc, old_caption)
        replace_picture_paragraph(doc.paragraphs[idx - 1], image_path)
        set_caption(doc.paragraphs[idx], new_caption)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = False

    doc.save(target_path)


def write_report() -> None:
    REPORT_PATH.write_text(
        "\n".join(
            [
                "# 校园物资智能管理系统论文改稿说明（第四轮图表重绘）",
                "",
                "## 工作文件",
                f"- working draft：`{WORKING_DRAFT.relative_to(ROOT)}`",
                f"- fresh backup：`{BACKUP_PATH.relative_to(ROOT)}`",
                "",
                "## 本轮处理",
                "- 重画了 `图3-2` 系统功能模块图，改为更规整的中心辐射式论文模块图，不再使用上一版交叉感过强的示意布局。",
                "- 重画了 `图3-3` 至 `图3-5`，统一改为“实体矩形 + 属性椭圆”的实体图风格，贴近你给出的参考样式。",
                "- 三张实体图的字段全部回到 `sql/schema.sql` 的真实表字段，没有沿用旧稿中的 `application_form`、`inventory_info`、`alert_log` 等旧命名。",
                "- 同步把题注从“实体关系图”收口为“实体图”，使图的表达方式与题注一致。",
                "",
                "## 新图对应关系",
                "- 图3-2：系统功能模块图",
                "- 图3-3：RBAC 与组织实体图",
                "- 图3-4：库存与批次实体图",
                "- 图3-5：业务单据、预警与通知实体图",
                "",
                "## 仍需本机 Word 复核",
                "- 打开 working draft 后检查四张新图的分页和缩放是否满足学校模板观感。",
                "- 如需进一步贴近导师常见风格，可在 Word 中把图 3-3 至图 3-5 的图片宽度再微调 1-2 个字号级别。",
            ]
        ),
        encoding="utf-8",
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=WORKING_DRAFT)
    parser.add_argument("--target", type=Path, default=WORKING_DRAFT)
    parser.add_argument("--skip-backup", action="store_true")
    args = parser.parse_args()

    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    if not args.skip_backup and args.source.resolve() == WORKING_DRAFT.resolve():
        shutil.copy2(WORKING_DRAFT, BACKUP_PATH)

    render_module_figure(FIGURE_DIR / "fig_3_2_modules.png")
    render_entity_figure(
        FIGURE_DIR / "fig_3_3_rbac_entity.png",
        [
            ((40, 30, 610, 560), "部门（sys_dept）", ["id", "dept_name", "parent_id", "deleted", "version"]),
            ((655, 30, 1245, 600), "用户（sys_user）", ["id", "username", "password", "real_name", "dept_id", "role_id", "status"]),
            ((1290, 30, 1860, 560), "角色（sys_role）", ["id", "role_code", "role_name", "description"]),
            ((140, 640, 840, 1200), "刷新令牌（auth_refresh_token）", ["id", "user_id", "token_id", "token_hash", "expire_at", "revoked"]),
            ((1030, 640, 1760, 1200), "登录日志（login_log）", ["id", "user_id", "username", "login_ip", "login_status", "login_time"]),
        ],
    )
    render_entity_figure(
        FIGURE_DIR / "fig_3_4_inventory_entity.png",
        [
            ((40, 30, 610, 560), "分类（material_category）", ["id", "category_name", "remark"]),
            ((655, 30, 1245, 600), "物资档案（material_info）", ["id", "material_code", "material_name", "category_id", "safety_stock", "supplier"]),
            ((1290, 30, 1860, 560), "仓库（warehouse）", ["id", "warehouse_name", "campus", "address", "manager"]),
            ((140, 640, 840, 1200), "库存（inventory）", ["id", "material_id", "warehouse_id", "current_qty", "locked_qty"]),
            ((1030, 640, 1760, 1200), "库存批次（inventory_batch）", ["id", "material_id", "warehouse_id", "batch_no", "in_qty", "remain_qty", "expire_date"]),
        ],
    )
    render_entity_figure(
        FIGURE_DIR / "fig_3_5_business_entity.png",
        [
            ((20, 20, 610, 560), "申领单（apply_order）", ["id", "dept_id", "applicant_id", "urgency_level", "status", "approver_id"]),
            ((650, 20, 1250, 560), "申领明细（apply_order_item）", ["id", "apply_order_id", "material_id", "apply_qty", "actual_qty"]),
            ((1290, 20, 1880, 560), "调拨单（transfer_order）", ["id", "from_warehouse_id", "to_warehouse_id", "status", "applicant_id", "approver_id"]),
            ((140, 640, 840, 1200), "预警记录（warning_record）", ["id", "warning_type", "material_id", "warehouse_id", "handle_status", "handler_id"]),
            ((1030, 640, 1760, 1200), "通知消息（notification）", ["id", "title", "msg_type", "target_user_id", "is_read", "biz_id"]),
        ],
    )

    replace_figures_in_doc(args.source, args.target)
    write_report()

    doc = Document(args.target)
    print(args.target)
    if not args.skip_backup and args.source.resolve() == WORKING_DRAFT.resolve():
        print(BACKUP_PATH)
    print(REPORT_PATH)
    print("images", len(doc.inline_shapes))
    print("tables", len(doc.tables))


if __name__ == "__main__":
    main()
