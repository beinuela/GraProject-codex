from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "Thesis Template" / "7 软件学院本科毕业设计（论文）模板 2026.docx"
OUT_DIR = ROOT / "output" / "doc"
FINAL_DOCX = OUT_DIR / "校园物资智能管理系统设计与实现-正式定稿版.docx"
REPORT = OUT_DIR / "校园物资智能管理系统设计与实现-定稿改写说明.md"
FIG_DIR = OUT_DIR / "figures"
SCREEN_DIR = OUT_DIR / "runtime-screenshots"

TITLE = "校园物资智能管理系统设计与实现"
EN_TITLE = "Design and Implementation of Campus Material Intelligent Management System"


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_run_font(run, size=12, bold=False, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_para_format(paragraph, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)


def add_para(doc: Document, text: str = "", first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12):
    p = doc.add_paragraph()
    set_para_format(p, first_line=first_line, align=align)
    r = p.add_run(text)
    set_run_font(r, size=size)
    return p


def add_center(doc: Document, text: str, size=16, bold=True, font="黑体"):
    p = doc.add_paragraph()
    set_para_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, font=font)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading("", level=level)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.page_break_before = True
        size = 16
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 15
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        size = 12
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, font="黑体")
    return p


def add_table_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, size=10.5, bold=True)
    return p


def apply_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = f"w:{edge}"
                element = borders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    borders.append(element)
                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), "4")
                element.set(qn("w:space"), "0")
                element.set(qn("w:color"), "000000")


def add_figure_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, size=10.5, bold=True)
    return p


def add_picture(doc: Document, path: Path, caption: str, width_cm=14.5):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        add_figure_caption(doc, caption)
    else:
        add_para(doc, f"{caption}所需图片文件未找到：{path.name}", first_line=False)


def add_table(doc: Document, caption: str, headers, rows, widths=None):
    add_table_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for style_name in ("Table Grid", "网格型", "网格表"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.2
                for r in p.runs:
                    set_run_font(r, size=9 if len(rows) > 8 else 10.5)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    apply_table_borders(table)
    return table


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_field(paragraph, instr: str):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    run._r.append(instr_text)

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char)

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def add_toc(doc: Document):
    add_center(doc, "目  录", size=16, bold=True, font="黑体")
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u')
    add_page_break(doc)


def clear_header_footer(section):
    for part in (section.header, section.footer):
        part.is_linked_to_previous = False
        for p in part.paragraphs:
            p.text = ""


def configure_doc(doc: Document):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)
        clear_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True


def set_body_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.text = ""
    run = header_p.add_run(TITLE)
    set_run_font(run, size=9)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer_p, "PAGE")
    for r in footer_p.runs:
        set_run_font(r, size=9, font="Times New Roman")


REFERENCES = [
    "李晓英, 高菲. 智慧校园背景下的高职院校物资管理系统建设创新研究[J]. 现代物业(中旬刊), 2019(09): 12.",
    "唐锡雷. 基于高职院校的后勤物资管理系统设计与实现[D]. 成都: 电子科技大学, 2015.",
    "曾文勇, 刘小玲. 数智赋能物资供应管理提质提效[N]. 中国石化报, 2026-04-24(002).",
    "张俊萌, 张华贵. 医疗机构物资流通管理信息系统的设计与实现[J]. 工业控制计算机, 2026, 39(04): 132-134.",
    "厉志安, 李笑, 葛江勤, 等. 基于物联网的高校实验室危化品溯源智能管理[J]. 实验室研究与探索, 2025, 44(03): 258-263.",
    "崔靖茹, 文华, 刘宏磊, 等. 基于Vue和SpringBoot框架的高校信息化项目管理系统的设计与实现[J]. 现代信息科技, 2025, 9(22): 77-81.",
    "薛楠楠, 张伟, 钟化雨, 等. 基于SpringBoot+Vue的事故报告数据库系统设计与实现[J]. 计算机仿真, 2025, 42(07): 456-462+570.",
    "李琴, 崔名扬, 钱奕文, 等. 基于SpringBoot的研究生学术档案管理系统开发[J]. 电脑知识与技术, 2023, 19(18): 46-48+51.",
    "沈露, 孙雨晨, 义智文宇, 等. 基于互联网的爱心校园服务平台设计与实现[J]. 电脑编程技巧与维护, 2023(06): 56-58.",
    "Kevin Santiago Rey Rodriguez, Julián David Avellaneda Galindo, Josep Tárrega Juan, et al. Secure Development Methodology for Full Stack Web Applications: Proof of the Methodology Applied to Vue.js, Spring Boot and MySQL[J]. Computers, Materials & Continua, 2025, 85(1).",
    "Yuanrun Zhu. Contract Management System Based on SpringBoot and Vue[J]. Advances in Computer, Signals and Systems, 2024, 8(5).",
    "Yixuan Liu. Design and Implementation of a Student Attendance Management System based on Springboot and Vue Technology[J]. Frontiers in Computing and Intelligent Systems, 2024, 8(1).",
    "Jian Chen, Chen Jian, Pan Hailan. Design of Man Hour Management Information System on SpringBoot Framework[J]. Journal of Physics: Conference Series, 2020, 1646(1).",
    "王晴. 基于Java和MySQL技术的供求信息网设计与实现[J]. 兰州石化职业技术大学学报, 2025, 25(04): 26-31.",
    "王阳坤, 陈宇峰, 向郑涛, 等. 基于Vue框架的动态路由监控系统设计[J]. 湖北汽车工业学院学报, 2026, 40(01): 8-11.",
    "陆岫昶, 林娜, 骆永翰, 等. Vue支持下的多层次协同可视化前端设计[J]. 电子设计工程, 2025, 33(19): 127-131.",
    "张文亮. MySQL应用实战与性能调优[M]. 北京: 机械工业出版社, 2022.",
    "袁子伯, 杨富源. 基于MySQL的数据库查询性能优化技术研究[N]. 农业科技报, 2026-02-02(006).",
    "卢万有. 基于JWT的RBAC在前后端分离项目中的设计与实现[J]. 电脑编程技巧与维护, 2025(01): 46-48.",
    "胡传甫, 王昕葳, 周宬, 等. 业务逻辑漏洞常见类型分析与JWT测试示例研究[J]. 计算机时代, 2025(02): 6-10.",
    "张迪, 刘小翔, 姚欢欢, 等. 基于B/S架构的入厂煤智能管理平台设计与开发[J]. 自动化应用, 2026, 67(07): 222-224+228.",
    "杨职远, 马海龙. 基于人工智能的智慧实验室管理系统的设计与应用[J]. 软件, 2025, 46(11): 28-30.",
]


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE, FINAL_DOCX)
    doc = Document(str(FINAL_DOCX))
    clear_document(doc)
    configure_doc(doc)

    # Cover
    add_center(doc, "郑州轻工业大学", size=16, bold=True, font="宋体")
    add_center(doc, "本科毕业设计（论文）", size=22, bold=True, font="黑体")
    for _ in range(4):
        add_para(doc, "", first_line=False)
    add_center(doc, f"题    目    {TITLE}", size=16, bold=True, font="宋体")
    for label, value in [
        ("学生姓名", "刘祎凯"),
        ("专业班级", "软件工程 22-06 班"),
        ("学    号", "542213340621"),
        ("学    院", "软件学院"),
        ("指导教师（职称）", "周开来"),
        ("完成时间", "2026年5月18日"),
    ]:
        add_center(doc, f"{label}    {value}", size=15, bold=False, font="宋体")
    add_page_break(doc)

    # Task book
    add_center(doc, "郑州轻工业大学", size=16, bold=True, font="宋体")
    add_center(doc, "毕业设计（论文）任务书", size=16, bold=True, font="黑体")
    add_para(doc, f"题目    {TITLE}", first_line=False)
    add_para(doc, "专业    软件工程    学号    542213340621    姓名    刘祎凯", first_line=False)
    add_para(doc, "主要内容：本课题围绕高校物资管理中台账分散、库存批次难追踪、申领审批链条不清晰和预警处置滞后等问题，设计并实现校园物资智能管理系统。系统采用前后端分离架构，后端基于 Spring Boot、Spring Security、JWT、MyBatis-Plus 和 MySQL，前端基于 Vue 3、Vite、Pinia、Element Plus 与 ECharts，完成基础数据、库存批次、申领审批、调拨协同、预警处理、通知日志和统计分析等功能。")
    add_para(doc, "基本要求：系统功能应与真实代码、数据库脚本和测试材料一致；论文应围绕需求分析、总体设计、数据库设计、详细实现和系统测试展开；智能化表述限于库存预警、效期规则、候选仓排序、补货建议和统计分析等可复核内容。")
    add_para(doc, "思政要求：课题应体现服务校园治理、规范物资流转、强化数据安全和审计追踪的责任意识，在系统设计与论文写作中坚持真实、严谨和可验证原则。")
    add_para(doc, "主要参考资料：")
    for i in [1, 4, 6, 10, 17]:
        add_para(doc, f"[{i}] {REFERENCES[i-1]}", first_line=False, size=10.5)
    add_para(doc, "完  成  期  限：2026年5月18日", first_line=False)
    add_para(doc, "指导教师签名：", first_line=False)
    add_para(doc, "专业负责人签名：", first_line=False)
    add_para(doc, "年    月    日", first_line=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_break(doc)

    # Abstracts
    add_center(doc, TITLE, size=16, bold=True, font="黑体")
    cn_abs = (
        "摘 要：针对高校物资管理中台账分散、审批流转不清晰、库存批次难以追踪和预警处理滞后等问题，本文基于当前项目仓库完成校园物资智能管理系统的设计与实现。系统采用前后端分离架构，前端使用 Vue 3、Vite、Pinia、Element Plus 和 ECharts 组织业务页面与统计图表，后端使用 Spring Boot、Spring Security、JWT、MyBatis-Plus 和 MySQL 实现认证授权、业务流程控制与数据持久化。系统围绕校园物资流转场景实现了基础数据维护、库存与批次管理、申领审批、出库执行、仓间调拨、预警扫描、通知消息、日志审计和统计分析等功能。本文所述智能能力限定为可解释的规则化辅助功能，包括低库存、积压、临期、过期和异常出库预警，按效期优先的批次扣减，基于固定校区图的候选仓排序，以及基于近 30 天出库量和近 6 个月历史的补货参考。测试结果显示，后端自动化测试、前端构建、前端单测和端到端流程验证均已通过，系统能够支撑校园物资核心业务流程的演示与验证。"
    )
    add_para(doc, cn_abs)
    add_para(doc, "关键词：校园物资管理; Spring Boot; Vue 3; 库存预警; 申领审批")
    add_page_break(doc)

    add_center(doc, EN_TITLE, size=16, bold=True, font="Times New Roman")
    en_abs = (
        "Abstract: This thesis designs and implements a campus material intelligent management system for problems such as scattered ledgers, unclear approval flows, weak batch traceability, and delayed warning handling in campus material management. The system adopts a front-end and back-end separation architecture. The front end uses Vue 3, Vite, Pinia, Element Plus, and ECharts to organize business pages and statistical charts, while the back end uses Spring Boot, Spring Security, JWT, MyBatis-Plus, and MySQL to implement authentication, authorization, workflow control, and data persistence. The implemented functions include master data maintenance, inventory and batch management, requisition approval, outbound execution, warehouse transfer, warning scanning, notification, audit logging, and statistical analysis. The intelligent capabilities described in this thesis are limited to explainable rule-based assistance, including stock warnings, expiry warnings, FEFO batch deduction, candidate warehouse ranking, replenishment suggestions, and moving-average-based reference forecasts. The latest verification shows that the back-end automated tests, front-end production build, front-end unit tests, and end-to-end workflow tests have passed. The system can support the demonstration and verification of core campus material management processes."
    )
    add_para(doc, en_abs)
    add_para(doc, "Keywords: campus material management; Spring Boot; Vue 3; inventory warning; requisition approval")
    add_page_break(doc)

    add_toc(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    clear_header_footer(doc.sections[0])
    set_body_header_footer(body_section)

    write_chapters(doc)
    write_references_and_appendix(doc)

    doc.save(str(FINAL_DOCX))
    write_report()


def write_chapters(doc: Document):
    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    for text in [
        "高校日常运行离不开实验耗材、办公用品、医疗防护、清洁消杀、应急照明和设备备件等物资。传统管理方式常依靠纸质台账、分散 Excel 表和人工沟通完成，数据更新滞后、审批记录不完整、库存批次不透明等问题会直接影响物资调配效率。智慧校园建设要求后勤保障业务逐步从单点记录转向流程协同和数据治理，物资管理系统因此需要承担主数据维护、库存追踪、审批流转和风险提醒等职责[1][2]。",
        "从物资管理研究和行业实践看，信息化系统通常将入库、领用、审核、库存监控等环节统一管理，以减少人工统计和重复沟通成本[3][4]。高校实验室和后勤物资还涉及有效期、批次和安全合规问题，相关研究强调全生命周期管理、身份认证、异常提醒和后续数据分析的重要性[5]。本课题正是在这一背景下，将校园物资从入库、申领、审批、出库、调拨到预警处置的核心链路转化为可追踪的数据流程。",
        "需要说明的是，本文所称“智能管理”并不等同于复杂人工智能模型。当前项目仓库能够证明的能力主要是规则化判断、候选仓排序、批次优先扣减、补货建议和图表统计。论文写作因此以真实实现为边界，把智能部分限定为对业务人员有帮助、规则可解释、结果可复核的辅助能力。"
    ]:
        add_para(doc, text)

    add_heading(doc, "1.2 研究意义", 2)
    for text in [
        "本课题的应用意义在于将校园物资管理中原本分散的角色、单据、库存和日志整理为统一流程。部门用户提交申领后，审批人员可以基于单据状态进行处理，仓库管理员再根据库存批次执行出库或调拨，最终形成可回溯的业务闭环。这种处理方式能够减少口头确认和线下台账之间的信息偏差。",
        "本课题的工程意义在于将软件工程中的需求分析、数据库设计、前后端分离、权限控制、事务处理和测试验证结合起来。系统不是单纯展示静态页面，而是围绕 `sys_user`、`material_info`、`inventory_batch`、`apply_order`、`transfer_order`、`warning_record` 等真实数据对象组织功能。论文各章也据此展开，使需求、设计、实现和测试能够相互对应。",
        "从安全和管理角度看，系统通过 JWT 登录认证、角色菜单、接口权限和日志记录形成基本的访问边界。相关研究表明，前后端分离项目中的 JWT 与 RBAC 能够降低会话依赖并提升权限管理清晰度[19]，但也需要关注业务逻辑漏洞和令牌使用风险[20]。本系统在实现中采用刷新令牌持久化、登录日志和操作日志等机制，为后续管理和排查提供依据。"
    ]:
        add_para(doc, text)

    add_heading(doc, "1.3 国内外研究现状", 2)
    add_heading(doc, "1.3.1 国内研究现状", 3)
    for text in [
        "国内有关高校信息化和管理系统的研究较多，基于 Spring Boot 与 Vue 的前后端分离方案已经广泛用于高校项目管理、事故报告数据库、研究生档案管理和校园服务平台等场景[6][7][8][9]。这些系统通常采用 B/S 架构，将用户管理、角色权限、业务单据、数据查询和统计展示拆分为若干模块，适合本科毕业设计中说明需求分析、总体设计、数据库设计和系统测试。",
        "与本课题更接近的研究集中在物资流通、后勤物资、实验室危化品和库存监控等方向。医疗机构物资流通系统强调入库、借用、领用、审核和库存监控的连续处理[4]，高校危化品溯源研究强调从采购、入库、使用到报废的全生命周期闭环[5]。这些研究为校园物资系统提供了流程建模参考，但本文不直接照搬其业务场景，而是以当前项目中已有代码和数据库为准。",
        "国内近年的智能管理研究也开始关注预警、流程自动化和数据分析。例如入厂煤智能管理平台将流程管控和预警模型结合起来[21]，智慧实验室管理研究讨论了更复杂的人工智能方法[22]。相比之下，本系统实现规模更适合本科毕业设计，采用固定阈值、时间窗口、移动平均和候选排序等轻量方法，避免把未实现的模型写成系统成果。"
    ]:
        add_para(doc, text)
    add_heading(doc, "1.3.2 国外相关实践启示", 3)
    for text in [
        "国外文献中与本课题直接相关的主要是 Web 应用系统工程实践和全栈安全方法。有关 Vue.js、Spring Boot 和 MySQL 的全栈安全方法研究强调身份认证、授权控制、错误处理、日志和输入校验等安全维度[10]。Spring Boot 和 Vue 也被用于合同管理、学生考勤和工时管理等信息系统实现[11][12][13]，说明该技术路线适合构建以业务流程为中心的管理系统。",
        "这些实践的启示在于，管理系统的价值不仅来自页面数量，也来自业务流程是否稳定、数据模型是否清楚、权限边界是否明确、测试是否可复核。本文后续章节因此不追求复杂理论模型，而是围绕本项目的真实功能说明系统如何支持校园物资的申领、审批、库存、调拨和预警。"
    ]:
        add_para(doc, text)
    add_heading(doc, "1.4 研究内容与论文结构", 2)
    add_para(doc, "本文主要研究内容包括：分析校园物资管理的角色和业务流程；基于前后端分离架构设计系统总体结构；根据 `schema.sql` 建立数据库表和索引说明；围绕认证授权、基础数据、库存批次、申领审批、调拨推荐、预警扫描、补货建议和统计分析说明详细实现；最后结合自动化测试、前端构建、E2E 和 k6 脚本进行测试验证。全文共分 7 章，依次为绪论、相关技术介绍、系统需求分析、系统总体设计、数据库设计、系统详细设计与实现、系统测试，最后给出结束语、致谢、参考文献和附录。")

    add_heading(doc, "2 相关技术介绍", 1)
    add_heading(doc, "2.1 技术选型概述", 2)
    for text in [
        "本系统采用前后端分离技术路线。前端负责页面展示、表单交互、路由跳转和图表呈现，后端负责身份认证、权限控制、业务规则、事务处理和数据访问。此类技术组合在高校管理系统和通用信息系统中较为常见，能够较好支持角色隔离、列表查询和流程审批[6][7][11][12]。",
        "根据 `frontend/package.json` 和 `backend/pom.xml`，系统实际使用的核心技术包括 Vue 3、Vite、Pinia、Element Plus、Axios、ECharts、Spring Boot 3.3.5、Spring Security、JWT、MyBatis-Plus、JdbcTemplate、MySQL 8 和 H2。本文只介绍这些真实使用的技术，不展开未在项目中承担核心功能的其他技术。"
    ]:
        add_para(doc, text)
    add_heading(doc, "2.2 前端开发技术", 2)
    for text in [
        "Vue 3 用于构建前端页面，配合 Vue Router 管理登录页、仪表盘、基础数据、库存、申领、调拨、预警、统计分析和系统工具等路由。Pinia 用于保存登录状态和用户信息，Axios 封装后端接口请求，Element Plus 提供表单、表格、对话框、按钮和分页等基础组件。Vue 框架在动态监控和可视化系统中具有组件化、响应式和交互效率方面的优势[15][16]。",
        "Vite 用于前端工程构建。它能够提供较快的本地开发启动和生产构建能力，适合本系统这类多页面管理端。ECharts 用于库存占比、出入库趋势、仓库分布、部门排行和效期统计等图表展示。本文在第 6 章只将图表说明为统计分析和可视化展示，不把它写成实时大数据分析平台。"
    ]:
        add_para(doc, text)
    add_heading(doc, "2.3 后端与数据访问技术", 2)
    for text in [
        "后端基于 Spring Boot 3.3.5 构建 REST API，使用 Controller、Service、Mapper 和 Entity 分层组织代码。Spring Boot 能够简化配置管理、依赖装配和测试集成，在管理信息系统开发中具有较好的工程适配性[13]。MyBatis-Plus 负责实体映射和通用 CRUD，JdbcTemplate 用于统计分析和补货建议中的聚合查询。",
        "数据层使用 MySQL 8 作为业务数据库，测试和截图场景使用 H2 模拟 MySQL 模式。MySQL 在中小型管理系统中具有成熟的事务、索引和查询能力，数据库设计与查询优化需要围绕业务查询场景进行[17][18]。本系统的索引主要覆盖用户角色、批次出库、预警查询、日志时间和通知时间等高频访问方向。"
    ]:
        add_para(doc, text)
    add_heading(doc, "2.4 安全认证与可视化技术", 2)
    for text in [
        "系统采用 Spring Security 与 JWT 实现无状态认证。用户登录后获得 access token 和 refresh token，前端在后续请求中通过 Authorization 头传递令牌；后端解析令牌并将当前用户、角色和权限信息放入安全上下文。JWT 与 RBAC 在前后端分离项目中常用于简化会话管理和权限控制[19]。",
        "安全设计不能只停留在登录成功。项目中还实现了 refresh token 持久化、刷新令牌轮换、登录日志、操作日志、安全响应头和部分高风险接口限流。业务逻辑漏洞研究提醒开发者要重视登录认证、业务流程和接口调用中的边界条件[20]，因此论文后续会把状态流转和权限边界作为实现章节的重要内容。"
    ]:
        add_para(doc, text)

    add_heading(doc, "3 系统需求分析", 1)
    add_heading(doc, "3.1 业务场景与角色划分", 2)
    add_para(doc, "系统面向高校物资管理场景，核心用户包括系统管理员、仓库管理员、审批人员和部门用户。系统管理员维护用户、角色、部门、校区、物资分类、物资档案、供应商、系统配置等基础数据；仓库管理员负责仓库、库位、库存、入库、出库、调拨执行和预警处理；审批人员负责申领单和调拨单审核；部门用户主要完成物资申领、查看流转状态和签收。该划分来自 `AuthService.buildMenusByRole`、前端路由和种子数据。")
    add_heading(doc, "3.2 功能需求分析", 2)
    need_rows = [
        ("基础数据管理", "维护用户、角色、部门、校区、仓库、库位、物资分类、物资档案、供应商和系统配置，为后续业务单据提供基础引用。"),
        ("库存与批次管理", "支持库存分页查询、入库建批次、批次余量查询、出库扣减和库存盘点。库存汇总与批次明细分表管理。"),
        ("申领审批与出库", "部门用户创建申领单，提交后锁定库存；审批通过后仓库管理员按申领单执行出库，部门用户完成签收。"),
        ("调拨协同", "仓库管理员创建调拨单，审批人员审核后执行跨仓批次搬移，目标仓生成镜像批次并更新库存汇总。"),
        ("预警统计与支撑", "支持低库存、积压、临期、过期、异常出库预警；支持补货建议、统计图表、通知消息、登录日志和操作日志。"),
    ]
    add_table(doc, "表3-1 系统功能需求概括", ["功能域", "需求说明"], need_rows, widths=[3.8, 11])
    add_heading(doc, "3.2.1 基础数据管理需求", 3)
    add_para(doc, "基础数据是系统运行的前提。物资档案中的物资编码、名称、分类、规格、单位、安全库存、保质期和供应商信息会被库存、预警、申领和补货建议反复引用；仓库和库位数据决定库存记录和调拨单的归属；部门和角色决定用户可以访问的菜单与业务流程。")
    add_heading(doc, "3.2.2 库存批次管理需求", 3)
    add_para(doc, "校园物资管理不能只依赖一个总库存数字。系统需要同时保存某物资在某仓库中的汇总库存和每个批次的余量、生产日期、到期日期。这样才能在出库时优先使用更早到期的批次，并在临期或过期时生成预警。")
    add_heading(doc, "3.2.3 申领审批需求", 3)
    add_para(doc, "申领流程要求部门用户先创建单据，再提交进入审批。提交时系统根据可用库存选择满足物资需求的仓库并锁定库存，避免多个申请同时提交造成库存超占。审批通过后单据进入可出库状态，出库完成后部门用户签收，形成业务闭环。")
    add_heading(doc, "3.2.4 调拨协同需求", 3)
    add_para(doc, "调拨流程用于处理跨仓库存不平衡。系统需要保存调出仓、调入仓、调拨物资和数量，并通过提交、审批、执行、签收等状态节点记录协同过程。推荐功能只根据固定校区图和库存余量对候选仓排序，不承担完整物流调度。")
    add_heading(doc, "3.2.5 预警统计需求", 3)
    add_para(doc, "预警需求包括低库存、库存积压、临期、过期和异常出库。统计需求包括运营总览、库存占比、出入库趋势、仓库分布、部门排行、效期统计等。相关结果用于辅助管理人员发现库存风险和了解物资流转情况。")
    add_heading(doc, "3.3 非功能需求", 2)
    add_table(doc, "表3-2 系统非功能需求", ["类别", "要求"], [
        ("安全性", "登录认证、JWT 校验、角色菜单、接口授权、安全响应头、登录日志和操作日志。"),
        ("一致性", "库存汇总、批次余量、申领实发数量和单据状态在同一事务内更新。"),
        ("可维护性", "后端按模块拆分 Controller、Service、Mapper、Entity，前端按视图和通用组件组织。"),
        ("可测试性", "提供后端单元与集成测试、前端单测、Playwright E2E 和 k6 脚本。"),
        ("可观测性", "提供 Actuator、Prometheus 指标、JSON 日志和基础观测配置。"),
    ], widths=[3, 12])
    add_heading(doc, "3.4 业务流程与用例", 2)
    add_picture(doc, FIG_DIR / "fig_2_1_apply_flow.png", "图3-1 申领审批闭环流程图")
    add_para(doc, "申领审批流程从部门用户创建申领单开始，经过提交、审批、出库和签收。系统在提交阶段锁定库存，在出库阶段回写实发数量，并在签收阶段释放未使用的锁定量。")
    add_picture(doc, FIG_DIR / "fig_2_2_transfer_flow.png", "图3-2 调拨执行流程图")
    add_para(doc, "调拨执行流程强调跨仓协同。调拨单审批通过后，服务层按 FEFO 顺序扣减来源仓批次，并为目标仓生成包含调拨单号的新批次记录。")
    add_picture(doc, FIG_DIR / "fig_2_3_warning_flow.png", "图3-3 预警处置流程图")
    add_para(doc, "预警处置流程包括定时或手动扫描、预警去重、记录生成和人工处理。系统只对未处理预警做去重，避免同一风险重复堆积。")

    add_heading(doc, "4 系统总体设计", 1)
    add_heading(doc, "4.1 系统总体架构", 2)
    add_para(doc, "系统采用 B/S 架构和前后端分离设计。用户通过浏览器访问 Vue 前端，前端通过 Axios 调用 Spring Boot API；后端负责认证、权限、业务规则和数据访问；MySQL 保存正式业务数据，H2 用于测试和截图环境。该设计与当前仓库 `docs/architecture.md`、`backend/pom.xml` 和 `frontend/package.json` 保持一致。")
    add_picture(doc, FIG_DIR / "fig_3_1_architecture.png", "图4-1 系统总体架构图")
    add_heading(doc, "4.2 功能模块设计", 2)
    add_para(doc, "系统功能模块按照实际代码目录和前端路由划分，包括认证授权、RBAC、基础数据、物资档案、仓库库位、库存出入库、申领审批、调拨管理、预警中心、统计分析、通知消息、事件管理和日志审计。模块之间通过业务单据和数据库主键关联，不采用微服务拆分。")
    add_picture(doc, FIG_DIR / "fig_3_2_modules.png", "图4-2 系统功能模块图")
    add_heading(doc, "4.3 接口与安全设计", 2)
    add_para(doc, "系统接口统一以 `/api` 为前缀，登录、刷新、当前用户和菜单由 `AuthController` 提供；申领、库存、调拨、预警、统计、通知和日志分别由对应 Controller 提供。接口响应采用 `ApiResponse` 统一包装，分页接口使用 `PageQuery` 和 `PageResult`，返回 records、total、page 和 size。")
    add_para(doc, "安全设计包括三层边界：第一层是登录认证，用户通过用户名和 BCrypt 密码校验获取令牌；第二层是 JWT 过滤器，后端从 Authorization 头解析用户身份；第三层是角色与方法级权限，控制不同角色可访问的菜单和接口。refresh token 以 hash 形式持久化，刷新成功后撤销旧令牌并签发新令牌。")
    add_heading(doc, "4.4 智能能力设计", 2)
    add_para(doc, "本文所述智能能力只指规则化和统计化辅助能力。入厂煤智能平台和智慧实验室研究中存在流程自动化、预警模型和人工智能应用[21][22]，但本系统没有实现复杂预测、图像识别或深度学习模型。系统当前可写的智能能力如表4-1所示。")
    add_table(doc, "表4-1 智能能力与实现方式对照表", ["能力", "实现依据", "写作边界"], [
        ("库存预警", "WarningService 扫描 currentQty 与 safetyStock", "写低库存、积压等规则，不写预测模型。"),
        ("效期提醒", "inventory_batch.expire_date 与当前日期比较", "写临期、过期提醒。"),
        ("FEFO 出库", "InventoryService 按 expire_date、id 排序扣减批次", "写批次优先规则。"),
        ("调拨推荐", "DijkstraUtil 固定校区图 + TransferService 库存筛选", "写候选仓排序，不写物流调度。"),
        ("补货建议", "SmartService 根据 30 天出库量和保障天数计算建议量", "写补货参考。"),
        ("移动平均预测", "SmartService 使用近 6 个月出库历史均值", "写参考预测，不写模型训练。"),
        ("统计分析", "AnalyticsController 聚合库存和出入库数据", "写可视化统计。"),
    ], widths=[3, 6, 6])

    add_heading(doc, "5 数据库设计", 1)
    add_heading(doc, "5.1 数据库设计原则", 2)
    add_para(doc, "数据库设计以 `sql/schema.sql` 为唯一依据。系统将用户角色、物资主数据、库存汇总、库存批次、业务单据、预警记录、通知消息和日志审计分别建表，既避免把所有业务混在一张表中，也避免设计当前代码不使用的冗余表。")
    add_heading(doc, "5.2 实体关系设计", 2)
    add_para(doc, "系统实体关系可分为三组。第一组是用户、角色和部门，用于支撑身份认证和菜单权限；第二组是物资、仓库、库存和批次，用于支撑物资存储与出入库；第三组是申领单、调拨单、预警、通知和日志，用于支撑业务闭环与审计。")
    add_picture(doc, FIG_DIR / "fig_3_3_rbac_er.png", "图5-1 用户角色与组织 E-R 图")
    add_picture(doc, FIG_DIR / "fig_3_4_inventory_er.png", "图5-2 库存与批次 E-R 图")
    add_picture(doc, FIG_DIR / "fig_3_5_business_er.png", "图5-3 业务单据、预警与通知 E-R 图")
    add_heading(doc, "5.3 关键数据表设计", 2)
    db_rows = [
        ("sys_user", "username、password、real_name、dept_id、role_id、status", "系统账号和角色归属。"),
        ("sys_role", "role_code、role_name、description", "角色编码与名称。"),
        ("sys_dept", "dept_name、parent_id", "组织部门层级。"),
        ("auth_refresh_token", "user_id、token_id、token_hash、expire_at、revoked", "刷新令牌持久化和撤销。"),
        ("material_info", "material_code、material_name、category_id、safety_stock、shelf_life_days", "物资档案与安全库存。"),
        ("warehouse", "warehouse_name、campus、address、manager", "仓库基础信息。"),
        ("inventory", "material_id、warehouse_id、current_qty、locked_qty", "物资仓库维度库存汇总。"),
        ("inventory_batch", "material_id、warehouse_id、batch_no、remain_qty、expire_date", "批次库存与效期。"),
        ("apply_order", "dept_id、applicant_id、urgency_level、status、reserved_warehouse_id", "申领主单。"),
        ("apply_order_item", "apply_order_id、material_id、apply_qty、actual_qty", "申领明细。"),
        ("stock_in / stock_out", "warehouse_id、operator_id、apply_order_id、remark", "入库和出库主单。"),
        ("transfer_order", "from_warehouse_id、to_warehouse_id、status、approver_id", "调拨主单。"),
        ("warning_record", "warning_type、material_id、warehouse_id、handle_status", "预警记录。"),
        ("notification", "title、content、msg_type、target_user_id、is_read", "站内通知。"),
        ("login_log / operation_log", "用户、模块、操作、时间、状态", "登录与操作审计。"),
    ]
    add_table(doc, "表5-1 关键数据表说明", ["数据表", "主要字段", "设计作用"], db_rows, widths=[3.2, 6.5, 5.3])
    add_heading(doc, "5.4 索引与约束设计", 2)
    add_para(doc, "系统索引围绕实际查询和状态过滤设计。`sys_user.username`、`sys_role.role_code`、`material_info.material_code` 和 `inventory(material_id, warehouse_id)` 使用唯一约束保证关键业务对象不重复。批次出库使用 `idx_batch_outbound_pick(material_id, warehouse_id, expire_date, remain_qty)` 支撑按物资、仓库和效期查找候选批次；预警查询使用 `idx_warning_type_material_warehouse_status` 支撑去重和状态筛选；日志和通知按时间与用户建立索引，便于分页查询。")

    add_heading(doc, "6 系统详细设计与实现", 1)
    add_heading(doc, "6.1 开发与运行环境", 2)
    add_table(doc, "表6-1 系统开发与运行环境", ["类别", "内容"], [
        ("后端", "Java 17，Spring Boot 3.3.5，Spring Security，MyBatis-Plus 3.5.7，JJWT，Bucket4j，Caffeine。"),
        ("前端", "Vue 3.5.13，Vite 6.2.0/6.4.1，Pinia 2.3.1，Element Plus 2.9.7，ECharts 5.6.0，Axios。"),
        ("数据库", "MySQL 8 用于业务数据，H2 用于测试、截图和 E2E 隔离环境。"),
        ("测试", "JUnit、Spring Boot Test、Vitest、Playwright、k6。"),
        ("运行接口", "后端 8080，管理端点 18080，前端开发/预览端口 5173 或 4173。"),
    ], widths=[3, 12])
    add_heading(doc, "6.2 认证授权与登录态管理实现", 2)
    add_para(doc, "认证授权由 `AuthController`、`AuthService`、`JwtTokenProvider`、`JwtAuthenticationFilter` 和 `SecurityConfig` 共同完成。登录时系统校验用户名、账号状态和 BCrypt 密码，成功后签发 access token 和 refresh token，并记录登录日志。refresh token 不以明文保存，而是保存 token_id、token_hash、过期时间和撤销状态。")
    add_picture(doc, FIG_DIR / "fig_4_1_auth_flow.png", "图6-1 登录认证与令牌续签流程图")
    add_picture(doc, SCREEN_DIR / "fig_6_2_login.png", "图6-2 系统登录界面")
    add_para(doc, "刷新令牌采用单次使用策略。服务端校验 refresh token 的类型、jti、hash 和过期时间，校验通过后先撤销旧记录，再生成新的 token 对。该设计能够减少刷新令牌重复使用风险。")
    add_heading(doc, "6.3 基础数据管理实现", 2)
    add_para(doc, "基础数据模块包括用户、角色、部门、校区、仓库、库位、物资分类、物资档案、供应商和系统配置。后端相关模块位于 `rbac`、`campus`、`warehouse`、`material`、`supplier` 和 `config` 目录，前端对应 `UserView`、`DeptView`、`CampusView`、`WarehouseView`、`MaterialView` 等页面。保存操作通过 MyBatis-Plus 完成实体持久化，删除操作采用逻辑删除字段。")
    add_heading(doc, "6.4 库存批次与出入库实现", 2)
    add_para(doc, "库存模块由 `InventoryService` 实现。入库时系统先创建 `stock_in` 主单和 `stock_in_item` 明细，再增加 `inventory` 汇总库存，并创建 `inventory_batch` 批次记录。出库时系统先校验汇总库存和锁定量，再按 `expire_date ASC, id ASC` 查询候选批次并扣减余量，最后更新汇总库存、出库明细和预警状态。")
    add_picture(doc, SCREEN_DIR / "fig_6_4_inventory.png", "图6-3 库存查询界面")
    add_para(doc, "申领出库与普通出库的差别在于，申领出库需要检查申领单是否已锁定到当前仓库，并在出库后回写 `apply_order_item.actual_qty`。这些操作位于同一事务中，避免库存汇总、批次余量和申领单明细不一致。")
    add_heading(doc, "6.5 申领审批与调拨实现", 2)
    add_para(doc, "申领服务由 `ApplyService` 实现。部门用户创建申领单时生成 DRAFT 状态；提交时系统聚合申领明细，选择可满足全部物资需求的仓库并锁定库存。普通申领进入 SUBMITTED 状态等待审批，紧急等级达到阈值时走快速通道并记录自动审批信息。审批通过后进入 APPROVED 状态，仓库管理员执行出库后进入 OUTBOUND，部门用户签收后进入 RECEIVED。")
    add_picture(doc, SCREEN_DIR / "fig_6_3_apply.png", "图6-4 部门用户申领界面")
    add_para(doc, "调拨服务由 `TransferService` 实现。调拨单创建后处于 DRAFT 状态，提交后由审批人员审核。执行调拨时，系统按 FEFO 顺序扣减来源仓批次，并在目标仓生成带有调拨单号后缀的镜像批次，随后分别更新来源仓和目标仓的库存汇总。")
    add_picture(doc, FIG_DIR / "fig_4_2_transfer_recommend_flow.png", "图6-5 调拨执行与候选仓排序流程图")
    add_picture(doc, SCREEN_DIR / "fig_6_6_transfer.png", "图6-6 调拨管理界面")
    add_heading(doc, "6.6 预警、补货建议与统计分析实现", 2)
    add_para(doc, "预警扫描由 `WarningService.scan` 实现，包含低库存、积压、临期、过期和异常出库五类规则。低库存规则判断当前库存是否低于安全库存；积压规则判断当前库存是否超过安全库存 3 倍；临期规则查找 30 天内到期且仍有余量的批次；过期规则查找已超过到期日的批次；异常出库规则比较近 7 天出库量与近 30 天折算周均出库量。")
    add_picture(doc, SCREEN_DIR / "fig_6_7_warning.png", "图6-7 预警管理界面")
    add_para(doc, "补货建议和移动平均预测由 `SmartService` 实现。补货建议根据物资安全库存、当前库存、近 30 天出库量和保障天数计算目标库存与建议补货量；预测接口按近 6 个月出库历史聚合月度数量，并用平均值生成未来月份参考值。该实现没有训练模型，因此论文中只写为补货参考和移动平均预测。")
    add_picture(doc, SCREEN_DIR / "fig_6_9_analytics.png", "图6-8 统计分析界面")
    add_heading(doc, "6.7 日志通知与事件管理实现", 2)
    add_para(doc, "日志模块包括登录日志和操作日志。登录日志记录用户、用户名、登录 IP、登录状态、登录时间和 user-agent；操作日志记录操作者、模块、操作类型和详情。通知模块提供未读数量、标记已读、全部已读和删除功能。事件管理模块用于记录校园突发事件的标题、类型、级别、校区、地点、描述、处理人和状态。")

    add_heading(doc, "7 系统测试", 1)
    add_heading(doc, "7.1 测试环境与方法", 2)
    add_para(doc, "系统测试采用自动化测试、构建验证、端到端流程验证和本地性能基线相结合的方法。后端测试使用 Maven、JUnit 和 Spring Boot Test；前端单测使用 Vitest；端到端流程使用 Playwright 启动 screenshot profile 和前端预览服务；性能基线使用 k6 脚本记录登录、库存分页、预警分页和操作日志分页的本地结果。")
    add_heading(doc, "7.2 后端自动化测试", 2)
    add_para(doc, "2026年4月26日重新执行 `mvn test`，后端共 49 项测试通过，失败 0 项、错误 0 项、跳过 0 项。测试范围包括统一响应、业务异常、全局异常处理、登录刷新、限流、申领出库主链路、并发申领库存锁定、调拨执行、预警处理和分页接口。")
    add_table(doc, "表7-1 后端自动化测试执行情况", ["测试对象", "覆盖内容", "结果"], [
        ("common", "ApiResponse、BizException、GlobalExceptionHandler", "通过"),
        ("auth", "登录、refresh token 轮换、令牌清理、安全响应头", "通过"),
        ("apply", "申领创建、提交、审批、驳回、签收", "通过"),
        ("inventory", "申领出库、库存锁定、批次扣减", "通过"),
        ("transfer", "调拨创建、审批、执行和库存搬移", "通过"),
        ("warning", "预警处理和高风险接口限流", "通过"),
        ("integration", "核心流程与分页响应契约", "通过"),
    ], widths=[3, 8, 3])
    add_heading(doc, "7.3 前端构建、单测与 E2E 验证", 2)
    add_para(doc, "前端执行 `npm run build` 后，Vite 生产构建成功；执行 `npm run test:unit` 后，3 个测试文件、8 项测试通过；执行 `npm run test:e2e` 后，4 个端到端场景通过。E2E 场景包括管理员登录查看库存页面、部门用户创建申领单、仓库管理员创建调拨单、仓库管理员处理预警。")
    add_table(doc, "表7-2 前端验证结果", ["命令", "验证内容", "结果"], [
        ("npm run build", "生产构建和资源打包", "通过"),
        ("npm run test:unit", "HTTP 封装、登录页、状态标签组件", "3 个文件 8 项通过"),
        ("npm run test:e2e", "登录、库存、申领、调拨、预警流程", "4 个场景通过"),
    ], widths=[4, 8, 3])
    add_heading(doc, "7.4 业务场景验证", 2)
    add_para(doc, "业务场景验证主要围绕论文中的核心流程。申领审批场景验证从部门用户创建申领单到提交、审批、出库和签收的状态变化；调拨场景验证来源仓扣减、目标仓增加和批次镜像生成；预警场景验证手动扫描和处理动作；认证场景验证登录、刷新、旧 refresh token 失效和安全响应头。")
    add_table(doc, "表7-3 典型业务场景验证说明", ["场景", "验证点", "证据"], [
        ("登录认证", "access token、refresh token、安全响应头、登录日志", "CoreFlowIntegrationTest、AuthServiceTest"),
        ("申领审批", "DRAFT、SUBMITTED、APPROVED、OUTBOUND、RECEIVED 状态链路", "CoreFlowIntegrationTest、ApplyServiceTest"),
        ("库存出库", "锁定量校验、FEFO 批次扣减、实发数量回写", "InventoryService、集成测试"),
        ("调拨执行", "来源仓扣减、目标仓增加、镜像批次生成", "TransferServiceTest、E2E"),
        ("预警处理", "未处理预警处理为已处理，高风险接口限流", "WarningControllerTest、E2E"),
    ], widths=[3, 8, 4])
    add_heading(doc, "7.5 性能基线与边界说明", 2)
    add_para(doc, "仓库中提供了 `tests/performance` 下的 k6 脚本和 `docs/performance-baseline.md` 基线记录。该基线记录采集于本地 screenshot profile 和隔离 H2 数据集，场景包括登录、库存分页列表、预警分页列表和操作日志分页列表。由于该数据只适用于同机型、同脚本和同数据集的相对对比，本文不将其写成生产 SLA 或大规模并发结论。")
    add_table(doc, "表7-4 本地 k6 基线记录", ["场景", "脚本", "错误率", "p95"], [
        ("登录", "tests/performance/login.js", "0.00%", "78.20 ms"),
        ("库存分页列表", "tests/performance/inventory-list.js", "0.00%", "96.51 ms"),
        ("预警分页列表", "tests/performance/warning-list.js", "0.00%", "20.42 ms"),
        ("操作日志分页列表", "tests/performance/operation-log-list.js", "0.00%", "17.41 ms"),
    ], widths=[4, 6, 2.5, 2.5])

    add_heading(doc, "结束语", 1)
    for text in [
        "本文围绕校园物资智能管理系统完成了需求分析、总体设计、数据库设计、详细实现和系统测试。系统采用前后端分离架构，围绕校园物资管理中的基础数据、库存批次、申领审批、调拨协同、预警处理、通知消息、日志审计和统计分析建立了较完整的业务流程。",
        "从实现结果看，系统已经能够支持多角色登录和菜单边界控制，能够完成库存入库、批次维护、按效期优先出库、申领审批闭环、跨仓调拨、库存风险预警和补货参考等功能。测试验证表明，当前后端自动化测试、前端构建、前端单测和端到端流程均已通过，能够支撑毕业设计演示和论文论述。",
        "系统仍存在后续完善空间。第一，可以继续细化采购计划、供应商协同和报废处置流程；第二，可以在更充分的数据积累后尝试更精细的需求预测和异常检测；第三，可以扩展移动端或扫码核验能力；第四，可以在真实部署环境中继续完善数据库备份、权限审计和运行监控。上述内容属于未来工作，本文不将其作为已经实现的系统功能。"
    ]:
        add_para(doc, text)

    add_heading(doc, "致  谢", 1)
    add_para(doc, "在毕业设计和论文撰写过程中，指导教师在选题确定、系统实现、论文结构和格式规范方面给予了耐心指导。软件学院提供的课程训练和毕业设计要求，使我能够将软件工程方法、数据库设计、前后端开发和测试验证结合起来完成本课题。")
    add_para(doc, "同时感谢同学和家人在资料整理、系统演示和论文修改过程中的支持。通过本次毕业设计，我进一步认识到工程项目必须以真实需求和可验证证据为基础，论文写作也必须坚持事实边界、规范表达和严谨态度。")


def write_references_and_appendix(doc: Document):
    add_heading(doc, "参考文献", 1)
    for idx, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)
        r = p.add_run(f"[{idx}] {ref}")
        set_run_font(r, size=10.5)

    add_heading(doc, "附录 1 关键接口与测试补充说明", 1)
    add_para(doc, "附录内容用于补充正文中不宜展开的接口入口与测试范围说明，所有条目均来源于当前项目的控制器、服务类和测试文件。")
    add_table(doc, "附表1-1 关键业务接口补充说明", ["功能域", "主要接口", "对应模块", "说明"], [
        ("认证授权", "/api/auth/login、/api/auth/refresh、/api/auth/me、/api/auth/menus", "AuthController、AuthService", "完成登录、续签、当前用户和角色菜单加载。"),
        ("库存与批次", "/api/inventory/list、/api/inventory/batches、/api/inventory/stock-in、/api/inventory/stock-out", "InventoryController、InventoryService", "支撑库存查询、批次查询、入库和出库。"),
        ("申领审批", "/api/apply/list、/api/apply/{id}/submit、/approve、/reject、/receive", "ApplyController、ApplyService", "支撑申领单提交、审批、驳回和签收。"),
        ("调拨管理", "/api/transfer/list、/api/transfer/recommend、/submit、/approve、/execute、/receive", "TransferController、TransferService", "支撑调拨推荐、审批、执行和签收。"),
        ("预警与统计", "/api/warning/list、/api/warning/scan、/api/smart/*、/api/analytics/*", "WarningService、SmartService、AnalyticsService", "支撑预警扫描、补货参考、移动平均预测和图表统计。"),
        ("日志通知", "/api/login-log/*、/api/log/list、/api/notification/*", "Log、Notification 模块", "支撑登录日志、操作日志和通知消息。"),
    ], widths=[2.5, 4.6, 4.1, 4.2])
    add_para(doc, "附录测试范围包括后端 Maven 自动化测试、前端 Vitest 单元测试、Vite 生产构建、Playwright E2E 和 k6 脚本。对应命令为 `mvn test`、`npm run test:unit`、`npm run build`、`npm run test:e2e` 和 `k6 run tests/performance/*.js`。")


def write_report():
    report = f"""# 校园物资智能管理系统设计与实现-定稿改写说明

## 交付文件
- 终稿 DOCX：`{FINAL_DOCX}`
- 改写说明：`{REPORT}`
- 运行截图目录：`{SCREEN_DIR}`

## 本轮执行依据
- 学校模板：`Thesis Template/7 软件学院本科毕业设计（论文）模板 2026.docx`
- 撰写规范：`Thesis Template/4 软件学院本科毕业设计（论文）撰写规范.docx`
- 现有 PDF 工作稿：`tmp/docs/table-4-1-format/校园物资智能管理系统设计与实现-正式定稿版.pdf`
- 参考文献源：`C:/Users/48195/Desktop/CNKI-20260426111148676.txt`
- 项目证据：`backend/`、`frontend/`、`sql/schema.sql`、`docs/`、`tests/`

## 主要改写动作
- 未找到可编辑 DOCX 工作稿，因此按计划使用 2026 模板新建终稿 DOCX，未覆盖模板、源码、SQL 或 CNKI 文件。
- 将论文重构为“绪论、相关技术介绍、系统需求分析、系统总体设计、数据库设计、系统详细设计与实现、系统测试、结束语、致谢、参考文献、附录”。
- 删除旧稿中未列入 CNKI 文件的参考文献体系，重选 22 篇 CNKI 来源文献并按正文首次引用顺序编号。
- 将“智能”统一收口为规则化能力：库存预警、效期提醒、FEFO、候选仓排序、补货建议、移动平均预测和统计分析。
- 数据库章节按 `sql/schema.sql` 重写，未新增不存在的表、字段、外键或存储过程。
- 插入现有 drawio 黑白工程图，并重新截取登录、申领、库存、调拨、预警和统计分析页面截图。

## 验证结果
- `mvn test`：49 项后端测试通过，失败 0，错误 0，跳过 0。
- `npm run build`：前端生产构建通过。
- `npm run test:unit`：3 个测试文件、8 项测试通过。
- `npm run test:e2e`：4 个 Playwright 场景通过。
- k6 性能表使用仓库已有 `docs/performance-baseline.md`，未写成生产 SLA。

## 保守处理项
- 未写 DeepSeek、Shiro、Redis、OAuth2、RocketMQ、小程序、移动端、二维码、硬件柜等无本项目证据功能。
- 未写机器学习、神经网络、复杂 AI 调度和大规模生产运行成效。
- 观测、Sentry、Grafana 只作为工程化支撑背景，不作为论文主线展开。

## 后续人工复核
- 在 Word 中打开终稿后执行“全选 -> F9”刷新目录和页码域。
- 检查封面完成日期、指导教师职称等个人信息是否需按学院最终通知调整。
- 视觉排版已用 python-docx 生成；本机未检测到 LibreOffice/soffice，未额外导出 PDF 进行逐页渲染复核。
"""
    REPORT.write_text(report, encoding="utf-8")


if __name__ == "__main__":
    build_doc()
    print(FINAL_DOCX)
    print(REPORT)
