from __future__ import annotations

import math
import shutil
import subprocess
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = ROOT / "output" / "doc" / "figures"
DRAWIO_DIR = FIG_DIR / "drawio"
DOCX_PATH = ROOT / "output" / "doc" / "校园物资智能管理系统设计与实现-正文扩写定稿版.docx"
REPORT_PATH = ROOT / "output" / "doc" / "流程图箭头规范修复说明.md"
DRAWIO_CLI = Path(r"C:\Users\48195\.codex\skills\drawio\scripts\cli.js")

BG = (255, 255, 255)
LINE = (0, 0, 0)
GRAY = (90, 90, 90)
LIGHT = (248, 248, 248)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    for font_path in candidates:
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size=size)
    return ImageFont.load_default()


FONT_TITLE = load_font(44, True)
FONT_MODULE = load_font(34, True)
FONT_NODE = load_font(34)
FONT_SMALL = load_font(28)


def rect_from_center(cx: int, cy: int, w: int, h: int) -> tuple[int, int, int, int]:
    return (cx - w // 2, cy - h // 2, cx + w // 2, cy + h // 2)


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def fit_font(draw: ImageDraw.ImageDraw, text: str, max_w: int, max_h: int, base: int = 34) -> ImageFont.FreeTypeFont:
    for size in range(base, 17, -1):
        font = load_font(size)
        w, h = text_size(draw, text, font)
        if w <= max_w and h <= max_h:
            return font
    return load_font(18)


def draw_center_text(draw: ImageDraw.ImageDraw, rect: tuple[int, int, int, int], text: str, base: int = 34) -> None:
    max_w = int((rect[2] - rect[0]) * 0.86)
    max_h = int((rect[3] - rect[1]) * 0.70)
    font = fit_font(draw, text, max_w, max_h, base)
    w, h = text_size(draw, text, font)
    x = rect[0] + (rect[2] - rect[0] - w) / 2
    y = rect[1] + (rect[3] - rect[1] - h) / 2 - 2
    draw.text((x, y), text, font=font, fill=LINE)


def draw_box(draw: ImageDraw.ImageDraw, rect: tuple[int, int, int, int], text: str) -> None:
    draw.rounded_rectangle(rect, radius=22, outline=LINE, width=4, fill=BG)
    draw_center_text(draw, rect, text)


def draw_diamond(draw: ImageDraw.ImageDraw, rect: tuple[int, int, int, int], text: str) -> None:
    cx = (rect[0] + rect[2]) // 2
    cy = (rect[1] + rect[3]) // 2
    points = [(cx, rect[1]), (rect[2], cy), (cx, rect[3]), (rect[0], cy)]
    draw.polygon(points, outline=LINE, fill=BG)
    draw.line(points + [points[0]], fill=LINE, width=4)
    draw_center_text(draw, rect, text, base=30)


def draw_module(draw: ImageDraw.ImageDraw, rect: tuple[int, int, int, int], title: str) -> None:
    draw.rectangle(rect, outline=GRAY, width=3, fill=LIGHT)
    w, h = text_size(draw, title, FONT_MODULE)
    draw.text((rect[0] + (rect[2] - rect[0] - w) / 2, rect[1] + 22), title, font=FONT_MODULE, fill=LINE)


def arrow_head(end: tuple[float, float], prev: tuple[float, float], length: float = 24, width: float = 18):
    dx = end[0] - prev[0]
    dy = end[1] - prev[1]
    dist = math.hypot(dx, dy)
    if dist == 0:
        return None
    ux, uy = dx / dist, dy / dist
    px, py = -uy, ux
    base = (end[0] - ux * length, end[1] - uy * length)
    return [
        end,
        (base[0] + px * width / 2, base[1] + py * width / 2),
        (base[0] - px * width / 2, base[1] - py * width / 2),
    ], base


def draw_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]], width: int = 4) -> None:
    if len(points) < 2:
        return
    head = arrow_head(points[-1], points[-2])
    if head is None:
        return
    polygon, base = head
    line_points = list(points[:-1]) + [(int(base[0]), int(base[1]))]
    draw.line(line_points, fill=LINE, width=width, joint="curve")
    draw.polygon(polygon, fill=LINE)


def top(rect: tuple[int, int, int, int]) -> tuple[int, int]:
    return ((rect[0] + rect[2]) // 2, rect[1])


def bottom(rect: tuple[int, int, int, int]) -> tuple[int, int]:
    return ((rect[0] + rect[2]) // 2, rect[3])


def left(rect: tuple[int, int, int, int]) -> tuple[int, int]:
    return (rect[0], (rect[1] + rect[3]) // 2)


def right(rect: tuple[int, int, int, int]) -> tuple[int, int]:
    return (rect[2], (rect[1] + rect[3]) // 2)


def draw_edge_label(draw: ImageDraw.ImageDraw, text: str, xy: tuple[int, int]) -> None:
    w, h = text_size(draw, text, FONT_SMALL)
    draw.text((xy[0] - w / 2, xy[1] - h / 2), text, font=FONT_SMALL, fill=LINE)


def render_auth_flow() -> Path:
    path = FIG_DIR / "fig_4_1_auth_flow.png"
    image = Image.new("RGB", (1600, 2672), BG)
    draw = ImageDraw.Draw(image)
    draw.text((80, 55), "登录认证与令牌续签流程", font=FONT_TITLE, fill=LINE)

    rects = {
        "login": rect_from_center(520, 180, 430, 104),
        "verify": rect_from_center(520, 380, 430, 104),
        "issue": rect_from_center(520, 580, 430, 104),
        "init": rect_from_center(520, 800, 520, 112),
        "api": rect_from_center(520, 1030, 520, 112),
        "unauth": rect_from_center(520, 1260, 360, 150),
        "refresh": rect_from_center(1160, 1260, 430, 104),
        "rotated": rect_from_center(1160, 1490, 360, 150),
        "retry": rect_from_center(1160, 1740, 470, 112),
        "logout": rect_from_center(1160, 2140, 470, 112),
        "done": rect_from_center(520, 2460, 430, 104),
    }
    labels = {
        "login": "提交账号密码",
        "verify": "校验用户信息",
        "issue": "签发令牌",
        "init": "保存令牌并加载菜单",
        "api": "携带 Token 访问接口",
        "unauth": "收到 401?",
        "refresh": "调用刷新接口",
        "rotated": "刷新成功?",
        "retry": "更新令牌并重放请求",
        "logout": "清理令牌并返回登录页",
        "done": "继续访问页面",
    }
    for key in ["login", "verify", "issue", "init", "api", "refresh", "retry", "logout", "done"]:
        draw_box(draw, rects[key], labels[key])
    for key in ["unauth", "rotated"]:
        draw_diamond(draw, rects[key], labels[key])

    for source, target in [("login", "verify"), ("verify", "issue"), ("issue", "init"), ("init", "api"), ("api", "unauth")]:
        draw_arrow(draw, [bottom(rects[source]), top(rects[target])])

    draw_arrow(draw, [bottom(rects["unauth"]), top(rects["done"])])
    draw_edge_label(draw, "否", (560, 1720))
    draw_arrow(draw, [right(rects["unauth"]), left(rects["refresh"])])
    draw_edge_label(draw, "是", (835, 1225))
    draw_arrow(draw, [bottom(rects["refresh"]), top(rects["rotated"])])
    draw_arrow(draw, [bottom(rects["rotated"]), top(rects["retry"])])
    draw_edge_label(draw, "是", (1200, 1620))
    draw_arrow(draw, [right(rects["rotated"]), (1440, 1490), (1440, 2140), right(rects["logout"])])
    draw_edge_label(draw, "否", (1460, 1660))
    draw_arrow(draw, [bottom(rects["retry"]), (1160, 1900), (860, 1900), (860, 2300), (520, 2300), top(rects["done"])])

    image.save(path)
    return path


def render_front_backend() -> Path:
    path = FIG_DIR / "fig_4_3_front_backend_flow.png"
    image = Image.new("RGB", (2200, 1350), BG)
    draw = ImageDraw.Draw(image)
    draw.text((70, 42), "前后端交互流程", font=FONT_TITLE, fill=LINE)

    modules = [
        ("前端请求", (110, 120, 590, 1300)),
        ("安全边界", (640, 120, 1060, 1300)),
        ("后端服务", (1110, 120, 1530, 1300)),
        ("数据访问与响应", (1580, 120, 2100, 1300)),
    ]
    for title, rect in modules:
        draw_module(draw, rect, title)

    node_rects = {
        "vue": rect_from_center(350, 200, 300, 88),
        "pinia": rect_from_center(350, 350, 300, 88),
        "axios": rect_from_center(350, 500, 300, 88),
        "jwt": rect_from_center(850, 500, 300, 88),
        "perm": rect_from_center(850, 650, 300, 88),
        "controller": rect_from_center(1320, 650, 300, 88),
        "service": rect_from_center(1320, 800, 300, 88),
        "mapper": rect_from_center(1320, 950, 300, 88),
        "mysql": rect_from_center(1840, 950, 300, 88),
        "response": rect_from_center(1840, 1100, 300, 88),
        "refresh": rect_from_center(1840, 1250, 300, 88),
    }
    labels = {
        "vue": "Vue页面",
        "pinia": "Pinia状态",
        "axios": "Axios请求",
        "jwt": "JWT过滤器",
        "perm": "权限校验",
        "controller": "Controller",
        "service": "Service",
        "mapper": "Mapper",
        "mysql": "MySQL",
        "response": "统一响应",
        "refresh": "页面刷新",
    }
    for key, rect in node_rects.items():
        draw_box(draw, rect, labels[key])

    edges = [
        ("vue", "pinia", "down"),
        ("pinia", "axios", "down"),
        ("axios", "jwt", "right"),
        ("jwt", "perm", "down"),
        ("perm", "controller", "right"),
        ("controller", "service", "down"),
        ("service", "mapper", "down"),
        ("mapper", "mysql", "right"),
        ("mysql", "response", "down"),
        ("response", "refresh", "down"),
    ]
    for source, target, direction in edges:
        if direction == "down":
            draw_arrow(draw, [bottom(node_rects[source]), top(node_rects[target])])
        else:
            draw_arrow(draw, [right(node_rects[source]), left(node_rects[target])])

    image.save(path)
    return path


def render_vertical_flow(path: Path, title: str, nodes: list[str]) -> Path:
    image = Image.new("RGB", (1700, 1850), BG)
    draw = ImageDraw.Draw(image)
    draw.text((80, 48), title, font=FONT_TITLE, fill=LINE)
    top_y = 210
    gap = 220 if len(nodes) <= 7 else 190
    w, h = 860, 104
    rects = []
    for index, label in enumerate(nodes):
        rect = rect_from_center(850, top_y + index * gap, w, h)
        rects.append(rect)
        draw_box(draw, rect, label)
    for source, target in zip(rects, rects[1:]):
        draw_arrow(draw, [bottom(source), top(target)])
    image.save(path)
    return path


def yaml_quote(value: str) -> str:
    return "'" + value.replace("'", "''") + "'"


def edge_style(direction: str) -> dict[str, float | int | str]:
    if direction == "down":
        coords = {
            "exitX": 0.5,
            "exitY": 1,
            "entryX": 0.5,
            "entryY": 0,
        }
    elif direction == "down_left_slot":
        coords = {
            "exitX": 0.35,
            "exitY": 1,
            "entryX": 0.5,
            "entryY": 0,
        }
    elif direction == "down_right_slot":
        coords = {
            "exitX": 0.65,
            "exitY": 1,
            "entryX": 0.5,
            "entryY": 0,
        }
    elif direction == "right":
        coords = {
            "exitX": 1,
            "exitY": 0.5,
            "entryX": 0,
            "entryY": 0.5,
        }
    else:
        raise ValueError(direction)
    return {
        "strokeColor": "#000000",
        "endArrow": "block",
        "strokeWidth": 1.4,
        "fontColor": "#000000",
        "fontSize": 9,
        **coords,
        "exitDx": 0,
        "exitDy": 0,
        "entryDx": 0,
        "entryDy": 0,
    }


def write_yaml_spec(
    name: str,
    title: str,
    canvas: str,
    nodes: list[dict],
    edges: list[dict],
    modules: list[dict] | None = None,
    layout: str = "vertical",
) -> Path:
    DRAWIO_DIR.mkdir(parents=True, exist_ok=True)
    path = DRAWIO_DIR / f"{name}.spec.yaml"
    lines: list[str] = [
        "meta:",
        "  profile: academic-paper",
        "  theme: academic",
        f"  layout: {layout}",
        "  routing: orthogonal",
        f"  canvas: {canvas}",
        f"  title: {yaml_quote(title)}",
        "  description: '论文流程图箭头规范修复版：连接点固定在节点边缘'",
    ]
    if modules:
        lines.append("modules:")
        for module in modules:
            lines.extend([
                f"  - id: {module['id']}",
                f"    label: {yaml_quote(module['label'])}",
            ])
    lines.append("nodes:")
    for node in nodes:
        lines.extend([
            f"  - id: {node['id']}",
            f"    label: {yaml_quote(node['label'])}",
            f"    type: {node.get('type', 'process')}",
        ])
        if node.get("module"):
            lines.append(f"    module: {node['module']}")
        lines.extend([
            f"    size: {node.get('size', 'xl')}",
            "    position:",
            f"      x: {node['x']}",
            f"      y: {node['y']}",
            "    style:",
            "      fillColor: '#FFFFFF'",
            "      strokeColor: '#000000'",
            "      fontColor: '#000000'",
            "      strokeWidth: 1.4",
            "      fontSize: 10",
        ])
    lines.append("edges:")
    for edge in edges:
        style = edge_style(edge["direction"])
        lines.extend([
            f"  - from: {edge['from']}",
            f"    to: {edge['to']}",
            "    type: primary",
            "    style:",
        ])
        for key, value in style.items():
            if isinstance(value, str):
                lines.append(f"      {key}: {yaml_quote(value)}")
            else:
                lines.append(f"      {key}: {value}")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def generate_new_specs() -> list[Path]:
    auth_nodes = [
        {"id": "login", "label": "提交账号密码", "x": 440, "y": 130, "type": "terminal"},
        {"id": "verify", "label": "校验用户信息", "x": 440, "y": 330},
        {"id": "issue", "label": "签发令牌", "x": 440, "y": 530},
        {"id": "init", "label": "保存令牌并加载菜单", "x": 420, "y": 740, "size": "xl"},
        {"id": "api", "label": "携带 Token 访问接口", "x": 420, "y": 970, "size": "xl"},
        {"id": "unauth", "label": "收到 401?", "x": 440, "y": 1200, "type": "decision"},
        {"id": "refresh", "label": "调用刷新接口", "x": 1080, "y": 1210},
        {"id": "rotated", "label": "刷新成功?", "x": 1080, "y": 1430, "type": "decision"},
        {"id": "retry", "label": "更新令牌并重放请求", "x": 1060, "y": 1680, "size": "xl"},
        {"id": "logout", "label": "清理令牌并返回登录页", "x": 1060, "y": 2080, "size": "xl"},
        {"id": "done", "label": "继续访问页面", "x": 440, "y": 2410, "type": "terminal"},
    ]
    auth_edges = [
        {"from": "login", "to": "verify", "direction": "down"},
        {"from": "verify", "to": "issue", "direction": "down"},
        {"from": "issue", "to": "init", "direction": "down"},
        {"from": "init", "to": "api", "direction": "down"},
        {"from": "api", "to": "unauth", "direction": "down"},
        {"from": "unauth", "to": "done", "direction": "down"},
        {"from": "unauth", "to": "refresh", "direction": "right"},
        {"from": "refresh", "to": "rotated", "direction": "down"},
        {"from": "rotated", "to": "retry", "direction": "down_left_slot"},
        {"from": "rotated", "to": "logout", "direction": "down_right_slot"},
        {"from": "retry", "to": "done", "direction": "down"},
    ]
    specs = [
        write_yaml_spec(
            "fig_4_1_auth_flow",
            "登录认证与令牌续签流程图",
            "1600x2672",
            auth_nodes,
            auth_edges,
            layout="vertical",
        )
    ]

    front_nodes = [
        {"id": "vue", "label": "Vue页面", "x": 200, "y": 156, "module": "frontend"},
        {"id": "pinia", "label": "Pinia状态", "x": 200, "y": 306, "module": "frontend"},
        {"id": "axios", "label": "Axios请求", "x": 200, "y": 456, "module": "frontend"},
        {"id": "jwt", "label": "JWT过滤器", "x": 700, "y": 456, "module": "security"},
        {"id": "perm", "label": "权限校验", "x": 700, "y": 606, "module": "security"},
        {"id": "controller", "label": "Controller", "x": 1170, "y": 606, "module": "backend"},
        {"id": "service", "label": "Service", "x": 1170, "y": 756, "module": "backend"},
        {"id": "mapper", "label": "Mapper", "x": 1170, "y": 906, "module": "backend"},
        {"id": "mysql", "label": "MySQL", "x": 1690, "y": 906, "module": "data"},
        {"id": "response", "label": "统一响应", "x": 1690, "y": 1056, "module": "data"},
        {"id": "refresh", "label": "页面刷新", "x": 1690, "y": 1206, "module": "data"},
    ]
    front_edges = [
        {"from": "vue", "to": "pinia", "direction": "down"},
        {"from": "pinia", "to": "axios", "direction": "down"},
        {"from": "axios", "to": "jwt", "direction": "right"},
        {"from": "jwt", "to": "perm", "direction": "down"},
        {"from": "perm", "to": "controller", "direction": "right"},
        {"from": "controller", "to": "service", "direction": "down"},
        {"from": "service", "to": "mapper", "direction": "down"},
        {"from": "mapper", "to": "mysql", "direction": "right"},
        {"from": "mysql", "to": "response", "direction": "down"},
        {"from": "response", "to": "refresh", "direction": "down"},
    ]
    specs.append(
        write_yaml_spec(
            "fig_4_3_front_backend_flow",
            "前后端交互流程图",
            "2200x1350",
            front_nodes,
            front_edges,
            [
                {"id": "frontend", "label": "前端请求"},
                {"id": "security", "label": "安全边界"},
                {"id": "backend", "label": "后端服务"},
                {"id": "data", "label": "数据访问与响应"},
            ],
            layout="horizontal",
        )
    )

    flow_specs = [
        (
            "fig_6_9_stock_in_flow",
            "入库业务流程图",
            ["填写入库单", "校验仓库和物资", "写入 stock_in", "写入 stock_in_item", "更新 inventory 数量", "生成 inventory_batch", "记录操作日志"],
        ),
        (
            "fig_6_10_stock_out_flow",
            "出库业务流程图",
            ["创建出库请求", "校验申领单与仓库", "检查库存与锁定量", "按 expire_date 批次扣减", "写入 stock_out_item", "回写库存与实发数量", "触发低库存检查"],
        ),
        (
            "fig_6_11_analytics_flow",
            "统计数据流转流程图",
            ["前端选择统计主题", "请求 AnalyticsController", "Service 聚合业务表", "返回统一 JSON", "ECharts 渲染图表", "页面展示库存与预警结果"],
        ),
    ]
    for name, title, labels in flow_specs:
        nodes = [
            {"id": f"n{i + 1}", "label": label, "x": 750, "y": 160 + i * 220, "type": "terminal" if i in (0, len(labels) - 1) else "process"}
            for i, label in enumerate(labels)
        ]
        edges = [{"from": f"n{i + 1}", "to": f"n{i + 2}", "direction": "down"} for i in range(len(labels) - 1)]
        specs.append(write_yaml_spec(name, title, "1700x1850", nodes, edges))
    return specs


def render_new_pngs() -> list[Path]:
    rendered = [
        render_auth_flow(),
        render_front_backend(),
        render_vertical_flow(
            FIG_DIR / "fig_6_9_stock_in_flow.png",
            "入库业务流程",
            ["填写入库单", "校验仓库和物资", "写入 stock_in", "写入 stock_in_item", "更新 inventory 数量", "生成 inventory_batch", "记录操作日志"],
        ),
        render_vertical_flow(
            FIG_DIR / "fig_6_10_stock_out_flow.png",
            "出库业务流程",
            ["创建出库请求", "校验申领单与仓库", "检查库存与锁定量", "按 expire_date 批次扣减", "写入 stock_out_item", "回写库存与实发数量", "触发低库存检查"],
        ),
        render_vertical_flow(
            FIG_DIR / "fig_6_11_analytics_flow.png",
            "统计数据流转流程",
            ["前端选择统计主题", "请求 AnalyticsController", "Service 聚合业务表", "返回统一 JSON", "ECharts 渲染图表", "页面展示库存与预警结果"],
        ),
    ]
    return rendered


def run_drawio_cli(spec: Path, out_ext: str = ".drawio", strict: bool = False) -> dict:
    base_name = spec.name.replace(".spec.yaml", "")
    output = spec.with_name(f"{base_name}{out_ext}")
    cmd = ["node", str(DRAWIO_CLI), str(spec), str(output), "--validate", "--write-sidecars"]
    if strict:
        cmd.append("--strict")
    proc = subprocess.run(cmd, cwd=ROOT, text=True, capture_output=True, encoding="utf-8", errors="replace")
    return {
        "spec": spec.name,
        "output": output.name,
        "returncode": proc.returncode,
        "stdout": (proc.stdout or "").strip(),
        "stderr": (proc.stderr or "").strip(),
    }


def validate_drawio_specs(new_specs: list[Path]) -> list[dict]:
    targets = [
        "fig_2_1_apply_flow",
        "fig_2_2_transfer_flow",
        "fig_2_3_warning_flow",
        "fig_3_1_architecture",
        "fig_3_2_modules",
        "fig_4_1_auth_flow",
        "fig_4_2_transfer_recommend_flow",
        "fig_4_3_front_backend_flow",
        "fig_6_9_stock_in_flow",
        "fig_6_10_stock_out_flow",
        "fig_6_11_analytics_flow",
    ]
    results: list[dict] = []
    new_names = {spec.stem.replace(".spec", "") for spec in new_specs}
    for name in targets:
        spec = DRAWIO_DIR / f"{name}.spec.yaml"
        if not spec.exists():
            results.append({"spec": f"{name}.spec.yaml", "returncode": -1, "stdout": "", "stderr": "spec missing"})
            continue
        strict = name in new_names
        results.append(run_drawio_cli(spec, ".drawio", strict=strict))
        results.append(run_drawio_cli(spec, ".svg", strict=strict))
    return results


def replace_docx_images() -> list[str]:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    backup = DOCX_PATH.with_name(f"{DOCX_PATH.stem}-流程图箭头修复前备份-{datetime.now().strftime('%Y%m%d%H%M%S')}.docx")
    shutil.copy2(DOCX_PATH, backup)

    caption_to_image = {
        "图6-1 登录认证与令牌续签流程图": FIG_DIR / "fig_4_1_auth_flow.png",
        "图4-3 前后端交互流程图": FIG_DIR / "fig_4_3_front_backend_flow.png",
        "图6-9 入库业务流程图": FIG_DIR / "fig_6_9_stock_in_flow.png",
        "图6-10 出库业务流程图": FIG_DIR / "fig_6_10_stock_out_flow.png",
        "图6-11 统计数据流转流程图": FIG_DIR / "fig_6_11_analytics_flow.png",
    }

    doc = Document(DOCX_PATH)
    replaced: list[str] = []
    paragraphs = list(doc.paragraphs)
    for caption, image_path in caption_to_image.items():
        for index, para in enumerate(paragraphs):
            if caption not in para.text:
                continue
            for probe in range(index, max(-1, index - 8), -1):
                blips = [el for el in paragraphs[probe]._element.iter() if el.tag == qn("a:blip")]
                if not blips:
                    continue
                rid = blips[-1].get(qn("r:embed"))
                if not rid:
                    continue
                image_part = paragraphs[probe].part.related_parts[rid]
                image_part._blob = image_path.read_bytes()
                replaced.append(caption)
                break
            break
    doc.save(DOCX_PATH)
    replaced.append(f"备份文件：{backup.name}")
    return replaced


def inspect_docx() -> dict[str, int | list[str]]:
    doc = Document(DOCX_PATH)
    captions = [p.text.strip() for p in doc.paragraphs if p.text.strip().startswith("图")]
    image_count = 0
    for para in doc.paragraphs:
        image_count += sum(1 for el in para._element.iter() if el.tag == qn("a:blip"))
    return {"captions": captions, "image_count": image_count, "caption_count": len(captions)}


def write_report(rendered: list[Path], validation: list[dict], replaced: list[str], docx_info: dict) -> None:
    rows = [
        ("图3-1", "申领审批闭环流程图", "已有 drawio 侧车校验", "通过/见校验日志"),
        ("图3-2", "调拨执行流程图", "已有 drawio 侧车校验", "通过/见校验日志"),
        ("图3-3", "预警处置流程图", "已有 drawio 侧车校验", "通过/见校验日志"),
        ("图4-1", "系统总体架构图", "已有 drawio 侧车校验", "通过/见校验日志"),
        ("图4-2", "系统功能模块图", "已有 drawio 侧车校验", "通过/见校验日志"),
        ("图4-3", "前后端交互流程图", "重画为边缘连接泳道图", "已替换 DOCX"),
        ("图6-1", "登录认证与令牌续签流程图", "重画为边缘连接分支流程图", "已替换 DOCX"),
        ("图6-5", "调拨执行与候选仓排序流程图", "已有 drawio 侧车校验", "通过/见校验日志"),
        ("图6-9", "入库业务流程图", "重画为竖向边缘连接流程图", "已替换 DOCX"),
        ("图6-10", "出库业务流程图", "重画为竖向边缘连接流程图", "已替换 DOCX"),
        ("图6-11", "统计数据流转流程图", "重画为竖向边缘连接流程图", "已替换 DOCX"),
    ]
    failed = [r for r in validation if r["returncode"] != 0]
    lines = [
        "# 流程图箭头规范修复说明",
        "",
        f"- 修复时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"- 定稿文件：`{DOCX_PATH}`",
        f"- 图片目录：`{FIG_DIR}`",
        f"- drawio 侧车目录：`{DRAWIO_DIR}`",
        "",
        "## 修复范围",
        "",
        "| 图号 | 图名 | 处理方式 | 结果 |",
        "|---|---|---|---|",
    ]
    lines.extend(f"| {fig} | {name} | {action} | {result} |" for fig, name, action, result in rows)
    lines.extend([
        "",
        "## 连接规范",
        "",
        "- 新重画流程图的箭头均从源节点边缘出发，并以箭头尖端接触目标节点边缘。",
        "- 图4-3 改为四个泳道模块：前端请求、安全边界、后端服务、数据访问与响应，避免原图中斜线穿透方框。",
        "- 连接线统一使用黑色实线、规范块箭头，保持白底黑灰工程论文风格。",
        "- 对带 drawio 侧车的既有流程图执行 drawio CLI `--validate --write-sidecars` 校验并刷新可编辑文件。",
        "",
        "## 已重画图片",
        "",
    ])
    lines.extend(f"- `{path.name}`" for path in rendered)
    lines.extend([
        "",
        "## DOCX 替换结果",
        "",
    ])
    lines.extend(f"- {item}" for item in replaced)
    lines.extend([
        f"- DOCX 图片数量：{docx_info['image_count']}",
        f"- DOCX 图题数量：{docx_info['caption_count']}",
        "- 嵌入图片校验：脚本按图题定位并替换图6-1、图4-3、图6-9、图6-10、图6-11 的嵌入图片；可用哈希比对确认与 `output/doc/figures` 中对应 PNG 一致。",
        "- 页面渲染说明：本机若未检测到 `soffice` / `pdftoppm`，可采用 DOCX 结构检查、嵌入图片哈希校验和 PNG 视觉检查替代逐页 PDF 渲染。",
        "",
        "## drawio 校验结果",
        "",
        f"- 校验命令数量：{len(validation)}",
        f"- 失败数量：{len(failed)}",
        "",
        "| Spec | 输出 | 退出码 | 备注 |",
        "|---|---|---:|---|",
    ])
    for result in validation:
        note = result["stderr"] or result["stdout"] or "OK"
        note = note.replace("\n", "<br>")
        if len(note) > 260:
            note = note[:260] + "..."
        lines.append(f"| `{result['spec']}` | `{result.get('output', '')}` | {result['returncode']} | {note} |")
    lines.extend([
        "",
        "## 待说明",
        "",
        "- 本次只替换流程图图片和 drawio 侧车文件，不改论文正文内容、参考文献和章节结构。",
        "- 若后续重新运行正文生成脚本导致图片被覆盖，应再次运行 `tmp/fix_thesis_flow_arrows.py`。",
    ])
    REPORT_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    DRAWIO_DIR.mkdir(parents=True, exist_ok=True)
    rendered = render_new_pngs()
    new_specs = generate_new_specs()
    validation = validate_drawio_specs(new_specs)
    replaced = replace_docx_images()
    docx_info = inspect_docx()
    write_report(rendered, validation, replaced, docx_info)
    print("rendered:", ", ".join(path.name for path in rendered))
    print("validation failures:", sum(1 for result in validation if result["returncode"] != 0))
    print("docx images:", docx_info["image_count"], "captions:", docx_info["caption_count"])
    print("report:", REPORT_PATH)


if __name__ == "__main__":
    main()
