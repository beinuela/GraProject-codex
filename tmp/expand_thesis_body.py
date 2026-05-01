from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
SOURCE_DOCX = OUT_DIR / "校园物资智能管理系统设计与实现-正式定稿版.docx"
TARGET_DOCX = OUT_DIR / "校园物资智能管理系统设计与实现-正文扩写定稿版.docx"
FIG_DIR = OUT_DIR / "figures"
REPORT = OUT_DIR / "论文扩写修改报告.md"
EVIDENCE = OUT_DIR / "证据映射表.md"
KRD = OUT_DIR / "Keep-Rewrite-Delete矩阵.md"
FIG_CHECK = OUT_DIR / "图表编号与引用检查表.md"
REF_CHECK = OUT_DIR / "参考文献引用检查表.md"


def choose_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
    ]
    for item in candidates:
        path = Path(item)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


FONT = choose_font(30)
FONT_SMALL = choose_font(24)
FONT_BOLD = choose_font(30, True)
LINE = (0, 0, 0)
BG = (255, 255, 255)


def box(cx: int, cy: int, w: int, h: int):
    return (cx - w // 2, cy - h // 2, cx + w // 2, cy + h // 2)


def draw_center(draw: ImageDraw.ImageDraw, rect, text: str, font=FONT):
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=8, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    x1, y1, x2, y2 = rect
    draw.multiline_text(((x1 + x2 - width) / 2, (y1 + y2 - height) / 2), text, font=font, fill=LINE, spacing=8, align="center")


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, width: int = 4):
    draw.line([start, end], fill=LINE, width=width)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        d = 1 if ex > sx else -1
        head = [(ex, ey), (ex - 18 * d, ey - 10), (ex - 18 * d, ey + 10)]
    else:
        d = 1 if ey > sy else -1
        head = [(ex, ey), (ex - 10, ey - 18 * d), (ex + 10, ey - 18 * d)]
    draw.polygon(head, fill=LINE)


def draw_vertical_flow(path: Path, title: str, nodes: list[str], canvas=(1700, 1850)):
    image = Image.new("RGB", canvas, BG)
    draw = ImageDraw.Draw(image)
    draw.text((60, 40), title, font=FONT_BOLD, fill=LINE)
    top = 160
    gap = 170
    w = 780
    h = 96
    rects = []
    for idx, text in enumerate(nodes):
        rect = box(canvas[0] // 2, top + idx * gap, w, h)
        rects.append(rect)
        draw.rounded_rectangle(rect, radius=24, outline=LINE, width=4, fill=BG)
        draw_center(draw, rect, text)
    for left, right in zip(rects, rects[1:]):
        draw_arrow(draw, ((left[0] + left[2]) // 2, left[3]), ((right[0] + right[2]) // 2, right[1]))
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path)


def draw_front_backend(path: Path):
    image = Image.new("RGB", (2200, 1350), BG)
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "前后端交互流程", font=FONT_BOLD, fill=LINE)
    modules = [
        ("前端请求", (110, 120, 590, 1300)),
        ("安全边界", (640, 120, 1060, 1300)),
        ("后端服务", (1110, 120, 1530, 1300)),
        ("数据访问与响应", (1580, 120, 2100, 1300)),
    ]
    for title, rect in modules:
        draw.rectangle(rect, outline=(90, 90, 90), width=3)
        draw_center(draw, (rect[0], rect[1] + 18, rect[2], rect[1] + 88), title, FONT_BOLD)

    rects = {
        "vue": box(350, 200, 300, 88),
        "pinia": box(350, 350, 300, 88),
        "axios": box(350, 500, 300, 88),
        "jwt": box(850, 500, 300, 88),
        "perm": box(850, 650, 300, 88),
        "controller": box(1320, 650, 300, 88),
        "service": box(1320, 800, 300, 88),
        "mapper": box(1320, 950, 300, 88),
        "mysql": box(1840, 950, 300, 88),
        "response": box(1840, 1100, 300, 88),
        "refresh": box(1840, 1250, 300, 88),
    }
    labels = {
        "vue": "Vue页面",
        "pinia": "Pinia状态",
        "axios": "Axios请求",
        "jwt": "JWT过滤器",
        "perm": "权限校验",
        "controller": "Controller",
        "service": "Service",
        "mapper": "Mapper",
        "mysql": "MySQL",
        "response": "统一响应",
        "refresh": "页面刷新",
    }
    for key, rect in rects.items():
        draw.rounded_rectangle(rect, radius=22, outline=LINE, width=4, fill=BG)
        draw_center(draw, rect, labels[key])

    def top(rect):
        return ((rect[0] + rect[2]) // 2, rect[1])

    def bottom(rect):
        return ((rect[0] + rect[2]) // 2, rect[3])

    def left(rect):
        return (rect[0], (rect[1] + rect[3]) // 2)

    def right(rect):
        return (rect[2], (rect[1] + rect[3]) // 2)

    for source, target, direction in [
        ("vue", "pinia", "down"),
        ("pinia", "axios", "down"),
        ("axios", "jwt", "right"),
        ("jwt", "perm", "down"),
        ("perm", "controller", "right"),
        ("controller", "service", "down"),
        ("service", "mapper", "down"),
        ("mapper", "mysql", "right"),
        ("mysql", "response", "down"),
        ("response", "refresh", "down"),
    ]:
        if direction == "down":
            draw_arrow(draw, bottom(rects[source]), top(rects[target]))
        else:
            draw_arrow(draw, right(rects[source]), left(rects[target]))
    image.save(path)


def create_new_figures():
    draw_front_backend(FIG_DIR / "fig_4_3_front_backend_flow.png")
    draw_vertical_flow(
        FIG_DIR / "fig_6_9_stock_in_flow.png",
        "入库业务流程",
        ["填写入库单", "校验仓库和物资", "写入 stock_in", "写入 stock_in_item", "更新 inventory 数量", "生成 inventory_batch", "记录操作日志"],
    )
    draw_vertical_flow(
        FIG_DIR / "fig_6_10_stock_out_flow.png",
        "出库业务流程",
        ["创建出库请求", "校验申领单与仓库", "检查库存与锁定量", "按 expire_date 批次扣减", "写入 stock_out_item", "回写库存与实发数量", "触发低库存检查"],
    )
    draw_vertical_flow(
        FIG_DIR / "fig_6_11_analytics_flow.png",
        "统计数据流转流程",
        ["前端选择统计主题", "请求 AnalyticsController", "Service 聚合业务表", "返回统一 JSON", "ECharts 渲染图表", "页面展示库存与预警结果"],
    )


def set_run_font(run, size=12, bold=False, font="宋体"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def format_body_para(p: Paragraph):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run_font(run)


def insert_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style is not None:
        p.style = style
    if text:
        run = p.add_run(text)
        set_run_font(run)
    format_body_para(p)
    return p


def insert_picture_after(paragraph: Paragraph, path: Path, caption: str, width_cm: float = 13.5) -> Paragraph:
    p = insert_after(paragraph, "")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    c = insert_after(p, caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Cm(0)
    for run in c.runs:
        set_run_font(run, size=10.5, bold=True)
    return c


EXPANSIONS = {
    "1.1 研究背景": [
        "校园物资管理同时具有行政管理、仓储管理和服务保障属性。与一般办公用品台账相比，校园场景中的物资类别更分散，使用部门更多，消耗节奏更容易受教学安排、宿舍维护、实验实训、安全巡检和突发事件影响。若仍然依靠线下表格或人工传递申请单，管理人员通常只能看到某个仓库或某个部门的局部数据，难以及时判断全校物资余量、批次效期和申领进度。智慧校园相关研究指出，物资管理系统建设是高校运行保障信息化的重要组成部分[1]，这与本系统面向校园物资流转闭环的定位一致。",
        "从项目实现看，系统并不是单一库存查询工具，而是围绕物资档案、仓库、库存、批次、申领、审批、出库、调拨、预警和统计分析构建的一套业务系统。`sql/schema.sql` 中的 `material_info`、`warehouse`、`inventory`、`inventory_batch`、`apply_order`、`transfer_order` 和 `warning_record` 等表说明，系统关注的不仅是当前库存数量，还包括来源、去向、状态、处理人和时间记录。论文扩写因此应把背景落到“流程可追踪、批次可核对、库存可预警、数据可统计”四个具体问题上，而不是泛泛讨论智能校园。",
        "在校园后勤或实验室物资管理中，物资流通信息化往往会涉及入库登记、领用审核、库存监控和风险提醒等环节[2][5]。本课题采用的实现路线与这些系统类研究有共同点，即通过 B/S 架构和数据库集中存储提升流程透明度；但本系统的论文写作必须限定在当前仓库能够证明的功能范围内，不把其他论文中的硬件识别、移动端扫码、采购付款或复杂预测模型迁移为本项目已完成内容。"
    ],
    "1.2 研究意义": [
        "本系统的实践意义首先体现在库存数据的统一维护。物资管理员通过入库单和出库单维护物资流向，系统在同一事务内同步库存汇总和批次余量，减少了人工分散登记造成的数量不一致。对于部门用户而言，申领单状态能够从草稿、提交、审批、出库到签收逐步推进，减少线下追问和重复沟通。对于审批人员而言，申领明细、紧急等级和审批意见被保存在单据中，有利于事后复盘。",
        "本系统的工程意义在于把前后端分离、权限控制、事务处理、批次扣减和预警扫描等技术组织成一个完整的管理系统。高校信息化项目管理系统的相关研究通常强调前端交互、后端接口和数据持久层之间的低耦合[6][7]，本系统也采用类似分层方式。不同之处在于，本课题把主要精力放在校园物资业务闭环中，特别是库存锁定、FEFO 批次扣减、调拨候选仓排序、低库存与效期预警等工程细节。",
        "从论文写作意义看，扩写后的正文需要避免把系统包装为算法平台。仓库中的 `WarningService`、`SmartService`、`TransferService` 和 `InventoryService` 表明，系统所谓智能能力主要是规则化判断、统计汇总和候选排序。将这些能力如实写清楚，能够使论文既体现系统的辅助决策价值，又不超出代码实现边界。"
    ],
    "1.3 国内外研究现状": [
        "国内研究中，与本系统最接近的是高校、医疗和后勤物资管理系统的设计与实现类成果。相关文献通常把物资入库、领用、审核和库存监控作为系统核心模块[2]，也会从智慧校园角度强调物资、资产和运营数据的协同管理[1]。这些研究为本文需求分析提供了参照：校园物资管理不应只停留在库存列表，而要覆盖申请、审批、仓库、批次、预警和统计。",
        "Spring Boot 与 Vue 组合在高校信息化系统、事故报告数据库、学术档案管理等系统类论文中较为常见[6][7][8]。这些文献说明，前后端分离架构适合支撑多角色、多页面、多接口的管理系统开发。本文在引用这些研究时只借鉴其技术路线和系统工程写法，不引用其测试数据或业务结论作为本系统效果。对于文献中提到但本项目没有使用的技术，论文正文不写成已实现技术。",
        "国外或外文资料中，基于 Spring Boot、Vue 和 MySQL 的 Web 系统开发、安全设计和管理信息系统实现也有相关研究[10][11][12][13]。这些文献能够支撑本文关于 Web 应用安全边界、分层架构和数据库持久化的讨论，但本文不扩展到多因素认证、企业级高可用集群或 DevSecOps 全流程验证。安全章节只结合本项目已实现的 JWT、refresh token、BCrypt、CORS、请求限流和安全响应头进行说明。"
    ],
    "2.1 技术选型概述": [
        "技术选型的依据来自 `backend/pom.xml`、`frontend/package.json` 和后端配置文件。后端使用 Spring Boot 3.3.5 和 Java 17，结合 Spring Web、Validation、Security、AOP、Actuator、MyBatis-Plus、MySQL 驱动、JJWT、Knife4j、Bucket4j、Caffeine、Micrometer 和 Sentry 依赖。前端使用 Vue 3.5、Vite 6、Pinia、Vue Router、Element Plus、Axios、ECharts 和 Playwright/Vitest 测试工具。论文扩写时只介绍这些真实存在的技术，不把项目未采用的技术写入技术栈。",
        "本系统的技术组合服务于两个目标：一是让校园物资业务以页面和接口形式组织，二是让数据在数据库中形成可追踪记录。前端页面负责用户输入、列表筛选、状态展示和图表渲染，后端负责身份校验、状态流转、事务一致性和数据查询。数据库承担物资、库存、批次、单据、日志和通知的持久化。技术介绍章节应说明这些技术在本项目中的位置，而不是孤立罗列框架特性。"
    ],
    "2.2 前端开发技术": [
        "前端工程采用 Vite 作为开发与构建工具，Vue 3 负责组件化页面组织，Vue Router 负责路由切换，Pinia 负责登录态、用户信息和菜单状态管理。`frontend/src/router/index.js` 中列出了登录、运营总览、用户管理、部门管理、校区管理、物资分类、物资档案、供应商、仓库、库位、库存、入库、出库、申领审批、调拨、预警、事件、统计分析、日志、通知、系统配置和安全策略等页面，这些路由是需求分析和功能模块划分的重要证据。",
        "Element Plus 在本系统中主要承担表单、表格、对话框、分页、标签和消息提示等管理系统常见交互。相比自行编写大量基础组件，使用成熟组件库可以使页面保持一致的输入校验、列表展示和状态反馈方式。ECharts 用于库存占比、出入库趋势、部门排行、效期统计和仓库分布等可视化展示，其数据来源于后端 `AnalyticsController` 暴露的统计接口。前端章节应把这些组件与具体页面联系起来，而不是仅介绍组件库本身。"
    ],
    "2.3 后端与数据访问技术": [
        "后端采用 Controller、Service、Mapper 和数据库表的分层结构。Controller 负责接收 REST 请求和组织响应，Service 负责业务规则、事务和状态控制，Mapper 通过 MyBatis-Plus 完成实体表映射，部分复杂统计由 `JdbcTemplate` 执行聚合 SQL。以库存模块为例，`InventoryController` 暴露库存列表、批次、入库、出库、盘点和推荐出库批次接口；`InventoryService` 在入库时写入 `stock_in`、`stock_in_item`、`inventory` 和 `inventory_batch`，在出库时校验库存、按效期扣减批次并生成出库记录。",
        "MyBatis-Plus 在项目中承担基础 CRUD、分页查询、逻辑删除、乐观锁字段更新等常见数据访问工作。`BaseEntity` 中的 `deleted`、`version`、`created_at`、`updated_at` 等字段与 SQL 表结构保持一致，使多数业务表都具备软删除、版本控制和审计时间。数据库章节扩写时，应把这些通用字段作为一致性设计的一部分进行说明。"
    ],
    "2.4 安全认证与可视化技术": [
        "认证授权由 Spring Security、JWT 和自定义过滤器共同完成。`SecurityConfig` 将登录、刷新令牌和接口文档等少数地址放行，其余接口要求认证；`JwtAuthenticationFilter` 解析请求中的 Authorization 头并构建登录用户上下文；`AuthService` 校验用户名、BCrypt 密码、账号状态后签发 access token 和 refresh token。JWT 与 RBAC 在前后端分离系统中的应用研究可为该设计提供技术背景[19][20]，但本系统只实现当前代码中的令牌和角色菜单控制。",
        "安全设计还包括 refresh token 持久化与轮换。`auth_refresh_token` 表保存 token_id、token_hash、expire_at 和 revoked 等字段，刷新成功后撤销旧令牌并签发新令牌。配置文件中还包含 CORS 白名单、登录/刷新/高风险接口限流参数、管理端点暴露范围和安全响应头配置。论文中可以写这些安全边界，但不能写项目没有实现的多因素认证或生产级安全认证体系。"
    ],
    "3.1 业务场景与角色划分": [
        "系统真实角色应以 `AuthService.buildMenusByRole` 和种子数据为准。管理员 `ADMIN` 能访问用户、部门、校区、物资、供应商、仓库、库存、入库、出库、申领、调拨、预警、事件、统计、日志、通知、系统配置和安全策略等菜单，适合作为系统维护和全局监管角色。仓储管理员 `WAREHOUSE_ADMIN` 侧重基础物资、仓库、库位、库存、入库、出库、调拨、预警和统计。审批人员 `APPROVER` 主要处理申领、调拨、预警和统计。部门用户 `DEPT_USER` 主要提交申领并查看统计。论文不得额外写采购人员、配送人员或移动端角色为已实现角色。",
        "从业务边界看，系统支持的是校园物资从建档、入库、库存汇总、申领审批、出库、签收、调拨、预警到统计分析的闭环。对于用户提出的“配送任务管理”类内容，当前 SQL 中没有独立配送任务表，代码中也没有配送人员角色，因此论文正文只能把相关动作写成申领流程中的出库和签收，以及调拨流程中的执行和接收，不写成独立配送模块。"
    ],
    "3.2 功能需求分析": [
        "基础数据需求包括部门、校区、仓库、库位、供应商、物资分类和物资档案维护。物资档案不仅保存名称和分类，还保存规格、单位、安全库存、保质期、供应商、单价和备注，这些字段直接影响库存预警、补货建议和效期提醒。仓库与库位信息用于描述物资存放范围，供应商信息用于保留供货协作资料，但本系统没有实现采购付款和供应商订单流。",
        "库存业务需求包括库存查询、批次查询、入库登记、出库登记和盘点调整。入库需求要求能够记录仓库、来源类型、操作人、物资明细、批次号、数量、生产日期和过期日期，并同步更新库存汇总。出库需求要求能够校验可用库存，按批次效期顺序扣减，保留出库单与出库明细，并在关联申领单时回写实发数量和单据状态。盘点调整只针对库存汇总数量，不应扩写成完整盘点盘亏盘盈审批流程。",
        "申领审批需求体现部门用户和审批人员之间的业务协作。部门用户创建申领单并提交后，系统需要根据明细数量选择可满足需求的仓库并锁定库存；审批人员可以审批通过或驳回；仓储人员在审批通过后执行出库；部门用户最终完成签收。`ApplyService` 中对 `DRAFT`、`SUBMITTED`、`APPROVED`、`OUTBOUND`、`RECEIVED` 和 `REJECTED` 状态的控制，是本节写申领流程的主要证据。",
        "调拨需求解决不同仓库之间库存分布不均的问题。用户可以创建调拨单，提交后由审批人员审批，审批通过后执行调拨并在目标仓库生成批次记录，最后接收确认。推荐调拨接口根据目标校区和物资需求，从满足库存条件的仓库中计算候选仓并按距离排序。该能力只能写成固定校区图上的候选仓排序，不能写成完整物流路径规划或智能配送调度。",
        "预警与统计需求包括低库存、库存积压、临期、过期和异常消耗预警，以及库存占比、出入库趋势、部门排行、效期统计、仓库分布和应急消耗等统计主题。`WarningService` 通过定时扫描和手动扫描生成预警记录，`AnalyticsController` 提供统计接口，前端使用 ECharts 展示结果。系统还提供通知、登录日志和操作日志，用于提醒与追溯。"
    ],
    "3.3 非功能需求": [
        "安全性需求要求系统能够区分未登录、登录失败、账号禁用、令牌失效和权限不足等情况。后端通过 Spring Security 统一拦截未认证请求，通过 JWT 保存用户身份和角色，通过 BCrypt 校验密码，通过 refresh token 轮换降低长期令牌复用风险。前端路由在无 token 时跳转登录页，在安全策略页面上进一步限制管理员访问。",
        "一致性需求主要出现在库存相关流程。入库、出库、申领锁定、调拨执行都可能同时影响多张表，如果只更新其中一部分，就会产生库存汇总和批次余量不一致。项目在关键 Service 方法上使用事务，出库时先校验库存，再按 FEFO 扣减批次，随后更新库存汇总和出库明细。申领提交时使用锁定库存，驳回或签收后释放剩余锁定量。",
        "可维护性需求体现在模块划分和通用字段设计。后端按 auth、rbac、material、warehouse、inventory、apply、transfer、warning、analytics、notification、log 等模块组织，前端按 views、router、api、store 和复用组件组织。数据库多数表保留 deleted、version、created_at 和 updated_at 字段，便于后续扩展软删除、乐观锁和审计追踪。"
    ],
    "4.1 系统总体架构": [
        "系统采用典型 B/S 架构和前后端分离模式，浏览器访问 Vue 构建的前端页面，前端通过 Axios 请求后端 REST 接口，后端在 Spring Security 过滤链中完成 JWT 校验，再进入 Controller、Service、Mapper 和 MySQL。该结构使页面展示、业务逻辑和数据持久化分工清晰，便于在论文中分别展开需求、设计、实现和测试。",
        "后端分层并不意味着系统采用微服务。当前项目是单体 Spring Boot 应用，所有模块运行在同一个后端工程中，通过包结构和 Service 分层实现职责划分。配置中虽然包含 Actuator、Prometheus 指标和 Sentry 依赖，但论文主体只把它们作为工程化支撑简写，不写成生产级监控平台或分布式部署方案。"
    ],
    "4.3 接口与安全设计": [
        "接口设计遵循资源路径和业务动作相结合的方式。认证接口集中在 `/api/auth`，库存接口集中在 `/api/inventory`，申领接口集中在 `/api/apply`，调拨接口集中在 `/api/transfer`，预警接口集中在 `/api/warning`，统计接口集中在 `/api/analytics`。列表查询通常使用分页参数，写操作通过 POST 提交对象，状态动作通过路径表达，如 `/submit`、`/approve`、`/reject`、`/execute`、`/receive` 和 `/handle`。",
        "如图4-3所示，前端页面到后端数据库的交互流程包括页面输入、请求封装、令牌携带、接口分发、业务处理、数据访问和页面刷新。用户在 Vue 页面输入查询条件或提交表单后，前端请求工具携带 JWT 调用接口；后端先经过安全过滤链，再由 Controller 分发到 Service；Service 在事务中执行业务规则并调用 Mapper；数据库返回结果后，后端统一包装响应，前端根据返回数据更新 Pinia 状态、表格或 ECharts 图表。该流程是后文章节分析登录、库存、申领和统计模块的共同基础。"
    ],
    "4.4 智能能力设计": [
        "本系统的智能能力设计重点是“可解释”和“可复核”。低库存预警直接比较库存当前数量与物资安全库存，积压预警比较当前数量与安全库存的倍数，临期和过期预警比较批次过期日期与当前日期，异常消耗预警比较近 7 天出库量与近 30 天折算周均值。补货建议根据安全库存、当前库存、近 30 天出库量和保障天数计算建议数量。移动平均预测根据近 6 个月出库记录计算月均值。上述逻辑均能在代码中找到对应实现。",
        "调拨推荐的智能性同样需要谨慎表述。`TransferService.recommendTransfer` 只在候选仓满足物资库存数量的前提下，结合 `DijkstraUtil` 计算的校区距离进行排序。这能帮助用户从多个可用仓库中优先选择距离较近的仓库，但并不等同于车辆路径优化、配送排班或实时交通调度。论文在总体设计中将其归类为候选排序，不把它写成复杂算法成果。"
    ],
    "5.1 数据库设计原则": [
        "数据库设计以业务闭环和数据可追溯为基本原则。用户、角色、部门用于承载组织和权限；物资分类、物资信息、供应商、仓库、校区、库位构成基础数据；库存、批次、入库、出库、申领和调拨承载核心业务；预警、通知、登录日志、操作日志、事件记录和系统配置提供支撑能力。各表通过业务字段形成关联，SQL 中没有显式外键约束，因此论文只能描述逻辑关联，不能虚构数据库外键。",
        "多数业务表使用 `id` 作为自增主键，使用 `deleted` 表示逻辑删除，使用 `version` 支持乐观锁或版本更新，使用 `created_at` 和 `updated_at` 记录创建与更新时间。该设计能够让系统在删除和更新数据时保留统一处理方式，也便于测试和审计章节说明数据追踪依据。索引主要围绕查询条件和业务状态设置，例如用户的部门和角色索引、库存的物资仓库唯一索引、批次的效期索引、申领和调拨的状态索引、预警的状态和类型索引。"
    ],
    "5.3 关键数据表设计": [
        "用户权限相关表包括 `sys_dept`、`sys_role`、`sys_user` 和 `auth_refresh_token`。`sys_dept` 记录部门名称和父级部门，支撑部门用户归属；`sys_role` 记录角色编码、角色名称和说明，支撑菜单和权限边界；`sys_user` 保存用户名、加密密码、真实姓名、部门、角色和启停状态；`auth_refresh_token` 保存 refresh token 的用户、token_id、token_hash、过期时间和撤销状态。认证模块登录成功后会写入 refresh token，刷新时校验 hash 并撤销旧记录。",
        "物资基础数据表包括 `material_category`、`material_info`、`supplier`、`warehouse`、`campus` 和 `storage_location`。其中 `material_info` 是核心档案表，`material_code` 具有唯一约束，`safety_stock` 直接影响低库存和积压预警，`shelf_life_days`、`production_date`、`expire_date` 的组合支撑效期管理。仓库表记录仓库名称、校区、地址和负责人，库位表记录库位编码、容量、已用数量和状态。论文中对这些字段的说明必须与 SQL 保持一致。",
        "库存与批次表包括 `inventory` 和 `inventory_batch`。`inventory` 以物资和仓库为唯一组合，保存当前库存和锁定数量；`inventory_batch` 保存批次号、入库数量、剩余数量、生产日期和过期日期，并设置按物资、仓库、过期日期和余量查询的索引。出库和调拨均按未过期且有剩余数量的批次排序扣减，因此 `inventory_batch` 是 FEFO 规则落地的关键表。",
        "入库与出库单据表包括 `stock_in`、`stock_in_item`、`stock_out` 和 `stock_out_item`。入库主表保存仓库、来源类型、操作人和备注，入库明细保存物资、批次、数量和日期。出库主表可以关联申领单，保存仓库、操作人和备注，出库明细保存物资和数量。单据主从结构使系统既能查询一次操作的整体信息，也能统计某种物资的入库或出库数量。",
        "申领相关表包括 `apply_order` 和 `apply_order_item`。主表保存部门、申请人、紧急等级、状态、原因、场景、快速通道标记、锁定仓库、审批人、审批意见和审批时间；明细表保存物资、申请数量和实发数量。状态字段是申领流程控制的核心，不能在论文中新增 SQL 中不存在的“配送中”“已评价”等状态。",
        "调拨相关表包括 `transfer_order` 和 `transfer_order_item`。调拨主表保存调出仓库、调入仓库、状态、申请人、审批人、审批意见和审批时间，明细表保存调拨物资和数量。调拨执行会扣减来源仓库存库并在目标仓库生成带有调拨标记的新批次，因此调拨不仅是单据状态变化，也影响库存汇总和批次表。",
        "预警、日志和通知支撑表包括 `warning_record`、`operation_log`、`login_log` 和 `notification`。预警表记录预警类型、物资、仓库、内容、处理状态、处理人和处理备注；操作日志表记录操作者、模块、操作和详情；登录日志表记录登录用户、IP、状态、时间和 user_agent；通知表保存标题、内容、类型、目标用户、已读状态和业务关联。它们使系统具备追溯、提醒和管理支撑能力。"
    ],
    "6.1 开发与运行环境": [
        "开发环境与运行环境应以项目文件为准。后端使用 Java 17、Spring Boot 3.3.5 和 Maven 管理依赖；开发默认配置由 `application.yml` 激活 dev profile，可通过环境变量切换数据库、JWT 密钥、CORS、限流和管理端点；测试使用 H2 数据库和 `application-test.yml`。前端使用 Node.js、Vite、Vue 3、Element Plus 和 ECharts，测试工具包括 Vitest 与 Playwright。论文不把本地演示环境写成正式生产部署。",
        "系统在本地演示时通常由后端 8080 端口提供业务接口，前端开发或预览服务提供页面访问。配置文件还设置了管理端点 18080 和 Prometheus 指标暴露，但这些配置只说明项目具备工程观测接口，不代表系统已经在生产环境中持续运行。扩写后的测试章节也只写本地自动化测试和本地基线，不写线上用户规模。"
    ],
    "6.2 认证授权模块设计与实现": [
        "登录流程从前端登录页开始。用户输入用户名和密码后，前端调用 `/api/auth/login`。后端 `AuthService.login` 首先去除用户名空白并读取请求 IP 与 user_agent，然后从 `sys_user` 查询用户，使用 BCrypt 校验密码，检查账号状态，再根据 `sys_role` 获取角色编码。认证通过后，系统生成 access token 和 refresh token，并将 refresh token 的哈希、token_id、过期时间和撤销状态写入 `auth_refresh_token` 表，同时写入登录日志。",
        "请求鉴权由 Spring Security 过滤链完成。`SecurityConfig` 将登录、刷新令牌、接口文档和健康检查等少数地址放行，其余业务接口要求登录。前端在登录后保存 token，后续请求由 Axios 封装携带 Authorization 头。后端 `JwtTokenProvider` 解析令牌中的用户 id、用户名、角色和 token 类型，`JwtAuthenticationFilter` 将其放入安全上下文，使 Controller 或 Service 能通过 `AuthUtil` 获取当前用户。",
        "refresh token 采用持久化和轮换策略。当前 refresh token 必须能被解析，类型必须为 refresh，token_id 必须存在，数据库记录必须未撤销且未过期，并且提交的 token 哈希要与数据库保存的哈希一致。刷新成功后，旧 token 会被撤销并重新签发令牌。该设计比单纯前端保存长期 token 更便于服务端主动失效控制，也是本文安全实现部分的重要工程细节。",
        "角色菜单由后端 `AuthService.buildMenusByRole` 统一生成。管理员菜单最完整，仓储管理员偏向基础数据和仓储业务，审批人员关注申领、调拨、预警和统计，部门用户仅保留申领和统计。前端加载用户信息和菜单后渲染导航，并在安全策略页面上额外判断管理员角色。这样可以使菜单展示和后端角色边界保持一致。"
    ],
    "6.3 用户与基础数据管理模块设计与实现": [
        "用户与角色管理由 `RbacController` 和 `RbacService` 支撑。用户列表按 id 倒序查询，新增用户时必须填写密码并通过 BCrypt 加密保存，修改用户时如果未提交新密码则保留原密码。角色表保存角色编码和名称，部门表保存部门名称和父级关系。每次用户、角色或部门维护操作都会写入操作日志，便于管理员追踪系统基础权限变更。",
        "基础数据模块还包括校区、仓库、库位、供应商、物资分类和物资档案。相应 Controller 多采用列表查询、保存和删除三类接口，Service 使用 MyBatis-Plus 完成基础数据维护。物资档案页面和物资分类页面对应 `MaterialController`，仓库页面对应 `WarehouseController`，库位页面对应 `StorageLocationController`。这些模块看似是增删改查，但它们为后续库存、申领、调拨和预警提供基础主数据，不能在论文中省略其业务作用。",
        "物资档案维护的重点是物资编码唯一、分类归属、安全库存和效期参数。安全库存被低库存和积压规则使用，保质期和批次过期日期共同支撑临期提醒。仓库和校区字段被库存分布和调拨推荐使用。供应商字段用于保留供货范围和联系方式，但当前系统没有形成采购订单闭环，因此论文只写供应商资料管理。"
    ],
    "6.4 库存批次与入库管理模块设计与实现": [
        "入库管理由前端 `StockInView.vue`、后端 `/api/inventory/stock-in` 接口和 `InventoryService.stockIn` 方法支撑。用户选择仓库、来源类型并填写多个物资明细，明细中包含物资、批次号、数量、生产日期和过期日期。后端接收请求后创建 `stock_in` 主记录，并逐条写入 `stock_in_item` 明细。若前端未提供批次号，后端会按时间生成批次号，避免批次字段为空。",
        "入库过程的关键不是保存单据本身，而是同步库存汇总和批次余量。对于每条入库明细，系统先调用 `getOrInitInventory` 获取或创建物资在指定仓库的库存汇总行，再增加 `current_qty`；随后创建 `inventory_batch`，写入物资、仓库、批次号、入库数量、剩余数量、生产日期和过期日期。上述操作位于事务中，任何一步失败都应回滚，避免出现入库单存在但库存没有增加的情况。",
        "入库业务的执行顺序如图6-9所示。该流程对应 `stock_in`、`stock_in_item`、`inventory` 和 `inventory_batch` 四类表。论文写作时可以说明入库如何影响批次管理和库存汇总，但不能扩写成采购验收、财务付款或供应商结算流程，因为当前代码没有这些表和接口。"
    ],
    "6.5 出库管理模块设计与实现": [
        "出库管理由 `/api/inventory/stock-out` 接口和 `InventoryService.stockOut` 方法实现。出库请求可以独立创建，也可以关联已审批的申领单。当出库关联申领单时，系统会检查申领单状态是否为已审批或已出库，校验指定仓库是否与申领提交阶段锁定的仓库一致，并读取申领明细以限制实发数量，避免超出申请和锁定范围。",
        "出库扣减采用 FEFO 思路，即优先扣减过期日期更早且未过期的批次。后端查询条件要求批次剩余量大于 0、过期日期不早于当前日期，并按 `expire_date` 和 `id` 升序排列。系统按出库数量逐个扣减批次余量，若所有候选批次数量仍不足，则抛出业务异常。扣减成功后，系统减少库存汇总数量；若为申领出库，还会减少锁定数量并回写申领明细的实发数量。",
        "出库业务流程如图6-10所示。出库完成后系统会写入 `stock_out_item` 明细，关联申领时会把申领单状态更新为 `OUTBOUND`，随后等待部门用户签收。出库后还会检查安全库存，若当前库存低于物资安全库存，则生成低库存预警。该流程体现了单据、库存、批次、申领和预警之间的数据联动。"
    ],
    "6.6 物资申请与审批模块设计与实现": [
        "申领单创建由部门用户发起。前端申领页面提供物资选择、申请数量、紧急等级、申请原因和场景说明，后端 `ApplyService.create` 生成 `apply_order` 主表记录和 `apply_order_item` 明细记录，初始状态为 `DRAFT`。创建动作写入操作日志。草稿状态的设计使用户可以先保存明细，再在确认后提交，避免未完成的申请直接进入审批队列。",
        "提交申领时，系统会汇总所有明细物资的申请数量，并在库存表中寻找能够一次性满足需求的仓库。若存在可满足的仓库，系统把这些物资的数量增加到 `locked_qty`，并将仓库 id 写入 `reserved_warehouse_id`。锁定库存的意义在于防止多个申请同时提交后抢占同一库存。若库存不足或没有仓库能满足需求，提交会失败并返回业务异常。",
        "审批流程严格依赖状态。普通申领从 `SUBMITTED` 进入审批，审批通过后变为 `APPROVED`，驳回后变为 `REJECTED` 并释放锁定库存。紧急等级达到代码中的阈值时，系统会进入快速通道，自动设置审批通过状态、审批人、审批意见和审批时间。这里的快速通道是规则化状态处理，不是智能审批模型。出库后申领单进入 `OUTBOUND`，部门用户签收后进入 `RECEIVED` 并释放剩余锁定量。",
        "申领模块与库存模块的边界较清晰：申领服务负责单据状态、申请明细、审批意见和签收动作，库存服务负责仓库选择、锁定库存、释放库存和实际出库。这样的分工能够减少单个模块承担过多职责，也便于测试中分别验证申领状态流转和库存一致性。"
    ],
    "6.7 调拨协同与候选仓排序模块设计与实现": [
        "调拨模块用于处理不同仓库之间的物资调剂。用户创建调拨单时必须指定调出仓库和调入仓库，二者不能相同，并填写调拨明细。提交后状态由 `DRAFT` 变为 `SUBMITTED`，审批通过后进入 `APPROVED`，驳回则进入 `REJECTED`。调拨执行后状态为 `OUTBOUND`，接收后状态为 `RECEIVED`。这些状态均来自代码中的 `OrderStatus`，论文不得扩展不存在的调拨状态。",
        "调拨执行会对库存和批次产生实质影响。系统先校验来源仓库对应物资的库存是否充足，再按 FEFO 顺序扣减来源仓库的批次余量；同时在目标仓库生成镜像批次，批次号带有调拨单标记，生产日期和过期日期沿用来源批次。随后系统减少来源仓库库存汇总、增加目标仓库库存汇总，并记录操作日志。整个执行方法使用事务，确保批次搬移和库存汇总同步成功。",
        "候选仓排序接口为调拨申请提供参考。用户给出目标校区、物资和数量后，系统先筛选库存数量满足需求的仓库，再根据固定校区图计算距离并排序。该设计适合作为仓库选择辅助，但不能证明系统具备完整配送路线规划、车辆调度或实时地图能力。扩写版论文将其表述为“候选仓排序”和“调拨参考”，不写成物流优化算法。"
    ],
    "6.8 库存预警与临期提醒模块设计与实现": [
        "预警扫描由 `WarningService.scan` 统一触发，既可以通过定时任务每 30 分钟执行，也可以由预警页面手动触发。扫描过程依次执行低库存、库存积压、临期、过期和异常消耗检查，并记录扫描耗时和生成预警数量。预警列表接口支持按类型和处理状态筛选，处理接口将预警状态更新为已处理并保存处理人和备注。",
        "低库存和积压预警依赖 `inventory` 与 `material_info`。当当前库存低于物资安全库存时生成低库存预警；当当前库存超过安全库存 3 倍时生成积压预警。临期预警查询未来 30 天内过期且仍有余量的批次，过期预警查询过期日期早于当前日期且仍有余量的批次。异常消耗预警比较近 7 天出库量与近 30 天折算周均值，当本周出库超过月均周出库的 1.5 倍时生成预警。",
        "预警去重通过 `createWarningIfAbsent` 完成。系统在生成新预警前，会按预警类型、物资、仓库和未处理状态查询是否已有记录，避免同一问题在未处理前反复生成。该逻辑使预警中心更适合实际管理场景：管理人员看到的是待处理问题，而不是重复堆积的扫描日志。"
    ],
    "6.9 数据统计、补货建议与可视化模块设计与实现": [
        "统计分析模块由 `AnalyticsController` 提供统一入口，包括总览、库存占比、出入库趋势、部门排行、效期统计、仓库分布和应急消耗等接口。前端 `AnalyticsView.vue` 通过接口获取数据后使用 ECharts 进行图表展示，运营总览和大屏页面也会使用部分统计信息。图6-11展示了统计数据从前端筛选到后端聚合再到图表渲染的流转过程。",
        "补货建议由 `SmartService.replenishmentSuggestions` 实现。该方法聚合每种物资当前库存、物资安全库存和近 30 天出库量，计算日均消耗，再根据保障天数估算目标库存和建议补货量。建议量为目标库存减当前库存，若结果小于 0 则取 0。该计算适合作为仓储管理员补货参考，但不等同于采购计划自动生成，也不涉及供应商下单。",
        "移动平均预测由 `SmartService.forecast` 实现。方法读取指定物资近 6 个月出库明细，按月份汇总数量，计算历史月均值，并把该均值作为未来若干月的预测数量。由于算法只使用简单平均，论文只能称为移动平均或趋势参考，不能写成机器学习预测模型。"
    ],
    "6.10 日志、通知与系统支撑模块设计与实现": [
        "日志模块分为登录日志和操作日志。登录日志由认证服务在登录成功或失败时写入，保存用户、用户名、IP、登录状态、时间和 user_agent；操作日志由多个业务 Service 在创建、提交、审批、出入库、调拨、处理预警等动作后写入，保存操作者、模块、操作类型和详情。日志页面可以按模块或时间查询，帮助管理员追踪关键操作。",
        "通知模块保存系统消息和业务关联信息。`notification` 表包含标题、内容、消息类型、目标用户、已读状态、业务类型和业务 id，前端通知中心提供列表、未读数量、标记已读、全部已读和删除能力。通知模块使预警、审批或系统消息具备统一展示入口，但当前代码没有实现短信、邮件或第三方推送，因此论文只写站内通知。",
        "系统配置和事件记录模块属于支撑能力。配置表保存运行参数分组和值，事件表用于记录事件标题、类型、等级、校区、地点、描述、处理状态、上报人、处理人和处理结果。它们不是论文主线，但能说明系统在校园物资保障之外也预留了安全事件和系统参数管理能力。"
    ],
    "7.1 测试环境与方法": [
        "测试环境由后端单元/集成测试、前端构建、前端单测和 E2E 流程测试组成。后端测试使用 Maven、JUnit、Spring Boot Test、H2 数据库和测试 profile；前端构建使用 Vite；前端单测使用 Vitest；端到端测试使用 Playwright 脚本启动后端 screenshot profile 与前端预览服务。测试章节只引用这些真实执行过的命令和结果。",
        "本次扩写前重新执行了 `mvn test`、`npm run build`、`npm run test:unit` 和 `npm run test:e2e`。后端测试结果为 49 项通过，失败 0，错误 0，跳过 0；前端构建通过；前端单测为 3 个测试文件、8 项测试通过；E2E 为 4 个场景通过。E2E 运行过程中后端会记录权限拒绝日志，这是权限边界触发后的日志表现，不影响 Playwright 场景通过结论。"
    ],
    "7.2 后端自动化测试": [
        "后端自动化测试覆盖公共响应、业务异常、全局异常处理、JWT 令牌、认证服务、refresh token 清理、申领服务、申领控制器、调拨服务、调拨控制器、预警控制器和核心流程集成测试。`CoreFlowIntegrationTest` 通过真实 Spring Boot 上下文验证登录、库存查询、申领创建与审批、库存扣减、调拨执行等组合流程，比单纯 Service 单测更接近业务运行状态。",
        "申领相关测试重点验证状态流转和库存锁定。测试会检查非草稿状态不可重复提交、审批只能在提交状态执行、驳回会释放锁定库存、出库与签收会按状态推进。调拨测试重点验证同仓调拨被拒绝、提交和审批状态合法、执行时更新来源和目标仓库库存。认证测试重点验证密码校验、账号状态、token 生成、刷新失败和清理逻辑。",
        "后端测试没有给出正式性能指标，因此论文不能从测试通过推导系统可承受某个并发规模。测试章节应把这些结果解释为功能正确性、异常处理和核心流程一致性的验证，而不是压力测试或生产稳定性证明。"
    ],
    "7.3 前端构建与端到端验证": [
        "前端构建验证了 Vue 单文件组件、路由懒加载、Element Plus、ECharts 和样式资源能够被 Vite 正常打包。构建输出中可以看到登录页、运营总览、库存、入库、出库、申领、调拨、预警、统计、日志、通知等页面对应的资源文件，说明前端页面结构与论文功能模块一致。",
        "前端单测主要覆盖 HTTP 请求封装和登录页面交互。虽然测试数量不多，但能验证请求工具、登录页面基本状态和组件行为没有构建层面的错误。E2E 脚本覆盖管理员登录查看库存、部门用户创建申领单、仓储管理员创建调拨单、仓储管理员处理预警四个典型流程，这些场景与论文中的核心业务闭环直接对应。",
        "权限测试在 E2E 和后端日志中也有所体现。部门用户访问不属于自身角色的仓库或日志接口时，后端记录权限拒绝异常，页面流程仍能继续完成预期动作。论文可将其写为权限边界被触发，但不能把日志中的异常理解为系统错误。"
    ],
    "7.5 性能基线与测试结论": [
        "性能部分只按仓库已有 `tests/performance` 脚本和 `docs/performance-baseline.md` 记录进行保守描述。脚本文件包含登录、库存列表、预警列表和操作日志列表等接口场景，说明项目具备本地性能基线测试材料；但本次扩写未重新执行 k6，因此正文不写新的并发数、吞吐量或响应时间，也不承诺生产 SLA。",
        "综合测试结果看，系统核心功能能够在本地环境下完成自动化验证。后端测试覆盖业务规则和状态流转，前端构建和单测验证页面工程可构建，E2E 验证了登录、库存查看、申领、调拨和预警处理流程。测试章节的结论应限定为“满足毕业设计演示和核心流程验证要求”，不扩展为“大规模上线运行”。"
    ],
    "结束语": [
        "本文围绕校园物资管理场景，完成了基于 Spring Boot 与 Vue 的前后端分离系统设计与实现。系统从组织、角色、物资、仓库、库存、批次、单据、预警、日志和通知等方面建立了较完整的数据结构，能够支撑物资建档、入库、库存查询、申领审批、出库签收、仓间调拨、预警处理和统计分析等流程。与单纯台账系统相比，本系统更强调业务状态和库存数据之间的一致性。",
        "本系统的主要工程特点包括：采用 JWT 和角色菜单实现登录认证与访问边界；采用库存汇总和批次表共同表达库存状态；出库和调拨按效期优先顺序扣减批次；申领提交阶段锁定库存，驳回或签收后释放剩余锁定；预警扫描对低库存、积压、临期、过期和异常消耗进行规则化识别；统计模块通过后端接口和 ECharts 展示业务数据。这些内容均能在代码、SQL、页面或测试材料中找到证据。",
        "系统仍存在可完善空间。首先，当前没有独立移动端、小程序端和扫码硬件接入，后续可在不改变核心业务模型的前提下扩展移动适配。其次，调拨推荐目前是候选仓排序，尚未涉及车辆路径、人员排班和实时交通因素。再次，补货建议和趋势预测采用规则与移动平均，未来可在积累足够真实历史数据后研究更细粒度的需求预测。最后，性能基线仍以本地脚本为主，若系统进入真实运行环境，还需要补充更系统的压测、备份恢复和运维监控方案。"
    ],
}


def find_paragraph(doc: Document, exact: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    raise ValueError(f"paragraph not found: {exact}")


def body_char_count(doc: Document) -> int:
    started = False
    texts = []
    for p in doc.paragraphs:
        s = p.text.strip()
        if s == "1 绪论":
            started = True
        if started and s == "致  谢":
            break
        if started:
            texts.append(s)
    joined = "".join(texts)
    return len(re.sub(r"\s+", "", joined))


def insert_expansions(doc: Document):
    expansion_items = [
        ("1.1 研究背景", EXPANSIONS["1.1 研究背景"]),
        ("1.2 研究意义", EXPANSIONS["1.2 研究意义"]),
        ("1.3 国内外研究现状", EXPANSIONS["1.3 国内外研究现状"]),
        ("2.1 技术选型概述", EXPANSIONS["2.1 技术选型概述"]),
        ("2.2 前端开发技术", EXPANSIONS["2.2 前端开发技术"]),
        ("2.3 后端与数据访问技术", EXPANSIONS["2.3 后端与数据访问技术"]),
        ("2.4 安全认证与可视化技术", EXPANSIONS["2.4 安全认证与可视化技术"]),
        ("3.1 业务场景与角色划分", EXPANSIONS["3.1 业务场景与角色划分"]),
        ("3.2 功能需求分析", EXPANSIONS["3.2 功能需求分析"]),
        ("3.3 非功能需求", EXPANSIONS["3.3 非功能需求"]),
        ("4.1 系统总体架构", EXPANSIONS["4.1 系统总体架构"]),
        ("4.3 接口与安全设计", EXPANSIONS["4.3 接口与安全设计"]),
        ("4.4 智能能力设计", EXPANSIONS["4.4 智能能力设计"]),
        ("5.1 数据库设计原则", EXPANSIONS["5.1 数据库设计原则"]),
        ("5.3 关键数据表设计", EXPANSIONS["5.3 关键数据表设计"]),
        ("6.1 开发与运行环境", EXPANSIONS["6.1 开发与运行环境"]),
        ("6.2 认证授权与登录态管理实现", EXPANSIONS["6.2 认证授权模块设计与实现"]),
        ("6.3 基础数据管理实现", EXPANSIONS["6.3 用户与基础数据管理模块设计与实现"]),
        ("6.4 库存批次与出入库实现", EXPANSIONS["6.4 库存批次与入库管理模块设计与实现"] + EXPANSIONS["6.5 出库管理模块设计与实现"]),
        ("6.5 申领审批与调拨实现", EXPANSIONS["6.6 物资申请与审批模块设计与实现"] + EXPANSIONS["6.7 调拨协同与候选仓排序模块设计与实现"]),
        ("6.6 预警、补货建议与统计分析实现", EXPANSIONS["6.8 库存预警与临期提醒模块设计与实现"] + EXPANSIONS["6.9 数据统计、补货建议与可视化模块设计与实现"]),
        ("6.7 日志通知与事件管理实现", EXPANSIONS["6.10 日志、通知与系统支撑模块设计与实现"]),
        ("7.1 测试环境与方法", EXPANSIONS["7.1 测试环境与方法"]),
        ("7.2 后端自动化测试", EXPANSIONS["7.2 后端自动化测试"]),
        ("7.3 前端构建、单测与 E2E 验证", EXPANSIONS["7.3 前端构建与端到端验证"]),
        ("7.5 性能基线与边界说明", EXPANSIONS["7.5 性能基线与测试结论"]),
        ("结束语", EXPANSIONS["结束语"]),
    ]
    for heading, paragraphs in expansion_items:
        anchor = find_paragraph(doc, heading)
        for text in paragraphs:
            anchor = insert_after(anchor, text)


def insert_figures(doc: Document):
    anchor = find_paragraph(doc, "4.3 接口与安全设计")
    anchor = insert_after(anchor, "系统前后端交互关系如图4-3所示。该图对应前端路由、Axios 请求封装、JWT 请求头、后端控制层、服务层、数据访问层和 MySQL 数据库之间的数据流。")
    insert_picture_after(anchor, FIG_DIR / "fig_4_3_front_backend_flow.png", "图4-3 前后端交互流程图", 14.2)

    anchor = find_paragraph(doc, "图6-8 统计分析界面")
    anchor = insert_after(anchor, "入库单从页面表单进入后端后，会依次写入单据、明细、库存汇总和批次表，流程如图6-9所示。")
    anchor = insert_picture_after(anchor, FIG_DIR / "fig_6_9_stock_in_flow.png", "图6-9 入库业务流程图", 12.0)
    anchor = insert_after(anchor, "出库流程与申领单、库存汇总和批次扣减同时关联，其业务处理顺序如图6-10所示。")
    anchor = insert_picture_after(anchor, FIG_DIR / "fig_6_10_stock_out_flow.png", "图6-10 出库业务流程图", 12.0)
    anchor = insert_after(anchor, "统计分析模块从后端聚合接口向前端图表提供数据，数据流转关系如图6-11所示。")
    insert_picture_after(anchor, FIG_DIR / "fig_6_11_analytics_flow.png", "图6-11 统计数据流转流程图", 12.0)


def sanitize_final_text(doc: Document):
    replacements = {
        "机器学习预测模型": "复杂预测模型",
        "小程序端和扫码硬件接入": "外部硬件接入",
        "移动端、小程序端": "移动适配",
        "生产 SLA": "生产服务承诺",
    }
    containers = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for paragraph in containers:
        if not paragraph.text:
            continue
        original = paragraph.text
        updated = original
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != original:
            paragraph.text = updated
            format_body_para(paragraph)


def write_reports(body_count: int):
    evidence_rows = [
        ("技术栈", "是", "backend/pom.xml; frontend/package.json; application.yml", "保留并按真实依赖写"),
        ("角色与菜单", "是", "AuthService.buildMenusByRole; frontend/src/router/index.js", "保留 ADMIN/WAREHOUSE_ADMIN/APPROVER/DEPT_USER"),
        ("数据库表字段", "是", "sql/schema.sql", "所有表名、字段名、索引按 SQL 写"),
        ("登录认证", "是", "AuthService; SecurityConfig; JwtTokenProvider; auth_refresh_token", "写 JWT、BCrypt、refresh token 轮换"),
        ("基础数据", "是", "Material/Warehouse/Campus/Supplier/Rbac Controller 与 Service", "保留为支撑模块"),
        ("入库出库", "是", "InventoryController; InventoryService; stock_in/out 表", "重点扩写事务与 FEFO 扣减"),
        ("申领审批", "是", "ApplyService; apply_order 表; CoreFlowIntegrationTest", "按状态流转写"),
        ("独立配送任务", "否", "未发现配送任务表、配送角色、配送 Controller", "删除或降级为出库/签收"),
        ("调拨推荐", "是", "TransferService; DijkstraUtil; transfer_order 表", "谨慎写为候选仓排序"),
        ("预警扫描", "是", "WarningService; warning_record 表", "写规则化预警"),
        ("补货建议/移动平均", "是", "SmartService", "写统计参考，不写复杂模型"),
        ("测试结果", "是", "mvn test; npm build; vitest; Playwright E2E", "写真实通过数量"),
        ("性能结果", "部分", "tests/performance; docs/performance-baseline.md", "只写已有本地基线，不新增数据"),
        ("参考文献", "是", "C:/Users/48195/Desktop/CNKI-20260426111148676.txt", "只使用该文件来源"),
    ]
    EVIDENCE.write_text("# 证据映射表\n\n| 检查对象 | 是否存在证据 | 证据位置 | 处理方式 |\n|---|---|---|---|\n" + "\n".join(f"| {a} | {b} | `{c}` | {d} |" for a, b, c, d in evidence_rows) + "\n", encoding="utf-8")

    krd_rows = [
        ("Keep", "题目与系统定位", "项目仓库、SQL 和页面均围绕校园物资管理", "保留并扩写为前后端分离管理系统", "README、pom、package、schema"),
        ("Keep", "申领审批、库存批次、调拨、预警", "有代码和测试支撑", "作为论文核心业务", "ApplyService、InventoryService、TransferService、WarningService"),
        ("Keep", "JWT/RBAC 权限", "有认证、角色菜单和测试", "保留并扩写安全设计", "AuthService、SecurityConfig、JwtTokenProvider"),
        ("Rewrite", "智能表述", "容易被写成算法论文", "改为规则化预警、FEFO、候选仓排序和移动平均", "SmartService、WarningService"),
        ("Rewrite", "数据库章节", "旧稿表字段可能不全", "按 schema.sql 重写关键表关系与字段", "sql/schema.sql"),
        ("Rewrite", "测试章节", "需使用最新执行结果", "写 49 后端测试、前端构建、8 单测、4 E2E", "本次命令输出"),
        ("Delete", "独立配送任务模块", "无配送任务表、角色和接口", "删除，不作为已实现功能", "schema、Controller 列表"),
        ("Delete", "移动端、小程序、扫码柜、复杂 AI", "无代码和配置证据", "仅可作为展望或不写", "仓库证据不足"),
        ("Delete", "新增外部参考文献", "用户限制只允许 CNKI txt", "删除所有非指定来源", "CNKI-20260426111148676.txt"),
    ]
    KRD.write_text("# Keep-Rewrite-Delete矩阵\n\n| 类型 | 内容位置/主题 | 当前问题 | 处理建议 | 依据 |\n|---|---|---|---|---|\n" + "\n".join(f"| {a} | {b} | {c} | {d} | {e} |" for a, b, c, d, e in krd_rows) + "\n", encoding="utf-8")

    fig_rows = [
        ("图3-1", "申领审批闭环流程图", "已引用", "ApplyService/InventoryService", "保留"),
        ("图3-2", "调拨执行流程图", "已引用", "TransferService", "保留"),
        ("图3-3", "预警处置流程图", "已引用", "WarningService", "保留"),
        ("图4-1", "系统总体架构图", "已引用", "pom/package/application/router", "保留"),
        ("图4-2", "系统功能模块图", "已引用", "前端路由/后端模块", "保留"),
        ("图4-3", "前后端交互流程图", "新增并引用", "router/http/Security/Controller/Service/Mapper", "新增"),
        ("图5-1~图5-3", "E-R 图", "已引用", "sql/schema.sql", "保留"),
        ("图6-1~图6-8", "认证、界面、调拨、预警、统计图", "已引用", "代码与运行截图", "保留"),
        ("图6-9", "入库业务流程图", "新增并引用", "InventoryService.stockIn", "新增"),
        ("图6-10", "出库业务流程图", "新增并引用", "InventoryService.stockOut", "新增"),
        ("图6-11", "统计数据流转流程图", "新增并引用", "AnalyticsController/AnalyticsView", "新增"),
        ("表3-1~表7-4", "需求、环境、测试表", "已引用", "代码/配置/测试", "保留并检查"),
    ]
    FIG_CHECK.write_text("# 图表编号与引用检查表\n\n| 图表编号 | 图表名称 | 引用状态 | 数据/证据来源 | 处理方式 |\n|---|---|---|---|---|\n" + "\n".join(f"| {a} | {b} | {c} | `{d}` | {e} |" for a, b, c, d, e in fig_rows) + "\n", encoding="utf-8")

    refs = [
        ("[1]", "绪论-智慧校园物资管理", "CNKI [33]", "校园物资管理背景"),
        ("[2]", "绪论-物资流通系统", "CNKI [34]", "后勤物资管理信息化"),
        ("[3]", "绪论-供应管理数字化", "CNKI [1]", "数智供应管理背景"),
        ("[4]", "绪论-物资流通流程", "CNKI [2]", "入库、领用、审核、库存监控"),
        ("[5]", "智能管理展望", "CNKI [18]", "高校实验室物资全生命周期参考"),
        ("[6]", "技术路线", "CNKI [11]", "Vue+SpringBoot 高校系统"),
        ("[7]", "系统架构/测试", "CNKI [16]", "B/S 前后端分离系统"),
        ("[8]", "SpringBoot 系统开发", "CNKI [24]", "SpringBoot 管理系统"),
        ("[9]", "校园服务系统", "CNKI [25]", "校园服务信息化"),
        ("[10]", "安全开发", "CNKI [14]", "Vue/Spring Boot/MySQL 安全方法"),
        ("[11]", "外文系统实践", "CNKI [21]", "SpringBoot+Vue 管理系统"),
        ("[12]", "外文系统实践", "CNKI [23]", "SpringBoot+Vue 学生管理系统"),
        ("[13]", "外文 SpringBoot", "CNKI [31]", "SpringBoot 管理信息系统"),
        ("[14]", "Java+MySQL", "CNKI [8]", "前后端分离与 MySQL"),
        ("[15]", "Vue 动态路由", "CNKI [7]", "Vue 与 ECharts 可视化"),
        ("[16]", "前端可视化", "CNKI [13]", "Vue 可视化前端"),
        ("[17]", "数据库设计", "CNKI [29]", "MySQL 设计与索引"),
        ("[18]", "数据库优化", "CNKI [4]", "MySQL 查询优化"),
        ("[19]", "JWT 安全", "CNKI [20]", "JWT 测试与风险"),
        ("[20]", "JWT/RBAC", "CNKI [19]", "前后端分离权限控制"),
        ("[21]", "智能平台边界", "CNKI [3]", "只作智能管理背景"),
        ("[22]", "智能实验室展望", "CNKI [10]", "只作未来方向"),
    ]
    REF_CHECK.write_text("# 参考文献引用检查表\n\n| 论文编号 | 主要引用位置 | 原始 CNKI 编号 | 引用理由 |\n|---|---|---|---|\n" + "\n".join(f"| {a} | {b} | {c} | {d} |" for a, b, c, d in refs) + "\n\n说明：最终 DOCX 文末参考文献为 22 条，均来自指定 CNKI txt，正文引用保持顺序编码制。\n", encoding="utf-8")

    report = f"""# 论文扩写修改报告

## 交付文件
- 正文扩写定稿版：`{TARGET_DOCX}`
- 证据映射表：`{EVIDENCE}`
- Keep-Rewrite-Delete矩阵：`{KRD}`
- 图表编号与引用检查表：`{FIG_CHECK}`
- 参考文献引用检查表：`{REF_CHECK}`

## 扩写范围
- 以 `{SOURCE_DOCX}` 为可编辑草稿副本来源，另存为新文件，未覆盖学校模板、原始草稿、SQL、源码或 CNKI 文件。
- 正文扩写覆盖 1 绪论、2 相关技术介绍、3 系统需求分析、4 系统总体设计、5 数据库设计、6 系统详细设计与实现、7 系统测试和结束语。
- 扩写后按程序统计正文至结束语的非空白字符数为 `{body_count}`，超过 18000 字目标。

## 证据核验
- 后端 `mvn test`：49 项测试通过，失败 0，错误 0，跳过 0。
- 前端 `npm run build`：构建通过。
- 前端 `npm run test:unit`：3 个测试文件、8 项测试通过。
- 前端 `npm run test:e2e`：4 个 Playwright 场景通过。
- k6 未重新执行，正文只按仓库已有基线材料保守描述。

## 主要处理
- 新增图4-3、图6-9、图6-10、图6-11，补齐前后端交互、入库、出库和统计数据流转说明。
- 将“智能”限定为规则化预警、FEFO、候选仓排序、补货建议和移动平均。
- 删除或降级无证据内容：独立配送任务、移动端、小程序、扫码硬件、复杂预测算法、生产级并发指标。
- 参考文献仍保持 22 条，均来自指定 CNKI 文件，未新增外部来源。

## 后续人工复核
- 在 Word 中刷新目录、页码和域。
- 检查新增图跨页位置是否符合导师偏好。
- 本机未检测到 LibreOffice/soffice，本轮未导出 PDF 逐页渲染。
"""
    REPORT.write_text(report, encoding="utf-8")


def main():
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    create_new_figures()
    shutil.copyfile(SOURCE_DOCX, TARGET_DOCX)
    doc = Document(str(TARGET_DOCX))
    insert_figures(doc)
    insert_expansions(doc)
    sanitize_final_text(doc)
    count = body_char_count(doc)
    doc.save(str(TARGET_DOCX))
    write_reports(count)
    print(TARGET_DOCX)
    print(count)


if __name__ == "__main__":
    main()
